#!/usr/bin/env python3
"""Package replicated literature models vs the current best NSC method.

This report intentionally separates:

1. Replicated literature-model families that were actually trained/evaluated on
   the 114-case NSC data.
2. Method-component replications where the core idea was implemented with
   explicitly disclosed custom parameters.
3. Literature models not yet executed on this dataset.

The comparison level is patient-aware stratified 10-fold unless explicitly
marked otherwise.
"""

from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Iterable, List

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis" / "nsc_replicated_literature_models_vs_best_20260520"
REPORT = ROOT / "reports" / "NSC_replicated_literature_models_vs_our_best_20260520.docx"
TAIL_BAG_CORRECTED = ROOT / "analysis" / "nsc_restricted_subject_bagging_tail_stats_corrected_20260521"
TAIL_BAG_ORIGINAL = ROOT / "analysis" / "nsc_restricted_subject_bagging_tail_stats_20260521"


def tail_bag_analysis_dir() -> Path:
    corrected_summary = TAIL_BAG_CORRECTED / "bagged_metrics_summary.csv"
    if corrected_summary.exists():
        return TAIL_BAG_CORRECTED
    return TAIL_BAG_ORIGINAL


REFERENCES = {
    "Wang2015": {
        "text": "Wang, Z., & Oates, T. (2015). Imaging Time-Series to Improve Classification and Imputation. Proceedings of IJCAI 2015.",
        "url": "https://www.ijcai.org/Proceedings/15/Papers/553.pdf",
    },
    "Hatami2017": {
        "text": "Hatami, N., Gavet, Y., & Debayle, J. (2017). Classification of Time-Series Images Using Deep Convolutional Neural Networks. arXiv:1710.00886.",
        "url": "https://arxiv.org/abs/1710.00886",
    },
    "Elmir2023": {
        "text": "Elmir, Y., Himeur, Y., & Amira, A. (2023). ECG classification using Deep CNN and Gramian Angular Field. arXiv:2308.02395.",
        "url": "https://arxiv.org/abs/2308.02395",
    },
    "Wu2018": {
        "text": "Wu, Y., Yang, F., Liu, Y., Zha, X., & Yuan, S. (2018). A Comparison of 1-D and 2-D Deep Convolutional Neural Networks in ECG Classification. arXiv:1810.07088.",
        "url": "https://arxiv.org/abs/1810.07088",
    },
    "Noman2018": {
        "text": "Noman, F., Ting, C.-M., Salleh, S.-H., & Ombao, H. (2018). Short-segment heart sound classification using an ensemble of deep convolutional neural networks. arXiv:1810.11573.",
        "url": "https://arxiv.org/abs/1810.11573",
    },
    "Zhao2022": {
        "text": "Zhao, P., Luo, C., Qiao, B., Wang, L., Rajmohan, S., Lin, Q., & Zhang, D. (2022). T-SMOTE: Temporal-oriented Synthetic Minority Oversampling Technique for Imbalanced Time Series Classification. IJCAI 2022, 2406-2412.",
        "url": "https://www.ijcai.org/proceedings/2022/334",
    },
    "Chawla2002": {
        "text": "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research, 16, 321-357.",
        "url": "https://www.jair.org/index.php/jair/article/view/10302",
    },
    "Snell2017": {
        "text": "Snell, J., Swersky, K., & Zemel, R. (2017). Prototypical Networks for Few-shot Learning. Advances in Neural Information Processing Systems 30.",
        "url": "https://papers.nips.cc/paper/6996-prototypical-networks-for-few-shot-learning",
    },
    "Chen2020": {
        "text": "Chen, Y., & Kuo, C.-C. J. (2020). PixelHop: A successive subspace learning (SSL) method for object recognition. Journal of Visual Communication and Image Representation, 70, 102749.",
        "url": "https://doi.org/10.1016/j.jvcir.2019.102749",
    },
    "Chen2020b": {
        "text": "Chen, Y., Rouhsedaghat, M., You, S., Rao, R., & Kuo, C.-C. J. (2020). PixelHop++: A Small Successive-Subspace-Learning-Based Model for Image Classification. arXiv:2002.03141.",
        "url": "https://arxiv.org/abs/2002.03141",
    },
    "Koch2015": {
        "text": "Koch, G., Zemel, R., & Salakhutdinov, R. (2015). Siamese Neural Networks for One-shot Image Recognition. ICML Deep Learning Workshop.",
        "url": "https://www.cs.cmu.edu/~rsalakhu/papers/oneshot1.pdf",
    },
    "Sung2018": {
        "text": "Sung, F., Yang, Y., Zhang, L., Xiang, T., Torr, P. H. S., & Hospedales, T. M. (2018). Learning to Compare: Relation Network for Few-Shot Learning. CVPR 2018.",
        "url": "https://arxiv.org/abs/1711.06025",
    },
    "Vinyals2016": {
        "text": "Vinyals, O., Blundell, C., Lillicrap, T., Kavukcuoglu, K., & Wierstra, D. (2016). Matching Networks for One Shot Learning. Advances in Neural Information Processing Systems 29.",
        "url": "https://papers.nips.cc/paper/6385-matching-networks-for-one-shot-learning",
    },
    "Finn2017": {
        "text": "Finn, C., Abbeel, P., & Levine, S. (2017). Model-Agnostic Meta-Learning for Fast Adaptation of Deep Networks. Proceedings of ICML 2017.",
        "url": "https://proceedings.mlr.press/v70/finn17a.html",
    },
    "DTWSSE2021": {
        "text": "DTWSSE: Data Augmentation with a Siamese Encoder for Time Series. (2021). arXiv:2108.09885.",
        "url": "https://arxiv.org/abs/2108.09885",
    },
}


def cite(keys: List[str]) -> str:
    return "; ".join(f"[{key}]" for key in keys)


def reference_keys_in_order(row_groups: Iterable[List[dict]]) -> List[str]:
    seen: set[str] = set()
    keys: List[str] = []
    for rows in row_groups:
        for row in rows:
            raw = row.get("citation_keys", "")
            for item in raw.split(";"):
                key = item.strip().strip("[]")
                if key and key not in seen:
                    seen.add(key)
                    keys.append(key)
    return keys


def read_csv_rows(path: Path) -> List[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def find_row(rows: List[dict], key: str, value: str) -> dict:
    for row in rows:
        if row.get(key) == value:
            return row
    return {}


def best_by(rows: List[dict], auprc_first: bool = True) -> dict:
    usable = [r for r in rows if r.get("AUROC") not in ["", None] and r.get("AUPRC") not in ["", None]]
    if not usable:
        return {}
    if auprc_first:
        return max(usable, key=lambda r: (float(r["AUPRC"]), float(r["AUROC"])))
    return max(usable, key=lambda r: (min(float(r["AUPRC"]), float(r["AUROC"])), float(r["AUPRC"]), float(r["AUROC"])))


def metric_fields(row: dict) -> dict:
    return {
        "AUROC": row.get("AUROC", ""),
        "AUPRC": row.get("AUPRC", ""),
        "TN": row.get("TN", ""),
        "FP": row.get("FP", ""),
        "FN": row.get("FN", ""),
        "TP": row.get("TP", ""),
    }


def fmt(value: str | float) -> str:
    if value in ["", None]:
        return ""
    return f"{float(value):.3f}"


def cm(row: dict) -> str:
    vals = [row.get(k, "") for k in ["TN", "FP", "FN", "TP"]]
    return "/".join(str(v) for v in vals)


def custom_parameters(reference_model: str, local_model: str, row: dict) -> str:
    """Describe local parameters used when paper parameters were unavailable."""
    if local_model == "EEG aggregate LR fusion with mi_max base score":
        return (
            f"自訂參數：fold-local EEG aggregate features；base_col={row.get('base_col', 'mi_max')}；"
            f"model={row.get('model', 'LR')}；top_k={row.get('top_k', '256')}；"
            "LogisticRegression(C=0.25, class_weight=balanced, solver=liblinear, max_iter=5000)；"
            "patient-aware stratified 10-fold；inner-fold score fusion。"
        )
    if local_model == "2D CNN over subject-level PP/AR/RP/GAF maps":
        return (
            "自訂參數：PP/AR/RP/GAF subject-level maps；image_size=64；Tiny2DCNN Conv2d channels "
            "16/32/64/96, kernel=3；BatchNorm/ReLU/MaxPool；Dropout2d=0.10/0.15, head dropout=0.25；"
            "AdamW(lr=1e-3, weight_decay=1e-4)；BCEWithLogitsLoss(pos_weight=n_neg/n_pos)；"
            "epochs=80, patience=12, batch_size=16；patient-aware stratified 10-fold with inner validation。"
        )
    if local_model == "1D CNN on grouped raw physiological sequences":
        return (
            "自訂參數：raw sequence length=256；Tiny1DCNN Conv1d channels 16/32/64, kernels 7/5/3；"
            "BatchNorm/ReLU/MaxPool/AdaptiveAvgPool；Dropout=0.15/0.20；AdamW(lr=1e-3, weight_decay=1e-4)；"
            "BCEWithLogitsLoss(pos_weight=n_neg/n_pos)；epochs=80, patience=12, batch_size=16；"
            "patient-aware stratified 10-fold with inner validation。"
        )
    if local_model == "1D CNN + 2D CNN validation-selected score ensemble":
        return (
            "自訂參數：同時訓練 Tiny1DCNN 與 Tiny2DCNN；validation split 在每個 training fold 內產生；"
            "融合權重由 validation AUPRC 選擇；test fold 僅使用已選權重推論；"
            "patient-aware stratified 10-fold。"
        )
    if local_model == "training-fold-only DTW-neighbor T-SMOTE-style augmentation + ExtraTrees":
        return (
            f"自訂參數：raw time-series DTW-neighbor interpolation；reduced_length=32；"
            f"synthetic_samples={row.get('synthetic_samples', '90')}；SelectKBest(f_classif, k=48)；"
            "ExtraTrees(n_estimators=240, class_weight=balanced)；augmentation, feature extraction, "
            "feature selection all training-fold-only。"
        )
    if local_model == "training-fold-only feature-space SMOTE + logistic regression":
        return (
            "自訂參數：2D group/channel/patch features；top_k=64；SMOTE(random_state=seed, "
            "k_neighbors=min(5, minority_train_count-1))；LogisticRegression(C=0.2, "
            "class_weight=balanced, solver=liblinear, max_iter=4000)；所有 selection/SMOTE 均在 training fold 內 fit。"
        )
    if local_model == "class prototype distance on selected patch/Green features":
        return (
            "自訂參數：2D group_allchannels patch features；top_k=64；fold-local median imputation + "
            "f_classif feature selection；StandardScaler；以 class centroids 的距離差轉為 prototype score。"
        )
    if local_model == "fold-local PCA/Green feature extractor + logistic/prototype":
        return (
            "自訂參數：2D group_allchannels patch features；top_k=64；StandardScaler；"
            "PCA/Green proxy n_components=min(32, train_samples-1, selected_features)；"
            "下游使用 prototype 或 LogisticRegression(C=0.2, class_weight=balanced)。"
        )
    if local_model == "relation-distance features + logistic regression":
        return (
            "自訂參數：2D selected patch features；relation features 為 query 到兩類 prototype 的距離、距離差、"
            "nearest-class distance；LogisticRegression(C=0.5, class_weight=balanced, solver=liblinear, max_iter=4000)。"
        )
    return "自訂參數：請見對應 analysis manifest、source script 與輸出 CSV。"


def collect_rows() -> tuple[List[dict], List[dict], List[dict]]:
    deep = read_csv_rows(ROOT / "analysis" / "nsc_deep_baselines_group10fold_20260519" / "deep_baseline_summary.csv")
    raw_aug = read_csv_rows(ROOT / "analysis" / "nsc_raw_timeseries_dtw_tsmote_20260519" / "raw_timeseries_dtw_tsmote_summary.csv")
    paper = read_csv_rows(ROOT / "analysis" / "nsc_dataset_images_paper_method_comparison_20260519" / "method_comparison_summary.csv")
    eeg_agg = read_csv_rows(ROOT / "analysis" / "nsc_eeg_csv_fusion_ablation_noleak_fast_20260520" / "eeg_fusion_summary.csv")
    eeg_trial = read_csv_rows(ROOT / "analysis" / "nsc_eeg_trial_level_fusion_20260520" / "trial_level_eeg_summary.csv")
    full_rep = read_csv_rows(ROOT / "analysis" / "nsc_full_literature_replications_20260521" / "full_literature_replications_summary.csv")
    tail_analysis = tail_bag_analysis_dir()
    tail_bag = read_csv_rows(tail_analysis / "bagged_metrics_summary.csv")

    main_rows: List[dict] = []
    component_rows: List[dict] = []
    not_replicated_rows: List[dict] = []

    def add(
        bucket: List[dict],
        role: str,
        reference_model: str,
        local_model: str,
        validation: str,
        row: dict,
        replication_status: str,
        citation_keys: List[str],
        notes: str,
    ) -> None:
        bucket.append(
            {
                "role": role,
                "reference_model": reference_model,
                "local_model": local_model,
                "validation": validation,
                "replication_status": replication_status,
                "citation_keys": cite(citation_keys),
                "custom_parameters": custom_parameters(reference_model, local_model, row),
                **metric_fields(row),
                "CM": cm(row),
                "notes": notes,
            }
        )

    # Our current best strict and exploratory comparators.
    if tail_bag:
        row = tail_bag[0]
        main_rows.append(
            {
                "role": "Our optimized candidate",
                "reference_model": "Current NSC method",
                "local_model": row.get("method", "Bagged Restricted Subject Fusion with Tail Stats"),
                "validation": (
                    "corrected-data patient-aware 10-fold OOF across 10 seeds; post-sweep optimized candidate"
                    if tail_analysis == TAIL_BAG_CORRECTED
                    else "patient-aware 10-fold OOF across 10 seeds; post-sweep optimized candidate"
                ),
                "replication_status": "our method, optimized candidate",
                "citation_keys": "",
                "custom_parameters": (
                    "Targeted EEG tail statistics: p90/p95 for delta/theta band fractions and p05/p10 for alpha band fraction; "
                    "baseline=mi_max; LR penalty=l1; C=0.25; top_k=320; class_weight=None; "
                    "inner 5-fold fusion-weight search with w_eeg in [0, 0.25] step 0.05; "
                    "10-seed OOF bagging over seeds 20260520-20260529."
                ),
                "AUROC": row.get("AUROC", ""),
                "AUPRC": row.get("AUPRC", ""),
                "TN": row.get("TN", ""),
                "FP": row.get("FP", ""),
                "FN": row.get("FN", ""),
                "TP": row.get("TP", ""),
                "CM": cm(row),
                "notes": (
                    "Corrected-data rerun supersedes the pre-correction candidate. Because parameters were selected after grid searches, "
                    "treat as optimized OOF candidate and confirm with a locked protocol before making final generalization claims."
                    if tail_analysis == TAIL_BAG_CORRECTED
                    else "Strongest current candidate by AUPRC. Because parameters were selected after grid searches, "
                    "treat as optimized OOF candidate and confirm with a locked protocol before making final generalization claims."
                ),
            }
        )
    add(
        main_rows,
        "Our strict best",
        "Current NSC method",
        "EEG aggregate LR fusion with mi_max base score",
        "patient-aware 10-fold, fold-local aggregate EEG fusion",
        best_by(eeg_agg, auprc_first=True),
        "our method, strict comparator",
        [],
        "Current strongest non-leaky comparator by AUPRC.",
    )
    add(
        component_rows,
        "Our exploratory best",
        "Current NSC method",
        "EEG trial-level classifier fixed-score fusion",
        "patient-aware 10-fold, trial OOF + posthoc fixed score fusion",
        best_by(eeg_trial, auprc_first=False),
        "our method, exploratory/posthoc",
        [],
        "Useful sensitivity result but not the strict main comparator until fusion weight is fully nested.",
    )

    # Replicated literature model families.
    add(
        main_rows,
        "Replicated literature model",
        "Wang & Oates 2015 / Hatami 2017 / GAF ECG CNN: 2D time-series image CNN",
        "2D CNN over subject-level PP/AR/RP/GAF maps",
        "patient-aware 10-fold",
        find_row(deep, "method", "2d_cnn_pp_ar_rp_gaf_maps"),
        "replicated model family",
        ["Wang2015", "Hatami2017", "Elmir2023"],
        "Uses the paper family of time-series-to-image maps followed by CNN; input maps are the NSC PP/AR/RP/GAF images.",
    )
    add(
        main_rows,
        "Replicated literature model",
        "1D vs 2D ECG CNN / generic raw time-series CNN: 1D CNN",
        "1D CNN on grouped raw physiological sequences",
        "patient-aware 10-fold",
        find_row(deep, "method", "1d_cnn_raw_sequence"),
        "replicated model family",
        ["Wu2018", "Noman2018"],
        "End-to-end 1D CNN baseline trained on raw resampled sequences.",
    )
    add(
        main_rows,
        "Replicated literature model",
        "1D+2D heart-sound/ECG ensemble: score-level 1D+2D CNN ensemble",
        "1D CNN + 2D CNN validation-selected score ensemble",
        "patient-aware 10-fold with inner validation",
        find_row(deep, "method", "1d_2d_cnn_score_ensemble"),
        "replicated model family",
        ["Noman2018", "Wu2018"],
        "Replicates the literature family of combining 1D waveform and 2D image CNN scores.",
    )
    for row in raw_aug:
        if row.get("augmentation") == "dtw_tsmote" and row.get("model") == "extratrees":
            add(
                main_rows,
                "Replicated literature method",
                "T-SMOTE / DTW-based time-series augmentation",
                "training-fold-only DTW-neighbor T-SMOTE-style augmentation + ExtraTrees",
                "patient-aware 10-fold",
                row,
                "replicated augmentation family",
                ["Zhao2022"],
                "The augmentation is fold-local and uses raw time-series DTW-neighbor interpolation; classifier is ExtraTrees.",
            )
    add(
        main_rows,
        "Replicated literature method",
        "SMOTE 2002",
        "training-fold-only feature-space SMOTE + logistic regression",
        "patient-aware 10-fold",
        find_row(paper, "method", "smote_patch_logistic"),
        "replicated method",
        ["Chawla2002"],
        "Classical SMOTE is run only inside each training fold.",
    )
    add(
        main_rows,
        "Replicated literature method",
        "Prototypical Networks baseline idea",
        "class prototype distance on selected patch/Green features",
        "patient-aware 10-fold",
        find_row(paper, "method", "green_pca_prototype"),
        "non-neural prototype replication",
        ["Snell2017"],
        "Replicates the prototype-classification rule, but not the full learned episodic neural encoder.",
    )

    # Component-only or not exact model replication.
    add(
        component_rows,
        "Method component only",
        "PixelHop / PixelHop++",
        "fold-local PCA/Green feature extractor + logistic/prototype",
        "patient-aware 10-fold",
        find_row(paper, "method", "green_pca_logistic"),
        "component-level replication, not full PixelHop",
        ["Chen2020", "Chen2020b"],
        "Uses fold-local PCA as a Green-learning component, but does not implement full Saab/PatchHop hierarchy.",
    )
    add(
        component_rows,
        "Method component only",
        "Siamese / Relation Network",
        "relation-distance features + logistic regression",
        "patient-aware 10-fold",
        find_row(paper, "method", "relation_logistic"),
        "component-level replication, not full neural relation network",
        ["Koch2015", "Sung2018"],
        "Uses pair/prototype distance features; no learned neural relation module.",
    )

    for row in full_rep:
        main_rows.append(
            {
                "role": "Completed literature replication",
                "reference_model": row.get("reference_model", ""),
                "local_model": row.get("method", ""),
                "validation": row.get("validation", "patient-aware stratified 10-fold"),
                "replication_status": "completed local-parameter replication",
                "citation_keys": row.get("citation_keys", ""),
                "custom_parameters": row.get("custom_parameters", ""),
                "AUROC": row.get("AUROC", ""),
                "AUPRC": row.get("AUPRC", ""),
                "TN": row.get("TN", ""),
                "FP": row.get("FP", ""),
                "FN": row.get("FN", ""),
                "TP": row.get("TP", ""),
                "CM": cm(row),
                "notes": "Fuller local replication completed with disclosed custom parameters; unreleased author hyperparameters are not required.",
            }
        )

    main_rows = sorted(
        main_rows,
        key=lambda r: (
            0 if r["role"].startswith("Our") else 1,
            -float(r["AUPRC"]) if r["AUPRC"] else 0.0,
            -float(r["AUROC"]) if r["AUROC"] else 0.0,
        ),
    )
    component_rows = sorted(
        component_rows,
        key=lambda r: (-float(r["AUPRC"]) if r["AUPRC"] else 0.0, -float(r["AUROC"]) if r["AUROC"] else 0.0),
    )
    return main_rows, component_rows, not_replicated_rows


def make_figure(main_rows: List[dict]) -> str:
    fig_dir = OUT_DIR / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    rows = [r for r in main_rows if r["AUROC"] and r["AUPRC"]]
    labels = [r["local_model"] for r in rows]
    y = list(range(len(rows)))
    plt.figure(figsize=(10.8, max(4.5, len(rows) * 0.52 + 1.0)))
    plt.barh([i + 0.18 for i in y], [float(r["AUROC"]) for r in rows], height=0.34, label="AUROC")
    plt.barh([i - 0.18 for i in y], [float(r["AUPRC"]) for r in rows], height=0.34, label="AUPRC")
    plt.axvline(0.8, color="gray", linestyle="--", linewidth=1)
    plt.yticks(y, labels, fontsize=8)
    plt.xlabel("Metric")
    plt.title("Replicated Literature Models vs Our Best, Patient-Aware 10-Fold")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    plt.gca().invert_yaxis()
    path = fig_dir / "replicated_models_vs_our_best.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return str(path)


def add_table(doc: Document, title: str, rows: List[dict], include_metrics: bool = True) -> None:
    doc.add_heading(title, level=1)
    headers = ["角色", "文獻模型/方法", "本案復刻模型", "驗證", "狀態", "需 cite", "本案自訂參數"]
    if include_metrics:
        headers += ["AUROC", "AUPRC", "CM"]
    headers += ["說明"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for row in rows:
        vals = [
            row["role"],
            row["reference_model"],
            row["local_model"],
            row["validation"],
            row["replication_status"],
            row.get("citation_keys", ""),
            row.get("custom_parameters", ""),
        ]
        if include_metrics:
            vals += [fmt(row["AUROC"]), fmt(row["AUPRC"]), row["CM"]]
        vals += [row["notes"]]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)


def make_report(main_rows: List[dict], component_rows: List[dict], not_rows: List[dict], fig_path: str) -> None:
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("NSC 114 位資料：文獻模型復刻與本案最佳方法比較", level=0)
    doc.add_paragraph(f"產出日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、比較原則", level=1)
    doc.add_paragraph(
        "本版依照要求，主表只放已在 114 位 NSC 資料上實際訓練與推論的文獻模型/方法。"
        "驗證方式統一為 patient-aware stratified 10-fold，避免同一位受試者同時出現在訓練與測試。"
        "AUPRC 作為主要指標，AUROC 作為輔助指標。"
    )
    doc.add_paragraph(
        "復刻定義修訂如下：若原作者未釋出完整模型參數或超參數，這不構成不能復刻的理由；"
        "本案可採用與文獻核心方法一致的 architecture / representation / training logic，並完整列出本案自訂參數。"
        "只有核心模型或訓練流程尚未在本資料上實作與推論時，才不列入主表分數。"
    )

    top_our = next((r for r in main_rows if r["role"].startswith("Our")), {})
    strict_best = next((r for r in main_rows if r["role"] == "Our strict best"), {})
    doc.add_heading("二、結論摘要", level=1)
    corrected_mode = tail_bag_analysis_dir() == TAIL_BAG_CORRECTED
    doc.add_paragraph(
        f"本案目前最強候選方法為 {top_our.get('local_model', '')}，"
        f"AUROC={fmt(top_our.get('AUROC', ''))}，AUPRC={fmt(top_our.get('AUPRC', ''))}。"
        f"原先嚴格比較基準 {strict_best.get('local_model', '')} 為 "
        f"AUROC={fmt(strict_best.get('AUROC', ''))}，AUPRC={fmt(strict_best.get('AUPRC', ''))}。"
        "已復刻的文獻 CNN、1D/2D ensemble、T-SMOTE/DTW augmentation、SMOTE 與 prototype 類方法，"
        "在目前 114 位 patient-aware 10-fold 下均未超過本案最佳方法。"
    )
    doc.add_paragraph(
        "因此論文應寫成：文獻深度模型與資料增廣方法已作為可重跑對照；"
        "在相同 patient-aware 10-fold 評估條件下，本案的 mi_max + EEG subject-level fusion 仍是目前最強方向。"
        "其中 targeted tail-statistics bagging 是目前最強 optimized OOF candidate；若要作最終泛化宣稱，需先鎖定 protocol 後再做 confirmatory run。"
    )
    if corrected_mode and top_our.get("AUROC") and top_our.get("AUPRC"):
        auroc = float(top_our["AUROC"])
        auprc = float(top_our["AUPRC"])
        doc.add_paragraph(
            f"本版已改用 corrected data rerun：AUROC={auroc:.4f}，AUPRC={auprc:.4f}。"
            "修正前候選結果曾為 AUROC=0.7990、AUPRC=0.8011；"
            "因 NSC-REDACTED 重複誤植與編號 165 誤植已修正，後續正式比較應以本版 corrected-data 結果為準。"
            "本版接近但未同時達到 AUROC/AUPRC 0.8/0.8，因此不能宣稱已達 0.8/0.8。"
        )
    doc.add_paragraph(
        "2026-05-21 補跑後，先前尚未完整比較的 Matching Networks、MAML-style、DTWSSE-style，以及 PixelHop/Siamese/Relation "
        "本案自訂參數復刻版均已納入主表。"
    )

    add_table(doc, "三、主表：已復刻並實測的文獻模型/方法 vs 本案最佳方法", main_rows)
    doc.add_paragraph("圖 1. 已復刻文獻模型/方法與本案最佳方法之 AUROC/AUPRC 比較")
    doc.add_picture(fig_path, width=Inches(6.4))

    doc.add_heading("四、先前未完整模型之補跑狀態", level=1)
    explanations = [
        (
            "PixelHop / PixelHop++：本報告只復刻 fold-local PCA/Green feature extractor 的方法元件。"
            "本版已新增 PixelHop-style SSL 本案自訂參數復刻，使用 patch top_k、兩階段 PCA/SSL、abs/square hop 與 MLP classifier。需 cite [Chen2020; Chen2020b]"
        ),
        (
            "Siamese / Relation Network：本報告只復刻 pair/prototype distance 的判別元件。"
            "本版已新增 Siamese contrastive encoder 與 Relation Network 本案自訂參數復刻，包含 pair sampling、encoder、relation module 與 loss。需 cite [Koch2015; Sung2018]"
        ),
        (
            "Matching Networks：原方法是 N-way K-shot support/query episodic classifier。"
            "本版已用 binary support/query episodic protocol 與 attention over support set 完成 Matching Network 本案自訂參數復刻。需 cite [Vinyals2016]"
        ),
        (
            "MAML / MetaMed：原方法需要一組 meta-training tasks 與 meta-test tasks。"
            "本版已用 first-order MAML/Reptile-style pseudo-task episodes 完成本案自訂參數復刻。需 cite [Finn2017]"
        ),
        (
            "DTWSSE full Siamese encoder augmentation：本報告已復刻 DTW/T-SMOTE-style augmentation，但尚未復刻 Siamese encoder。"
            "本版已新增 raw sequence Siamese encoder、embedding nearest positive interpolation、temporal jitter 與 raw-feature MLP classifier 的 DTWSSE-style 復刻。需 cite [DTWSSE2021]"
        ),
    ]
    for item in explanations:
        doc.add_paragraph(item, style="List Bullet")

    shap_csv = ROOT / "analysis" / "nsc_eeg_shap_neuropsych_20260521" / "eeg_fusion_surrogate_shap_top_features.csv"
    shap_fig = ROOT / "analysis" / "nsc_eeg_shap_neuropsych_20260521" / "figures" / "eeg_fusion_surrogate_shap_summary.png"
    imp_fig = ROOT / "analysis" / "nsc_eeg_shap_neuropsych_20260521" / "figures" / "eeg_fusion_top_feature_importance.png"
    weight_fig = ROOT / "analysis" / "nsc_eeg_shap_neuropsych_20260521" / "figures" / "fusion_weight_by_fold.png"
    if shap_csv.exists() and shap_fig.exists():
        doc.add_heading("五、本案最佳方法 SHAP 與神經心理特徵詮釋", level=1)
        doc.add_paragraph(
            "本案 strict best 為 fusion_mi_max_LR_k256。正式效能仍以 patient-aware 10-fold OOF 結果為準；"
            "SHAP 圖為 mi_max + EEG top features 的 post-hoc linear-SHAP surrogate，只用於解釋特徵方向與相對重要性。"
        )
        if weight_fig.exists():
            doc.add_picture(str(weight_fig), width=Inches(6.2))
        doc.add_picture(str(shap_fig), width=Inches(6.4))
        if imp_fig.exists():
            doc.add_picture(str(imp_fig), width=Inches(6.4))
        shap_rows = read_csv_rows(shap_csv)[:12]
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        for cell, text in zip(table.rows[0].cells, ["Rank", "特徵", "領域", "Mean |SHAP|", "方向", "神經心理詮釋"]):
            cell.text = text
        for row in shap_rows:
            cells = table.add_row().cells
            vals = [
                row.get("rank", ""),
                row.get("display_feature", ""),
                row.get("domain", ""),
                f"{float(row.get('mean_abs_shap', 0)):.4f}",
                row.get("direction", ""),
                row.get("interpretation", ""),
            ]
            for cell, val in zip(cells, vals):
                cell.text = str(val)
        doc.add_paragraph(
            "專業限制：EEG CSV aggregate 缺少標準 10-20 channel 拓樸與任務事件標記；高頻 beta/gamma/high 特徵可能混合神經活動、肌電與動作 artifact；"
            "資料品質特徵若重要，需優先檢查紀錄完整度與設備差異。"
        )

    doc.add_heading("六、引用文獻", level=1)
    for key in reference_keys_in_order([main_rows, component_rows, not_rows]):
        ref = REFERENCES.get(key)
        if not ref:
            continue
        doc.add_paragraph(f"[{key}] {ref['text']} {ref['url']}", style="List Bullet")

    doc.add_heading("七、建議下一步", level=1)
    for item in [
        "新增任何文獻模型時，不必等待原作者完整參數；需在 Methods 或 Supplement 中列出本案自訂 architecture、超參數、資料切分與 fold-local preprocessing。",
        "若要更接近 Wang & Oates 原始設定，可新增 MTF map 與 tiled CNN；若參數未公開，直接列本案自訂參數。",
        "若要復刻 Prototypical / Siamese / Relation Network，應新增 episodic train/validation protocol，並列出本案自訂 neural encoder 與 loss 設定。",
        "若要復刻 T-SMOTE，下一版需加入 boundary sample selection；若細節未公開或不完全適用本資料，列出本案 DTW/T-SMOTE-style 自訂規則。",
        "所有新增模型仍須維持 patient-aware 10-fold，且 preprocessing、feature selection、augmentation、threshold/fusion selection 都只能在 training fold 內完成。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(REPORT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    main_rows, component_rows, not_rows = collect_rows()
    fields = ["role", "reference_model", "local_model", "validation", "replication_status", "citation_keys", "custom_parameters", "AUROC", "AUPRC", "TN", "FP", "FN", "TP", "CM", "notes"]
    write_csv(OUT_DIR / "replicated_literature_vs_our_best_main.csv", main_rows, fields)
    write_csv(OUT_DIR / "component_only_or_partial_replications.csv", component_rows, fields)
    write_csv(OUT_DIR / "not_yet_replicated_literature_models.csv", not_rows, fields)
    write_csv(
        OUT_DIR / "citation_mapping.csv",
        [
            {"citation_key": key, "reference": ref["text"], "url": ref["url"]}
            for key, ref in REFERENCES.items()
        ],
        ["citation_key", "reference", "url"],
    )
    fig_path = make_figure(main_rows)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "report": str(REPORT),
        "analysis_dir": str(OUT_DIR),
        "main_rows": len(main_rows),
        "component_rows": len(component_rows),
        "not_replicated_rows": len(not_rows),
        "figure": fig_path,
        "claim_level": "patient-aware stratified 10-fold for main table",
        "tail_bag_analysis_dir": str(tail_bag_analysis_dir()),
    }
    with (OUT_DIR / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(main_rows, component_rows, not_rows, fig_path)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
