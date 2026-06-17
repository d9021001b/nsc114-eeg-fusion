#!/usr/bin/env python3
"""NSC 114-case EEG CSV fusion ablation.

Inputs
------
- `nsc_dataset_images/manifest.csv` supplies the binary patient label.
- `eeg-csv-data-by-class/{0..4}/*.csv` supplies newly received EEG raw trials.
- Existing best OOF scores from `nsc_uncertain_band_patch_refinement_20260520`
  supply the current 2D/raw baseline (`mi_max`, etc.).

Validation
----------
Patient-level stratified 10-fold. Feature selection, imputation, model fitting,
EEG branch calibration, and score-fusion weights are all selected inside each
outer training fold. The script keeps all 114 patients: 106 have EEG CSV, 8 are
represented by missing EEG features and explicit availability counts.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import sys
import warnings
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier, RandomForestClassifier
from sklearn.feature_selection import f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)


EEG_FOLDERS = ["0", "1", "2", "3", "4"]
# User-provided mapping: folder 0 is class 0; folders 1/2/3/4 are all class 1.
# Keep the folder/subclass counts as optional audit features, but aggregate signal
# statistics at binary EEG class level for modeling.
EEG_FOLDER_TO_BINARY = {"0": "0", "1": "1", "2": "1", "3": "1", "4": "1"}
EEG_BINARY_CLASSES = ["0", "1"]
BANDS = [
    ("delta", 0.5, 4.0),
    ("theta", 4.0, 8.0),
    ("alpha", 8.0, 13.0),
    ("beta", 13.0, 30.0),
    ("gamma", 30.0, 80.0),
    ("high", 80.0, 120.0),
]


def natural_key(text: str) -> Tuple:
    parts = re.split(r"(\d+)", str(text))
    return tuple(int(p) if p.isdigit() else p for p in parts)


def write_csv(path: Path, rows: Iterable[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def dump_json(path: Path, payload: dict | list) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def load_labels(manifest_path: Path) -> Dict[str, int]:
    labels: Dict[str, int] = {}
    with manifest_path.open("r", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            labels[str(row["subject_id"])] = int(row["label"])
    return dict(sorted(labels.items(), key=lambda kv: natural_key(kv[0])))


def load_base_scores(path: Path, subjects: List[str]) -> Dict[str, np.ndarray]:
    table = pd.read_csv(path, encoding="utf-8-sig")
    table["subject_id"] = table["subject_id"].astype(str)
    table = table.set_index("subject_id")
    out: Dict[str, np.ndarray] = {}
    for col in ["mi_max", "mi_top3mean", "patient_refined_fpfn_family", "raw_only"]:
        if col in table.columns:
            out[col] = table.loc[subjects, col].to_numpy(dtype=float)
    return out


def file_cache_key(path: Path) -> str:
    stat = path.stat()
    raw = f"{path.resolve()}|{stat.st_size}|{int(stat.st_mtime)}"
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()[:24]


def safe_percentiles(values: np.ndarray, qs: List[float]) -> Dict[str, float]:
    if values.size == 0:
        return {f"p{int(q):02d}": math.nan for q in qs}
    vals = np.nanpercentile(values, qs)
    return {f"p{int(q):02d}": float(v) for q, v in zip(qs, vals)}


def downsample_even(values: np.ndarray, max_points: int = 8192) -> np.ndarray:
    values = values[np.isfinite(values)]
    if values.size <= max_points:
        return values.astype(float)
    idx = np.linspace(0, values.size - 1, max_points).astype(int)
    return values[idx].astype(float)


def extract_one_eeg_csv(path: Path, cache_dir: Path) -> dict:
    cache_dir.mkdir(parents=True, exist_ok=True)
    cache_path = cache_dir / f"{file_cache_key(path)}_{path.stem}.json"
    if cache_path.exists():
        with cache_path.open("r", encoding="utf-8") as f:
            return json.load(f)

    header = pd.read_csv(path, nrows=0)
    cols = list(header.columns)
    usecols = [c for c in ["time_sec", "value"] if c in cols]
    if "value" not in usecols:
        numeric = [c for c in cols if c not in {"time_sec", "sample_index"}]
        usecols = [c for c in ["time_sec"] if c in cols] + numeric[:1]
    data = pd.read_csv(path, usecols=usecols)
    value_col = "value" if "value" in data.columns else [c for c in data.columns if c != "time_sec"][0]
    values = data[value_col].to_numpy(dtype=float)
    if "time_sec" in data.columns:
        times = data["time_sec"].to_numpy(dtype=float)
    else:
        times = np.arange(len(values), dtype=float)
    del data

    finite = np.isfinite(values)
    values = values[finite]
    times = times[finite]
    if values.size == 0:
        feats = {"n": 0, "finite_ratio": 0.0}
        dump_json(cache_path, feats)
        return feats

    duration = float(times[-1] - times[0]) if times.size > 1 else 0.0
    sample_rate = float((times.size - 1) / duration) if duration > 0 else math.nan
    q = safe_percentiles(values, [1, 5, 10, 25, 50, 75, 90, 95, 99])
    med = q["p50"]
    centered = values - med
    diffs = np.diff(values) if values.size > 1 else np.asarray([], dtype=float)
    abs_diffs = np.abs(diffs)
    sampled = downsample_even(values)
    sampled_centered = sampled - float(np.nanmean(sampled))

    feats: Dict[str, float] = {
        "n": int(values.size),
        "duration_s": duration,
        "sample_rate_hz": sample_rate,
        "finite_ratio": float(np.mean(finite)),
        "mean": float(np.nanmean(values)),
        "std": float(np.nanstd(values)),
        "min": float(np.nanmin(values)),
        "max": float(np.nanmax(values)),
        "range": float(np.nanmax(values) - np.nanmin(values)),
        "median": float(med),
        "iqr": float(q["p75"] - q["p25"]),
        "mad": float(np.nanmedian(np.abs(centered))),
        "rms": float(np.sqrt(np.nanmean(values**2))),
        "abs_mean": float(np.nanmean(np.abs(values))),
        "center_abs_mean": float(np.nanmean(np.abs(centered))),
        "skew_proxy": float(np.nanmean(centered**3) / (np.nanstd(values) ** 3 + 1e-12)),
        "kurtosis_proxy": float(np.nanmean(centered**4) / (np.nanstd(values) ** 4 + 1e-12)),
        "diff_abs_mean": float(np.nanmean(abs_diffs)) if abs_diffs.size else 0.0,
        "diff_std": float(np.nanstd(diffs)) if diffs.size else 0.0,
        "diff_p95_abs": float(np.nanpercentile(abs_diffs, 95)) if abs_diffs.size else 0.0,
        "zero_cross_rate_centered": float(np.mean(np.diff(np.signbit(centered)) != 0)) if centered.size > 1 else 0.0,
    }
    feats.update(q)

    if sampled.size >= 16:
        x = np.linspace(0.0, 1.0, sampled.size)
        feats["trend_slope"] = float(np.polyfit(x, sampled, deg=1)[0])
        fft = np.abs(np.fft.rfft(sampled_centered))
        power = fft**2
        total_power = float(np.sum(power) + 1e-12)
        # FFT freq-axis fix (2026-06-17): use the EFFECTIVE rate after downsample_even,
        # not the full-resolution sample_rate. For trials decimated to the 8192 cap, the
        # original sample_rate mislabels the frequency axis (affects only the >8192-sample
        # trials). eff_rate = (sampled.size - 1) / duration keeps Hz bins correct.
        eff_rate = (sampled.size - 1) / duration if duration and duration > 0 else sample_rate
        freqs = np.fft.rfftfreq(sampled.size, d=1.0 / eff_rate) if eff_rate and np.isfinite(eff_rate) and eff_rate > 0 else np.arange(len(power))
        feats["spectral_entropy"] = float(-(power / total_power * np.log(power / total_power + 1e-12)).sum())
        feats["spectral_centroid_hz"] = float(np.sum(freqs * power) / total_power)
        cdf = np.cumsum(power) / total_power
        feats["spectral_edge95_hz"] = float(freqs[min(np.searchsorted(cdf, 0.95), len(freqs) - 1)])
        feats["fft_peak_frac"] = float(np.max(power) / total_power)
        for name, lo, hi in BANDS:
            mask = (freqs >= lo) & (freqs < hi)
            feats[f"band_{name}_frac"] = float(np.sum(power[mask]) / total_power) if np.any(mask) else 0.0
        bands = np.array_split(power, 6)
        for i, band in enumerate(bands, start=1):
            feats[f"fft_split{i}_frac"] = float(np.sum(band) / total_power)
    else:
        for name in ["trend_slope", "spectral_entropy", "spectral_centroid_hz", "spectral_edge95_hz", "fft_peak_frac"]:
            feats[name] = math.nan
        for name, _lo, _hi in BANDS:
            feats[f"band_{name}_frac"] = math.nan
        for i in range(1, 7):
            feats[f"fft_split{i}_frac"] = math.nan

    segments = np.array_split(values, 5)
    for i, seg in enumerate(segments, start=1):
        feats[f"seg{i}_mean"] = float(np.nanmean(seg)) if seg.size else math.nan
        feats[f"seg{i}_std"] = float(np.nanstd(seg)) if seg.size else math.nan
        feats[f"seg{i}_median"] = float(np.nanmedian(seg)) if seg.size else math.nan
    if segments[0].size and segments[-1].size:
        feats["last_minus_first_seg_mean"] = float(np.nanmean(segments[-1]) - np.nanmean(segments[0]))
        feats["last_over_first_seg_std"] = float((np.nanstd(segments[-1]) + 1e-12) / (np.nanstd(segments[0]) + 1e-12))

    dump_json(cache_path, feats)
    return feats


def collect_eeg_file_index(eeg_root: Path) -> Dict[str, Dict[str, List[Path]]]:
    out: Dict[str, Dict[str, List[Path]]] = defaultdict(lambda: defaultdict(list))
    for folder in EEG_FOLDERS:
        for path in sorted((eeg_root / folder).glob("*.csv"), key=lambda p: natural_key(p.name)):
            m = re.match(r"^(\d+)_", path.name)
            if not m:
                continue
            out[m.group(1)][EEG_FOLDER_TO_BINARY[folder]].append(path)
            out[m.group(1)][f"folder{folder}"].append(path)
    return out


def aggregate_trial_features(paths: List[Path], cache_dir: Path, prefix: str, agg_funcs: dict) -> Dict[str, float]:
    feats: Dict[str, float] = {}
    if not paths:
        return feats
    trial_feats = [extract_one_eeg_csv(path, cache_dir) for path in paths]
    numeric_names = sorted({k for item in trial_feats for k, v in item.items() if isinstance(v, (int, float))})
    for name in numeric_names:
        vals = np.asarray([float(item.get(name, math.nan)) for item in trial_feats], dtype=float)
        vals = vals[np.isfinite(vals)]
        if vals.size == 0:
            continue
        for agg_name, func in agg_funcs.items():
            feats[f"{prefix}/{name}/{agg_name}"] = float(func(vals))
    return feats


def build_subject_eeg_features(
    subjects: List[str],
    eeg_root: Path,
    cache_dir: Path,
    include_counts: bool,
    class_aware_features: bool,
) -> Tuple[np.ndarray, List[str], dict]:
    index = collect_eeg_file_index(eeg_root)
    rows: Dict[str, Dict[str, float]] = {}
    audit = {
        "subjects_with_eeg": 0,
        "subjects_without_eeg": 0,
        "files_by_class": Counter(),
        "files_by_subject": {},
    }
    agg_funcs = {
        "mean": np.nanmean,
        "std": np.nanstd,
        "min": np.nanmin,
        "max": np.nanmax,
        "p25": lambda x: np.nanpercentile(x, 25),
        "p50": lambda x: np.nanpercentile(x, 50),
        "p75": lambda x: np.nanpercentile(x, 75),
    }
    for subject in subjects:
        feats: Dict[str, float] = {}
        total_files = 0
        binary_paths: Dict[str, List[Path]] = {cls: index.get(subject, {}).get(cls, []) for cls in EEG_BINARY_CLASSES}
        all_paths = [p for cls in EEG_BINARY_CLASSES for p in binary_paths[cls]]
        for cls, paths in binary_paths.items():
            total_files += len(paths)
            audit["files_by_class"][cls] += len(paths)
            if class_aware_features:
                if include_counts:
                    feats[f"eeg_class{cls}/trial_count"] = float(len(paths))
                    feats[f"eeg_class{cls}/has_trials"] = float(len(paths) > 0)
                feats.update(aggregate_trial_features(paths, cache_dir, f"eeg_class{cls}", agg_funcs))
        if not class_aware_features:
            feats.update(aggregate_trial_features(all_paths, cache_dir, "eeg_all", agg_funcs))
        if include_counts:
            feats["eeg_all/total_trial_count"] = float(total_files)
            feats["eeg_all/has_any_eeg"] = float(total_files > 0)
            if class_aware_features:
                for folder in EEG_FOLDERS:
                    folder_paths = index.get(subject, {}).get(f"folder{folder}", [])
                    feats[f"folder{folder}/trial_count"] = float(len(folder_paths))
                    feats[f"folder{folder}/trial_fraction"] = float(len(folder_paths) / max(total_files, 1))
                for cls in EEG_BINARY_CLASSES:
                    denom = max(total_files, 1)
                    feats[f"eeg_class{cls}/trial_fraction"] = float(len(index.get(subject, {}).get(cls, [])) / denom)
        if total_files:
            audit["subjects_with_eeg"] += 1
        else:
            audit["subjects_without_eeg"] += 1
        audit["files_by_subject"][subject] = int(total_files)
        rows[subject] = feats

    feature_names = sorted({name for row in rows.values() for name in row})
    X = np.full((len(subjects), len(feature_names)), np.nan, dtype=float)
    idx = {name: i for i, name in enumerate(feature_names)}
    for i, subject in enumerate(subjects):
        for name, value in rows[subject].items():
            X[i, idx[name]] = value
    audit["files_by_class"] = dict(audit["files_by_class"])
    return X, feature_names, audit


def select_topk(X_train: np.ndarray, y_train: np.ndarray, k: int) -> Tuple[SimpleImputer, np.ndarray]:
    imputer = SimpleImputer(strategy="median")
    X_imp = imputer.fit_transform(X_train)
    if X_imp.shape[1] == 0:
        raise ValueError("No features.")
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        scores, _ = f_classif(X_imp, y_train)
    scores = np.nan_to_num(scores, nan=0.0, posinf=0.0, neginf=0.0)
    selected = np.argsort(scores)[::-1][: min(k, X_imp.shape[1])]
    return imputer, selected


def fit_model(model_name: str, X_train: np.ndarray, y_train: np.ndarray, X_test: np.ndarray, seed: int) -> Tuple[np.ndarray, object]:
    if model_name == "ExtraTrees":
        model = ExtraTreesClassifier(
            n_estimators=500,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=1,
            n_jobs=-1,
        )
    elif model_name == "RF":
        model = RandomForestClassifier(
            n_estimators=500,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=1,
            n_jobs=-1,
        )
    elif model_name == "HistGB":
        model = HistGradientBoostingClassifier(max_iter=120, learning_rate=0.035, l2_regularization=0.05, random_state=seed)
    elif model_name == "LR":
        model = make_pipeline(StandardScaler(), LogisticRegression(max_iter=5000, C=0.25, class_weight="balanced", solver="liblinear"))
    else:
        raise ValueError(f"Unknown model: {model_name}")
    model.fit(X_train, y_train)
    return model.predict_proba(X_test)[:, 1], model


def metrics(y_true: np.ndarray, score: np.ndarray, threshold: float = 0.5) -> dict:
    pred = (score >= threshold).astype(int)
    tn, fp, fn, tp = [int(x) for x in confusion_matrix(y_true, pred, labels=[0, 1]).ravel()]
    return {
        "AUROC": float(roc_auc_score(y_true, score)),
        "AUPRC": float(average_precision_score(y_true, score)),
        "accuracy": float((tn + tp) / max(len(y_true), 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": tn,
        "FP": fp,
        "FN": fn,
        "TP": tp,
    }


def robust_unit_fit(train_scores: np.ndarray, test_scores: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
    lo, hi = np.nanpercentile(train_scores, [5, 95])
    if not np.isfinite(lo) or not np.isfinite(hi) or hi <= lo:
        lo, hi = float(np.nanmin(train_scores)), float(np.nanmax(train_scores))
    if hi <= lo:
        return np.full_like(train_scores, 0.5, dtype=float), np.full_like(test_scores, 0.5, dtype=float)
    return np.clip((train_scores - lo) / (hi - lo), 0, 1), np.clip((test_scores - lo) / (hi - lo), 0, 1)


def fit_eeg_oof(
    X: np.ndarray,
    y: np.ndarray,
    top_k: int,
    model_name: str,
    seed: int,
    n_splits: int,
) -> Tuple[np.ndarray, List[dict], List[dict]]:
    splitter = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=seed)
    score = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    feature_rows: List[dict] = []
    for fold, (train_idx, test_idx) in enumerate(splitter.split(X, y), start=1):
        imputer, selected = select_topk(X[train_idx], y[train_idx], top_k)
        X_train = imputer.transform(X[train_idx])[:, selected]
        X_test = imputer.transform(X[test_idx])[:, selected]
        fold_score, model = fit_model(model_name, X_train, y[train_idx], X_test, seed + fold)
        score[test_idx] = fold_score
        fold_rows.append({"fold": fold, "top_k": top_k, "model": model_name, "train_cases": len(train_idx), "test_cases": len(test_idx)})
        for rank, feat_idx in enumerate(selected[:30], start=1):
            feature_rows.append({"fold": fold, "model": model_name, "top_k": top_k, "rank": rank, "feature_index": int(feat_idx)})
    return score, fold_rows, feature_rows


def inner_select_weight(
    X: np.ndarray,
    y: np.ndarray,
    base_score: np.ndarray,
    outer_train_idx: np.ndarray,
    top_k: int,
    model_name: str,
    seed: int,
    inner_splits: int,
    objective: str,
) -> Tuple[float, List[dict]]:
    inner = StratifiedKFold(n_splits=inner_splits, shuffle=True, random_state=seed)
    eeg_oof = np.zeros(len(outer_train_idx), dtype=float)
    for inner_fold, (tr_local, va_local) in enumerate(inner.split(X[outer_train_idx], y[outer_train_idx]), start=1):
        tr_idx = outer_train_idx[tr_local]
        va_idx = outer_train_idx[va_local]
        imputer, selected = select_topk(X[tr_idx], y[tr_idx], top_k)
        Xtr = imputer.transform(X[tr_idx])[:, selected]
        Xva = imputer.transform(X[va_idx])[:, selected]
        sc, _model = fit_model(model_name, Xtr, y[tr_idx], Xva, seed + inner_fold)
        eeg_oof[va_local] = sc
    base_train, eeg_train = robust_unit_fit(base_score[outer_train_idx], eeg_oof)
    rows = []
    best = (float("-inf"), 0.0)
    for w_eeg in np.linspace(0.0, 1.0, 21):
        fused = (1.0 - w_eeg) * base_train + w_eeg * eeg_train
        met = metrics(y[outer_train_idx], fused)
        if objective == "min_metric":
            obj = min(met["AUROC"], met["AUPRC"])
        else:
            obj = met["AUPRC"] + 0.05 * met["AUROC"]
        rows.append({"w_eeg": float(w_eeg), "objective": float(obj), **met})
        if obj > best[0]:
            best = (obj, float(w_eeg))
    return best[1], rows


def nested_fusion_oof(
    X: np.ndarray,
    y: np.ndarray,
    base_scores: Dict[str, np.ndarray],
    base_col: str,
    top_k: int,
    model_name: str,
    seed: int,
    n_splits: int,
    inner_splits: int,
    objective: str,
) -> Tuple[np.ndarray, np.ndarray, List[dict], List[dict]]:
    splitter = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=seed)
    eeg_score = np.zeros(len(y), dtype=float)
    fused_score = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    weight_rows: List[dict] = []
    base = base_scores[base_col]
    for fold, (train_idx, test_idx) in enumerate(splitter.split(X, y), start=1):
        w_eeg, inner_rows = inner_select_weight(
            X,
            y,
            base,
            train_idx,
            top_k,
            model_name,
            seed + fold * 100,
            inner_splits,
            objective,
        )
        for row in inner_rows:
            weight_rows.append({"fold": fold, "base_col": base_col, "top_k": top_k, "model": model_name, **row})
        imputer, selected = select_topk(X[train_idx], y[train_idx], top_k)
        Xtr = imputer.transform(X[train_idx])[:, selected]
        Xte = imputer.transform(X[test_idx])[:, selected]
        sc_test, _model = fit_model(model_name, Xtr, y[train_idx], Xte, seed + fold)
        # Calibrate base and EEG using the outer training distribution.
        # EEG train score uses model's in-sample score only for scaling, not for weight selection.
        sc_train, _ = fit_model(model_name, Xtr, y[train_idx], Xtr, seed + fold + 7000)
        base_train_unit, base_test_unit = robust_unit_fit(base[train_idx], base[test_idx])
        eeg_train_unit, eeg_test_unit = robust_unit_fit(sc_train, sc_test)
        eeg_score[test_idx] = eeg_test_unit
        fused_score[test_idx] = (1.0 - w_eeg) * base_test_unit + w_eeg * eeg_test_unit
        fold_rows.append(
            {
                "fold": fold,
                "base_col": base_col,
                "top_k": top_k,
                "model": model_name,
                "w_eeg": float(w_eeg),
                "train_cases": len(train_idx),
                "test_cases": len(test_idx),
            }
        )
    return eeg_score, fused_score, fold_rows, weight_rows


def make_figures(out_dir: Path, y: np.ndarray, rows: List[dict], predictions: List[dict]) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    top = rows[:16]
    labels = [r["method"] for r in top]
    ypos = np.arange(len(top))
    plt.figure(figsize=(10.5, max(4.8, 0.38 * len(top))))
    plt.barh(ypos + 0.18, [r["AUROC"] for r in top], height=0.34, label="AUROC")
    plt.barh(ypos - 0.18, [r["AUPRC"] for r in top], height=0.34, label="AUPRC")
    plt.axvline(0.8, color="gray", linestyle="--")
    plt.yticks(ypos, labels, fontsize=8)
    plt.gca().invert_yaxis()
    plt.xlabel("Metric")
    plt.title("NSC EEG CSV fusion ablation, patient-aware 10-fold")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    path = fig_dir / "eeg_fusion_ablation_bar.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["bar"] = str(path)

    if rows:
        best_method = rows[0]["method"]
        pred = [r for r in predictions if r["method"] == best_method]
        pred = sorted(pred, key=lambda r: natural_key(r["subject_id"]))
        score = np.asarray([float(r["score"]) for r in pred], dtype=float)
        fpr, tpr, _ = roc_curve(y, score)
        plt.figure(figsize=(5.4, 4.3))
        plt.plot(fpr, tpr, label=f"AUROC={rows[0]['AUROC']:.3f}")
        plt.plot([0, 1], [0, 1], "--", color="gray")
        plt.xlabel("False Positive Rate")
        plt.ylabel("True Positive Rate")
        plt.title("Best EEG fusion ROC")
        plt.grid(alpha=0.25)
        plt.legend(loc="lower right")
        path = fig_dir / "best_eeg_fusion_roc.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        paths["roc"] = str(path)

        precision, recall, _ = precision_recall_curve(y, score)
        plt.figure(figsize=(5.4, 4.3))
        plt.plot(recall, precision, label=f"AUPRC={rows[0]['AUPRC']:.3f}")
        plt.axhline(float(np.mean(y)), linestyle="--", color="gray", label=f"prevalence={np.mean(y):.3f}")
        plt.xlabel("Recall")
        plt.ylabel("Precision")
        plt.title("Best EEG fusion PRC")
        plt.grid(alpha=0.25)
        plt.legend(loc="lower left")
        path = fig_dir / "best_eeg_fusion_prc.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        paths["prc"] = str(path)
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC 114 位 EEG CSV 融合 Ablation 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、資料與驗證", level=1)
    doc.add_paragraph(
        "本報告將新取得的 EEG CSV trial 資料轉為 patient-level 統計與頻譜特徵，並與既有 raw/2D uncertain-band baseline 做 score-level fusion。"
        "所有模型採 patient-aware stratified 10-fold；EEG 特徵選擇、imputation、模型與 fusion 權重均只在 outer training fold 內決定。"
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "數值"
    for label, value in [
        ("subjects", manifest["subjects"]),
        ("label 0 / 1", f"{manifest['cases_by_label'].get('0', 0)} / {manifest['cases_by_label'].get('1', 0)}"),
        ("subjects with EEG", manifest["eeg_audit"]["subjects_with_eeg"]),
        ("subjects without EEG", manifest["eeg_audit"]["subjects_without_eeg"]),
        ("EEG feature count", manifest["eeg_feature_count"]),
        ("validation", manifest["validation"]),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    doc.add_heading("二、結果", level=1)
    table = doc.add_table(rows=1, cols=12)
    table.style = "Table Grid"
    headers = ["rank", "method", "AUROC", "AUPRC", "ACC", "Sens", "Spec", "PPV", "NPV", "CM", "top_k", "model"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for i, row in enumerate(rows[:24], start=1):
        cm = f"{row['TN']}/{row['FP']}/{row['FN']}/{row['TP']}"
        vals = [
            i,
            row["method"],
            f"{row['AUROC']:.3f}",
            f"{row['AUPRC']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['sensitivity']:.3f}",
            f"{row['specificity']:.3f}",
            f"{row['PPV']:.3f}",
            f"{row['NPV']:.3f}",
            cm,
            row.get("top_k", ""),
            row.get("model", ""),
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、圖表", level=1)
    for key in ["bar", "roc", "prc"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.4 if key == "bar" else 5.2))

    doc.add_heading("四、判讀", level=1)
    best = rows[0]
    doc.add_paragraph(
        f"本輪最佳方法為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。"
    )
    if best["AUROC"] >= 0.8 and best["AUPRC"] >= 0.8:
        doc.add_paragraph("此版本在 patient-aware 10-fold 已達 0.8/0.8，建議下一步做 seed repeat、錯誤案例檢查與特徵穩定性分析。")
    else:
        doc.add_paragraph(
            "此版本尚未同時達 0.8/0.8。若 EEG branch 未提升 baseline，代表目前 EEG CSV 的 trial-level 統計/頻譜摘要仍不足；"
            "下一步應檢查 EEG trial class 與原 outcome 的對應意義，並考慮以 raw sequence encoder 或 class-specific prototype 取代全域統計聚合。"
        )
    doc.add_heading("五、輸出檔案", level=1)
    for k, v in manifest["outputs"].items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--eeg-root", default="eeg-csv-data-by-class")
    parser.add_argument("--base-predictions", default="analysis/nsc_uncertain_band_patch_refinement_20260520/uncertain_band_predictions.csv")
    parser.add_argument("--out-dir", default="analysis/nsc_eeg_csv_fusion_ablation_20260520")
    parser.add_argument("--report", default="reports/NSC_EEG_CSV_fusion_ablation_patient_aware_10fold_20260520.docx")
    parser.add_argument("--cache-dir", default="analysis/nsc_eeg_csv_fusion_ablation_20260520/eeg_feature_cache")
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--random-state", type=int, default=20260520)
    parser.add_argument("--top-ks", nargs="+", type=int, default=[32, 64, 96, 128, 192, 256, 384, 512])
    parser.add_argument("--models", nargs="+", default=["ExtraTrees", "RF", "HistGB", "LR"])
    parser.add_argument("--base-cols", nargs="+", default=["mi_max", "raw_only", "patient_refined_fpfn_family"])
    parser.add_argument("--objective", default="min_metric", choices=["min_metric", "auprc"])
    parser.add_argument("--include-count-features", action="store_true")
    parser.add_argument(
        "--class-aware-features",
        action="store_true",
        help="Use folder-derived class-specific EEG aggregates. This is exploratory only and may leak labels.",
    )
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    labels = load_labels(Path(args.manifest).resolve())
    subjects = list(labels)
    y = np.asarray([labels[s] for s in subjects], dtype=int)
    X_eeg, feature_names, eeg_audit = build_subject_eeg_features(
        subjects,
        Path(args.eeg_root).resolve(),
        Path(args.cache_dir).resolve(),
        include_counts=args.include_count_features,
        class_aware_features=args.class_aware_features,
    )
    base_scores = load_base_scores(Path(args.base_predictions).resolve(), subjects)

    rows: List[dict] = []
    predictions: List[dict] = []
    fold_rows: List[dict] = []
    weight_rows: List[dict] = []
    feature_rows: List[dict] = []

    # Existing baseline scores.
    for name, score in base_scores.items():
        row = metrics(y, score)
        row.update({"method": f"baseline_{name}", "kind": "baseline", "model": "", "top_k": ""})
        rows.append(row)
        for subject, yy, ss in zip(subjects, y, score):
            predictions.append({"subject_id": subject, "true_label": int(yy), "method": f"baseline_{name}", "score": float(ss)})

    for model_name in args.models:
        for top_k in args.top_ks:
            eeg_score, eeg_fold_rows, eeg_feature_rows = fit_eeg_oof(X_eeg, y, top_k, model_name, args.random_state, args.n_splits)
            row = metrics(y, eeg_score)
            row.update({"method": f"eeg_only_{model_name}_k{top_k}", "kind": "eeg_only", "model": model_name, "top_k": top_k})
            rows.append(row)
            fold_rows.extend([{"method": row["method"], **r} for r in eeg_fold_rows])
            feature_rows.extend([{"method": row["method"], **r} for r in eeg_feature_rows])
            for subject, yy, ss in zip(subjects, y, eeg_score):
                predictions.append({"subject_id": subject, "true_label": int(yy), "method": row["method"], "score": float(ss)})

            for base_col in args.base_cols:
                if base_col not in base_scores:
                    continue
                eeg_branch, fused, fusion_folds, fusion_weights = nested_fusion_oof(
                    X_eeg,
                    y,
                    base_scores,
                    base_col,
                    top_k,
                    model_name,
                    args.random_state,
                    args.n_splits,
                    args.inner_splits,
                    args.objective,
                )
                method = f"fusion_{base_col}_{model_name}_k{top_k}"
                row = metrics(y, fused)
                row.update({"method": method, "kind": "fusion", "model": model_name, "top_k": top_k, "base_col": base_col})
                rows.append(row)
                fold_rows.extend([{"method": method, **r} for r in fusion_folds])
                weight_rows.extend([{"method": method, **r} for r in fusion_weights])
                for subject, yy, ss, ee in zip(subjects, y, fused, eeg_branch):
                    predictions.append(
                        {"subject_id": subject, "true_label": int(yy), "method": method, "score": float(ss), "eeg_branch_score": float(ee)}
                    )

    rows.sort(key=lambda r: (min(r["AUROC"], r["AUPRC"]), r["AUPRC"], r["AUROC"]), reverse=True)
    result_fields = sorted({k for r in rows for k in r})
    pred_fields = sorted({k for r in predictions for k in r})
    fold_fields = sorted({k for r in fold_rows for k in r})
    weight_fields = sorted({k for r in weight_rows for k in r})
    feature_fields = sorted({k for r in feature_rows for k in r})
    write_csv(out_dir / "eeg_fusion_summary.csv", rows, result_fields)
    write_csv(out_dir / "eeg_fusion_predictions.csv", predictions, pred_fields)
    write_csv(out_dir / "eeg_fusion_fold_details.csv", fold_rows, fold_fields)
    write_csv(out_dir / "eeg_fusion_inner_weight_search.csv", weight_rows, weight_fields)
    write_csv(out_dir / "eeg_selected_features_by_fold.csv", feature_rows, feature_fields)

    figures = make_figures(out_dir, y, rows, predictions)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in Counter(y).items()},
        "eeg_root": str(Path(args.eeg_root).resolve()),
        "base_predictions": str(Path(args.base_predictions).resolve()),
        "eeg_feature_count": len(feature_names),
        "include_count_features": args.include_count_features,
        "class_aware_features": args.class_aware_features,
        "folder_mapping": EEG_FOLDER_TO_BINARY,
        "eeg_audit": eeg_audit,
        "validation": f"patient-aware stratified {args.n_splits}-fold with inner {args.inner_splits}-fold fusion weight selection",
        "objective": args.objective,
        "best": rows[0] if rows else {},
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "eeg_fusion_summary.csv"),
            "predictions": str(out_dir / "eeg_fusion_predictions.csv"),
            "fold_details": str(out_dir / "eeg_fusion_fold_details.csv"),
            "inner_weight_search": str(out_dir / "eeg_fusion_inner_weight_search.csv"),
            "selected_features": str(out_dir / "eeg_selected_features_by_fold.csv"),
            "report": str(Path(args.report).resolve()),
        },
    }
    dump_json(out_dir / "manifest.json", manifest)
    make_report(Path(args.report).resolve(), manifest, rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
