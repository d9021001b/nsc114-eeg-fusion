#!/usr/bin/env python3
"""Patient-aware multi-instance 2D visual feature fusion for NSC.

Earlier patient-level 2D branches averaged all images for a patient before
modelling. That can dilute local visual evidence. This script keeps every
session/channel/plot image as an instance, trains image-level visual models
inside the training patients only, then aggregates image scores back to the
patient level with mean/max/p75/top-k rules. The patient-level 2D score is
then fused with the raw time-series branch by inner-fold search.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
import warnings
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from PIL import Image
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier, RandomForestClassifier
from sklearn.feature_selection import SelectKBest, VarianceThreshold, f_classif
from sklearn.impute import SimpleImputer
from sklearn.metrics import confusion_matrix, precision_recall_curve, roc_curve
from sklearn.model_selection import StratifiedGroupKFold

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences, natural_key, signal_group  # noqa: E402
from score_level_fusion_2d_raw_nsc import fit_raw_dtw_et, metric_dict, write_csv  # noqa: E402

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)

PLOTS = ("AR", "PP", "RP", "GAF")


def image_visual_features(path: Path, grid: int) -> Dict[str, float]:
    img = Image.open(path).convert("L").resize((128, 128))
    arr = np.asarray(img, dtype=np.float32) / 255.0
    ink = 1.0 - arr
    h, w = ink.shape
    gy = np.abs(np.diff(ink, axis=0))
    gx = np.abs(np.diff(ink, axis=1))
    hist, _ = np.histogram(ink, bins=16, range=(0, 1), density=False)
    prob = hist.astype(np.float64) / max(hist.sum(), 1)
    feats: Dict[str, float] = {
        "global_ink_mean": float(ink.mean()),
        "global_ink_std": float(ink.std()),
        "global_ink_p10": float(np.percentile(ink, 10)),
        "global_ink_p90": float(np.percentile(ink, 90)),
        "global_edge_x": float(gx.mean()) if gx.size else 0.0,
        "global_edge_y": float(gy.mean()) if gy.size else 0.0,
        "global_entropy": float(-(prob * np.log(prob + 1e-12)).sum()),
        "global_left_minus_right": float(ink[:, : w // 2].mean() - ink[:, w // 2 :].mean()),
        "global_top_minus_bottom": float(ink[: h // 2, :].mean() - ink[h // 2 :, :].mean()),
        "global_tl_minus_br": float(ink[: h // 2, : w // 2].mean() - ink[h // 2 :, w // 2 :].mean()),
        "global_tr_minus_bl": float(ink[: h // 2, w // 2 :].mean() - ink[h // 2 :, : w // 2].mean()),
    }
    patch_h = h // grid
    patch_w = w // grid
    patch_values = np.zeros((grid, grid), dtype=np.float32)
    for r in range(grid):
        for c in range(grid):
            patch = ink[r * patch_h : (r + 1) * patch_h, c * patch_w : (c + 1) * patch_w]
            value = float(patch.mean())
            patch_values[r, c] = value
            feats[f"patch_g{grid}_r{r}c{c}_mean"] = value
            feats[f"patch_g{grid}_r{r}c{c}_std"] = float(patch.std())
    for r in range(grid):
        feats[f"row_g{grid}_r{r}_mean"] = float(patch_values[r, :].mean())
    for c in range(grid):
        feats[f"col_g{grid}_c{c}_mean"] = float(patch_values[:, c].mean())
    half = grid // 2
    feats[f"quad_g{grid}_tl"] = float(patch_values[:half, :half].mean())
    feats[f"quad_g{grid}_tr"] = float(patch_values[:half, half:].mean())
    feats[f"quad_g{grid}_bl"] = float(patch_values[half:, :half].mean())
    feats[f"quad_g{grid}_br"] = float(patch_values[half:, half:].mean())
    feats[f"quad_g{grid}_tl_minus_br"] = feats[f"quad_g{grid}_tl"] - feats[f"quad_g{grid}_br"]
    feats[f"quad_g{grid}_tr_minus_bl"] = feats[f"quad_g{grid}_tr"] - feats[f"quad_g{grid}_bl"]
    return feats


def collect_image_instances(root: Path, grid: int) -> Tuple[List[dict], List[str]]:
    instances: List[dict] = []
    feature_names = set()
    for path in sorted(root.rglob("*.png"), key=lambda p: natural_key(str(p))):
        rel = path.relative_to(root).parts
        if len(rel) < 4 or rel[0] not in {"0", "1"} or path.stem not in PLOTS:
            continue
        label = int(rel[0])
        session = rel[1]
        subject = session.split("_Sess", 1)[0]
        channel = rel[2]
        plot = path.stem
        feats = image_visual_features(path, grid)
        feats[f"meta_plot_{plot}"] = 1.0
        feats[f"meta_group_{signal_group(channel)}"] = 1.0
        feats[f"meta_channel_{channel}"] = 1.0
        for key in feats:
            feature_names.add(key)
        instances.append(
            {
                "subject_id": subject,
                "session": session,
                "channel": channel,
                "plot": plot,
                "label": label,
                "path": str(path),
                "features": feats,
            }
        )
    names = sorted(feature_names)
    return instances, names


def instances_to_matrix(instances: List[dict], names: List[str]) -> Tuple[np.ndarray, np.ndarray, np.ndarray]:
    idx = {name: i for i, name in enumerate(names)}
    X = np.zeros((len(instances), len(names)), dtype=np.float32)
    y = np.zeros(len(instances), dtype=int)
    subjects = np.empty(len(instances), dtype=object)
    for i, inst in enumerate(instances):
        y[i] = int(inst["label"])
        subjects[i] = str(inst["subject_id"])
        for key, value in inst["features"].items():
            X[i, idx[key]] = float(value)
    return X, y, subjects


def aggregate_instance_scores(subject_order: List[str], inst_subjects: np.ndarray, inst_scores: np.ndarray, mode: str) -> np.ndarray:
    buckets: Dict[str, List[float]] = defaultdict(list)
    for subject, score in zip(inst_subjects, inst_scores):
        buckets[str(subject)].append(float(score))
    out = []
    for subject in subject_order:
        vals = np.asarray(buckets[str(subject)], dtype=float)
        if len(vals) == 0:
            out.append(0.5)
            continue
        vals = np.sort(vals)
        if mode == "mean":
            out.append(float(vals.mean()))
        elif mode == "max":
            out.append(float(vals.max()))
        elif mode == "p75":
            out.append(float(np.percentile(vals, 75)))
        elif mode == "top3mean":
            out.append(float(vals[-min(3, len(vals)) :].mean()))
        elif mode == "mean_max":
            out.append(float(0.5 * vals.mean() + 0.5 * vals.max()))
        else:
            raise ValueError(mode)
    return np.asarray(out, dtype=float)


def fit_2d_multi_instance(
    Xinst: np.ndarray,
    yinst: np.ndarray,
    inst_subjects: np.ndarray,
    subject_order: List[str],
    y_subject: np.ndarray,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    seed: int,
    top_k: int,
    model_name: str,
    aggregation: str,
) -> Tuple[np.ndarray, np.ndarray, dict]:
    train_subjects = set(str(subject_order[i]) for i in train_idx)
    test_subjects = set(str(subject_order[i]) for i in test_idx)
    inst_train_mask = np.asarray([str(s) in train_subjects for s in inst_subjects], dtype=bool)
    inst_test_mask = np.asarray([str(s) in test_subjects for s in inst_subjects], dtype=bool)
    Xtr = Xinst[inst_train_mask]
    ytr = yinst[inst_train_mask]
    Xte = Xinst[inst_test_mask]
    train_inst_subjects = inst_subjects[inst_train_mask]
    test_inst_subjects = inst_subjects[inst_test_mask]

    imputer = SimpleImputer(strategy="median")
    Xtr_i = imputer.fit_transform(Xtr)
    Xte_i = imputer.transform(Xte)
    var = VarianceThreshold(threshold=1e-8)
    Xtr_v = var.fit_transform(Xtr_i)
    Xte_v = var.transform(Xte_i)
    if Xtr_v.shape[1] == 0:
        Xtr_v, Xte_v = Xtr_i, Xte_i
    k = min(top_k, Xtr_v.shape[1])
    selector = SelectKBest(f_classif, k=k)
    Xtr_s = selector.fit_transform(Xtr_v, ytr)
    Xte_s = selector.transform(Xte_v)

    if model_name == "ExtraTrees":
        model = ExtraTreesClassifier(
            n_estimators=620,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=2,
            n_jobs=-1,
        )
    elif model_name == "RandomForest":
        model = RandomForestClassifier(
            n_estimators=520,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=2,
            n_jobs=-1,
        )
    elif model_name == "HistGB":
        model = HistGradientBoostingClassifier(max_iter=160, learning_rate=0.04, max_leaf_nodes=17, l2_regularization=0.05, random_state=seed)
    else:
        raise ValueError(model_name)

    counts = Counter(str(s) for s in train_inst_subjects)
    sample_weight = np.asarray([1.0 / counts[str(s)] for s in train_inst_subjects], dtype=float)
    try:
        model.fit(Xtr_s, ytr, sample_weight=sample_weight)
    except TypeError:
        model.fit(Xtr_s, ytr)

    train_inst_score = model.predict_proba(Xtr_s)[:, 1]
    test_inst_score = model.predict_proba(Xte_s)[:, 1]
    train_subject_order = [subject_order[i] for i in train_idx]
    test_subject_order = [subject_order[i] for i in test_idx]
    train_score = aggregate_instance_scores(train_subject_order, train_inst_subjects, train_inst_score, aggregation)
    test_score = aggregate_instance_scores(test_subject_order, test_inst_subjects, test_inst_score, aggregation)
    stats = {
        "train_instances": int(Xtr.shape[0]),
        "test_instances": int(Xte.shape[0]),
        "input_features": int(Xinst.shape[1]),
        "after_variance_features": int(Xtr_v.shape[1]),
        "selected_features": int(k),
        "model": model_name,
        "aggregation": aggregation,
    }
    return train_score, test_score, stats


def search_fusion(scores: Dict[str, np.ndarray], y_true: np.ndarray, branches: Tuple[str, ...], step: float = 0.1) -> Tuple[dict, List[dict]]:
    rows: List[dict] = []
    if len(branches) == 1:
        branch = branches[0]
        score = scores[branch]
        row = {
            "branches": branch,
            "weight_raw": 1.0 if branch == "raw" else 0.0,
            "weight_2d": 1.0 if branch == "2d" else 0.0,
            "weight_vlm": 0.0,
            **metric_dict(y_true, score),
        }
        return row, [row]
    grid = np.arange(0, 1.00001, step)
    best = None
    for weight_2d in grid:
        weight_2d = float(weight_2d)
        weight_raw = 1.0 - weight_2d
        score = weight_raw * scores["raw"] + weight_2d * scores["2d"]
        row = {
            "branches": "+".join(branches),
            "weight_raw": weight_raw,
            "weight_2d": weight_2d,
            "weight_vlm": 0.0,
            **metric_dict(y_true, score),
        }
        rows.append(row)
        if best is None or (row["AUPRC"], row["AUROC"]) > (best["AUPRC"], best["AUROC"]):
            best = row
    return best, rows


def make_figures(out_dir: Path, rows: List[dict], y: np.ndarray, best_score: np.ndarray, best: dict) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    top = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)[:16]
    labels = [f"{r['model']} {r['aggregation']} {r['fusion']}" for r in top]
    ypos = np.arange(len(top))
    plt.figure(figsize=(12, max(5, 0.38 * len(top) + 1.2)))
    plt.barh(ypos + 0.18, [r["AUROC"] for r in top], height=0.34, label="AUROC")
    plt.barh(ypos - 0.18, [r["AUPRC"] for r in top], height=0.34, label="AUPRC")
    plt.axvline(0.8, linestyle="--", color="gray")
    plt.yticks(ypos, labels, fontsize=8)
    plt.xlabel("Metric")
    plt.title("Multi-instance 2D visual fusion")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    plt.gca().invert_yaxis()
    path = fig_dir / "multi_instance_metric_bar.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["bar"] = str(path)

    fpr, tpr, _ = roc_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={best['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Best multi-instance ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    path = fig_dir / "best_multi_instance_roc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["roc"] = str(path)

    precision, recall, _ = precision_recall_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={best['AUPRC']:.3f}")
    plt.axhline(float(y.mean()), linestyle="--", color="gray", label=f"prevalence={y.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Best multi-instance PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    path = fig_dir / "best_multi_instance_prc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["prc"] = str(path)

    cm = confusion_matrix(y, (best_score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Best multi-instance confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    path = fig_dir / "best_multi_instance_cm.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["cm"] = str(path)
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC 2D Multi-instance Visual Feature Fusion 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、整合策略", level=1)
    doc.add_paragraph(
        "本版不再把同一病人的所有 2D 圖先平均，而是把每張 session/channel/plot 圖視為一個 visual instance。"
        "模型在 training patients 的 image instances 上學局部 patch/紋理/位置特徵，之後用 mean、max、p75 或 top3mean 聚合成 patient-level 分數。"
    )
    doc.add_paragraph(
        "目的在保留局部視覺訊號：若 label 1 的可分辨現象只出現在特定 channel 或特定 plot，max/p75/top-k 聚合比 patient 平均更不容易被沖淡。"
    )
    doc.add_heading("二、驗證設計", level=1)
    doc.add_paragraph(
        f"{manifest['validation']}。同一 patient-id 不跨 train/test；image-level model、feature selection、"
        "raw DTW/T-SMOTE 與 fusion weight search 都只在 training fold 內完成。"
    )
    doc.add_paragraph(f"Subjects: {manifest['subjects']}；Image instances: {manifest['image_instances']}；Target reached: {manifest['target_reached']}")
    doc.add_heading("三、結果", level=1)
    table = doc.add_table(rows=1, cols=11)
    table.style = "Table Grid"
    headers = ["model", "agg", "fusion", "AUROC", "AUPRC", "ACC", "Sens", "Spec", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for row in rows[:20]:
        vals = [
            row["model"],
            row["aggregation"],
            row["fusion"],
            f"{row['AUROC']:.3f}",
            f"{row['AUPRC']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['sensitivity']:.3f}",
            f"{row['specificity']:.3f}",
            f"{row['PPV']:.3f}",
            f"{row['NPV']:.3f}",
            f"{row['TN']}/{row['FP']}/{row['FN']}/{row['TP']}",
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("四、圖表", level=1)
    for key in ["bar", "roc", "prc", "cm"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.3 if key == "bar" else 5.2))
    doc.add_heading("五、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(f"最佳組合為 {best['model']} / {best['aggregation']} / {best['fusion']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。")
    if manifest["target_reached"]:
        doc.add_paragraph("本版已在 patient-aware group validation 達到 0.8/0.8，可進一步做重複 random seed 穩定性檢查與錯誤案例人工 review。")
    else:
        doc.add_paragraph(
            "本版若仍未達 0.8/0.8，代表目前資料中的局部 2D 訊號仍不足以完全修正 raw branch 錯誤。"
            "下一步應針對 FP/FN 的高分 image instances 進行人工/VLM 規則審查，建立少量高可信 patch family，而非再盲目增加所有 patch。"
        )
    doc.add_heading("六、輸出", level=1)
    for k, v in manifest["outputs"].items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_multi_instance_2d_raw_fusion_20260519")
    parser.add_argument("--report", default="reports/NSC_multi_instance_2D_raw_fusion_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--grid", type=int, default=8)
    parser.add_argument("--top-k", type=int, default=192)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--models", default="ExtraTrees,RandomForest,HistGB")
    parser.add_argument("--aggregations", default="mean,max,p75,top3mean,mean_max")
    parser.add_argument("--fusion-step", type=float, default=0.1)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.max_rows_per_csv,
    )
    instances, feature_names = collect_image_instances(Path(args.image_root).resolve(), args.grid)
    Xinst, yinst, inst_subjects = instances_to_matrix(instances, feature_names)
    allowed_subjects = set(subjects)
    keep = np.asarray([str(s) in allowed_subjects for s in inst_subjects], dtype=bool)
    Xinst = Xinst[keep]
    yinst = yinst[keep]
    inst_subjects = inst_subjects[keep]

    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)
    models = [m.strip() for m in args.models.split(",") if m.strip()]
    aggs = [a.strip() for a in args.aggregations.split(",") if a.strip()]
    fusion_defs = [("2d_mi_only", ("2d",)), ("raw_plus_2d_mi", ("raw", "2d"))]
    results: List[dict] = []
    result_scores: Dict[str, np.ndarray] = {}
    fold_rows: List[dict] = []
    inner_rows: List[dict] = []

    for model_name in models:
        for agg in aggs:
            scores_by_fusion = {f: np.zeros(len(y), dtype=float) for f, _b in fusion_defs}
            for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
                y_train = y[train_idx]
                groups_train = groups[train_idx]
                raw_oof = np.zeros(len(train_idx), dtype=float)
                two_oof = np.zeros(len(train_idx), dtype=float)
                inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 101)
                for inner_fold, (tr_local, va_local) in enumerate(inner.split(np.zeros(len(train_idx)), y_train, groups=groups_train), start=1):
                    tr_idx = train_idx[tr_local]
                    va_idx = train_idx[va_local]
                    _tr_raw, raw_va, _n_syn = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, args.random_state + fold * 1000 + inner_fold)
                    _tr_2d, two_va, _st = fit_2d_multi_instance(
                        Xinst, yinst, inst_subjects, subjects, y, tr_idx, va_idx, args.random_state + fold * 2000 + inner_fold, args.top_k, model_name, agg
                    )
                    raw_oof[va_local] = raw_va
                    two_oof[va_local] = two_va
                _tr_raw, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
                _tr_2d, two_test, st = fit_2d_multi_instance(
                    Xinst, yinst, inst_subjects, subjects, y, train_idx, test_idx, args.random_state + fold + 500, args.top_k, model_name, agg
                )
                branch_oof = {"raw": raw_oof, "2d": two_oof}
                branch_test = {"raw": raw_test, "2d": two_test}
                for fusion_name, branches in fusion_defs:
                    best, candidates = search_fusion(branch_oof, y_train, branches, step=args.fusion_step)
                    for cand in candidates:
                        inner_rows.append({"model": model_name, "aggregation": agg, "fold": fold, "fusion": fusion_name, **cand})
                    fused = np.zeros(len(test_idx), dtype=float)
                    if "raw" in branches:
                        fused += float(best.get("weight_raw", 0.0)) * branch_test["raw"]
                    if "2d" in branches:
                        fused += float(best.get("weight_2d", 0.0)) * branch_test["2d"]
                    scores_by_fusion[fusion_name][test_idx] = fused
                    fold_rows.append(
                        {
                            "model": model_name,
                            "aggregation": agg,
                            "fold": fold,
                            "fusion": fusion_name,
                            "weight_raw": best.get("weight_raw", 0.0),
                            "weight_2d": best.get("weight_2d", 0.0),
                            "inner_AUROC": best["AUROC"],
                            "inner_AUPRC": best["AUPRC"],
                            "raw_synthetic_samples": n_syn,
                            **st,
                        }
                    )
            for fusion_name, score in scores_by_fusion.items():
                key = f"{model_name}::{agg}::{fusion_name}"
                result_scores[key] = score
                results.append({"model": model_name, "aggregation": agg, "fusion": fusion_name, **metric_dict(y, score)})

    results = sorted(results, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    write_csv(out_dir / "multi_instance_fusion_summary.csv", results, ["model", "aggregation", "fusion", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"])
    write_csv(out_dir / "fold_details.csv", fold_rows, ["model", "aggregation", "fold", "fusion", "weight_raw", "weight_2d", "inner_AUROC", "inner_AUPRC", "raw_synthetic_samples", "train_instances", "test_instances", "input_features", "after_variance_features", "selected_features"])
    write_csv(out_dir / "inner_fusion_search.csv", inner_rows, ["model", "aggregation", "fold", "fusion", "branches", "weight_raw", "weight_2d", "weight_vlm", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"])
    pred_rows = []
    for i, subject in enumerate(subjects):
        row = {"subject_id": subject, "true_label": int(y[i])}
        for key, score in result_scores.items():
            row[key] = float(score[i])
        pred_rows.append(row)
    write_csv(out_dir / "multi_instance_predictions.csv", pred_rows, ["subject_id", "true_label"] + sorted(result_scores))
    best = results[0]
    best_key = f"{best['model']}::{best['aggregation']}::{best['fusion']}"
    figures = make_figures(out_dir, results, y, result_scores[best_key], best)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold fusion search",
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "image_instances": int(len(instances)),
        "usable_image_instances": int(Xinst.shape[0]),
        "image_feature_count": int(Xinst.shape[1]),
        "grid": args.grid,
        "top_k": args.top_k,
        "models": models,
        "aggregations": aggs,
        "best": best,
        "target_reached": bool(best["AUROC"] >= 0.8 and best["AUPRC"] >= 0.8),
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "multi_instance_fusion_summary.csv"),
            "predictions": str(out_dir / "multi_instance_predictions.csv"),
            "fold_details": str(out_dir / "fold_details.csv"),
            "inner_fusion_search": str(out_dir / "inner_fusion_search.csv"),
            "report": str(Path(args.report).resolve()),
        },
        "figures": figures,
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, results, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
