#!/usr/bin/env python3
"""Package the NSC targeted tail-statistics bagged fusion candidate report."""

from __future__ import annotations

import argparse
import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_ANALYSIS = ROOT / "analysis" / "nsc_restricted_subject_bagging_tail_stats_20260521"
DEFAULT_REPORT = ROOT / "reports" / "NSC_targeted_tail_stats_bagged_fusion_candidate_20260521.docx"


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--analysis", type=Path, default=DEFAULT_ANALYSIS, help="Analysis output directory.")
    parser.add_argument("--report", type=Path, default=DEFAULT_REPORT, help="Output docx path.")
    parser.add_argument(
        "--corrected-data",
        action="store_true",
        help="Mark the report as the corrected-data rerun that supersedes the pre-correction result.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    analysis = args.analysis
    report = args.report
    if not analysis.is_absolute():
        analysis = ROOT / analysis
    if not report.is_absolute():
        report = ROOT / report

    with (analysis / "manifest.json").open("r", encoding="utf-8") as f:
        manifest = json.load(f)
    summary = read_csv(analysis / "bagged_metrics_summary.csv")[0]
    fold_rows = read_csv(analysis / "seed_fold_details.csv")

    report.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title = "NSC 114 位資料：Targeted Tail-statistics Bagged Fusion 候選方法報告"
    if args.corrected_data:
        title += "（corrected data rerun）"
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"產出日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、定位", level=1)
    if args.corrected_data:
        doc.add_paragraph(
            "本報告使用 corrected data 重新執行 targeted tail-statistics subject-level EEG fusion 實驗。"
            "corrected data 已刪除 NSC-REDACTED 的重複誤植列，並將編號 165 的 NSC-REDACTED 修正為 NSC-REDACTED。"
            "本報告結果取代修正前資料的候選結果。"
        )
    else:
        doc.add_paragraph(
            "本報告整理 Antigravity 新增的 targeted tail-statistics subject-level EEG fusion 實驗。"
        )
    doc.add_paragraph(
        "此方法在 patient-aware 10-fold OOF 架構下，使用 10 個 random seeds 產生 OOF 分數後做 bagging。"
    )
    doc.add_paragraph(
        "重要判讀：這是 optimized OOF candidate。由於參數組合經過多輪 grid search 探索，"
        "正式論文若要宣稱最終泛化能力，應先鎖定本 protocol 後再跑 confirmatory evaluation。"
    )

    doc.add_heading("二、方法參數", level=1)
    params = manifest["parameters"]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "內容"
    for key, value in [
        ("個案數", manifest["subjects"]),
        ("Seeds", ", ".join(str(x) for x in manifest["seeds_run"])),
        ("Base score", params["base_col"]),
        ("模型", f"Logistic Regression, penalty={params['penalty']}, C={params['C']}, class_weight=None"),
        ("特徵數", params["top_k"]),
        ("特徵工程", "delta/theta: p90/p95；alpha: p05/p10；其餘 EEG feature 使用標準 7 aggregation"),
        ("Fusion weight search", f"w_eeg <= {params['w_eeg_max']}, step={params['w_eeg_step']}, objective={params['objective']}"),
    ]:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    doc.add_heading("三、結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["Method", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    cells = table.add_row().cells
    vals = [
        summary["method"],
        f"{float(summary['AUROC']):.4f}",
        f"{float(summary['AUPRC']):.4f}",
        f"{float(summary['ACC']):.4f}",
        f"{float(summary['Sens']):.4f}",
        f"{float(summary['Spec']):.4f}",
        f"{float(summary['PPV']):.4f}",
        f"{float(summary['NPV']):.4f}",
        f"{summary['TN']}/{summary['FP']}/{summary['FN']}/{summary['TP']}",
    ]
    for cell, value in zip(cells, vals):
        cell.text = str(value)

    auroc = float(summary["AUROC"])
    auprc = float(summary["AUPRC"])
    if auroc >= 0.8 and auprc >= 0.8:
        result_text = (
            f"本結果 AUROC={auroc:.4f}、AUPRC={auprc:.4f}，兩項指標均跨過 0.80。"
            "相較 baseline mi_max AUROC=0.7835、AUPRC=0.7820，兩項指標均提升。"
        )
    else:
        result_text = (
            f"本結果 AUROC={auroc:.4f}、AUPRC={auprc:.4f}，接近但未同時跨過 0.80。"
            "相較 baseline mi_max AUROC=0.7835、AUPRC=0.7820，兩項指標仍有提升，"
            "但正式報告不可再寫成已達 0.8/0.8。"
        )
    if args.corrected_data:
        result_text += (
            " 修正前候選結果曾為 AUROC=0.7990、AUPRC=0.8011；"
            "因資料編號誤植已修正，後續應以本 corrected-data rerun 為準。"
        )
    doc.add_paragraph(result_text)

    doc.add_heading("四、圖表", level=1)
    roc = analysis / "figures" / "bagged_fusion_roc.png"
    prc = analysis / "figures" / "bagged_fusion_prc.png"
    if roc.exists():
        doc.add_picture(str(roc), width=Inches(5.5))
    if prc.exists():
        doc.add_picture(str(prc), width=Inches(5.5))

    doc.add_heading("五、Fold/Seed 穩定性摘要", level=1)
    selected_weights = [float(r["selected_w"]) for r in fold_rows if r.get("selected_w", "") != ""]
    if selected_weights:
        doc.add_paragraph(
            f"共 {len(selected_weights)} 個 seed-fold decisions；"
            f"平均 selected w_eeg={sum(selected_weights)/len(selected_weights):.3f}，"
            f"最小={min(selected_weights):.3f}，最大={max(selected_weights):.3f}。"
        )
    doc.add_paragraph(
        "這表示 EEG branch 以受限權重作為輔助分數來源，不是取代 mi_max baseline；"
        "tail statistics 的主要價值是捕捉 transient slow-wave 或 alpha suppression 類候選訊號。"
    )

    doc.add_heading("六、限制與下一步", level=1)
    for item in [
        "目前 AUROC 尚未正式超過 0.80；若 reviewer 要求 AUROC/AUPRC 同時超過 0.80，仍需下一步鎖定 protocol 後再優化或確認。",
        "此配置由多輪 grid search 找到，需標註為 optimized OOF candidate，不能直接寫成 untouched confirmatory result。",
        "下一步建議固定此參數組合，新增 bootstrap confidence interval、seed-wise variability 與 locked rerun 報告。",
        "若要補神經心理詮釋，需把 targeted tail features 納入 SHAP/feature importance 報告，而不是只解釋舊版 fusion_mi_max_LR_k256。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("七、輸出檔案", level=1)
    for key, value in manifest["outputs"].items():
        doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    doc.save(report)
    print(report)


if __name__ == "__main__":
    main()
