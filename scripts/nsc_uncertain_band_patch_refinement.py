#!/usr/bin/env python3
"""Raw uncertain-band gating with FP/FN patch-family refinement for NSC.

This script implements the next modelling step:

1. Train a raw initial-segment branch inside patient-aware folds.
2. Use only the outer training patients to identify raw-branch FP/FN patch
   families.
3. Train refined 2D branches on the selected patch families.
4. Search an uncertain-band gate on the outer training OOF scores:
      if raw_score is inside [low, high], blend in the 2D branch.
      otherwise keep the raw score.
5. Freeze the band/weight and evaluate on the outer test patients.

The outer test fold is never used for patch-family selection, band selection,
feature selection, model training, or augmentation.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import sys
import warnings
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.metrics import confusion_matrix, precision_recall_curve, roc_curve
from sklearn.model_selection import StratifiedGroupKFold

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_rule_learning_nsc_dataset_images import parse_feature_key  # noqa: E402
from nsc_multi_instance_2d_raw_fusion import collect_image_instances, fit_2d_multi_instance, instances_to_matrix  # noqa: E402
from patch_error_analysis_nsc import feature_stats, group_labels  # noqa: E402
from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences  # noqa: E402
from score_level_fusion_2d_raw_nsc import build_2d_matrix, fit_raw_dtw_et, metric_dict, write_csv  # noqa: E402
from stacked_multiscale_raw2d_nsc import transform_with_fold_local_filter  # noqa: E402

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)


def fit_patient_2d_et(
    X2d: np.ndarray,
    y: np.ndarray,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    seed: int,
    top_k: int,
) -> Tuple[np.ndarray, np.ndarray, dict]:
    Xtr = X2d[train_idx]
    Xte = X2d[test_idx]
    Xtr_s, Xtr_orig_s, Xte_s, stats = transform_with_fold_local_filter(Xtr, y[train_idx], Xtr, Xte, top_k=top_k)
    model = ExtraTreesClassifier(
        n_estimators=420,
        random_state=seed,
        class_weight="balanced",
        max_features="sqrt",
        min_samples_leaf=2,
        n_jobs=-1,
    )
    model.fit(Xtr_s, y[train_idx])
    return model.predict_proba(Xtr_orig_s)[:, 1], model.predict_proba(Xte_s)[:, 1], stats


def subset_matrix(X: np.ndarray, names: List[str], selected: Iterable[str]) -> Tuple[np.ndarray, List[str]]:
    wanted = set(selected)
    keep_idx = [i for i, name in enumerate(names) if name in wanted]
    if not keep_idx:
        keep_idx = list(range(len(names)))
    return X[:, keep_idx], [names[i] for i in keep_idx]


def parse_patch(name: str) -> Tuple[str, str, int, int] | None:
    try:
        return parse_feature_key(name)
    except Exception:
        return None


def expand_patch_family(seed_features: Iterable[str], all_names: List[str], radius: int) -> List[str]:
    parsed_all = []
    for name in all_names:
        parsed = parse_patch(name)
        if parsed is not None:
            parsed_all.append((name, *parsed))
    seeds = []
    for feat in seed_features:
        parsed = parse_patch(feat)
        if parsed is not None:
            seeds.append((feat, *parsed))
    selected = set()
    for _feat, channel, plot, row, col in seeds:
        for name, ch, pl, r, c in parsed_all:
            if ch == channel and pl == plot and abs(r - row) <= radius and abs(c - col) <= radius:
                selected.add(name)
    return sorted(selected)


def select_refined_patch_family(
    X2d: np.ndarray,
    names: List[str],
    y_train: np.ndarray,
    raw_oof: np.ndarray,
    top_error: int,
    top_class: int,
    radius: int,
    min_support: int,
) -> Tuple[List[str], List[dict]]:
    groups = group_labels(y_train, raw_oof, 0.5)
    stats = feature_stats(X2d, names, y_train, groups, min_support=min_support)
    if not stats:
        return list(names), []
    error_rank = sorted(
        stats,
        key=lambda r: (
            float(r["error_alignment_score"]),
            max(float(r["fp_toward_label1_score"]), float(r["fn_toward_label0_score"])),
            abs(float(r["class_effect_label1_minus_label0"])),
        ),
        reverse=True,
    )
    class_rank = sorted(stats, key=lambda r: abs(float(r["class_effect_label1_minus_label0"])), reverse=True)
    seed_features = [r["feature"] for r in error_rank[:top_error]] + [r["feature"] for r in class_rank[:top_class]]
    selected = expand_patch_family(seed_features, names, radius=radius)
    if not selected:
        selected = seed_features
    selected_set = set(selected)
    rows = []
    for r in stats:
        if r["feature"] in selected_set:
            rows.append({**r, "selected": 1})
    return selected, rows


def search_uncertain_gate(
    raw_score: np.ndarray,
    branch_score: np.ndarray,
    y_true: np.ndarray,
    low_grid: List[float],
    high_grid: List[float],
    weight_grid: List[float],
    min_band_cases: int,
) -> Tuple[dict, List[dict]]:
    rows: List[dict] = []
    raw_met = metric_dict(y_true, raw_score)
    best = {"gate_low": math.nan, "gate_high": math.nan, "weight_2d": 0.0, "band_cases": 0, "mode": "raw_only", **raw_met}
    rows.append(best)
    branch_met = metric_dict(y_true, branch_score)
    rows.append({"gate_low": -math.inf, "gate_high": math.inf, "weight_2d": 1.0, "band_cases": int(len(y_true)), "mode": "branch_only", **branch_met})
    if (branch_met["AUPRC"], branch_met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
        best = rows[-1]
    for low in low_grid:
        for high in high_grid:
            if high <= low:
                continue
            mask = (raw_score >= low) & (raw_score <= high)
            band_cases = int(mask.sum())
            if band_cases < min_band_cases:
                continue
            for weight in weight_grid:
                fused = raw_score.copy()
                fused[mask] = (1.0 - weight) * raw_score[mask] + weight * branch_score[mask]
                met = metric_dict(y_true, fused)
                row = {
                    "gate_low": float(low),
                    "gate_high": float(high),
                    "weight_2d": float(weight),
                    "band_cases": band_cases,
                    "mode": "uncertain_band",
                    **met,
                }
                rows.append(row)
                if (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
                    best = row
    return best, rows


def apply_gate(raw_score: np.ndarray, branch_score: np.ndarray, gate: dict) -> np.ndarray:
    if gate["mode"] == "raw_only":
        return raw_score.copy()
    if gate["mode"] == "branch_only":
        return branch_score.copy()
    low = float(gate["gate_low"])
    high = float(gate["gate_high"])
    weight = float(gate["weight_2d"])
    fused = raw_score.copy()
    mask = (raw_score >= low) & (raw_score <= high)
    fused[mask] = (1.0 - weight) * raw_score[mask] + weight * branch_score[mask]
    return fused


def make_figures(out_dir: Path, rows: List[dict], y: np.ndarray, best_score: np.ndarray, best: dict, feature_freq: List[dict]) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}

    top = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    labels = [f"{r['branch']} | {r['gate_mode']}" for r in top]
    ypos = np.arange(len(top))
    plt.figure(figsize=(11.5, max(4.8, 0.38 * len(top) + 1.2)))
    plt.barh(ypos + 0.18, [r["AUROC"] for r in top], height=0.34, label="AUROC")
    plt.barh(ypos - 0.18, [r["AUPRC"] for r in top], height=0.34, label="AUPRC")
    plt.axvline(0.8, linestyle="--", color="gray")
    plt.yticks(ypos, labels, fontsize=8)
    plt.xlabel("Metric")
    plt.title("Raw uncertain-band gating + 2D patch refinement")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    plt.gca().invert_yaxis()
    path = fig_dir / "uncertain_band_gating_metric_bar.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["bar"] = str(path)

    fpr, tpr, _ = roc_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={best['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Best uncertain-band ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    path = fig_dir / "best_uncertain_band_roc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["roc"] = str(path)

    precision, recall, _ = precision_recall_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={best['AUPRC']:.3f}")
    plt.axhline(float(y.mean()), linestyle="--", color="gray", label=f"prevalence={y.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Best uncertain-band PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    path = fig_dir / "best_uncertain_band_prc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["prc"] = str(path)

    cm = confusion_matrix(y, (best_score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Best uncertain-band confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    path = fig_dir / "best_uncertain_band_cm.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["cm"] = str(path)

    if feature_freq:
        top_feat = feature_freq[:20]
        labels = [r["feature"] for r in top_feat]
        vals = [int(r["selected_folds"]) for r in top_feat]
        plt.figure(figsize=(12, 5.8))
        plt.bar(np.arange(len(top_feat)), vals, color="#637a9f")
        plt.xticks(np.arange(len(top_feat)), labels, rotation=45, ha="right", fontsize=8)
        plt.ylabel("Selected fold count")
        plt.title("Most frequent refined patch-family features")
        plt.grid(axis="y", alpha=0.25)
        path = fig_dir / "refined_patch_family_frequency.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        paths["feature_freq"] = str(path)
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], fold_rows: List[dict], feature_freq: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC Raw Uncertain-band Gating 與 FP/FN Patch-family 精修報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版把 raw branch 視為主要訊號；只有 raw score 落在不確定區間時，才引入 2D branch 修正。"
        "不確定區間 [low, high] 與 2D 權重都在 outer training fold 的 inner OOF scores 上選定，"
        "再凍結套用到 outer test fold。"
    )
    doc.add_paragraph(
        "FP/FN patch-family 精修使用 outer training fold 的 raw OOF 錯誤：找出 FP 往 label 1 紋理偏移、"
        "FN 往 label 0 紋理偏移的 channel/plot/patch，並展開鄰近 patch family 作為 refined 2D branch。"
        "outer test fold 不參與 patch selection。"
    )
    doc.add_paragraph(f"驗證設計：{manifest['validation']}；Subjects={manifest['subjects']}；label counts={manifest['cases_by_label']}")

    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=13)
    table.style = "Table Grid"
    headers = ["branch", "gate", "AUROC", "AUPRC", "ACC", "Sens", "Spec", "PPV", "NPV", "CM", "avg low", "avg high", "avg w2d"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for row in rows:
        vals = [
            row["branch"],
            row["gate_mode"],
            f"{row['AUROC']:.3f}",
            f"{row['AUPRC']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['sensitivity']:.3f}",
            f"{row['specificity']:.3f}",
            f"{row['PPV']:.3f}",
            f"{row['NPV']:.3f}",
            f"{row['TN']}/{row['FP']}/{row['FN']}/{row['TP']}",
            row["avg_gate_low"],
            row["avg_gate_high"],
            row["avg_weight_2d"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、常見 refined patch-family", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["feature", "selected folds", "channel", "plot", "position"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for row in feature_freq[:24]:
        vals = [row["feature"], row["selected_folds"], row["channel"], row["plot"], row["position"]]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、圖表", level=1)
    for key in ["bar", "roc", "prc", "cm", "feature_freq"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.3 if key in {"bar", "feature_freq"} else 5.1))

    doc.add_heading("五、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(f"最佳結果為 {best['branch']} / {best['gate_mode']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。")
    if manifest["target_reached"]:
        doc.add_paragraph("此版本已在 patient-aware group validation 達到 0.8/0.8，下一步應做 random seed 重複與錯誤案例人工審查。")
    else:
        doc.add_paragraph(
            "此版本仍未達 0.8/0.8。若 gating 有提升但未達標，代表 2D branch 有局部互補訊號，但 false positive 控制仍不足；"
            "下一步應針對被選中的 patch family 做人工/VLM 紋理規則審查，或改用代表性 holdout 作 proof-of-signal。"
        )
    doc.add_heading("六、輸出檔案", level=1)
    for k, v in manifest["outputs"].items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def summarize_gate_folds(fold_rows: List[dict], branch: str) -> Tuple[str, str, str]:
    vals = [r for r in fold_rows if r["branch"] == branch]
    def avg(key: str) -> str:
        xs = [float(r[key]) for r in vals if r[key] not in {"", "nan", "inf", "-inf"} and math.isfinite(float(r[key]))]
        return f"{float(np.mean(xs)):.3f}" if xs else ""
    return avg("gate_low"), avg("gate_high"), avg("weight_2d")


def feature_frequency(selected_rows_by_fold: List[dict]) -> List[dict]:
    counts = Counter()
    meta = {}
    for row in selected_rows_by_fold:
        feat = row["feature"]
        counts[feat] += 1
        if feat not in meta:
            parsed = parse_patch(feat)
            if parsed is None:
                meta[feat] = ("", "", "")
            else:
                channel, plot, r, c = parsed
                meta[feat] = (channel, plot, f"r{r + 1}c{c + 1}")
    out = []
    for feat, count in counts.most_common():
        channel, plot, pos = meta[feat]
        out.append({"feature": feat, "selected_folds": int(count), "channel": channel, "plot": plot, "position": pos})
    return out


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="data/physiology-csv")
    parser.add_argument("--manifest", default="data/nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="data/nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_uncertain_band_patch_refinement_20260520")
    parser.add_argument("--report", default="reports/NSC_uncertain_band_patch_family_refinement_20260520.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--random-state", type=int, default=20260520)
    parser.add_argument("--outer-random-state", type=int, default=42)
    parser.add_argument("--patient-top-k", type=int, default=96)
    parser.add_argument("--mi-top-k", type=int, default=192)
    parser.add_argument("--top-error", type=int, default=24)
    parser.add_argument("--top-class", type=int, default=24)
    parser.add_argument("--family-radius", type=int, default=1)
    parser.add_argument("--fine-grid", action="store_true", help="Use finer uncertain-band and weight search grids.")
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(Path(args.csv_dir).resolve(), Path(args.manifest).resolve(), args.sequence_length, args.max_rows_per_csv)
    X2d, names2d = build_2d_matrix(Path(args.image_root).resolve(), subjects)
    instances, inst_names = collect_image_instances(Path(args.image_root).resolve(), grid=8)
    Xinst, yinst, inst_subjects = instances_to_matrix(instances, inst_names)
    allowed = set(subjects)
    keep = np.asarray([str(s) in allowed for s in inst_subjects], dtype=bool)
    Xinst, yinst, inst_subjects = Xinst[keep], yinst[keep], inst_subjects[keep]
    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(
        n_splits=args.n_splits,
        shuffle=True,
        random_state=args.outer_random_state,
    )

    branch_names = ["patient_all_2d", "patient_refined_fpfn_family", "mi_max", "mi_top3mean"]
    scores_by_branch = {branch: np.zeros(len(y), dtype=float) for branch in branch_names}
    raw_scores = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    gate_search_rows: List[dict] = []
    selected_feature_rows: List[dict] = []
    if args.fine_grid:
        low_grid = [round(x, 3) for x in np.arange(0.15, 0.501, 0.025)]
        high_grid = [round(x, 3) for x in np.arange(0.40, 0.801, 0.025)]
        weight_grid = [round(x, 3) for x in np.arange(0.05, 1.001, 0.05)]
    else:
        low_grid = [0.20, 0.25, 0.30, 0.35, 0.40, 0.45]
        high_grid = [0.45, 0.50, 0.55, 0.60, 0.65, 0.70, 0.75]
        weight_grid = [round(x, 2) for x in np.linspace(0.1, 1.0, 10)]

    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        y_train = y[train_idx]
        groups_train = groups[train_idx]
        raw_oof = np.zeros(len(train_idx), dtype=float)
        branch_oof = {branch: np.zeros(len(train_idx), dtype=float) for branch in branch_names}
        inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 101)
        # First pass: raw OOF only, used for fold-local FP/FN patch selection.
        for inner_fold, (tr_local, va_local) in enumerate(inner.split(np.zeros(len(train_idx)), y_train, groups=groups_train), start=1):
            tr_idx = train_idx[tr_local]
            va_idx = train_idx[va_local]
            _raw_tr, raw_va, _n_syn = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, args.random_state + fold * 1000 + inner_fold)
            raw_oof[va_local] = raw_va
        selected, selected_rows = select_refined_patch_family(
            X2d[train_idx],
            names2d,
            y_train,
            raw_oof,
            top_error=args.top_error,
            top_class=args.top_class,
            radius=args.family_radius,
            min_support=max(6, int(len(train_idx) * 0.25)),
        )
        for row in selected_rows:
            selected_feature_rows.append({"fold": fold, **row})
        X2d_refined, names_refined = subset_matrix(X2d, names2d, selected)
        # Second pass: branch OOF scores for gate selection.
        inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 103)
        for inner_fold, (tr_local, va_local) in enumerate(inner.split(np.zeros(len(train_idx)), y_train, groups=groups_train), start=1):
            tr_idx = train_idx[tr_local]
            va_idx = train_idx[va_local]
            _tr, te, _st = fit_patient_2d_et(X2d, y, tr_idx, va_idx, args.random_state + fold * 2000 + inner_fold, top_k=args.patient_top_k)
            branch_oof["patient_all_2d"][va_local] = te
            _tr, te, _st = fit_patient_2d_et(X2d_refined, y, tr_idx, va_idx, args.random_state + fold * 3000 + inner_fold, top_k=args.patient_top_k)
            branch_oof["patient_refined_fpfn_family"][va_local] = te
            _tr, te, _st = fit_2d_multi_instance(Xinst, yinst, inst_subjects, subjects, y, tr_idx, va_idx, args.random_state + fold * 4000 + inner_fold, args.mi_top_k, "ExtraTrees", "max")
            branch_oof["mi_max"][va_local] = te
            _tr, te, _st = fit_2d_multi_instance(Xinst, yinst, inst_subjects, subjects, y, tr_idx, va_idx, args.random_state + fold * 5000 + inner_fold, args.mi_top_k, "ExtraTrees", "top3mean")
            branch_oof["mi_top3mean"][va_local] = te

        raw_train, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
        raw_scores[test_idx] = raw_test
        branch_test: Dict[str, np.ndarray] = {}
        _tr, branch_test["patient_all_2d"], st_all = fit_patient_2d_et(X2d, y, train_idx, test_idx, args.random_state + fold + 110, top_k=args.patient_top_k)
        _tr, branch_test["patient_refined_fpfn_family"], st_refined = fit_patient_2d_et(X2d_refined, y, train_idx, test_idx, args.random_state + fold + 210, top_k=args.patient_top_k)
        _tr, branch_test["mi_max"], st_mi_max = fit_2d_multi_instance(Xinst, yinst, inst_subjects, subjects, y, train_idx, test_idx, args.random_state + fold + 310, args.mi_top_k, "ExtraTrees", "max")
        _tr, branch_test["mi_top3mean"], st_mi_top3 = fit_2d_multi_instance(Xinst, yinst, inst_subjects, subjects, y, train_idx, test_idx, args.random_state + fold + 410, args.mi_top_k, "ExtraTrees", "top3mean")
        stats_map = {
            "patient_all_2d": st_all,
            "patient_refined_fpfn_family": st_refined,
            "mi_max": st_mi_max,
            "mi_top3mean": st_mi_top3,
        }

        for branch in branch_names:
            best_gate, candidates = search_uncertain_gate(
                raw_oof,
                branch_oof[branch],
                y_train,
                low_grid=low_grid,
                high_grid=high_grid,
                weight_grid=weight_grid,
                min_band_cases=max(8, int(len(train_idx) * 0.10)),
            )
            for cand in candidates:
                gate_search_rows.append({"fold": fold, "branch": branch, **cand})
            fused_test = apply_gate(raw_test, branch_test[branch], best_gate)
            scores_by_branch[branch][test_idx] = fused_test
            fold_rows.append(
                {
                    "fold": fold,
                    "branch": branch,
                    "gate_mode": best_gate["mode"],
                    "gate_low": best_gate["gate_low"],
                    "gate_high": best_gate["gate_high"],
                    "weight_2d": best_gate["weight_2d"],
                    "band_cases": best_gate["band_cases"],
                    "inner_AUROC": best_gate["AUROC"],
                    "inner_AUPRC": best_gate["AUPRC"],
                    "raw_synthetic_samples": n_syn,
                    "refined_selected_features": len(names_refined),
                    **{f"branch_{k}": v for k, v in stats_map[branch].items()},
                }
            )

    results = [{"branch": "raw_only", "gate_mode": "raw_only", **metric_dict(y, raw_scores)}]
    result_scores = {"raw_only": raw_scores}
    for branch, score in scores_by_branch.items():
        rows = [r for r in fold_rows if r["branch"] == branch]
        gate_modes = Counter(r["gate_mode"] for r in rows).most_common()
        gate_mode = ";".join(f"{m}:{c}" for m, c in gate_modes)
        low, high, weight = summarize_gate_folds(fold_rows, branch)
        row = {
            "branch": branch,
            "gate_mode": gate_mode,
            **metric_dict(y, score),
            "avg_gate_low": low,
            "avg_gate_high": high,
            "avg_weight_2d": weight,
        }
        results.append(row)
        result_scores[branch] = score
    # Add missing summary columns for raw baseline.
    results[0].update({"avg_gate_low": "", "avg_gate_high": "", "avg_weight_2d": ""})
    results = sorted(results, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    feature_freq = feature_frequency(selected_feature_rows)
    best = results[0]
    best_score = result_scores[best["branch"]]
    figures = make_figures(out_dir, results, y, best_score, best, feature_freq)

    write_csv(
        out_dir / "uncertain_band_summary.csv",
        results,
        ["branch", "gate_mode", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP", "avg_gate_low", "avg_gate_high", "avg_weight_2d"],
    )
    write_csv(
        out_dir / "fold_details.csv",
        fold_rows,
        [
            "fold",
            "branch",
            "gate_mode",
            "gate_low",
            "gate_high",
            "weight_2d",
            "band_cases",
            "inner_AUROC",
            "inner_AUPRC",
            "raw_synthetic_samples",
            "refined_selected_features",
            "branch_input_features",
            "branch_after_variance_features",
            "branch_selected_features",
            "branch_train_instances",
            "branch_test_instances",
            "branch_model",
            "branch_aggregation",
        ],
    )
    write_csv(
        out_dir / "gate_search.csv",
        gate_search_rows,
        ["fold", "branch", "mode", "gate_low", "gate_high", "weight_2d", "band_cases", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"],
    )
    write_csv(out_dir / "selected_patch_family_frequency.csv", feature_freq, ["feature", "selected_folds", "channel", "plot", "position"])
    pred_rows = []
    for i, subject in enumerate(subjects):
        row = {"subject_id": subject, "true_label": int(y[i])}
        for key, score in result_scores.items():
            row[key] = float(score[i])
        pred_rows.append(row)
    write_csv(out_dir / "uncertain_band_predictions.csv", pred_rows, ["subject_id", "true_label"] + sorted(result_scores))
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold OOF raw/2D branch scores for gate selection",
        "random_state": int(args.random_state),
        "outer_random_state": int(args.outer_random_state),
        "subjects": int(len(subjects)),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "branch_names": branch_names,
        "best": best,
        "target_reached": bool(best["AUROC"] >= 0.8 and best["AUPRC"] >= 0.8),
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "uncertain_band_summary.csv"),
            "predictions": str(out_dir / "uncertain_band_predictions.csv"),
            "fold_details": str(out_dir / "fold_details.csv"),
            "gate_search": str(out_dir / "gate_search.csv"),
            "selected_patch_family_frequency": str(out_dir / "selected_patch_family_frequency.csv"),
            "report": str(Path(args.report).resolve()),
        },
        "figures": figures,
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, results, fold_rows, feature_freq, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
