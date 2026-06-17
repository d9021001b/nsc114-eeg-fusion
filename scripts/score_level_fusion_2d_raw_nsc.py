#!/usr/bin/env python3
"""Nested score-level fusion of 2D Green prototype and raw DTW/T-SMOTE branch.

This script is designed to keep 2D features central while using raw 1D
time-series as a complementary branch:

- 2D branch: group/channel/all-channel patch features -> fold-local selection
  -> fold-local PCA/Green proxy -> prototype score.
- Raw branch: initial raw time-series segment -> DTW-neighbor T-SMOTE-style
  augmentation inside the training fold -> ExtraTrees score.
- Fusion: raw and 2D branches are trained separately. Fusion weight is selected
  by inner validation inside each outer training fold, then frozen for the outer
  test fold.

No outer test sample is used for feature selection, augmentation, PCA, model
training, or fusion-weight selection.
"""

from __future__ import annotations

import argparse
import csv
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from sklearn.decomposition import PCA
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.impute import SimpleImputer
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedGroupKFold
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_rule_learning_nsc_dataset_images import build_case_features, collect_case_images, natural_key  # noqa: E402
from raw_timeseries_dtw_tsmote_nsc_dataset_images import dtw_tsmote_augment, extract_features, load_case_sequences  # noqa: E402


def write_csv(path: Path, rows: Iterable[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def build_2d_matrix(image_root: Path, subjects: List[str]) -> Tuple[np.ndarray, List[str]]:
    case_features = build_case_features(collect_case_images(image_root))
    by_subject = {v["case_id"]: v for v in case_features.values()}
    feature_names = sorted(
        {
            f
            for subject in subjects
            if subject in by_subject
            for f in by_subject[subject]["features"]
            if f.startswith("GROUP_") or f.startswith("ALL_CHANNELS/")
        }
    )
    idx = {f: i for i, f in enumerate(feature_names)}
    X = np.full((len(subjects), len(feature_names)), np.nan, dtype=float)
    for i, subject in enumerate(subjects):
        if subject not in by_subject:
            continue
        for f, v in by_subject[subject]["features"].items():
            if f in idx:
                X[i, idx[f]] = v
    return X, feature_names


def robust_unit(train_scores: np.ndarray, test_scores: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
    lo, hi = np.percentile(train_scores, [5, 95])
    if hi <= lo:
        lo, hi = float(np.min(train_scores)), float(np.max(train_scores))
    if hi <= lo:
        return np.full_like(train_scores, 0.5, dtype=float), np.full_like(test_scores, 0.5, dtype=float)
    return (
        np.clip((train_scores - lo) / (hi - lo), 0, 1),
        np.clip((test_scores - lo) / (hi - lo), 0, 1),
    )


def fit_2d_green_proto(X2d: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, top_k: int = 64) -> Tuple[np.ndarray, np.ndarray]:
    imp = SimpleImputer(strategy="median")
    Xtr = imp.fit_transform(X2d[train_idx])
    Xte = imp.transform(X2d[test_idx])
    selector = SelectKBest(f_classif, k=min(top_k, Xtr.shape[1]))
    Xtr = selector.fit_transform(Xtr, y[train_idx])
    Xte = selector.transform(Xte)
    scaler = StandardScaler()
    Xtr = scaler.fit_transform(Xtr)
    Xte = scaler.transform(Xte)
    n_components = min(16, Xtr.shape[0] - 2, Xtr.shape[1])
    pca = PCA(n_components=n_components, random_state=seed)
    Ztr = pca.fit_transform(Xtr)
    Zte = pca.transform(Xte)
    c0 = Ztr[y[train_idx] == 0].mean(axis=0)
    c1 = Ztr[y[train_idx] == 1].mean(axis=0)
    train_raw = np.linalg.norm(Ztr - c0, axis=1) - np.linalg.norm(Ztr - c1, axis=1)
    test_raw = np.linalg.norm(Zte - c0, axis=1) - np.linalg.norm(Zte - c1, axis=1)
    return robust_unit(train_raw, test_raw)


def fit_raw_dtw_et(Xseq: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int) -> Tuple[np.ndarray, np.ndarray, int]:
    X_train_seq, y_train, n_syn = dtw_tsmote_augment(Xseq[train_idx], y[train_idx], target_label=1, random_state=seed)
    Xtr = extract_features(X_train_seq)
    Xte = extract_features(Xseq[test_idx])
    Xtr_orig = extract_features(Xseq[train_idx])
    pipe = [
        ("imputer", SimpleImputer(strategy="median")),
        ("selector", SelectKBest(f_classif, k=min(48, Xtr.shape[1]))),
    ]
    imputer = pipe[0][1]
    selector = pipe[1][1]
    Xtr_i = imputer.fit_transform(Xtr)
    Xtr_s = selector.fit_transform(Xtr_i, y_train)
    Xte_s = selector.transform(imputer.transform(Xte))
    Xtr_orig_s = selector.transform(imputer.transform(Xtr_orig))
    model = ExtraTreesClassifier(
        n_estimators=260,
        random_state=seed,
        class_weight="balanced",
        max_features="sqrt",
        min_samples_leaf=2,
        n_jobs=-1,
    )
    model.fit(Xtr_s, y_train)
    train_score = model.predict_proba(Xtr_orig_s)[:, 1]
    test_score = model.predict_proba(Xte_s)[:, 1]
    return train_score, test_score, n_syn


def metric_dict(y_true: np.ndarray, score: np.ndarray, threshold: float = 0.5) -> dict:
    pred = (score >= threshold).astype(int)
    tn, fp, fn, tp = confusion_matrix(y_true, pred, labels=[0, 1]).ravel()
    return {
        "AUROC": float(roc_auc_score(y_true, score)),
        "AUPRC": float(average_precision_score(y_true, score)),
        "accuracy": float((tp + tn) / max(tp + tn + fp + fn, 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": int(tn),
        "FP": int(fp),
        "FN": int(fn),
        "TP": int(tp),
    }


def choose_weight_inner(
    Xseq: np.ndarray,
    X2d: np.ndarray,
    y: np.ndarray,
    groups: np.ndarray,
    outer_train_idx: np.ndarray,
    seed: int,
    inner_splits: int,
) -> Tuple[float, List[dict]]:
    weights = np.linspace(0, 1, 11)  # weight assigned to 2D branch.
    inner = StratifiedGroupKFold(n_splits=inner_splits, shuffle=True, random_state=seed)
    raw_scores = np.zeros(len(outer_train_idx), dtype=float)
    two_scores = np.zeros(len(outer_train_idx), dtype=float)
    inner_rows = []
    outer_train_groups = groups[outer_train_idx]
    for inner_fold, (tr_local, va_local) in enumerate(
        inner.split(np.zeros(len(outer_train_idx)), y[outer_train_idx], groups=outer_train_groups),
        start=1,
    ):
        tr_idx = outer_train_idx[tr_local]
        va_idx = outer_train_idx[va_local]
        _, raw_val, _ = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, seed + inner_fold)
        _, two_val = fit_2d_green_proto(X2d, y, tr_idx, va_idx, seed + inner_fold)
        raw_scores[va_local] = raw_val
        two_scores[va_local] = two_val
    best = None
    for w in weights:
        fused = w * two_scores + (1.0 - w) * raw_scores
        met = metric_dict(y[outer_train_idx], fused)
        row = {"weight_2d": float(w), **met}
        inner_rows.append(row)
        if best is None or (met["AUPRC"], met["AUROC"]) > (best[1]["AUPRC"], best[1]["AUROC"]):
            best = (float(w), met)
    return best[0], inner_rows


def make_figures(out_dir: Path, y_true: np.ndarray, score: np.ndarray, met: dict) -> Dict[str, Path]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    fpr, tpr, _ = roc_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={met['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Nested score-level fusion ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = fig_dir / "nested_score_fusion_roc.png"
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()
    precision, recall, _ = precision_recall_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={met['AUPRC']:.3f}")
    plt.axhline(float(y_true.mean()), linestyle="--", color="gray", label=f"prevalence={y_true.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Nested score-level fusion PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = fig_dir / "nested_score_fusion_prc.png"
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()
    cm = confusion_matrix(y_true, (score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Nested score-level fusion confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = fig_dir / "nested_score_fusion_cm.png"
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], fold_rows: List[dict], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC 2D + Raw Score-level Late Fusion 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版明確兼顧 2D 特徵：2D branch 使用 group/channel/all-channel patch features + Green/PCA prototype；"
        "raw branch 使用 initial-segment raw time-series + training-fold-only DTW/T-SMOTE + ExtraTrees。"
        "兩個 branch 各自訓練，最後只在 score level 做 late fusion。"
    )
    doc.add_paragraph(
        f"本版驗證設計為 {manifest['validation']}。每個 outer fold 的 fusion weight 由該 fold 的 training data "
        f"內部 {manifest['inner_splits']}-fold validation 選出，再凍結用於 outer test，避免用測試折調權重。"
    )
    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = table.add_row().cells
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("三、各折選出的 2D 權重", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for c, h in zip(table.rows[0].cells, ["fold", "weight_2d", "synthetic_samples", "train_cases", "test_cases", "group_overlap"]):
        c.text = h
    for r in fold_rows:
        cells = table.add_row().cells
        for c, v in zip(cells, [r["fold"], r["weight_2d"], r["synthetic_samples"], r["train_cases"], r["test_cases"], r["group_overlap"]]):
            c.text = str(v)
    doc.add_heading("四、Fusion 圖表", level=1)
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))
    doc.add_heading("五、判讀", level=1)
    doc.add_paragraph(
        "這版避免將 raw synthetic samples 與不存在的 synthetic 2D images 混在 feature-level 訓練，"
        "改用 score-level fusion 保留 2D branch 的獨立貢獻。若 fusion 仍未達 0.8/0.8，"
        f"代表 strict patient-aware group {manifest['outer_splits']}-fold 仍受個案異質性限制；"
        "但可以清楚比較 2D-only、raw-only 與 2D+raw 是否有互補。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_score_level_fusion_2d_raw_20260519")
    parser.add_argument("--report", default="reports/NSC_score_level_fusion_2D_raw_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, meta = load_case_sequences(Path(args.csv_dir).resolve(), Path(args.manifest).resolve(), args.sequence_length, args.max_rows_per_csv)
    X2d, feature_names_2d = build_2d_matrix(Path(args.image_root).resolve(), subjects)
    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)
    raw_all = np.zeros(len(y), dtype=float)
    two_all = np.zeros(len(y), dtype=float)
    fusion_all = np.zeros(len(y), dtype=float)
    fold_rows = []
    inner_rows_all = []
    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        weight_2d, inner_rows = choose_weight_inner(Xseq, X2d, y, groups, train_idx, args.random_state + fold * 100, args.inner_splits)
        for r in inner_rows:
            r["outer_fold"] = fold
            inner_rows_all.append(r)
        _, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
        _, two_test = fit_2d_green_proto(X2d, y, train_idx, test_idx, args.random_state + fold)
        raw_all[test_idx] = raw_test
        two_all[test_idx] = two_test
        fusion_all[test_idx] = weight_2d * two_test + (1.0 - weight_2d) * raw_test
        train_groups = set(groups[train_idx].tolist())
        test_groups = set(groups[test_idx].tolist())
        fold_rows.append(
            {
                "fold": fold,
                "weight_2d": float(weight_2d),
                "synthetic_samples": int(n_syn),
                "train_cases": int(len(train_idx)),
                "test_cases": int(len(test_idx)),
                "group_overlap": int(len(train_groups & test_groups)),
            }
        )
    rows = [
        {"method": "raw_initial_segment_dtw_tsmote_extratrees", **metric_dict(y, raw_all)},
        {"method": "2d_green_patch_prototype", **metric_dict(y, two_all)},
        {"method": "nested_score_fusion_raw_plus_2d", **metric_dict(y, fusion_all)},
    ]
    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    pred_rows = []
    for i, subject in enumerate(subjects):
        pred_rows.append({"subject_id": subject, "true_label": int(y[i]), "raw_score": float(raw_all[i]), "two_d_score": float(two_all[i]), "fusion_score": float(fusion_all[i])})
    row_fields = ["method", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "score_level_fusion_summary.csv", rows, row_fields)
    write_csv(out_dir / "score_level_fusion_predictions.csv", pred_rows, ["subject_id", "true_label", "raw_score", "two_d_score", "fusion_score"])
    write_csv(out_dir / "fusion_fold_weights.csv", fold_rows, ["fold", "weight_2d", "synthetic_samples", "train_cases", "test_cases", "group_overlap"])
    write_csv(out_dir / "inner_weight_search.csv", inner_rows_all, ["outer_fold", "weight_2d", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"])
    fusion_met = next(r for r in rows if r["method"] == "nested_score_fusion_raw_plus_2d")
    figures = make_figures(out_dir, y, fusion_all, fusion_met)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold fusion-weight selection",
        "outer_splits": args.n_splits,
        "inner_splits": args.inner_splits,
        "sequence_length": args.sequence_length,
        "max_rows_per_csv": args.max_rows_per_csv,
        "feature_names_2d": len(feature_names_2d),
        "best": rows[0],
        "outputs": {"analysis_dir": str(out_dir), "summary": str(out_dir / "score_level_fusion_summary.csv"), "predictions": str(out_dir / "score_level_fusion_predictions.csv"), "fold_weights": str(out_dir / "fusion_fold_weights.csv"), "inner_search": str(out_dir / "inner_weight_search.csv"), "report": str(Path(args.report).resolve()), "figures": {k: str(v) for k, v in figures.items()}},
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, fold_rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
