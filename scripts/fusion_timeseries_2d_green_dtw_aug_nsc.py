#!/usr/bin/env python3
"""Late fusion of segmented raw time-series, 2D patch/green features, and DTW augmentation.

This is the v2 experiment after the initial raw time-series DTW/T-SMOTE proxy:
- reads multiple evenly-spaced chunks from large raw CSV files;
- extracts raw time-series summary features;
- extracts 2D group/channel patch features with fold-local Green/PCA transform;
- applies train-fold-only DTW boundary/barycenter-style augmentation on raw sequences;
- compares raw, 2D green, and late-fusion models under patient-level 10-fold.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.decomposition import PCA
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_rule_learning_nsc_dataset_images import build_case_features, collect_case_images, natural_key  # noqa: E402
from raw_timeseries_dtw_tsmote_nsc_dataset_images import (  # noqa: E402
    GROUPS,
    dtw_distance,
    extract_features,
    flattened_for_dtw,
    normalize_sequence,
    robust_resample,
    signal_group,
    value_column,
)


def write_csv(path: Path, rows: Iterable[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def load_labels(manifest_path: Path) -> Dict[str, int]:
    labels = {}
    with manifest_path.open("r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            labels[str(row["subject_id"])] = int(row["label"])
    return labels


def count_rows_fast(path: Path) -> int:
    with path.open("rb") as f:
        return max(sum(1 for _ in f) - 1, 0)


def read_segmented_csv(path: Path, chunks: int, rows_per_chunk: int) -> pd.DataFrame:
    header = pd.read_csv(path, nrows=0)
    cols = list(header.columns)
    read_cols = ["time_sec"]
    for candidate in ["value", "ibi_ms", "resp_period_ms"]:
        if candidate in cols:
            read_cols.append(candidate)
            break
    if len(read_cols) == 1:
        read_cols = cols[:3]
    n_rows = count_rows_fast(path)
    if n_rows <= chunks * rows_per_chunk:
        return pd.read_csv(path, usecols=read_cols)
    starts = np.linspace(0, max(n_rows - rows_per_chunk, 0), chunks).astype(int)
    parts = []
    for start in starts:
        skip = range(1, int(start) + 1)
        part = pd.read_csv(path, usecols=read_cols, skiprows=skip, nrows=rows_per_chunk)
        parts.append(part)
    return pd.concat(parts, ignore_index=True)


def load_segmented_case_sequences(csv_dir: Path, manifest_path: Path, length: int, chunks: int, rows_per_chunk: int):
    labels = load_labels(manifest_path)
    grouped: Dict[str, Dict[str, List[np.ndarray]]] = defaultdict(lambda: defaultdict(list))
    file_counts: Dict[str, Counter] = defaultdict(Counter)
    pattern = re.compile(r"^(\d+)_Sess\d+_(.+)\.csv$")
    for path in sorted(csv_dir.glob("*.csv"), key=lambda p: natural_key(p.name)):
        match = pattern.match(path.name)
        if not match:
            continue
        subject, signal = match.group(1), match.group(2)
        if subject not in labels:
            continue
        group = signal_group(signal)
        if group not in GROUPS:
            continue
        try:
            df = read_segmented_csv(path, chunks=chunks, rows_per_chunk=rows_per_chunk)
            col = value_column(df)
            seq = robust_resample(df["time_sec"].to_numpy(dtype=float), df[col].to_numpy(dtype=float), length)
            grouped[subject][group].append(normalize_sequence(seq))
            file_counts[subject][group] += 1
        except Exception:
            continue
    subjects = sorted(grouped.keys(), key=natural_key)
    X_seq = np.zeros((len(subjects), len(GROUPS), length), dtype=np.float32)
    meta = {}
    for i, subject in enumerate(subjects):
        for g, group in enumerate(GROUPS):
            seqs = grouped[subject].get(group, [])
            if seqs:
                X_seq[i, g] = np.mean(np.stack(seqs), axis=0)
        meta[subject] = {"label": labels[subject], "file_counts": dict(file_counts[subject])}
    y = np.asarray([labels[s] for s in subjects], dtype=int)
    return subjects, X_seq, y, meta


def build_2d_patch_matrix(dataset_root: Path, subjects: List[str]):
    case_features = build_case_features(collect_case_images(dataset_root))
    by_subject = {v["case_id"]: v for v in case_features.values()}
    feature_names = sorted(
        {
            f
            for subject in subjects
            if subject in by_subject
            for f in by_subject[subject]["features"]
            if f.startswith("GROUP_") or f.startswith("ALL_CHANNELS/")
        }
    )
    idx = {f: i for i, f in enumerate(feature_names)}
    X = np.full((len(subjects), len(feature_names)), np.nan, dtype=float)
    for i, subject in enumerate(subjects):
        if subject not in by_subject:
            continue
        for f, v in by_subject[subject]["features"].items():
            if f in idx:
                X[i, idx[f]] = v
    return X, feature_names


def dtw_aug_boundary_barycenter(X_seq: np.ndarray, y: np.ndarray, seed: int):
    rng = np.random.default_rng(seed)
    counts = Counter(y.tolist())
    n_to_make = max(0, counts[0] - counts[1])
    if n_to_make == 0:
        return X_seq, y, 0
    pos_idx = np.where(y == 1)[0]
    neg_idx = np.where(y == 0)[0]
    pos_flat = [flattened_for_dtw(X_seq[i], reduced_length=32) for i in pos_idx]
    neg_flat = [flattened_for_dtw(X_seq[i], reduced_length=32) for i in neg_idx]
    # Boundary positives are closest to any negative.
    boundary_scores = []
    for i, pf in enumerate(pos_flat):
        ds = [dtw_distance(pf, nf) for nf in neg_flat]
        boundary_scores.append((min(ds), i))
    boundary_scores.sort(key=lambda x: x[0])
    boundary_pool = [idx for _, idx in boundary_scores[: max(4, len(boundary_scores) // 2)]]
    # Positive-positive DTW neighbor map within boundary pool.
    synthetic = []
    for _ in range(n_to_make):
        local_i = int(rng.choice(boundary_pool))
        ds = []
        for local_j in boundary_pool:
            if local_j == local_i:
                continue
            ds.append((dtw_distance(pos_flat[local_i], pos_flat[local_j]), local_j))
        ds.sort(key=lambda x: x[0])
        local_j = ds[0][1] if ds else local_i
        a = X_seq[pos_idx[local_i]]
        b = X_seq[pos_idx[local_j]]
        lam = float(rng.uniform(0.35, 0.65))
        synth = (1 - lam) * a + lam * b
        # Barycenter-style local smoothing.
        smooth = synth.copy()
        smooth[:, 1:-1] = 0.2 * synth[:, :-2] + 0.6 * synth[:, 1:-1] + 0.2 * synth[:, 2:]
        shift = int(rng.integers(-4, 5))
        smooth = np.roll(smooth, shift=shift, axis=1)
        synthetic.append(smooth.astype(np.float32))
    return np.concatenate([X_seq, np.stack(synthetic)], axis=0), np.concatenate([y, np.ones(len(synthetic), dtype=int)]), len(synthetic)


def select_transform_2d_green(X_train: np.ndarray, y_train: np.ndarray, X_test: np.ndarray, top_k: int, seed: int):
    imp = SimpleImputer(strategy="median")
    Xtr = imp.fit_transform(X_train)
    Xte = imp.transform(X_test)
    selector = SelectKBest(f_classif, k=min(top_k, Xtr.shape[1]))
    with np.errstate(invalid="ignore"):
        Xtr = selector.fit_transform(Xtr, y_train)
    Xte = selector.transform(Xte)
    scaler = StandardScaler()
    Xtr = scaler.fit_transform(Xtr)
    Xte = scaler.transform(Xte)
    n_components = min(16, Xtr.shape[0] - 2, Xtr.shape[1])
    pca = PCA(n_components=n_components, random_state=seed)
    return pca.fit_transform(Xtr), pca.transform(Xte)


def select_raw_features(X_train: np.ndarray, y_train: np.ndarray, X_test: np.ndarray):
    pipe = make_pipeline(SimpleImputer(strategy="median"), SelectKBest(f_classif, k=min(48, X_train.shape[1])), StandardScaler())
    with np.errstate(invalid="ignore"):
        Xtr = pipe.fit_transform(X_train, y_train)
    Xte = pipe.transform(X_test)
    return Xtr, Xte


def train_score(model_name: str, X_train: np.ndarray, y_train: np.ndarray, X_test: np.ndarray, seed: int):
    if model_name == "logistic":
        model = LogisticRegression(max_iter=4000, C=0.35, class_weight="balanced", solver="liblinear")
    elif model_name == "extratrees":
        model = ExtraTreesClassifier(
            n_estimators=300,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=2,
            n_jobs=-1,
        )
    else:
        raise ValueError(model_name)
    model.fit(X_train, y_train)
    score = model.predict_proba(X_test)[:, 1]
    return score, (score >= 0.5).astype(int)


def metrics(y_true, score, pred):
    tn, fp, fn, tp = confusion_matrix(y_true, pred, labels=[0, 1]).ravel()
    return {
        "AUROC": float(roc_auc_score(y_true, score)),
        "AUPRC": float(average_precision_score(y_true, score)),
        "accuracy": float((tp + tn) / max(tp + tn + fp + fn, 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": int(tn),
        "FP": int(fp),
        "FN": int(fn),
        "TP": int(tp),
    }


def plot_best(out_dir: Path, y, score, pred, title, met):
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    fpr, tpr, _ = roc_curve(y, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={met['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.title(f"Fusion ROC: {title}")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = fig_dir / "fusion_best_roc.png"
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()
    precision, recall, _ = precision_recall_curve(y, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={met['AUPRC']:.3f}")
    plt.axhline(float(np.mean(y)), linestyle="--", color="gray", label=f"prevalence={np.mean(y):.3f}")
    plt.title(f"Fusion PRC: {title}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = fig_dir / "fusion_best_prc.png"
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()
    cm = confusion_matrix(y, pred, labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title(f"Fusion CM: {title}")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = fig_dir / "fusion_best_confusion_matrix.png"
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest, rows, best, figures):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC Late Fusion + DTW Boundary Augmentation v2 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版同時使用 segmented raw time-series features、2D patch Green/PCA features，以及 late fusion。"
        "原始 CSV 採多段抽樣讀取，不再只讀取開頭片段；augmentation 採 training-fold-only DTW boundary/barycenter-style synthesis。"
    )
    doc.add_heading("二、結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["特徵", "augmentation", "模型", "合成樣本", "AUROC", "AUPRC", "ACC", "Sensitivity", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = table.add_row().cells
        vals = [
            r["feature_set"],
            r["augmentation"],
            r["model"],
            r["synthetic_samples"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("三、最佳圖表", level=1)
    doc.add_paragraph(f"最佳方法：{best['feature_set']} + {best['augmentation']} + {best['model']}")
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))
    doc.add_heading("四、判讀", level=1)
    doc.add_paragraph(
        "本版檢查完整方法路線是否能推升 strict patient-level 10-fold。若仍未達 0.8/0.8，"
        "可保留 random holdout 作 proof-of-signal，並等待新增個案後再做外部 frozen testing。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_fusion_timeseries_2d_green_dtw_aug_20260519")
    parser.add_argument("--report", default="reports/NSC_fusion_timeseries_2d_green_DTW_aug_v2_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--chunks", type=int, default=5)
    parser.add_argument("--rows-per-chunk", type=int, default=5000)
    parser.add_argument("--random-state", type=int, default=20260519)
    args = parser.parse_args()
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, X_seq, y, meta = load_segmented_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.chunks,
        args.rows_per_chunk,
    )
    X_img, img_feature_names = build_2d_patch_matrix(Path(args.image_root).resolve(), subjects)
    splitter = StratifiedKFold(n_splits=10, shuffle=True, random_state=42)
    configs = [
        ("raw", "none", "extratrees"),
        ("raw", "dtw_boundary", "extratrees"),
        ("green2d", "none", "logistic"),
        ("fusion", "none", "extratrees"),
        ("fusion", "dtw_boundary", "extratrees"),
        ("fusion", "dtw_boundary", "logistic"),
    ]
    rows = []
    long_preds = []
    best_payload = None
    for feature_set, augmentation, model_name in configs:
        all_score = np.zeros(len(y), dtype=float)
        all_pred = np.zeros(len(y), dtype=int)
        syn_total = 0
        for fold, (train_idx, test_idx) in enumerate(splitter.split(X_seq, y), start=1):
            X_train_seq, y_train = X_seq[train_idx], y[train_idx]
            X_test_seq = X_seq[test_idx]
            if augmentation == "dtw_boundary":
                X_train_seq_aug, y_train_aug, n_syn = dtw_aug_boundary_barycenter(X_train_seq, y_train, args.random_state + fold)
                syn_total += n_syn
            else:
                X_train_seq_aug, y_train_aug, n_syn = X_train_seq, y_train, 0
            raw_train = extract_features(X_train_seq_aug)
            raw_test = extract_features(X_test_seq)
            raw_train_t, raw_test_t = select_raw_features(raw_train, y_train_aug, raw_test)
            # 2D features cannot be synthesized by raw augmentation, so repeat selected 2D feature rows for synthetic samples by nearest raw anchor is not defensible.
            # For augmented fusion, train image branch on real rows only and append zeros for synthetic image part.
            img_train_green, img_test_green = select_transform_2d_green(X_img[train_idx], y_train, X_img[test_idx], 64, args.random_state + fold)
            if feature_set == "raw":
                train_mat, test_mat, train_y = raw_train_t, raw_test_t, y_train_aug
            elif feature_set == "green2d":
                train_mat, test_mat, train_y = img_train_green, img_test_green, y_train
            else:
                if augmentation == "dtw_boundary":
                    real_raw_count = len(train_idx)
                    pad_img = np.zeros((len(y_train_aug) - real_raw_count, img_train_green.shape[1]), dtype=float)
                    img_train_aug = np.vstack([img_train_green, pad_img])
                    train_mat = np.hstack([raw_train_t, img_train_aug])
                    train_y = y_train_aug
                else:
                    train_mat = np.hstack([raw_train_t, img_train_green])
                    train_y = y_train
                test_mat = np.hstack([raw_test_t, img_test_green])
            score, pred = train_score(model_name, train_mat, train_y, test_mat, args.random_state + fold)
            all_score[test_idx] = score
            all_pred[test_idx] = pred
        met = metrics(y, all_score, all_pred)
        row = {"feature_set": feature_set, "augmentation": augmentation, "model": model_name, "synthetic_samples": syn_total, **met}
        rows.append(row)
        for i, subject in enumerate(subjects):
            long_preds.append({"feature_set": feature_set, "augmentation": augmentation, "model": model_name, "subject_id": subject, "true_label": int(y[i]), "score": float(all_score[i]), "pred_label": int(all_pred[i])})
        if best_payload is None or (row["AUPRC"], row["AUROC"]) > (best_payload["row"]["AUPRC"], best_payload["row"]["AUROC"]):
            best_payload = {"row": row, "score": all_score.copy(), "pred": all_pred.copy()}
    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    row_fields = ["feature_set", "augmentation", "model", "synthetic_samples", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    pred_fields = ["feature_set", "augmentation", "model", "subject_id", "true_label", "score", "pred_label"]
    write_csv(out_dir / "fusion_method_summary.csv", rows, row_fields)
    write_csv(out_dir / "fusion_predictions_long.csv", long_preds, pred_fields)
    best = best_payload["row"]
    figures = plot_best(out_dir, y, best_payload["score"], best_payload["pred"], f"{best['feature_set']}+{best['augmentation']}+{best['model']}", best)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": "patient-level stratified 10-fold",
        "sequence_length": args.sequence_length,
        "chunks": args.chunks,
        "rows_per_chunk": args.rows_per_chunk,
        "best": best,
        "outputs": {"analysis_dir": str(out_dir), "summary": str(out_dir / "fusion_method_summary.csv"), "predictions": str(out_dir / "fusion_predictions_long.csv"), "report": str(Path(args.report).resolve()), "figures": {k: str(v) for k, v in figures.items()}},
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, best, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
