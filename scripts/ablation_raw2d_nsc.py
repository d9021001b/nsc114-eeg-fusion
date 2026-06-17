#!/usr/bin/env python3
"""Ablation for NSC raw + 2D patient-aware group 10-fold modelling."""

from __future__ import annotations

import argparse
import csv
import json
import math
import sys
import warnings
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from sklearn.decomposition import PCA
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, VarianceThreshold, f_classif
from sklearn.impute import SimpleImputer
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedGroupKFold
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences  # noqa: E402
from score_level_fusion_2d_raw_nsc import (  # noqa: E402
    build_2d_matrix,
    fit_2d_green_proto,
    fit_raw_dtw_et,
    metric_dict,
    robust_unit,
    write_csv,
)
from stacked_multiscale_raw2d_nsc import (  # noqa: E402
    build_2d_multiscale_matrix,
    fit_2d_multiscale_et,
    fit_meta_model,
    fit_raw_multiwindow_et,
    stack_features,
    transform_with_fold_local_filter,
)

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)


BRANCHES = [
    "raw_initial",
    "raw_multiwindow",
    "2d_green8",
    "2d_green8_var",
    "2d_var_et8",
    "2d_multiscale_et",
]

PAIR_DEFS = [
    ("baseline_grid_raw_initial_green8", "raw_initial", "2d_green8", "grid"),
    ("raw_multiwindow_only_change_grid", "raw_multiwindow", "2d_green8", "grid"),
    ("2d_variance_only_change_grid", "raw_initial", "2d_green8_var", "grid"),
    ("2d_et8_change_grid", "raw_initial", "2d_var_et8", "grid"),
    ("stacking_only_change", "raw_initial", "2d_green8", "stacking"),
    ("raw_multiwindow_plus_2d_var_grid", "raw_multiwindow", "2d_green8_var", "grid"),
    ("raw_multiwindow_plus_2d_et8_grid", "raw_multiwindow", "2d_var_et8", "grid"),
    ("raw_initial_plus_multiscale_grid", "raw_initial", "2d_multiscale_et", "grid"),
    ("raw_multiwindow_plus_multiscale_grid", "raw_multiwindow", "2d_multiscale_et", "grid"),
    ("raw_multiwindow_plus_multiscale_stacking", "raw_multiwindow", "2d_multiscale_et", "stacking"),
]


def fit_2d_green_proto_var(X2d: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, top_k: int = 64) -> Tuple[np.ndarray, np.ndarray, dict]:
    imp = SimpleImputer(strategy="median")
    Xtr = imp.fit_transform(X2d[train_idx])
    Xte = imp.transform(X2d[test_idx])
    var = VarianceThreshold(threshold=1e-8)
    Xtr_v = var.fit_transform(Xtr)
    Xte_v = var.transform(Xte)
    if Xtr_v.shape[1] == 0:
        Xtr_v, Xte_v = Xtr, Xte
    selector = SelectKBest(f_classif, k=min(top_k, Xtr_v.shape[1]))
    Xtr_s = selector.fit_transform(Xtr_v, y[train_idx])
    Xte_s = selector.transform(Xte_v)
    scaler = StandardScaler()
    Xtr_s = scaler.fit_transform(Xtr_s)
    Xte_s = scaler.transform(Xte_s)
    n_components = min(16, Xtr_s.shape[0] - 2, Xtr_s.shape[1])
    pca = PCA(n_components=n_components, random_state=seed)
    Ztr = pca.fit_transform(Xtr_s)
    Zte = pca.transform(Xte_s)
    c0 = Ztr[y[train_idx] == 0].mean(axis=0)
    c1 = Ztr[y[train_idx] == 1].mean(axis=0)
    train_raw = np.linalg.norm(Ztr - c0, axis=1) - np.linalg.norm(Ztr - c1, axis=1)
    test_raw = np.linalg.norm(Zte - c0, axis=1) - np.linalg.norm(Zte - c1, axis=1)
    train_score, test_score = robust_unit(train_raw, test_raw)
    stats = {
        "input_features": int(X2d.shape[1]),
        "after_variance_features": int(Xtr_v.shape[1]),
        "selected_features": int(min(top_k, Xtr_v.shape[1])),
    }
    return train_score, test_score, stats


def fit_2d_var_et8(X2d: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, top_k: int = 96) -> Tuple[np.ndarray, np.ndarray, dict]:
    Xtr = X2d[train_idx]
    Xte = X2d[test_idx]
    Xtr_s, Xtr_orig_s, Xte_s, stats = transform_with_fold_local_filter(Xtr, y[train_idx], Xtr, Xte, top_k=top_k)
    model = ExtraTreesClassifier(
        n_estimators=360,
        random_state=seed,
        class_weight="balanced",
        max_features="sqrt",
        min_samples_leaf=2,
        n_jobs=-1,
    )
    model.fit(Xtr_s, y[train_idx])
    return model.predict_proba(Xtr_orig_s)[:, 1], model.predict_proba(Xte_s)[:, 1], stats


def fit_all_branches(
    Xseq: np.ndarray,
    X2d_base: np.ndarray,
    X2d_multi: np.ndarray,
    y: np.ndarray,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    seed: int,
    raw_top_k: int,
    two_d_top_k: int,
) -> Tuple[Dict[str, np.ndarray], Dict[str, np.ndarray], Dict[str, dict]]:
    train_scores: Dict[str, np.ndarray] = {}
    test_scores: Dict[str, np.ndarray] = {}
    stats: Dict[str, dict] = {}

    tr, te, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, seed)
    train_scores["raw_initial"], test_scores["raw_initial"] = tr, te
    stats["raw_initial"] = {"synthetic_samples": int(n_syn)}

    tr, te, n_syn, st = fit_raw_multiwindow_et(Xseq, y, train_idx, test_idx, seed + 11, raw_top_k)
    train_scores["raw_multiwindow"], test_scores["raw_multiwindow"] = tr, te
    stats["raw_multiwindow"] = {"synthetic_samples": int(n_syn), **st}

    tr, te = fit_2d_green_proto(X2d_base, y, train_idx, test_idx, seed + 21)
    train_scores["2d_green8"], test_scores["2d_green8"] = tr, te
    stats["2d_green8"] = {"input_features": int(X2d_base.shape[1]), "selected_features": 64}

    tr, te, st = fit_2d_green_proto_var(X2d_base, y, train_idx, test_idx, seed + 31)
    train_scores["2d_green8_var"], test_scores["2d_green8_var"] = tr, te
    stats["2d_green8_var"] = st

    tr, te, st = fit_2d_var_et8(X2d_base, y, train_idx, test_idx, seed + 41, top_k=96)
    train_scores["2d_var_et8"], test_scores["2d_var_et8"] = tr, te
    stats["2d_var_et8"] = st

    tr, te, st = fit_2d_multiscale_et(X2d_multi, y, train_idx, test_idx, seed + 51, top_k=two_d_top_k)
    train_scores["2d_multiscale_et"], test_scores["2d_multiscale_et"] = tr, te
    stats["2d_multiscale_et"] = st

    return train_scores, test_scores, stats


def search_grid(score_a: np.ndarray, score_b: np.ndarray, y_true: np.ndarray) -> Tuple[dict, List[dict]]:
    rows: List[dict] = []
    best = None
    for weight_b in np.linspace(0, 1, 11):
        score = (1.0 - float(weight_b)) * score_a + float(weight_b) * score_b
        met = metric_dict(y_true, score)
        row = {"weight_a": 1.0 - float(weight_b), "weight_b": float(weight_b), **met}
        rows.append(row)
        if best is None or (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
            best = row
    return best, rows


def make_ablation_figures(out_dir: Path, rows: List[dict], y_true: np.ndarray, best_score: np.ndarray, best: dict) -> Dict[str, Path]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, Path] = {}

    top = sorted(rows, key=lambda r: r["AUPRC"], reverse=True)[:12]
    labels = [r["method"].replace("_", "\n") for r in top]
    x = np.arange(len(top))
    width = 0.38
    plt.figure(figsize=(13, 5.8))
    plt.bar(x - width / 2, [r["AUROC"] for r in top], width, label="AUROC")
    plt.bar(x + width / 2, [r["AUPRC"] for r in top], width, label="AUPRC")
    plt.axhline(0.8, linestyle="--", color="gray", linewidth=1)
    plt.xticks(x, labels, rotation=0, fontsize=8)
    plt.ylim(0, 1)
    plt.ylabel("Metric")
    plt.title("Raw + 2D ablation metrics")
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    paths["bar"] = fig_dir / "ablation_metric_bar.png"
    plt.tight_layout()
    plt.savefig(paths["bar"], dpi=180)
    plt.close()

    fpr, tpr, _ = roc_curve(y_true, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={best['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title(f"Best ablation ROC: {best['method']}")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = fig_dir / "best_ablation_roc.png"
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()

    precision, recall, _ = precision_recall_curve(y_true, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={best['AUPRC']:.3f}")
    plt.axhline(float(y_true.mean()), linestyle="--", color="gray", label=f"prevalence={y_true.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title(f"Best ablation PRC: {best['method']}")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = fig_dir / "best_ablation_prc.png"
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()

    cm = confusion_matrix(y_true, (best_score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Best ablation confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = fig_dir / "best_ablation_cm.png"
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], fold_rows: List[dict], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC Raw + 2D Ablation 技術報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、目的與驗證設計", level=1)
    doc.add_paragraph(
        "本次 ablation 拆解前一版下降的原因，固定採 patient-aware StratifiedGroupKFold 10-fold。"
        "同一 case_id/subject_id 不跨 train/test；所有 imputation、variance filtering、feature selection、DTW/T-SMOTE、"
        "branch model 與 score-level fusion/stacking 均只在 training fold 內完成。"
    )
    doc.add_paragraph(
        "比較因素包括：raw initial vs raw multi-window、2D Green/PCA prototype vs 2D fold-local variance filter/ExtraTrees/"
        "multi-scale patch、以及 grid score fusion vs calibrated logistic stacking。"
    )

    doc.add_heading("二、Ablation 整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in rows[:14]:
        cells = table.add_row().cells
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、各折 Fusion/Stacking 摘要", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["fold", "method", "strategy", "branch_a", "branch_b", "weight_b/meta", "test", "overlap"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in fold_rows[:40]:
        meta = r.get("weight_b", "")
        if r["strategy"] == "stacking":
            meta = f"{r.get('meta_kind', '')}; C={r.get('meta_C', '')}; {r.get('meta_class_weight', '')}"
        vals = [r["fold"], r["method"], r["strategy"], r["branch_a"], r["branch_b"], meta, r["test_cases"], r["group_overlap"]]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、圖表", level=1)
    doc.add_picture(str(figures["bar"]), width=Inches(6.6))
    doc.add_picture(str(figures["roc"]), width=Inches(5.2))
    doc.add_picture(str(figures["prc"]), width=Inches(5.2))
    doc.add_picture(str(figures["cm"]), width=Inches(4.7))

    doc.add_heading("五、結論", level=1)
    best = manifest["best"]
    doc.add_paragraph(
        f"最佳方法為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。"
        "本次 ablation 的重點不是只追求單一最高分，而是定位造成下降的因素。"
    )
    doc.add_paragraph(
        "若 raw multi-window 或 2D multi-scale 未優於基準，代表額外特徵在目前樣本數下增加了不穩定性；"
        "若 stacking 未優於 grid fusion，代表目前 branch score 的 out-of-fold 訊號不足以支撐較複雜的校準器。"
        "後續應保留能穩定提升 AUPRC 的單一改動，再逐步組合。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_raw2d_ablation_group10fold_20260519")
    parser.add_argument("--report", default="reports/NSC_raw_2D_ablation_patient_aware_group10fold_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--raw-top-k", type=int, default=96)
    parser.add_argument("--two-d-top-k", type=int, default=160)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.max_rows_per_csv,
    )
    X2d_base, _feature_names_base = build_2d_matrix(Path(args.image_root).resolve(), subjects)
    X2d_multi, feature_names_multi, meta_2d = build_2d_multiscale_matrix(Path(args.image_root).resolve(), subjects)
    groups = np.asarray(subjects)

    branch_scores = {name: np.zeros(len(y), dtype=float) for name in BRANCHES}
    pair_scores = {name: np.zeros(len(y), dtype=float) for name, _a, _b, _strategy in PAIR_DEFS}
    fold_rows: List[dict] = []
    branch_fold_rows: List[dict] = []
    inner_rows: List[dict] = []

    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)
    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        branch_oof = {name: np.zeros(len(train_idx), dtype=float) for name in BRANCHES}
        inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 100)
        for inner_fold, (tr_local, va_local) in enumerate(
            inner.split(np.zeros(len(train_idx)), y[train_idx], groups=groups[train_idx]),
            start=1,
        ):
            tr_idx = train_idx[tr_local]
            va_idx = train_idx[va_local]
            _tr_scores, va_scores, _stats = fit_all_branches(
                Xseq,
                X2d_base,
                X2d_multi,
                y,
                tr_idx,
                va_idx,
                args.random_state + fold * 100 + inner_fold,
                args.raw_top_k,
                args.two_d_top_k,
            )
            for name in BRANCHES:
                branch_oof[name][va_local] = va_scores[name]

        _train_scores, test_scores, branch_stats = fit_all_branches(
            Xseq,
            X2d_base,
            X2d_multi,
            y,
            train_idx,
            test_idx,
            args.random_state + fold,
            args.raw_top_k,
            args.two_d_top_k,
        )
        for name in BRANCHES:
            branch_scores[name][test_idx] = test_scores[name]
            st = branch_stats.get(name, {})
            branch_fold_rows.append(
                {
                    "fold": fold,
                    "branch": name,
                    "train_cases": len(train_idx),
                    "test_cases": len(test_idx),
                    "synthetic_samples": st.get("synthetic_samples", ""),
                    "input_features": st.get("input_features", ""),
                    "after_variance_features": st.get("after_variance_features", ""),
                    "selected_features": st.get("selected_features", ""),
                }
            )

        train_groups = set(groups[train_idx].tolist())
        test_groups = set(groups[test_idx].tolist())
        for method, branch_a, branch_b, strategy in PAIR_DEFS:
            if strategy == "grid":
                best, candidates = search_grid(branch_oof[branch_a], branch_oof[branch_b], y[train_idx])
                for row in candidates:
                    inner_rows.append({"outer_fold": fold, "method": method, "strategy": strategy, "branch_a": branch_a, "branch_b": branch_b, **row})
                score = (1.0 - best["weight_b"]) * test_scores[branch_a] + best["weight_b"] * test_scores[branch_b]
                pair_scores[method][test_idx] = score
                fold_rows.append(
                    {
                        "fold": fold,
                        "method": method,
                        "strategy": strategy,
                        "branch_a": branch_a,
                        "branch_b": branch_b,
                        "weight_b": best["weight_b"],
                        "meta_kind": "",
                        "meta_C": "",
                        "meta_class_weight": "",
                        "test_cases": len(test_idx),
                        "group_overlap": int(len(train_groups & test_groups)),
                    }
                )
            elif strategy == "stacking":
                model, best_meta, candidates = fit_meta_model(branch_oof[branch_a], branch_oof[branch_b], y[train_idx])
                for row in candidates:
                    inner_rows.append({"outer_fold": fold, "method": method, "strategy": strategy, "branch_a": branch_a, "branch_b": branch_b, **row})
                X_test_meta = stack_features(test_scores[branch_a], test_scores[branch_b], best_meta["meta_kind"])
                pair_scores[method][test_idx] = model.predict_proba(X_test_meta)[:, 1]
                fold_rows.append(
                    {
                        "fold": fold,
                        "method": method,
                        "strategy": strategy,
                        "branch_a": branch_a,
                        "branch_b": branch_b,
                        "weight_b": "",
                        "meta_kind": best_meta["meta_kind"],
                        "meta_C": best_meta["meta_C"],
                        "meta_class_weight": best_meta["meta_class_weight"],
                        "test_cases": len(test_idx),
                        "group_overlap": int(len(train_groups & test_groups)),
                    }
                )
            else:
                raise ValueError(strategy)

    result_rows = []
    for name in BRANCHES:
        result_rows.append({"method": name, "kind": "branch", **metric_dict(y, branch_scores[name])})
    for name, _a, _b, strategy in PAIR_DEFS:
        result_rows.append({"method": name, "kind": strategy, **metric_dict(y, pair_scores[name])})
    result_rows = sorted(result_rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)

    pred_rows = []
    for i, subject in enumerate(subjects):
        row = {"subject_id": subject, "true_label": int(y[i])}
        for name in BRANCHES:
            row[name] = float(branch_scores[name][i])
        for name in pair_scores:
            row[name] = float(pair_scores[name][i])
        pred_rows.append(row)

    row_fields = ["method", "kind", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "ablation_summary.csv", result_rows, row_fields)
    write_csv(out_dir / "ablation_predictions.csv", pred_rows, ["subject_id", "true_label"] + BRANCHES + list(pair_scores.keys()))
    write_csv(
        out_dir / "fold_fusion_details.csv",
        fold_rows,
        ["fold", "method", "strategy", "branch_a", "branch_b", "weight_b", "meta_kind", "meta_C", "meta_class_weight", "test_cases", "group_overlap"],
    )
    write_csv(
        out_dir / "branch_fold_details.csv",
        branch_fold_rows,
        ["fold", "branch", "train_cases", "test_cases", "synthetic_samples", "input_features", "after_variance_features", "selected_features"],
    )
    write_csv(
        out_dir / "inner_fusion_search.csv",
        inner_rows,
        ["outer_fold", "method", "strategy", "branch_a", "branch_b", "weight_a", "weight_b", "meta_kind", "meta_C", "meta_class_weight", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"],
    )

    best = result_rows[0]
    best_score = branch_scores[best["method"]] if best["method"] in branch_scores else pair_scores[best["method"]]
    figures = make_ablation_figures(out_dir, result_rows, y, best_score, best)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold fusion/stacking selection",
        "outer_splits": args.n_splits,
        "inner_splits": args.inner_splits,
        "branches": BRANCHES,
        "pairs": [p[0] for p in PAIR_DEFS],
        "base_2d_feature_count": int(X2d_base.shape[1]),
        "multiscale_2d_feature_count": len(feature_names_multi),
        "two_d_meta": meta_2d,
        "best": best,
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "ablation_summary.csv"),
            "predictions": str(out_dir / "ablation_predictions.csv"),
            "fold_fusion_details": str(out_dir / "fold_fusion_details.csv"),
            "branch_fold_details": str(out_dir / "branch_fold_details.csv"),
            "inner_fusion_search": str(out_dir / "inner_fusion_search.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": {k: str(v) for k, v in figures.items()},
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, result_rows, fold_rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
