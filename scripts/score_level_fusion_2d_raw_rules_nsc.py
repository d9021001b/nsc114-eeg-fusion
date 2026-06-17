#!/usr/bin/env python3
"""Patient-aware score-level fusion of raw, 2D, and heuristic-rule branches."""

from __future__ import annotations

import argparse
import csv
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import numpy as np
from sklearn.model_selection import StratifiedGroupKFold

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_5fold_validation_nsc_dataset_images import learn_rules_within_train  # noqa: E402
from heuristic_rule_learning_nsc_dataset_images import build_case_features, classify, collect_case_images, natural_key  # noqa: E402
from score_level_fusion_2d_raw_nsc import (  # noqa: E402
    build_2d_matrix,
    fit_2d_green_proto,
    fit_raw_dtw_et,
    make_figures,
    metric_dict,
    write_csv,
)
from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences  # noqa: E402


def build_rule_context(image_root: Path, subjects: List[str]) -> Tuple[Dict[str, dict], np.ndarray]:
    case_features = build_case_features(collect_case_images(image_root))
    subject_to_key = {v["case_id"]: k for k, v in case_features.items()}
    missing = [s for s in subjects if s not in subject_to_key]
    if missing:
        raise ValueError(f"Missing image-derived rule features for subjects: {missing[:10]}")
    return case_features, np.asarray([subject_to_key[s] for s in subjects])


def rule_scores(case_features: Dict[str, dict], case_keys: np.ndarray, rules: List) -> Tuple[np.ndarray, np.ndarray, np.ndarray]:
    scores = np.zeros(len(case_keys), dtype=float)
    confidences = np.zeros(len(case_keys), dtype=float)
    used_rules = np.zeros(len(case_keys), dtype=int)
    for i, key in enumerate(case_keys):
        pred = classify(case_features[key]["features"], rules)
        scores[i] = float(np.clip((pred["score"] + 1.0) / 2.0, 0.0, 1.0))
        confidences[i] = float(pred["confidence"])
        used_rules[i] = int(pred["used_rules"])
    return scores, confidences, used_rules


def fit_rule_branch(
    case_features: Dict[str, dict],
    case_keys: np.ndarray,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    confidence_threshold: float,
    max_rules: int,
) -> Tuple[np.ndarray, np.ndarray, dict]:
    train_keys = [str(case_keys[i]) for i in train_idx]
    train_case_features = {k: case_features[k] for k in train_keys}
    rules, _train_trials, stats = learn_rules_within_train(
        train_case_features,
        train_keys,
        confidence_threshold=confidence_threshold,
        max_rules=max_rules,
    )
    train_scores, train_conf, train_used = rule_scores(case_features, case_keys[train_idx], rules)
    test_scores, test_conf, test_used = rule_scores(case_features, case_keys[test_idx], rules)
    stats = {
        **stats,
        "mean_train_rule_confidence": float(train_conf.mean()) if len(train_conf) else 0.0,
        "mean_test_rule_confidence": float(test_conf.mean()) if len(test_conf) else 0.0,
        "mean_train_used_rules": float(train_used.mean()) if len(train_used) else 0.0,
        "mean_test_used_rules": float(test_used.mean()) if len(test_used) else 0.0,
    }
    return train_scores, test_scores, stats


def search_two_branch(raw_scores: np.ndarray, two_scores: np.ndarray, y_true: np.ndarray) -> Tuple[dict, List[dict]]:
    best = None
    rows: List[dict] = []
    for weight_2d in np.linspace(0, 1, 11):
        weight_raw = 1.0 - float(weight_2d)
        fused = weight_raw * raw_scores + float(weight_2d) * two_scores
        met = metric_dict(y_true, fused)
        row = {
            "fusion_family": "raw_plus_2d",
            "weight_raw": weight_raw,
            "weight_2d": float(weight_2d),
            "weight_rule": 0.0,
            **met,
        }
        rows.append(row)
        if best is None or (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
            best = row
    return best, rows


def clean_weight(value: float) -> float:
    value = float(value)
    if abs(value) < 1e-12:
        return 0.0
    return value


def search_three_branch(raw_scores: np.ndarray, two_scores: np.ndarray, rule_scores_: np.ndarray, y_true: np.ndarray) -> Tuple[dict, List[dict]]:
    best = None
    rows: List[dict] = []
    grid = np.linspace(0, 1, 11)
    for weight_raw in grid:
        for weight_2d in grid:
            if weight_raw + weight_2d > 1.000001:
                continue
            weight_raw_clean = clean_weight(weight_raw)
            weight_2d_clean = clean_weight(weight_2d)
            weight_rule = clean_weight(1.0 - float(weight_raw) - float(weight_2d))
            fused = weight_raw_clean * raw_scores + weight_2d_clean * two_scores + weight_rule * rule_scores_
            met = metric_dict(y_true, fused)
            row = {
                "fusion_family": "raw_plus_2d_plus_rules",
                "weight_raw": weight_raw_clean,
                "weight_2d": weight_2d_clean,
                "weight_rule": weight_rule,
                **met,
            }
            rows.append(row)
            if best is None or (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
                best = row
    return best, rows


def choose_weights_inner(
    Xseq: np.ndarray,
    X2d: np.ndarray,
    y: np.ndarray,
    groups: np.ndarray,
    case_features: Dict[str, dict],
    case_keys: np.ndarray,
    outer_train_idx: np.ndarray,
    seed: int,
    inner_splits: int,
    confidence_threshold: float,
    max_rules: int,
) -> Tuple[dict, dict, List[dict]]:
    inner = StratifiedGroupKFold(n_splits=inner_splits, shuffle=True, random_state=seed)
    raw_scores = np.zeros(len(outer_train_idx), dtype=float)
    two_scores = np.zeros(len(outer_train_idx), dtype=float)
    rule_scores_ = np.zeros(len(outer_train_idx), dtype=float)
    for inner_fold, (tr_local, va_local) in enumerate(
        inner.split(np.zeros(len(outer_train_idx)), y[outer_train_idx], groups=groups[outer_train_idx]),
        start=1,
    ):
        tr_idx = outer_train_idx[tr_local]
        va_idx = outer_train_idx[va_local]
        _, raw_val, _ = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, seed + inner_fold)
        _, two_val = fit_2d_green_proto(X2d, y, tr_idx, va_idx, seed + inner_fold)
        _, rule_val, _ = fit_rule_branch(case_features, case_keys, tr_idx, va_idx, confidence_threshold, max_rules)
        raw_scores[va_local] = raw_val
        two_scores[va_local] = two_val
        rule_scores_[va_local] = rule_val

    best_two, rows_two = search_two_branch(raw_scores, two_scores, y[outer_train_idx])
    best_three, rows_three = search_three_branch(raw_scores, two_scores, rule_scores_, y[outer_train_idx])
    return best_two, best_three, rows_two + rows_three


def make_report(report_path: Path, manifest: dict, rows: List[dict], fold_rows: List[dict], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC Raw + 2D + Heuristic Rule Nested Fusion 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版將 heuristic rule branch 納入 raw + 2D nested score-level fusion。"
        "raw branch 使用 initial-segment raw time-series + training-fold-only DTW/T-SMOTE + ExtraTrees；"
        "2D branch 使用 group/channel/all-channel patch features + Green/PCA prototype；"
        "rule branch 在每個 training fold 內重新建立啟發式規則庫，輸出 rule score。"
    )
    doc.add_paragraph(
        f"驗證設計：{manifest['validation']}。三分支 fusion 權重只在 outer training fold 的 inner validation 搜尋，"
        "outer test patient 完全不參與規則建立、特徵選擇、augmentation 或權重調整。"
    )

    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = table.add_row().cells
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、各折權重與規則統計", level=1)
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = ["fold", "w_raw", "w_2d", "w_rule", "2way_w_2d", "synthetic", "rules", "train", "test", "overlap"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in fold_rows:
        vals = [
            r["fold"],
            f"{r['weight_raw']:.1f}",
            f"{r['weight_2d']:.1f}",
            f"{r['weight_rule']:.1f}",
            f"{r['two_way_weight_2d']:.1f}",
            r["synthetic_samples"],
            r["final_rule_count"],
            r["train_cases"],
            r["test_cases"],
            r["group_overlap"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、三分支 Fusion 圖表", level=1)
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))

    doc.add_heading("五、判讀", level=1)
    doc.add_paragraph(
        "若 raw+2D+rule 較 raw+2D 提升，代表規則庫補上了部分可解釋的 2D/形態訊號；"
        "若未提升，則表示目前規則庫與既有 2D prototype branch 的資訊重疊，或規則門檻仍需進一步從錯誤樣本中修訂。"
        "本報告採 patient-aware group split，因此可作為比 random holdout 更嚴格的個案層級證據。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_score_level_fusion_2d_raw_rules_group5fold_20260519")
    parser.add_argument("--report", default="reports/NSC_score_level_fusion_2D_raw_rules_patient_aware_group5fold_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--n-splits", type=int, default=5)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--confidence-threshold", type=float, default=0.15)
    parser.add_argument("--max-rules", type=int, default=48)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.max_rows_per_csv,
    )
    X2d, feature_names_2d = build_2d_matrix(Path(args.image_root).resolve(), subjects)
    case_features, case_keys = build_rule_context(Path(args.image_root).resolve(), subjects)
    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)

    raw_all = np.zeros(len(y), dtype=float)
    two_all = np.zeros(len(y), dtype=float)
    rule_all = np.zeros(len(y), dtype=float)
    fusion_two_all = np.zeros(len(y), dtype=float)
    fusion_three_all = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    inner_rows_all: List[dict] = []

    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        best_two, best_three, inner_rows = choose_weights_inner(
            Xseq,
            X2d,
            y,
            groups,
            case_features,
            case_keys,
            train_idx,
            args.random_state + fold * 100,
            args.inner_splits,
            args.confidence_threshold,
            args.max_rules,
        )
        for r in inner_rows:
            r["outer_fold"] = fold
            inner_rows_all.append(r)

        _, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
        _, two_test = fit_2d_green_proto(X2d, y, train_idx, test_idx, args.random_state + fold)
        _, rule_test, rule_stats = fit_rule_branch(
            case_features,
            case_keys,
            train_idx,
            test_idx,
            args.confidence_threshold,
            args.max_rules,
        )
        raw_all[test_idx] = raw_test
        two_all[test_idx] = two_test
        rule_all[test_idx] = rule_test
        fusion_two_all[test_idx] = best_two["weight_raw"] * raw_test + best_two["weight_2d"] * two_test
        fusion_three_all[test_idx] = (
            best_three["weight_raw"] * raw_test
            + best_three["weight_2d"] * two_test
            + best_three["weight_rule"] * rule_test
        )
        train_groups = set(groups[train_idx].tolist())
        test_groups = set(groups[test_idx].tolist())
        fold_rows.append(
            {
                "fold": fold,
                "weight_raw": float(best_three["weight_raw"]),
                "weight_2d": float(best_three["weight_2d"]),
                "weight_rule": float(best_three["weight_rule"]),
                "two_way_weight_2d": float(best_two["weight_2d"]),
                "synthetic_samples": int(n_syn),
                "final_rule_count": int(rule_stats["final_rule_count"]),
                "train_cases": int(len(train_idx)),
                "test_cases": int(len(test_idx)),
                "group_overlap": int(len(train_groups & test_groups)),
                "mean_test_rule_confidence": float(rule_stats["mean_test_rule_confidence"]),
                "mean_test_used_rules": float(rule_stats["mean_test_used_rules"]),
            }
        )

    rows = [
        {"method": "raw_initial_segment_dtw_tsmote_extratrees", **metric_dict(y, raw_all)},
        {"method": "2d_green_patch_prototype", **metric_dict(y, two_all)},
        {"method": "heuristic_rule_branch", **metric_dict(y, rule_all)},
        {"method": "nested_score_fusion_raw_plus_2d", **metric_dict(y, fusion_two_all)},
        {"method": "nested_score_fusion_raw_plus_2d_plus_rules", **metric_dict(y, fusion_three_all)},
    ]
    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)

    pred_rows = []
    for i, subject in enumerate(subjects):
        pred_rows.append(
            {
                "subject_id": subject,
                "true_label": int(y[i]),
                "raw_score": float(raw_all[i]),
                "two_d_score": float(two_all[i]),
                "rule_score": float(rule_all[i]),
                "fusion_raw_2d_score": float(fusion_two_all[i]),
                "fusion_raw_2d_rule_score": float(fusion_three_all[i]),
            }
        )

    row_fields = ["method", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "score_level_fusion_summary.csv", rows, row_fields)
    write_csv(
        out_dir / "score_level_fusion_predictions.csv",
        pred_rows,
        ["subject_id", "true_label", "raw_score", "two_d_score", "rule_score", "fusion_raw_2d_score", "fusion_raw_2d_rule_score"],
    )
    write_csv(
        out_dir / "fusion_fold_weights.csv",
        fold_rows,
        [
            "fold",
            "weight_raw",
            "weight_2d",
            "weight_rule",
            "two_way_weight_2d",
            "synthetic_samples",
            "final_rule_count",
            "train_cases",
            "test_cases",
            "group_overlap",
            "mean_test_rule_confidence",
            "mean_test_used_rules",
        ],
    )
    write_csv(
        out_dir / "inner_weight_search.csv",
        inner_rows_all,
        ["outer_fold", "fusion_family", "weight_raw", "weight_2d", "weight_rule", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"],
    )

    fusion_met = next(r for r in rows if r["method"] == "nested_score_fusion_raw_plus_2d_plus_rules")
    figures = make_figures(out_dir, y, fusion_three_all, fusion_met)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold fusion-weight selection",
        "outer_splits": args.n_splits,
        "inner_splits": args.inner_splits,
        "sequence_length": args.sequence_length,
        "max_rows_per_csv": args.max_rows_per_csv,
        "feature_names_2d": len(feature_names_2d),
        "confidence_threshold": args.confidence_threshold,
        "max_rules": args.max_rules,
        "best": rows[0],
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "score_level_fusion_summary.csv"),
            "predictions": str(out_dir / "score_level_fusion_predictions.csv"),
            "fold_weights": str(out_dir / "fusion_fold_weights.csv"),
            "inner_search": str(out_dir / "inner_weight_search.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": {k: str(v) for k, v in figures.items()},
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, fold_rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
