#!/usr/bin/env python3
"""Patch-level error analysis for NSC 2D time-series images.

This is a post-hoc interpretability analysis over out-of-fold predictions.
It does not claim a new validation score. The output is a candidate
channel/plot/patch subset and VLM-review manifest for the next fold-local
modeling round.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_rule_learning_nsc_dataset_images import (  # noqa: E402
    PATCH_GRID,
    build_case_features,
    channel_group,
    collect_case_images,
    natural_key,
    parse_feature_key,
)


def read_csv(path: Path) -> List[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def write_jsonl(path: Path, rows: Iterable[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def build_feature_matrix(case_features: Dict[str, dict], subjects: List[str], feature_mode: str) -> Tuple[np.ndarray, List[str], Dict[str, str]]:
    by_subject = {v["case_id"]: k for k, v in case_features.items()}
    subject_key = {s: by_subject[s] for s in subjects if s in by_subject}
    if feature_mode == "model_compatible":
        allowed = lambda f: f.startswith("GROUP_") or f.startswith("ALL_CHANNELS/")
    elif feature_mode == "all":
        allowed = lambda f: True
    else:
        raise ValueError(feature_mode)
    feature_names = sorted(
        {
            f
            for s in subjects
            if s in subject_key
            for f in case_features[subject_key[s]]["features"]
            if allowed(f)
        }
    )
    X = np.full((len(subjects), len(feature_names)), np.nan, dtype=np.float32)
    idx = {f: i for i, f in enumerate(feature_names)}
    for i, subject in enumerate(subjects):
        if subject not in subject_key:
            continue
        for f, value in case_features[subject_key[subject]]["features"].items():
            if f in idx:
                X[i, idx[f]] = float(value)
    return X, feature_names, subject_key


def group_labels(y: np.ndarray, score: np.ndarray, threshold: float) -> np.ndarray:
    pred = (score >= threshold).astype(int)
    out = []
    for yi, pi in zip(y, pred):
        if yi == 0 and pi == 0:
            out.append("TN")
        elif yi == 0 and pi == 1:
            out.append("FP")
        elif yi == 1 and pi == 0:
            out.append("FN")
        elif yi == 1 and pi == 1:
            out.append("TP")
        else:
            out.append("NA")
    return np.asarray(out, dtype=object)


def safe_mean(values: np.ndarray) -> float:
    if len(values) == 0 or np.all(~np.isfinite(values)):
        return math.nan
    return float(np.nanmean(values))


def safe_std(values: np.ndarray) -> float:
    if len(values) <= 1 or np.all(~np.isfinite(values)):
        return 0.0
    return float(np.nanstd(values, ddof=1))


def effect(mean_a: float, mean_b: float, std_a: float, std_b: float) -> float:
    pooled = math.sqrt((std_a * std_a + std_b * std_b) / 2.0)
    pooled = max(pooled, 0.02)
    return float((mean_a - mean_b) / pooled)


def feature_stats(X: np.ndarray, feature_names: List[str], y: np.ndarray, groups: np.ndarray, min_support: int) -> List[dict]:
    rows: List[dict] = []
    masks = {g: groups == g for g in ["TN", "FP", "FN", "TP"]}
    label0 = y == 0
    label1 = y == 1
    for j, feature in enumerate(feature_names):
        col = X[:, j]
        finite = np.isfinite(col)
        if int(finite.sum()) < min_support:
            continue
        try:
            channel, plot, row, col_idx = parse_feature_key(feature)
        except Exception:
            continue
        vals = {}
        for g in ["TN", "FP", "FN", "TP"]:
            vals[g] = col[masks[g] & finite]
        if len(vals["TN"]) < 3 or len(vals["FP"]) < 3 or len(vals["FN"]) < 3 or len(vals["TP"]) < 3:
            continue
        mean0 = safe_mean(col[label0 & finite])
        mean1 = safe_mean(col[label1 & finite])
        class_sd0 = safe_std(col[label0 & finite])
        class_sd1 = safe_std(col[label1 & finite])
        class_effect = effect(mean1, mean0, class_sd1, class_sd0)
        means = {g: safe_mean(vals[g]) for g in vals}
        sds = {g: safe_std(vals[g]) for g in vals}
        fp_effect = effect(means["FP"], means["TN"], sds["FP"], sds["TN"])
        fn_effect = effect(means["FN"], means["TP"], sds["FN"], sds["TP"])
        class_sign = 1 if class_effect >= 0 else -1
        fp_aligned = fp_effect * class_sign
        fn_aligned = -fn_effect * class_sign
        error_alignment = max(fp_aligned, 0.0) + max(fn_aligned, 0.0)
        rows.append(
            {
                "feature": feature,
                "channel": channel,
                "plot": plot,
                "patch_row": row,
                "patch_col": col_idx,
                "support": int(finite.sum()),
                "coverage": float(finite.mean()),
                "mean_label0": mean0,
                "mean_label1": mean1,
                "class_effect_label1_minus_label0": class_effect,
                "mean_TN": means["TN"],
                "mean_FP": means["FP"],
                "mean_FN": means["FN"],
                "mean_TP": means["TP"],
                "fp_effect_FP_minus_TN": fp_effect,
                "fn_effect_FN_minus_TP": fn_effect,
                "fp_toward_label1_score": fp_aligned,
                "fn_toward_label0_score": fn_aligned,
                "error_alignment_score": error_alignment,
                "interpretation": describe_feature(channel, plot, row, col_idx, class_effect, fp_aligned, fn_aligned),
            }
        )
    return rows


def describe_feature(channel: str, plot: str, row: int, col: int, class_effect: float, fp_aligned: float, fn_aligned: float) -> str:
    label1_dir = "較深/較密" if class_effect > 0 else "較淺/較疏"
    pieces = [f"{channel}/{plot} 第{row + 1}列第{col + 1}欄，label 1 傾向{label1_dir}"]
    if fp_aligned > 0.4:
        pieces.append("FP 的 label 0 個案在此區域往 label 1 紋理偏移")
    if fn_aligned > 0.4:
        pieces.append("FN 的 label 1 個案在此區域往 label 0 紋理偏移")
    return "；".join(pieces)


def aggregate_combo_scores(rows: List[dict], score_key: str) -> List[dict]:
    buckets: Dict[Tuple[str, str], List[float]] = defaultdict(list)
    for row in rows:
        buckets[(row["channel"], row["plot"])].append(float(row[score_key]))
    out = []
    for (channel, plot), vals in buckets.items():
        arr = np.asarray(vals, dtype=float)
        out.append(
            {
                "channel": channel,
                "plot": plot,
                "mean_abs_score": float(np.nanmean(np.abs(arr))),
                "max_abs_score": float(np.nanmax(np.abs(arr))),
                "patch_count": int(len(arr)),
            }
        )
    return sorted(out, key=lambda r: (r["mean_abs_score"], r["max_abs_score"]), reverse=True)


def feature_heatmap(rows: List[dict], channel: str, plot: str, value_key: str) -> np.ndarray:
    mat = np.full((PATCH_GRID, PATCH_GRID), np.nan, dtype=float)
    for row in rows:
        if row["channel"] == channel and row["plot"] == plot:
            mat[int(row["patch_row"]), int(row["patch_col"])] = float(row[value_key])
    return mat


def make_heatmaps(out_dir: Path, rows: List[dict], top_fp: List[dict], top_fn: List[dict]) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    for label, combos, value_key, title in [
        ("fp", top_fp[:6], "fp_effect_FP_minus_TN", "FP minus TN patch effect"),
        ("fn", top_fn[:6], "fn_effect_FN_minus_TP", "FN minus TP patch effect"),
    ]:
        n = max(len(combos), 1)
        fig, axes = plt.subplots(1, n, figsize=(3.2 * n, 3.1), squeeze=False)
        vals = [abs(float(r["mean_abs_score"])) for r in combos]
        vmax = max(vals + [1.0])
        for ax, combo in zip(axes[0], combos):
            mat = feature_heatmap(rows, combo["channel"], combo["plot"], value_key)
            im = ax.imshow(mat, cmap="coolwarm", vmin=-vmax, vmax=vmax)
            ax.set_title(f"{combo['channel']}\n{combo['plot']}", fontsize=9)
            ax.set_xticks(range(PATCH_GRID))
            ax.set_yticks(range(PATCH_GRID))
            ax.tick_params(labelsize=7)
            for r in range(PATCH_GRID):
                for c in range(PATCH_GRID):
                    if np.isfinite(mat[r, c]) and abs(mat[r, c]) >= 0.6:
                        ax.text(c, r, "*", ha="center", va="center", fontsize=10, color="black")
        fig.suptitle(title)
        fig.colorbar(im, ax=axes.ravel().tolist(), shrink=0.72)
        path = fig_dir / f"{label}_patch_error_heatmaps.png"
        fig.savefig(path, dpi=180, bbox_inches="tight")
        plt.close(fig)
        paths[f"{label}_heatmap"] = str(path)
    return paths


def make_bar(out_dir: Path, rows: List[dict], key: str, name: str, top_n: int = 16) -> str:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    top = sorted(rows, key=lambda r: abs(float(r[key])), reverse=True)[:top_n]
    labels = [f"{r['channel']}/{r['plot']}/r{int(r['patch_row'])+1}c{int(r['patch_col'])+1}" for r in top]
    values = [float(r[key]) for r in top]
    plt.figure(figsize=(12, 5.2))
    plt.bar(np.arange(len(top)), values, color=["#9c3b3b" if v > 0 else "#3b6ea8" for v in values])
    plt.axhline(0, color="black", linewidth=0.8)
    plt.xticks(np.arange(len(top)), labels, rotation=45, ha="right", fontsize=8)
    plt.ylabel(key)
    plt.title(name)
    plt.grid(axis="y", alpha=0.25)
    path = fig_dir / f"{name.lower().replace(' ', '_')}.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return str(path)


def collect_image_lookup(case_features: Dict[str, dict]) -> Dict[str, List[Path]]:
    lookup: Dict[str, List[Path]] = defaultdict(list)
    for row in case_features.values():
        for path in row.get("image_paths", []):
            lookup[row["case_id"]].append(Path(path))
    # build_case_features discards image_paths, so this function is unused for it.
    return lookup


def find_case_images(image_root: Path) -> Dict[str, List[Path]]:
    out: Dict[str, List[Path]] = defaultdict(list)
    for path in image_root.rglob("*.png"):
        rel = path.relative_to(image_root).parts
        if len(rel) < 4:
            continue
        case_id = rel[1].split("_Sess", 1)[0]
        out[case_id].append(path)
    for case_id in out:
        out[case_id] = sorted(out[case_id], key=lambda p: natural_key(str(p)))
    return out


def image_matches_feature(path: Path, feature: dict) -> bool:
    channel = path.parent.name
    plot = path.stem
    if plot != feature["plot"]:
        return False
    feat_channel = feature["channel"]
    if feat_channel == "ALL_CHANNELS":
        return True
    if feat_channel.startswith("GROUP_"):
        return channel_group(channel) == feat_channel
    return channel == feat_channel


def crop_patch(path: Path, row: int, col: int, out_size: int = 72) -> Image.Image:
    img = Image.open(path).convert("RGB").resize((128, 128))
    patch = 128 // PATCH_GRID
    crop = img.crop((col * patch, row * patch, (col + 1) * patch, (row + 1) * patch))
    return crop.resize((out_size, out_size), Image.Resampling.NEAREST)


def make_contact_sheet(
    out_dir: Path,
    case_images: Dict[str, List[Path]],
    subjects: List[str],
    y: np.ndarray,
    score: np.ndarray,
    groups: np.ndarray,
    feature: dict,
    sheet_name: str,
    max_per_group: int = 5,
) -> Tuple[str, List[dict]]:
    fig_dir = out_dir / "vlm_contact_sheets"
    fig_dir.mkdir(parents=True, exist_ok=True)
    row = int(feature["patch_row"])
    col = int(feature["patch_col"])
    selected: List[Tuple[str, str, Path]] = []
    manifest_rows: List[dict] = []
    for group in ["TN", "FP", "TP", "FN"]:
        group_idx = [i for i, g in enumerate(groups.tolist()) if g == group]
        if group in {"FP", "TP"}:
            group_idx = sorted(group_idx, key=lambda i: score[i], reverse=True)
        else:
            group_idx = sorted(group_idx, key=lambda i: score[i])
        count = 0
        for i in group_idx:
            subject = subjects[i]
            matches = [p for p in case_images.get(subject, []) if image_matches_feature(p, feature)]
            if not matches:
                continue
            selected.append((group, subject, matches[0]))
            manifest_rows.append(
                {
                    "subject_id": subject,
                    "true_label": int(y[i]),
                    "prediction_group": group,
                    "score": float(score[i]),
                    "feature": feature["feature"],
                    "channel": feature["channel"],
                    "plot": feature["plot"],
                    "patch_row_1based": row + 1,
                    "patch_col_1based": col + 1,
                    "image_path": str(matches[0]),
                    "vlm_prompt": (
                        "Compare local texture density, line concentration, diagonal/periodic structure, "
                        "and whether this patch visually resembles label 0 or label 1 exemplars."
                    ),
                }
            )
            count += 1
            if count >= max_per_group:
                break
    cell_w, cell_h = 112, 102
    header_h = 50
    sheet = Image.new("RGB", (max_per_group * cell_w, header_h + 4 * cell_h), "white")
    draw = ImageDraw.Draw(sheet)
    draw.text((4, 4), f"{feature['feature']}  r{row + 1}c{col + 1}", fill=(0, 0, 0))
    group_offsets = {"TN": 0, "FP": 1, "TP": 2, "FN": 3}
    per_group_counter = Counter()
    for group, subject, path in selected:
        x = per_group_counter[group] * cell_w
        y0 = header_h + group_offsets[group] * cell_h
        try:
            crop = crop_patch(path, row, col)
            sheet.paste(crop, (x + 20, y0 + 4))
            draw.text((x + 4, y0 + 78), f"{group} {subject}", fill=(0, 0, 0))
        except Exception:
            draw.text((x + 4, y0 + 4), f"{group} {subject}\nload failed", fill=(150, 0, 0))
        per_group_counter[group] += 1
    path = fig_dir / f"{sheet_name}.png"
    sheet.save(path)
    return str(path), manifest_rows


def make_report(
    report_path: Path,
    manifest: dict,
    count_rows: List[dict],
    top_features: List[dict],
    fp_features: List[dict],
    fn_features: List[dict],
    candidate_subset: List[dict],
    figures: Dict[str, str],
    contact_sheets: List[str],
) -> None:
    doc = Document()
    doc.add_heading("NSC Patch-level Error Analysis 與 VLM 啟發式特徵篩選報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、分析定位", level=1)
    doc.add_paragraph(
        "本報告是 post-hoc error analysis，不宣稱新的 AUROC/AUPRC。目的在於將文獻中的 time-series imaging、"
        "local texture、patch/prototype/green learning 思路收斂成可解釋的 channel/plot/patch feature subset。"
    )
    doc.add_paragraph(
        f"使用 out-of-fold 預測方法：{manifest['prediction_method']}，threshold={manifest['threshold']}。"
        "錯誤類型定義為 FP=label 0 但預測為 1，FN=label 1 但預測為 0。"
    )

    doc.add_heading("二、錯誤案例分布", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for c, h in zip(table.rows[0].cells, ["group", "count", "description"]):
        c.text = h
    for r in count_rows:
        cells = table.add_row().cells
        for c, v in zip(cells, [r["group"], r["count"], r["description"]]):
            c.text = str(v)

    doc.add_heading("三、最重要的 Error-aligned Patch Features", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["feature", "class_effect", "FP->label1", "FN->label0", "support", "TN/FP/FN/TP means", "interpretation", "subset"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    subset_features = {r["feature"] for r in candidate_subset}
    for r in top_features[:18]:
        vals = [
            r["feature"],
            f"{float(r['class_effect_label1_minus_label0']):.2f}",
            f"{float(r['fp_toward_label1_score']):.2f}",
            f"{float(r['fn_toward_label0_score']):.2f}",
            r["support"],
            f"{float(r['mean_TN']):.3f}/{float(r['mean_FP']):.3f}/{float(r['mean_FN']):.3f}/{float(r['mean_TP']):.3f}",
            r["interpretation"],
            "Y" if r["feature"] in subset_features else "",
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、FP 與 FN 專屬 Patch", level=1)
    doc.add_paragraph("FP features：label 0 錯分為 label 1 時，局部紋理往 label 1 方向偏移的 patch。")
    for r in fp_features[:8]:
        doc.add_paragraph(f"{r['feature']}：{r['interpretation']}", style="List Bullet")
    doc.add_paragraph("FN features：label 1 錯分為 label 0 時，局部紋理往 label 0 方向偏移的 patch。")
    for r in fn_features[:8]:
        doc.add_paragraph(f"{r['feature']}：{r['interpretation']}", style="List Bullet")

    doc.add_heading("五、圖表", level=1)
    for key in ["fp_heatmap", "fn_heatmap", "top_error_bar", "top_fp_bar", "top_fn_bar"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.3))

    doc.add_heading("六、VLM 啟發式學習輸出", level=1)
    doc.add_paragraph(
        "已將 top local patch 裁切成 contact sheet，並輸出 VLM review manifest。VLM 或人工標註時，應比較局部紋理密度、"
        "線條集中區域、對角/週期結構、以及該 patch 是否更像 label 0 或 label 1。"
    )
    for sheet in contact_sheets[:6]:
        doc.add_picture(sheet, width=Inches(6.1))

    doc.add_heading("七、下一步", level=1)
    doc.add_paragraph(
        "下一輪建模不可直接用本 post-hoc 全資料 subset 宣稱泛化提升。正確做法是：在每個 training fold 內重新執行 "
        "patch-level feature selection，再用該 fold 選出的 subset 訓練 2D ET/fusion，最後評估 outer test。"
    )
    doc.add_paragraph(
        "若 fold-local error-guided subset 能穩定提升 AUPRC，才可納入 paper 主方法；否則只作為可解釋錯誤分析與 VLM 啟發式規則庫。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--predictions", default="analysis/nsc_raw2d_ablation_group10fold_20260519/ablation_predictions.csv")
    parser.add_argument("--method", default="2d_et8_change_grid")
    parser.add_argument("--threshold", type=float, default=0.5)
    parser.add_argument("--feature-mode", choices=["model_compatible", "all"], default="model_compatible")
    parser.add_argument("--min-support", type=int, default=80)
    parser.add_argument("--subset-size", type=int, default=48)
    parser.add_argument("--out-dir", default="analysis/nsc_patch_error_analysis_20260519")
    parser.add_argument("--report", default="reports/NSC_patch_level_error_analysis_VLM_feature_subset_20260519.docx")
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    preds = read_csv(Path(args.predictions))
    subjects = [str(r["subject_id"]) for r in preds]
    y = np.asarray([int(r["true_label"]) for r in preds], dtype=int)
    score = np.asarray([float(r[args.method]) for r in preds], dtype=float)
    groups = group_labels(y, score, args.threshold)

    cases_raw = collect_case_images(Path(args.image_root).resolve())
    case_features = build_case_features(cases_raw)
    # Preserve image paths for contact sheets.
    by_subject_images = find_case_images(Path(args.image_root).resolve())
    X, feature_names, _subject_key = build_feature_matrix(case_features, subjects, args.feature_mode)
    rows = feature_stats(X, feature_names, y, groups, args.min_support)
    rows = sorted(rows, key=lambda r: (float(r["error_alignment_score"]), abs(float(r["class_effect_label1_minus_label0"]))), reverse=True)
    fp_rows = sorted([r for r in rows if float(r["fp_toward_label1_score"]) > 0], key=lambda r: float(r["fp_toward_label1_score"]), reverse=True)
    fn_rows = sorted([r for r in rows if float(r["fn_toward_label0_score"]) > 0], key=lambda r: float(r["fn_toward_label0_score"]), reverse=True)

    candidate_subset: List[dict] = []
    seen = set()
    for source_rows, quota in [(rows, args.subset_size // 2), (fp_rows, args.subset_size // 4), (fn_rows, args.subset_size // 4)]:
        for r in source_rows:
            if r["feature"] in seen:
                continue
            candidate_subset.append({**r, "selection_reason": "error_alignment" if source_rows is rows else ("FP_signature" if source_rows is fp_rows else "FN_signature")})
            seen.add(r["feature"])
            if sum(1 for x in candidate_subset if x["selection_reason"] == candidate_subset[-1]["selection_reason"]) >= quota:
                break
    for r in rows:
        if len(candidate_subset) >= args.subset_size:
            break
        if r["feature"] not in seen:
            candidate_subset.append({**r, "selection_reason": "fill_by_error_alignment"})
            seen.add(r["feature"])

    count_rows = [
        {"group": "TN", "count": int((groups == "TN").sum()), "description": "label 0 correctly predicted as 0"},
        {"group": "FP", "count": int((groups == "FP").sum()), "description": "label 0 incorrectly predicted as 1"},
        {"group": "FN", "count": int((groups == "FN").sum()), "description": "label 1 incorrectly predicted as 0"},
        {"group": "TP", "count": int((groups == "TP").sum()), "description": "label 1 correctly predicted as 1"},
    ]

    fields = [
        "feature",
        "channel",
        "plot",
        "patch_row",
        "patch_col",
        "support",
        "coverage",
        "mean_label0",
        "mean_label1",
        "class_effect_label1_minus_label0",
        "mean_TN",
        "mean_FP",
        "mean_FN",
        "mean_TP",
        "fp_effect_FP_minus_TN",
        "fn_effect_FN_minus_TP",
        "fp_toward_label1_score",
        "fn_toward_label0_score",
        "error_alignment_score",
        "interpretation",
    ]
    write_csv(out_dir / "patch_error_feature_stats.csv", rows, fields)
    write_csv(out_dir / "fp_signature_features.csv", fp_rows, fields)
    write_csv(out_dir / "fn_signature_features.csv", fn_rows, fields)
    write_csv(out_dir / "candidate_patch_feature_subset.csv", candidate_subset, fields + ["selection_reason"])

    top_fp_combos = aggregate_combo_scores(rows, "fp_effect_FP_minus_TN")
    top_fn_combos = aggregate_combo_scores(rows, "fn_effect_FN_minus_TP")
    write_csv(out_dir / "fp_channel_plot_patch_summary.csv", top_fp_combos, ["channel", "plot", "mean_abs_score", "max_abs_score", "patch_count"])
    write_csv(out_dir / "fn_channel_plot_patch_summary.csv", top_fn_combos, ["channel", "plot", "mean_abs_score", "max_abs_score", "patch_count"])

    figures = make_heatmaps(out_dir, rows, top_fp_combos, top_fn_combos)
    figures["top_error_bar"] = make_bar(out_dir, rows, "error_alignment_score", "Top Error Alignment Features")
    figures["top_fp_bar"] = make_bar(out_dir, fp_rows, "fp_toward_label1_score", "Top FP Toward Label1 Features")
    figures["top_fn_bar"] = make_bar(out_dir, fn_rows, "fn_toward_label0_score", "Top FN Toward Label0 Features")

    manifest_rows: List[dict] = []
    contact_sheets: List[str] = []
    for i, feature in enumerate(candidate_subset[:8], start=1):
        sheet, sheet_rows = make_contact_sheet(out_dir, by_subject_images, subjects, y, score, groups, feature, f"vlm_patch_{i:02d}")
        contact_sheets.append(sheet)
        manifest_rows.extend(sheet_rows)
    write_jsonl(out_dir / "vlm_patch_review_manifest.jsonl", manifest_rows)

    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "prediction_method": args.method,
        "threshold": args.threshold,
        "feature_mode": args.feature_mode,
        "min_support": args.min_support,
        "subjects": len(subjects),
        "counts": {r["group"]: r["count"] for r in count_rows},
        "feature_count_analyzed": len(rows),
        "candidate_subset_size": len(candidate_subset),
        "outputs": {
            "analysis_dir": str(out_dir),
            "feature_stats": str(out_dir / "patch_error_feature_stats.csv"),
            "candidate_subset": str(out_dir / "candidate_patch_feature_subset.csv"),
            "vlm_manifest": str(out_dir / "vlm_patch_review_manifest.jsonl"),
            "report": str(Path(args.report).resolve()),
            "figures": figures,
            "contact_sheets": contact_sheets,
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, count_rows, rows, fp_rows, fn_rows, candidate_subset, figures, contact_sheets)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
