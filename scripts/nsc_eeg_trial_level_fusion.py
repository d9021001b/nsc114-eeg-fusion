#!/usr/bin/env python3
"""Patient-aware EEG trial-level classifier ablation for NSC.

Folder mapping is user-defined:
    folder 0 -> EEG class 0
    folders 1/2/3/4 -> EEG class 1

This script uses that mapping only as *training trial labels* inside each
patient-level outer fold. For test patients, it predicts every raw EEG trial
from signal-derived features and aggregates trial probabilities into a patient
score. It does not use test folder identities as model features.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier, RandomForestClassifier
from sklearn.feature_selection import f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from nsc_eeg_csv_fusion_ablation import (  # noqa: E402
    EEG_FOLDER_TO_BINARY,
    EEG_FOLDERS,
    dump_json,
    extract_one_eeg_csv,
    load_base_scores,
    load_labels,
    natural_key,
    robust_unit_fit,
    write_csv,
)


def collect_trial_table(eeg_root: Path, cache_dir: Path, allowed_subjects: set[str]) -> Tuple[List[dict], List[str]]:
    rows: List[dict] = []
    feature_names: set[str] = set()
    for folder in EEG_FOLDERS:
        binary = int(EEG_FOLDER_TO_BINARY[folder])
        for path in sorted((eeg_root / folder).glob("*.csv"), key=lambda p: natural_key(p.name)):
            m = re.match(r"^(\d+)_", path.name)
            if not m:
                continue
            subject = m.group(1)
            if subject not in allowed_subjects:
                continue
            feats = extract_one_eeg_csv(path, cache_dir)
            numeric = {k: float(v) for k, v in feats.items() if isinstance(v, (int, float)) and math.isfinite(float(v))}
            feature_names.update(numeric)
            rows.append({"subject_id": subject, "trial_label": binary, "folder": folder, "path": str(path), "features": numeric})
    return rows, sorted(feature_names)


def trial_matrix(rows: List[dict], feature_names: List[str]) -> Tuple[np.ndarray, np.ndarray, np.ndarray]:
    idx = {f: i for i, f in enumerate(feature_names)}
    X = np.full((len(rows), len(feature_names)), np.nan, dtype=float)
    y = np.asarray([int(r["trial_label"]) for r in rows], dtype=int)
    subjects = np.asarray([str(r["subject_id"]) for r in rows], dtype=object)
    for i, row in enumerate(rows):
        for name, value in row["features"].items():
            if name in idx:
                X[i, idx[name]] = value
    return X, y, subjects


def select_topk(X_train: np.ndarray, y_train: np.ndarray, k: int) -> Tuple[SimpleImputer, np.ndarray]:
    imputer = SimpleImputer(strategy="median")
    X_imp = imputer.fit_transform(X_train)
    scores, _ = f_classif(X_imp, y_train)
    scores = np.nan_to_num(scores, nan=0.0, posinf=0.0, neginf=0.0)
    selected = np.argsort(scores)[::-1][: min(k, X_imp.shape[1])]
    return imputer, selected


def make_model(name: str, seed: int):
    if name == "LR":
        return make_pipeline(StandardScaler(), LogisticRegression(max_iter=5000, C=0.35, class_weight="balanced", solver="liblinear"))
    if name == "ExtraTrees":
        return ExtraTreesClassifier(n_estimators=500, random_state=seed, class_weight="balanced", max_features="sqrt", min_samples_leaf=2, n_jobs=-1)
    if name == "RF":
        return RandomForestClassifier(n_estimators=500, random_state=seed, class_weight="balanced", max_features="sqrt", min_samples_leaf=2, n_jobs=-1)
    if name == "HistGB":
        return HistGradientBoostingClassifier(max_iter=160, learning_rate=0.035, l2_regularization=0.05, random_state=seed)
    raise ValueError(f"Unknown model: {name}")


def aggregate_scores(scores: np.ndarray, mode: str) -> float:
    if scores.size == 0:
        return 0.5
    if mode == "mean":
        return float(np.mean(scores))
    if mode == "median":
        return float(np.median(scores))
    if mode == "max":
        return float(np.max(scores))
    if mode == "q75":
        return float(np.quantile(scores, 0.75))
    if mode == "q90":
        return float(np.quantile(scores, 0.90))
    if mode == "top3mean":
        return float(np.mean(np.sort(scores)[-min(3, len(scores)) :]))
    if mode == "top5mean":
        return float(np.mean(np.sort(scores)[-min(5, len(scores)) :]))
    if mode == "mean_max":
        return float(0.6 * np.mean(scores) + 0.4 * np.max(scores))
    raise ValueError(mode)


def metrics(y_true: np.ndarray, score: np.ndarray, threshold: float = 0.5) -> dict:
    pred = (score >= threshold).astype(int)
    tn, fp, fn, tp = [int(x) for x in confusion_matrix(y_true, pred, labels=[0, 1]).ravel()]
    return {
        "AUROC": float(roc_auc_score(y_true, score)),
        "AUPRC": float(average_precision_score(y_true, score)),
        "accuracy": float((tn + tp) / max(len(y_true), 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": tn,
        "FP": fp,
        "FN": fn,
        "TP": tp,
    }


def trial_oof(
    subjects: List[str],
    y_patient: np.ndarray,
    trial_rows: List[dict],
    feature_names: List[str],
    top_k: int,
    model_name: str,
    agg_modes: List[str],
    seed: int,
    n_splits: int,
) -> Tuple[Dict[str, np.ndarray], List[dict]]:
    X_trial, y_trial, trial_subjects = trial_matrix(trial_rows, feature_names)
    subject_to_idx = {s: i for i, s in enumerate(subjects)}
    scores = {mode: np.full(len(subjects), 0.5, dtype=float) for mode in agg_modes}
    fold_rows: List[dict] = []
    outer = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=seed)
    subjects_arr = np.asarray(subjects, dtype=object)
    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(subjects)), y_patient), start=1):
        train_subjects = set(subjects_arr[train_idx])
        test_subjects = set(subjects_arr[test_idx])
        tr = np.asarray([s in train_subjects for s in trial_subjects], dtype=bool)
        te = np.asarray([s in test_subjects for s in trial_subjects], dtype=bool)
        if len(np.unique(y_trial[tr])) < 2:
            continue
        imputer, selected = select_topk(X_trial[tr], y_trial[tr], top_k)
        Xtr = imputer.transform(X_trial[tr])[:, selected]
        Xte = imputer.transform(X_trial[te])[:, selected]
        model = make_model(model_name, seed + fold)
        model.fit(Xtr, y_trial[tr])
        trial_score = model.predict_proba(Xte)[:, 1]
        te_subjects = trial_subjects[te]
        by_subject: Dict[str, List[float]] = defaultdict(list)
        for s, sc in zip(te_subjects, trial_score):
            by_subject[str(s)].append(float(sc))
        for s in test_subjects:
            idx = subject_to_idx[s]
            arr = np.asarray(by_subject.get(s, []), dtype=float)
            for mode in agg_modes:
                scores[mode][idx] = aggregate_scores(arr, mode)
        fold_rows.append(
            {
                "fold": fold,
                "top_k": top_k,
                "model": model_name,
                "train_subjects": len(train_idx),
                "test_subjects": len(test_idx),
                "train_trials": int(np.sum(tr)),
                "test_trials": int(np.sum(te)),
                "selected_features": int(len(selected)),
            }
        )
    return scores, fold_rows


def make_figures(out_dir: Path, y: np.ndarray, rows: List[dict], predictions: List[dict]) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    top = rows[:18]
    labels = [r["method"] for r in top]
    ypos = np.arange(len(top))
    plt.figure(figsize=(11, max(5, 0.36 * len(top))))
    plt.barh(ypos + 0.18, [r["AUROC"] for r in top], height=0.34, label="AUROC")
    plt.barh(ypos - 0.18, [r["AUPRC"] for r in top], height=0.34, label="AUPRC")
    plt.axvline(0.8, color="gray", linestyle="--")
    plt.yticks(ypos, labels, fontsize=8)
    plt.gca().invert_yaxis()
    plt.xlabel("Metric")
    plt.title("EEG trial-level classifier patient-aware 10-fold")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    path = fig_dir / "trial_level_eeg_results.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["bar"] = str(path)

    if rows:
        best = rows[0]["method"]
        pred = sorted([r for r in predictions if r["method"] == best], key=lambda r: natural_key(r["subject_id"]))
        score = np.asarray([float(r["score"]) for r in pred])
        fpr, tpr, _ = roc_curve(y, score)
        plt.figure(figsize=(5.4, 4.3))
        plt.plot(fpr, tpr, label=f"AUROC={rows[0]['AUROC']:.3f}")
        plt.plot([0, 1], [0, 1], "--", color="gray")
        plt.xlabel("False Positive Rate")
        plt.ylabel("True Positive Rate")
        plt.title("Best trial-level EEG ROC")
        plt.grid(alpha=0.25)
        plt.legend(loc="lower right")
        path = fig_dir / "best_trial_level_eeg_roc.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        paths["roc"] = str(path)
        precision, recall, _ = precision_recall_curve(y, score)
        plt.figure(figsize=(5.4, 4.3))
        plt.plot(recall, precision, label=f"AUPRC={rows[0]['AUPRC']:.3f}")
        plt.axhline(float(np.mean(y)), linestyle="--", color="gray")
        plt.xlabel("Recall")
        plt.ylabel("Precision")
        plt.title("Best trial-level EEG PRC")
        plt.grid(alpha=0.25)
        plt.legend(loc="lower left")
        path = fig_dir / "best_trial_level_eeg_prc.png"
        plt.tight_layout()
        plt.savefig(path, dpi=180)
        plt.close()
        paths["prc"] = str(path)
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC EEG Trial-Level Classifier Fusion 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "資料夾 mapping 採使用者指定：folder 0 為 EEG class 0，folders 1/2/3/4 合併為 EEG class 1。"
        "本報告只在 outer training fold 內使用這些 trial labels 訓練 EEG trial-level classifier；"
        "outer test patient 的資料夾 label 不作為模型輸入特徵，僅使用 raw EEG CSV 萃取出的統計與頻譜特徵。"
    )
    doc.add_paragraph(f"驗證設計：{manifest['validation']}；subjects={manifest['subjects']}；label counts={manifest['cases_by_label']}")
    doc.add_heading("二、結果", level=1)
    table = doc.add_table(rows=1, cols=12)
    table.style = "Table Grid"
    headers = ["rank", "method", "AUROC", "AUPRC", "ACC", "Sens", "Spec", "PPV", "NPV", "CM", "model", "top_k"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for i, row in enumerate(rows[:24], start=1):
        vals = [
            i,
            row["method"],
            f"{row['AUROC']:.3f}",
            f"{row['AUPRC']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['sensitivity']:.3f}",
            f"{row['specificity']:.3f}",
            f"{row['PPV']:.3f}",
            f"{row['NPV']:.3f}",
            f"{row['TN']}/{row['FP']}/{row['FN']}/{row['TP']}",
            row.get("model", ""),
            row.get("top_k", ""),
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("三、圖表", level=1)
    for key in ["bar", "roc", "prc"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.4 if key == "bar" else 5.2))
    doc.add_heading("四、判讀", level=1)
    best = rows[0]
    doc.add_paragraph(f"最佳方法為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。")
    doc.add_heading("五、輸出檔案", level=1)
    for k, v in manifest["outputs"].items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--eeg-root", default="eeg-csv-data-by-class")
    parser.add_argument("--base-predictions", default="analysis/nsc_uncertain_band_patch_refinement_20260520/uncertain_band_predictions.csv")
    parser.add_argument("--cache-dir", default="analysis/nsc_eeg_csv_fusion_ablation_20260520/eeg_feature_cache")
    parser.add_argument("--out-dir", default="analysis/nsc_eeg_trial_level_fusion_20260520")
    parser.add_argument("--report", default="reports/NSC_EEG_trial_level_fusion_patient_aware_10fold_20260520.docx")
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--random-state", type=int, default=20260520)
    parser.add_argument("--top-ks", nargs="+", type=int, default=[16, 24, 32, 48, 64])
    parser.add_argument("--models", nargs="+", default=["LR", "ExtraTrees", "HistGB"])
    parser.add_argument("--agg-modes", nargs="+", default=["mean", "median", "q75", "q90", "max", "top3mean", "top5mean", "mean_max"])
    parser.add_argument("--base-cols", nargs="+", default=["mi_max"])
    args = parser.parse_args()

    labels = load_labels(Path(args.manifest).resolve())
    subjects = list(labels)
    y = np.asarray([labels[s] for s in subjects], dtype=int)
    trial_rows, feature_names = collect_trial_table(Path(args.eeg_root).resolve(), Path(args.cache_dir).resolve(), set(subjects))
    base_scores = load_base_scores(Path(args.base_predictions).resolve(), subjects)
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    rows: List[dict] = []
    predictions: List[dict] = []
    fold_rows: List[dict] = []
    for base_col, score in base_scores.items():
        row = metrics(y, score)
        row.update({"method": f"baseline_{base_col}", "kind": "baseline", "model": "", "top_k": "", "agg": ""})
        rows.append(row)
        for s, yy, sc in zip(subjects, y, score):
            predictions.append({"subject_id": s, "true_label": int(yy), "method": row["method"], "score": float(sc)})

    for model in args.models:
        for top_k in args.top_ks:
            agg_scores, folds = trial_oof(subjects, y, trial_rows, feature_names, top_k, model, args.agg_modes, args.random_state, args.n_splits)
            fold_rows.extend([{**r, "model": model, "top_k": top_k} for r in folds])
            for agg, score in agg_scores.items():
                method = f"eeg_trial_{model}_k{top_k}_{agg}"
                row = metrics(y, score)
                row.update({"method": method, "kind": "eeg_trial", "model": model, "top_k": top_k, "agg": agg})
                rows.append(row)
                for s, yy, sc in zip(subjects, y, score):
                    predictions.append({"subject_id": s, "true_label": int(yy), "method": method, "score": float(sc)})
                for base_col in args.base_cols:
                    base = base_scores[base_col]
                    btr, _ = robust_unit_fit(base, base)
                    etr, _ = robust_unit_fit(score, score)
                    for w_eeg in [0.05, 0.10, 0.15, 0.20, 0.25]:
                        fused = (1.0 - w_eeg) * btr + w_eeg * etr
                        method = f"fusion_{base_col}_{model}_k{top_k}_{agg}_w{w_eeg:.2f}"
                        row = metrics(y, fused)
                        row.update({"method": method, "kind": "fixed_fusion_posthoc", "model": model, "top_k": top_k, "agg": agg, "w_eeg": w_eeg})
                        rows.append(row)
                        for s, yy, sc in zip(subjects, y, fused):
                            predictions.append({"subject_id": s, "true_label": int(yy), "method": method, "score": float(sc)})

    rows.sort(key=lambda r: (min(r["AUROC"], r["AUPRC"]), r["AUPRC"], r["AUROC"]), reverse=True)
    write_csv(out_dir / "trial_level_eeg_summary.csv", rows, sorted({k for r in rows for k in r}))
    write_csv(out_dir / "trial_level_eeg_predictions.csv", predictions, sorted({k for r in predictions for k in r}))
    write_csv(out_dir / "trial_level_eeg_fold_details.csv", fold_rows, sorted({k for r in fold_rows for k in r}))
    figures = make_figures(out_dir, y, rows, predictions)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in Counter(y).items()},
        "trial_count": len(trial_rows),
        "subjects_with_eeg": len({r["subject_id"] for r in trial_rows}),
        "eeg_feature_count": len(feature_names),
        "folder_mapping": EEG_FOLDER_TO_BINARY,
        "validation": f"patient-aware stratified {args.n_splits}-fold; trial classifier trained only on training subjects",
        "best": rows[0] if rows else {},
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "trial_level_eeg_summary.csv"),
            "predictions": str(out_dir / "trial_level_eeg_predictions.csv"),
            "fold_details": str(out_dir / "trial_level_eeg_fold_details.csv"),
            "report": str(Path(args.report).resolve()),
        },
    }
    dump_json(out_dir / "manifest.json", manifest)
    make_report(Path(args.report).resolve(), manifest, rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
