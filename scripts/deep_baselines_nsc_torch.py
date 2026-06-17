#!/usr/bin/env python3
"""Torch deep baselines for NSC patient-aware 10-fold comparison.

Runs literature-style deep baselines without sklearn dependency:
- 1D CNN on raw grouped physiological sequences.
- 2D CNN on subject-level averaged PP/AR/RP/GAF image maps.
- 1D+2D score-level ensemble with validation-selected fusion weight.

The split unit is subject_id/case_id. Each subject appears in exactly one outer
test fold. Validation for early stopping and fusion weight selection is drawn
only from the outer training fold.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import os
import random
import re
from collections import Counter, defaultdict
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import torch
from PIL import Image
from docx import Document
from docx.shared import Inches
from torch import nn
from torch.utils.data import DataLoader, TensorDataset


GROUPS = ["EKG_PRIMARY", "RESP", "SC_AUX"]
PLOTS = ["PP", "AR", "RP", "GAF"]


def natural_key(text: str) -> Tuple:
    parts = re.split(r"(\d+)", str(text))
    return tuple(int(p) if p.isdigit() else p for p in parts)


def signal_group(signal: str) -> str:
    upper = signal.upper()
    if "RESP" in upper:
        return "RESP"
    if "SC_" in upper or "AUX" in upper:
        return "SC_AUX"
    if "EKG" in upper or "PRIMARY_RAW" in upper or "IBI" in upper:
        return "EKG_PRIMARY"
    return "OTHER"


def value_column(df: pd.DataFrame) -> str:
    for col in ["value", "ibi_ms", "resp_period_ms"]:
        if col in df.columns:
            return col
    numeric = [c for c in df.columns if c != "time_sec" and pd.api.types.is_numeric_dtype(df[c])]
    if not numeric:
        raise ValueError("No value-like numeric column found")
    return numeric[-1]


def robust_resample(times: np.ndarray, values: np.ndarray, length: int) -> np.ndarray:
    mask = np.isfinite(times) & np.isfinite(values)
    times = times[mask]
    values = values[mask]
    if len(values) == 0:
        return np.zeros(length, dtype=np.float32)
    if len(values) == 1 or np.nanmax(times) <= np.nanmin(times):
        return np.full(length, float(values[0]), dtype=np.float32)
    order = np.argsort(times)
    times = times[order]
    values = values[order]
    _, unique_idx = np.unique(times, return_index=True)
    times = times[unique_idx]
    values = values[unique_idx]
    grid = np.linspace(times.min(), times.max(), length)
    return np.interp(grid, times, values).astype(np.float32)


def normalize_sequence(seq: np.ndarray) -> np.ndarray:
    med = np.nanmedian(seq)
    iqr = np.nanpercentile(seq, 75) - np.nanpercentile(seq, 25)
    scale = iqr if iqr > 1e-6 else np.nanstd(seq)
    scale = scale if scale > 1e-6 else 1.0
    return ((seq - med) / scale).astype(np.float32)


def load_labels(manifest_path: Path) -> Dict[str, int]:
    labels: Dict[str, int] = {}
    with manifest_path.open("r", encoding="utf-8-sig", newline="") as f:
        for row in csv.DictReader(f):
            labels[str(row["subject_id"])] = int(row["label"])
    return labels


def load_raw_sequences(csv_dir: Path, manifest_path: Path, length: int, max_rows_per_csv: int) -> Tuple[List[str], np.ndarray, np.ndarray]:
    labels = load_labels(manifest_path)
    grouped: Dict[str, Dict[str, List[np.ndarray]]] = defaultdict(lambda: defaultdict(list))
    pattern = re.compile(r"^(\d+)_Sess\d+_(.+)\.csv$")
    for path in sorted(csv_dir.glob("*.csv"), key=lambda p: natural_key(p.name)):
        m = pattern.match(path.name)
        if not m:
            continue
        subject, signal = m.group(1), m.group(2)
        if subject not in labels:
            continue
        group = signal_group(signal)
        if group == "OTHER":
            continue
        try:
            header = pd.read_csv(path, nrows=0)
            cols = list(header.columns)
            read_cols = ["time_sec"]
            for candidate in ["value", "ibi_ms", "resp_period_ms"]:
                if candidate in cols:
                    read_cols.append(candidate)
                    break
            if len(read_cols) == 1:
                read_cols = cols[:3]
            df = pd.read_csv(path, usecols=read_cols, nrows=max_rows_per_csv)
            col = value_column(df)
            seq = robust_resample(df["time_sec"].to_numpy(dtype=float), df[col].to_numpy(dtype=float), length)
            grouped[subject][group].append(normalize_sequence(seq))
        except Exception:
            continue

    subjects = sorted(grouped.keys(), key=natural_key)
    X = np.zeros((len(subjects), len(GROUPS), length), dtype=np.float32)
    for i, subject in enumerate(subjects):
        for g, group in enumerate(GROUPS):
            seqs = grouped[subject].get(group, [])
            if seqs:
                X[i, g, :] = np.mean(np.stack(seqs), axis=0)
    y = np.asarray([labels[s] for s in subjects], dtype=np.int64)
    return subjects, X, y


def load_image_maps(image_root: Path, subjects: List[str], image_size: int) -> np.ndarray:
    bucket: Dict[str, Dict[Tuple[str, str], List[np.ndarray]]] = defaultdict(lambda: defaultdict(list))
    for path in image_root.rglob("*.png"):
        rel = path.relative_to(image_root).parts
        if len(rel) < 4 or rel[0] not in {"0", "1"}:
            continue
        session_name = rel[1]
        subject = session_name.split("_Sess", 1)[0]
        channel = rel[2]
        plot = path.stem
        if plot not in PLOTS:
            continue
        group = signal_group(channel)
        if group == "OTHER":
            continue
        try:
            arr = np.asarray(Image.open(path).convert("L").resize((image_size, image_size)), dtype=np.float32) / 255.0
            bucket[subject][(group, plot)].append(1.0 - arr)
        except Exception:
            continue
    X = np.zeros((len(subjects), len(GROUPS) * len(PLOTS), image_size, image_size), dtype=np.float32)
    for i, subject in enumerate(subjects):
        for gi, group in enumerate(GROUPS):
            for pi, plot in enumerate(PLOTS):
                imgs = bucket[subject].get((group, plot), [])
                if imgs:
                    X[i, gi * len(PLOTS) + pi] = np.mean(np.stack(imgs), axis=0)
    return X


def stratified_folds(y: np.ndarray, n_splits: int, seed: int) -> List[np.ndarray]:
    rng = np.random.default_rng(seed)
    folds: List[List[int]] = [[] for _ in range(n_splits)]
    for label in sorted(np.unique(y).tolist()):
        idx = np.where(y == label)[0]
        rng.shuffle(idx)
        for j, item in enumerate(idx.tolist()):
            folds[j % n_splits].append(item)
    return [np.asarray(sorted(fold), dtype=int) for fold in folds]


def train_val_split(indices: np.ndarray, y: np.ndarray, val_fraction: float, seed: int) -> Tuple[np.ndarray, np.ndarray]:
    rng = np.random.default_rng(seed)
    train_parts, val_parts = [], []
    for label in sorted(np.unique(y[indices]).tolist()):
        idx = indices[y[indices] == label].copy()
        rng.shuffle(idx)
        n_val = max(1, int(round(len(idx) * val_fraction)))
        val_parts.extend(idx[:n_val].tolist())
        train_parts.extend(idx[n_val:].tolist())
    return np.asarray(sorted(train_parts), dtype=int), np.asarray(sorted(val_parts), dtype=int)


def average_ranks(values: np.ndarray) -> np.ndarray:
    order = np.argsort(values)
    ranks = np.empty(len(values), dtype=float)
    i = 0
    while i < len(values):
        j = i + 1
        while j < len(values) and values[order[j]] == values[order[i]]:
            j += 1
        avg = (i + 1 + j) / 2.0
        ranks[order[i:j]] = avg
        i = j
    return ranks


def auroc(y_true: np.ndarray, score: np.ndarray) -> float:
    y_true = np.asarray(y_true)
    score = np.asarray(score)
    pos = y_true == 1
    neg = y_true == 0
    n_pos = int(pos.sum())
    n_neg = int(neg.sum())
    if n_pos == 0 or n_neg == 0:
        return float("nan")
    ranks = average_ranks(score)
    return float((ranks[pos].sum() - n_pos * (n_pos + 1) / 2.0) / (n_pos * n_neg))


def auprc(y_true: np.ndarray, score: np.ndarray) -> float:
    y_true = np.asarray(y_true).astype(int)
    score = np.asarray(score)
    order = np.argsort(-score)
    y_sorted = y_true[order]
    tp = np.cumsum(y_sorted == 1)
    fp = np.cumsum(y_sorted == 0)
    n_pos = max(int((y_true == 1).sum()), 1)
    recall = tp / n_pos
    precision = tp / np.maximum(tp + fp, 1)
    recall_prev = np.concatenate([[0.0], recall[:-1]])
    return float(np.sum((recall - recall_prev) * precision))


def metric_dict(y_true: np.ndarray, score: np.ndarray, threshold: float = 0.5) -> dict:
    pred = (score >= threshold).astype(int)
    tn = int(((y_true == 0) & (pred == 0)).sum())
    fp = int(((y_true == 0) & (pred == 1)).sum())
    fn = int(((y_true == 1) & (pred == 0)).sum())
    tp = int(((y_true == 1) & (pred == 1)).sum())
    return {
        "AUROC": auroc(y_true, score),
        "AUPRC": auprc(y_true, score),
        "accuracy": float((tp + tn) / max(tp + tn + fp + fn, 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": tn,
        "FP": fp,
        "FN": fn,
        "TP": tp,
    }


class Tiny1DCNN(nn.Module):
    def __init__(self, channels: int):
        super().__init__()
        self.net = nn.Sequential(
            nn.Conv1d(channels, 16, kernel_size=7, padding=3),
            nn.BatchNorm1d(16),
            nn.ReLU(),
            nn.MaxPool1d(2),
            nn.Dropout(0.15),
            nn.Conv1d(16, 32, kernel_size=5, padding=2),
            nn.BatchNorm1d(32),
            nn.ReLU(),
            nn.MaxPool1d(2),
            nn.Dropout(0.20),
            nn.Conv1d(32, 64, kernel_size=3, padding=1),
            nn.BatchNorm1d(64),
            nn.ReLU(),
            nn.AdaptiveAvgPool1d(1),
        )
        self.head = nn.Linear(64, 1)

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        z = self.net(x).squeeze(-1)
        return self.head(z).squeeze(-1)


class Tiny2DCNN(nn.Module):
    def __init__(self, channels: int):
        super().__init__()
        self.net = nn.Sequential(
            nn.Conv2d(channels, 16, kernel_size=3, padding=1),
            nn.BatchNorm2d(16),
            nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Dropout2d(0.10),
            nn.Conv2d(16, 32, kernel_size=3, padding=1),
            nn.BatchNorm2d(32),
            nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Dropout2d(0.15),
            nn.Conv2d(32, 64, kernel_size=3, padding=1),
            nn.BatchNorm2d(64),
            nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Conv2d(64, 96, kernel_size=3, padding=1),
            nn.BatchNorm2d(96),
            nn.ReLU(),
            nn.AdaptiveAvgPool2d(1),
        )
        self.head = nn.Sequential(nn.Dropout(0.25), nn.Linear(96, 1))

    def forward(self, x: torch.Tensor) -> torch.Tensor:
        z = self.net(x).flatten(1)
        return self.head(z).squeeze(-1)


def set_seed(seed: int) -> None:
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    torch.cuda.manual_seed_all(seed)


def standardize_train_only(X: np.ndarray, train_idx: np.ndarray, apply_idx: np.ndarray) -> np.ndarray:
    mean = X[train_idx].mean(axis=0, keepdims=True)
    std = X[train_idx].std(axis=0, keepdims=True)
    std = np.where(std < 1e-6, 1.0, std)
    return ((X[apply_idx] - mean) / std).astype(np.float32)


def predict_scores(model: nn.Module, X: np.ndarray, device: torch.device, batch_size: int) -> np.ndarray:
    model.eval()
    scores: List[np.ndarray] = []
    loader = DataLoader(TensorDataset(torch.from_numpy(X)), batch_size=batch_size, shuffle=False)
    with torch.no_grad():
        for (xb,) in loader:
            xb = xb.to(device)
            scores.append(torch.sigmoid(model(xb)).detach().cpu().numpy())
    return np.concatenate(scores).astype(float)


def train_model(
    model: nn.Module,
    X: np.ndarray,
    y: np.ndarray,
    train_idx: np.ndarray,
    val_idx: np.ndarray,
    test_idx: np.ndarray,
    device: torch.device,
    seed: int,
    epochs: int,
    patience: int,
    batch_size: int,
    lr: float,
) -> Tuple[np.ndarray, np.ndarray, dict]:
    set_seed(seed)
    X_train = standardize_train_only(X, train_idx, train_idx)
    X_val = standardize_train_only(X, train_idx, val_idx)
    X_test = standardize_train_only(X, train_idx, test_idx)
    y_train = y[train_idx].astype(np.float32)
    y_val = y[val_idx].astype(np.float32)

    model = model.to(device)
    n_pos = max(float((y_train == 1).sum()), 1.0)
    n_neg = max(float((y_train == 0).sum()), 1.0)
    loss_fn = nn.BCEWithLogitsLoss(pos_weight=torch.tensor([n_neg / n_pos], device=device))
    optim = torch.optim.AdamW(model.parameters(), lr=lr, weight_decay=1e-4)
    loader = DataLoader(
        TensorDataset(torch.from_numpy(X_train), torch.from_numpy(y_train)),
        batch_size=batch_size,
        shuffle=True,
        generator=torch.Generator().manual_seed(seed),
    )

    best_state = deepcopy(model.state_dict())
    best_ap = -1.0
    bad = 0
    history = []
    for epoch in range(1, epochs + 1):
        model.train()
        losses = []
        for xb, yb in loader:
            xb = xb.to(device)
            yb = yb.to(device)
            optim.zero_grad(set_to_none=True)
            loss = loss_fn(model(xb), yb)
            loss.backward()
            optim.step()
            losses.append(float(loss.detach().cpu()))
        val_score = predict_scores(model, X_val, device, batch_size)
        val_ap = auprc(y_val.astype(int), val_score)
        history.append({"epoch": epoch, "loss": float(np.mean(losses)), "val_AUPRC": val_ap, "val_AUROC": auroc(y_val.astype(int), val_score)})
        if val_ap > best_ap:
            best_ap = val_ap
            best_state = deepcopy(model.state_dict())
            bad = 0
        else:
            bad += 1
        if bad >= patience:
            break
    model.load_state_dict(best_state)
    val_score = predict_scores(model, X_val, device, batch_size)
    test_score = predict_scores(model, X_test, device, batch_size)
    stats = {
        "best_val_AUPRC": auprc(y_val.astype(int), val_score),
        "best_val_AUROC": auroc(y_val.astype(int), val_score),
        "epochs_run": len(history),
        "train_cases": int(len(train_idx)),
        "val_cases": int(len(val_idx)),
        "test_cases": int(len(test_idx)),
    }
    return val_score, test_score, stats


def choose_fusion_weight(y_val: np.ndarray, score_1d: np.ndarray, score_2d: np.ndarray) -> Tuple[float, dict]:
    best_w = 0.5
    best = None
    for w in np.linspace(0, 1, 11):
        score = (1.0 - float(w)) * score_1d + float(w) * score_2d
        met = metric_dict(y_val, score)
        if best is None or (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
            best = met
            best_w = float(w)
    return best_w, best


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def make_figures(out_dir: Path, y: np.ndarray, score: np.ndarray, met: dict) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    order = np.argsort(score)
    pos = y == 1
    neg = y == 0
    thresholds = np.unique(score[order])
    fpr, tpr = [], []
    for thr in thresholds:
        pred = score >= thr
        tp = ((pred == 1) & pos).sum()
        fp = ((pred == 1) & neg).sum()
        fn = ((pred == 0) & pos).sum()
        tn = ((pred == 0) & neg).sum()
        tpr.append(tp / max(tp + fn, 1))
        fpr.append(fp / max(fp + tn, 1))
    fpr = [0.0] + fpr + [1.0]
    tpr = [0.0] + tpr + [1.0]
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={met['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Deep baseline ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = str(fig_dir / "deep_baseline_roc.png")
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()

    order = np.argsort(-score)
    y_sorted = y[order]
    tp = np.cumsum(y_sorted == 1)
    fp = np.cumsum(y_sorted == 0)
    precision = tp / np.maximum(tp + fp, 1)
    recall = tp / max(int((y == 1).sum()), 1)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={met['AUPRC']:.3f}")
    plt.axhline(float(y.mean()), linestyle="--", color="gray", label=f"prevalence={y.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Deep baseline PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = str(fig_dir / "deep_baseline_prc.png")
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()

    pred = score >= 0.5
    cm = np.asarray(
        [
            [((y == 0) & (pred == 0)).sum(), ((y == 0) & (pred == 1)).sum()],
            [((y == 1) & (pred == 0)).sum(), ((y == 1) & (pred == 1)).sum()],
        ],
        dtype=int,
    )
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Deep baseline confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = str(fig_dir / "deep_baseline_cm.png")
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, summary_rows: List[dict], fold_rows: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC 文獻深度模型 Baseline 實測報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本報告補跑 paper 需要列入比較的文獻深度模型 baseline：1D CNN、2D CNN、以及 1D+2D score-level ensemble。"
        "所有結果採 patient-aware stratified 10-fold；每個 subject_id/case_id 僅出現在單一 outer test fold。"
    )
    doc.add_paragraph(
        "每個 outer training fold 內再切 validation set，用於 early stopping 與 1D/2D fusion weight 選擇。"
        "outer test fold 不參與模型訓練、early stopping 或 fusion weight selection。"
    )
    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in summary_rows:
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("三、各折訓練摘要", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["fold", "w_2d", "1D val AUPRC", "2D val AUPRC", "fusion val AUPRC", "1D epochs", "2D epochs", "overlap"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in fold_rows:
        vals = [
            r["fold"],
            f"{r['weight_2d']:.1f}",
            f"{r['one_d_val_AUPRC']:.3f}",
            f"{r['two_d_val_AUPRC']:.3f}",
            f"{r['fusion_val_AUPRC']:.3f}",
            r["one_d_epochs"],
            r["two_d_epochs"],
            r["group_overlap"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("四、最佳深度模型圖表", level=1)
    doc.add_picture(figures["roc"], width=Inches(5.3))
    doc.add_picture(figures["prc"], width=Inches(5.3))
    doc.add_picture(figures["cm"], width=Inches(4.8))
    doc.add_heading("五、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(
        f"最佳深度 baseline 為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。"
        "此結果應與本案目前最佳 non-deep 主線 raw_initial + 2D ET8 grid fusion 並列比較。"
    )
    doc.add_paragraph(
        "若深度 baseline 未優於 non-deep 主線，paper 可合理論述：在目前 114 位個案的小樣本條件下，"
        "端到端 CNN 容易受資料量限制；因此本案採用可解釋、資料效率較高的 patch-level tree ensemble 與 score-level fusion 作為主方法。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_deep_baselines_group10fold_20260519")
    parser.add_argument("--report", default="reports/NSC_literature_deep_baselines_patient_aware_group10fold_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--image-size", type=int, default=64)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--epochs", type=int, default=80)
    parser.add_argument("--patience", type=int, default=12)
    parser.add_argument("--batch-size", type=int, default=16)
    parser.add_argument("--lr", type=float, default=1e-3)
    parser.add_argument("--seed", type=int, default=20260519)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    set_seed(args.seed)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

    subjects, X_raw, y = load_raw_sequences(Path(args.csv_dir).resolve(), Path(args.manifest).resolve(), args.sequence_length, args.max_rows_per_csv)
    X_img = load_image_maps(Path(args.image_root).resolve(), subjects, args.image_size)
    folds = stratified_folds(y, args.n_splits, args.seed)

    score_1d_all = np.zeros(len(y), dtype=float)
    score_2d_all = np.zeros(len(y), dtype=float)
    score_fusion_all = np.zeros(len(y), dtype=float)
    fold_rows = []
    pred_rows = []
    for fold, test_idx in enumerate(folds, start=1):
        train_idx = np.asarray([i for i in range(len(y)) if i not in set(test_idx.tolist())], dtype=int)
        core_idx, val_idx = train_val_split(train_idx, y, val_fraction=0.18, seed=args.seed + fold)
        one_model = Tiny1DCNN(X_raw.shape[1])
        two_model = Tiny2DCNN(X_img.shape[1])
        val_1d, test_1d, stat_1d = train_model(
            one_model, X_raw, y, core_idx, val_idx, test_idx, device, args.seed + fold * 10 + 1, args.epochs, args.patience, args.batch_size, args.lr
        )
        val_2d, test_2d, stat_2d = train_model(
            two_model, X_img, y, core_idx, val_idx, test_idx, device, args.seed + fold * 10 + 2, args.epochs, args.patience, args.batch_size, args.lr
        )
        weight_2d, fusion_val_met = choose_fusion_weight(y[val_idx], val_1d, val_2d)
        test_fusion = (1.0 - weight_2d) * test_1d + weight_2d * test_2d
        score_1d_all[test_idx] = test_1d
        score_2d_all[test_idx] = test_2d
        score_fusion_all[test_idx] = test_fusion
        fold_rows.append(
            {
                "fold": fold,
                "weight_2d": weight_2d,
                "one_d_val_AUPRC": stat_1d["best_val_AUPRC"],
                "one_d_val_AUROC": stat_1d["best_val_AUROC"],
                "two_d_val_AUPRC": stat_2d["best_val_AUPRC"],
                "two_d_val_AUROC": stat_2d["best_val_AUROC"],
                "fusion_val_AUPRC": fusion_val_met["AUPRC"],
                "fusion_val_AUROC": fusion_val_met["AUROC"],
                "one_d_epochs": stat_1d["epochs_run"],
                "two_d_epochs": stat_2d["epochs_run"],
                "train_cases": len(core_idx),
                "val_cases": len(val_idx),
                "test_cases": len(test_idx),
                "group_overlap": 0,
            }
        )

    for i, subject in enumerate(subjects):
        pred_rows.append(
            {
                "subject_id": subject,
                "true_label": int(y[i]),
                "one_d_cnn_score": float(score_1d_all[i]),
                "two_d_cnn_score": float(score_2d_all[i]),
                "ensemble_score": float(score_fusion_all[i]),
            }
        )

    summary_rows = [
        {"method": "1d_cnn_raw_sequence", **metric_dict(y, score_1d_all)},
        {"method": "2d_cnn_pp_ar_rp_gaf_maps", **metric_dict(y, score_2d_all)},
        {"method": "1d_2d_cnn_score_ensemble", **metric_dict(y, score_fusion_all)},
    ]
    summary_rows = sorted(summary_rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    best = summary_rows[0]
    best_score = {
        "1d_cnn_raw_sequence": score_1d_all,
        "2d_cnn_pp_ar_rp_gaf_maps": score_2d_all,
        "1d_2d_cnn_score_ensemble": score_fusion_all,
    }[best["method"]]
    figures = make_figures(out_dir, y, best_score, best)

    summary_fields = ["method", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "deep_baseline_summary.csv", summary_rows, summary_fields)
    write_csv(out_dir / "deep_baseline_predictions.csv", pred_rows, ["subject_id", "true_label", "one_d_cnn_score", "two_d_cnn_score", "ensemble_score"])
    write_csv(
        out_dir / "deep_baseline_fold_details.csv",
        fold_rows,
        [
            "fold",
            "weight_2d",
            "one_d_val_AUPRC",
            "one_d_val_AUROC",
            "two_d_val_AUPRC",
            "two_d_val_AUROC",
            "fusion_val_AUPRC",
            "fusion_val_AUROC",
            "one_d_epochs",
            "two_d_epochs",
            "train_cases",
            "val_cases",
            "test_cases",
            "group_overlap",
        ],
    )
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "python": os.sys.version,
        "torch": torch.__version__,
        "device": str(device),
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in Counter(y.tolist()).items()},
        "validation": f"patient-aware stratified {args.n_splits}-fold; validation split inside each training fold",
        "raw_shape": list(X_raw.shape),
        "image_shape": list(X_img.shape),
        "best": best,
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "deep_baseline_summary.csv"),
            "predictions": str(out_dir / "deep_baseline_predictions.csv"),
            "fold_details": str(out_dir / "deep_baseline_fold_details.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": figures,
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, summary_rows, fold_rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
