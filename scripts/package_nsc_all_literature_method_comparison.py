#!/usr/bin/env python3
"""Package all downloaded/referenced NSC literature-method comparisons.

The output is a report-ready comparison table. It separates:

* strict patient-level 10-fold results,
* representative random-holdout proof-of-signal,
* exact implementation vs proxy vs not-yet-implemented literature methods.
"""

from __future__ import annotations

import csv
import json
import math
from datetime import datetime
from pathlib import Path
from typing import Iterable, List

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis" / "nsc_all_literature_method_comparison_20260520"
REPORT = ROOT / "reports" / "NSC_all_downloaded_literature_method_comparison_20260520.docx"


def read_csv_rows(path: Path) -> List[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def metric_lookup() -> dict:
    lookup: dict[str, dict] = {}

    def as_float(value: str, default: float = -1.0) -> float:
        try:
            if value == "":
                return default
            return float(value)
        except Exception:
            return default

    def add(key: str, source: str, row: dict, method_col: str = "method") -> None:
        if not row:
            return
        lookup[key] = {
            "source": source,
            "method": row.get(method_col, key),
            "AUROC": row.get("AUROC", ""),
            "AUPRC": row.get("AUPRC", ""),
            "CM": "/".join(str(row.get(k, "")) for k in ["TN", "FP", "FN", "TP"]),
        }

    def find(rows: List[dict], key_col: str, value: str) -> dict:
        for row in rows:
            if row.get(key_col) == value:
                return row
        return {}

    def best_min_metric(rows: List[dict]) -> dict:
        usable = [r for r in rows if r.get("AUROC", "") != "" and r.get("AUPRC", "") != ""]
        if not usable:
            return {}
        return max(usable, key=lambda r: (min(as_float(r.get("AUROC", "")), as_float(r.get("AUPRC", ""))), as_float(r.get("AUPRC", "")), as_float(r.get("AUROC", ""))))

    def best_auroc(rows: List[dict]) -> dict:
        usable = [r for r in rows if r.get("AUROC", "") != "" and r.get("AUPRC", "") != ""]
        return max(usable, key=lambda r: (as_float(r.get("AUROC", "")), as_float(r.get("AUPRC", "")))) if usable else {}

    def best_auprc(rows: List[dict]) -> dict:
        usable = [r for r in rows if r.get("AUROC", "") != "" and r.get("AUPRC", "") != ""]
        return max(usable, key=lambda r: (as_float(r.get("AUPRC", "")), as_float(r.get("AUROC", "")))) if usable else {}

    paper = read_csv_rows(ROOT / "analysis" / "nsc_dataset_images_paper_method_comparison_20260519" / "method_comparison_summary.csv")
    deep = read_csv_rows(ROOT / "analysis" / "nsc_deep_baselines_group10fold_20260519" / "deep_baseline_summary.csv")
    raw_aug = read_csv_rows(ROOT / "analysis" / "nsc_raw_timeseries_dtw_tsmote_20260519" / "raw_timeseries_dtw_tsmote_summary.csv")
    fusion_aug = read_csv_rows(ROOT / "analysis" / "nsc_fusion_timeseries_2d_green_dtw_aug_20260519" / "fusion_method_summary.csv")
    uncertain = read_csv_rows(ROOT / "analysis" / "nsc_uncertain_band_patch_refinement_20260520" / "uncertain_band_summary.csv")
    ablation = read_csv_rows(ROOT / "analysis" / "nsc_raw2d_ablation_group10fold_20260519" / "ablation_summary.csv")
    vlm = read_csv_rows(ROOT / "analysis" / "nsc_vlm_feature_integration_strategy_20260519" / "vlm_feature_integration_summary.csv")
    artifact = read_csv_rows(ROOT / "analysis" / "nsc_artifact_family_rules_20260520" / "artifact_family_summary.csv")
    random_hits = read_csv_rows(ROOT / "analysis" / "nsc_dataset_images_random_holdout_signal_20260520" / "target_reaching_random_holdout_results.csv")
    eeg_agg = read_csv_rows(ROOT / "analysis" / "nsc_eeg_csv_fusion_ablation_noleak_fast_20260520" / "eeg_fusion_summary.csv")
    eeg_agg_count = read_csv_rows(ROOT / "analysis" / "nsc_eeg_csv_fusion_ablation_noleak_count_lr_20260520" / "eeg_fusion_summary.csv")
    eeg_trial = read_csv_rows(ROOT / "analysis" / "nsc_eeg_trial_level_fusion_20260520" / "trial_level_eeg_summary.csv")

    add("proposed_uncertain_band", "patient-aware 10-fold", find(uncertain, "branch", "mi_max"), method_col="branch")
    add("raw_initial", "patient-aware 10-fold", find(ablation, "method", "raw_initial"))
    add("raw_multiwindow", "patient-aware 10-fold", find(ablation, "method", "raw_multiwindow"))
    add("2d_patch_et", "patient-aware 10-fold", find(ablation, "method", "2d_et8_change_grid"))
    add("2d_var_et8", "patient-aware 10-fold", find(ablation, "method", "2d_var_et8"))
    add("raw_2d_fusion", "patient-aware 10-fold", find(ablation, "method", "2d_et8_change_grid"))
    add("green_pca_prototype", "patient-aware 10-fold", find(paper, "method", "green_pca_prototype"))
    add("green_pca_logistic", "patient-aware 10-fold", find(paper, "method", "green_pca_logistic"))
    add("patch_prototype", "patient-aware 10-fold", find(paper, "method", "patch_prototype"))
    add("patch_logistic", "patient-aware 10-fold", find(paper, "method", "patch_logistic"))
    add("patch_extratrees", "patient-aware 10-fold", find(paper, "method", "patch_extratrees"))
    add("relation_logistic", "patient-aware 10-fold", find(paper, "method", "relation_logistic"))
    add("smote_patch_logistic", "patient-aware 10-fold", find(paper, "method", "smote_patch_logistic"))
    add("2d_cnn", "patient-aware 10-fold", find(deep, "method", "2d_cnn_pp_ar_rp_gaf_maps"))
    add("1d_cnn", "patient-aware 10-fold", find(deep, "method", "1d_cnn_raw_sequence"))
    add("1d_2d_cnn_ensemble", "patient-aware 10-fold", find(deep, "method", "1d_2d_cnn_score_ensemble"))
    for row in raw_aug:
        if row.get("augmentation") == "dtw_tsmote" and row.get("model") == "extratrees":
            add("dtw_tsmote_extratrees", "patient-aware 10-fold", row)
    for row in fusion_aug:
        if row.get("feature_set") == "fusion" and row.get("augmentation") == "dtw_boundary" and row.get("model") == "extratrees":
            add("dtw_boundary_fusion", "patient-aware 10-fold", row)
    add("vlm_raw_plus_local", "patient-aware 10-fold", vlm[0] if vlm else {})
    add("artifact_family_suppressed", "patient-aware 10-fold", find(artifact, "method", "artifact_family_suppressed"))
    add("eeg_aggregate_fusion", "patient-aware 10-fold, fold-local aggregate EEG fusion", best_min_metric(eeg_agg))
    add("eeg_aggregate_count_fusion", "patient-aware 10-fold, fold-local aggregate EEG fusion", best_min_metric(eeg_agg_count))
    add("eeg_trial_fusion_best_balanced", "patient-aware 10-fold, trial OOF + posthoc fixed score fusion", best_min_metric(eeg_trial))
    add("eeg_trial_fusion_best_auroc", "patient-aware 10-fold, trial OOF + posthoc fixed score fusion", best_auroc(eeg_trial))
    add("eeg_trial_fusion_best_auprc", "patient-aware 10-fold, trial OOF + posthoc fixed score fusion", best_auprc(eeg_trial))
    if random_hits:
        row = random_hits[0]
        add("random_holdout_best", "representative random holdout, proof-of-signal only", row)
    return lookup


def f3(value: str) -> str:
    try:
        if value == "":
            return ""
        return f"{float(value):.3f}"
    except Exception:
        return str(value)


def literature_rows(metrics: dict) -> List[dict]:
    pdf_dir = ROOT / "references" / "nsc_2d_vs_1d_timeseries_20260519"
    downloaded = {
        "Wang & Oates 2015": str(pdf_dir / "01_wang_oates_ijcai2015_imaging_timeseries.pdf"),
        "Hatami et al. 2017": str(pdf_dir / "02_hatami_2017_timeseries_images_cnn.pdf"),
        "GAF ECG CNN 2023": str(pdf_dir / "03_ecg_deep_cnn_gaf_2023.pdf"),
        "1D vs 2D ECG CNN 2018": str(pdf_dir / "04_ecg_1d_vs_2d_cnn_2018.pdf"),
        "Mix Time-Series Imaging 2022": str(pdf_dir / "05_mix_timeseries_imaging_ecg_2022.pdf"),
        "1D/2D CNN seismic 2022": str(pdf_dir / "06_1d_2d_cnn_seismic_comparison_2022.pdf"),
        "1D+2D heart sound ensemble 2018": str(pdf_dir / "07_heart_sound_1d_2d_ensemble_2018.pdf"),
        "Gliner ECG image interpretability 2025": str(ROOT / "downloaded_papers" / "gliner-2025-clinically-meaningful-interpretability-ecg-images.pdf"),
    }

    def m(key: str) -> tuple[str, str, str]:
        row = metrics.get(key, {})
        return f3(row.get("AUROC", "")), f3(row.get("AUPRC", "")), row.get("CM", "")

    rows = []

    def add(
        category: str,
        reference: str,
        literature_method: str,
        local_status: str,
        local_proxy: str,
        metric_key: str,
        conclusion: str,
        evidence_level: str = "patient-aware 10-fold",
        local_file: str = "",
    ) -> None:
        auroc, auprc, cm = m(metric_key)
        rows.append(
            {
                "category": category,
                "reference": reference,
                "literature_method": literature_method,
                "downloaded_or_linked": "PDF downloaded" if local_file and Path(local_file).exists() else "linked/reference only",
                "local_file": local_file if local_file and Path(local_file).exists() else "",
                "local_status": local_status,
                "local_proxy_or_exact_method": local_proxy,
                "validation_level": evidence_level,
                "AUROC": auroc,
                "AUPRC": auprc,
                "CM": cm,
                "paper_conclusion": conclusion,
            }
        )

    add("Time-series imaging", "Wang & Oates 2015", "GAF/MTF/RP time-series images + CNN/tiled CNN", "proxy + baseline implemented", "PP/AR/RP/GAF 2D patch ET, 2D CNN, mi_max uncertain-band", "proposed_uncertain_band", "Supports converting 1D signals into 2D images; our strict best uses 2D image-instance score in uncertain-band fusion.", local_file=downloaded["Wang & Oates 2015"])
    add("Time-series imaging", "Eckmann et al. 1987", "Recurrence plot for dynamical systems", "proxy implemented", "RP local texture features in patch/instance models", "2d_patch_et", "RP is used as a representation family; not a standalone classifier.", local_file="")
    add("Time-series imaging", "Huang et al. 2025", "RP + GAF fusion with optimization", "proxy implemented", "multi-plot/multiscale fusion and nested grid search", "raw_2d_fusion", "Fusion is helpful only when constrained; high-dimensional multi-plot fusion was unstable in low-N.", local_file="")
    add("2D deep image", "GAF ECG CNN 2023", "GAF image + deep CNN", "lightweight baseline implemented", "2D CNN over PP/AR/RP/GAF maps", "2d_cnn", "Deep CNN did not outperform small-data tree/patch fusion under patient-aware 10-fold.", local_file=downloaded["GAF ECG CNN 2023"])
    add("1D vs 2D deep", "1D vs 2D ECG CNN 2018", "1D CNN, 2D CNN, transfer/ensemble comparison", "exact family baseline implemented", "1D CNN, 2D CNN, 1D+2D CNN score ensemble", "1d_2d_cnn_ensemble", "CNN ensemble was weaker than non-deep fusion; keep as literature baseline, not main model.", local_file=downloaded["1D vs 2D ECG CNN 2018"])
    add("2D deep image", "Hatami et al. 2017", "RP + deep CNN", "lightweight baseline/proxy implemented", "RP included in 2D CNN and patch ET branches", "2d_cnn", "Direct RP-CNN family tested as lightweight 2D CNN; low-N limits deep training.", local_file=downloaded["Hatami et al. 2017"])
    add("Multi-image fusion", "Mix Time-Series Imaging 2022", "GAF + RP multi-channel/tiled image fusion", "proxy implemented", "multi-plot/multiscale 2D fusion, score-level fusion", "raw_2d_fusion", "Multi-image representation is useful motivation; our strict best uses selective fusion rather than large deep fusion.", local_file=downloaded["Mix Time-Series Imaging 2022"])
    add("1D/2D comparison", "1D/2D CNN seismic 2022", "1D waveform vs 2D representation CNN comparison", "proxy implemented", "raw-only, 2D-only, raw+2D late fusion comparison", "raw_2d_fusion", "Supports reporting both 1D and 2D branches; domain differs, so no direct score comparison.", local_file=downloaded["1D/2D CNN seismic 2022"])
    add("Score-level ensemble", "1D+2D heart sound ensemble 2018", "1D-CNN + 2D-CNN score-level ensemble", "exact family baseline implemented", "1D+2D CNN ensemble and non-deep score-level fusion", "1d_2d_cnn_ensemble", "Deep ensemble baseline exists but is not strongest; non-deep uncertain-band fusion is better.", local_file=downloaded["1D+2D heart sound ensemble 2018"])
    add("Interpretability", "Gliner ECG image interpretability 2025", "Clinically meaningful interpretability of ECG images", "proxy implemented", "image-instance artifact review, local artifact family rules, VLM-style patch interpretation", "artifact_family_suppressed", "Interpretability/artifact review is useful for feature design, but direct suppression reduced AUPRC.", local_file=downloaded["Gliner ECG image interpretability 2025"])

    add("Few-shot prototype", "Snell et al. 2017 Prototypical Networks", "class prototype distance", "proxy implemented", "patch_prototype / green_pca_prototype", "green_pca_prototype", "Prototype/Green prototype is stable but below proposed fusion in strict 10-fold.")
    add("Few-shot metric", "Koch et al. 2015 Siamese Networks", "pairwise similarity learning", "shallow proxy implemented", "relation_logistic / distance features", "relation_logistic", "Metric-learning proxy underperformed; neural Siamese deferred until more cases.")
    add("Few-shot matching", "Vinyals et al. 2016 Matching Networks", "support/query episodic classifier", "not directly implemented", "prototype/relation proxy", "patch_prototype", "Can be referenced as future episodic-learning baseline; current N is small for stable neural episodic training.")
    add("Meta-learning", "Finn et al. 2017 MAML", "gradient-based fast adaptation", "not implemented", "no direct model; validation risk too high", "", "Keep as future work after additional patients; not suitable as current main comparison.")
    add("Few-shot relation", "Sung et al. 2018 Relation Network", "learned relation module", "shallow proxy implemented", "relation_logistic", "relation_logistic", "Relation proxy is reported; full neural relation network is not promoted due low-N overfitting risk.")
    add("Medical few-shot", "ProtoMed 2024", "prototype + auxiliary regularization for medical images", "proxy implemented", "green_pca_prototype / patch_prototype", "green_pca_prototype", "Supports prototype regularization direction; proxy did not exceed fusion.")
    add("Medical meta-learning", "MetaMed 2021", "gradient-based meta-learning for medical images", "not implemented", "no direct model; future after more cases", "", "Mention as related work, not as current result.")
    add("Medical few-shot", "Skin lesion prototypical network 2024", "medical image prototype classifier", "proxy implemented", "patch_prototype", "patch_prototype", "Comparable idea implemented as patch prototype; below proposed fusion.")
    add("Spatial attention", "Collateral location coding + spatial attention 2022", "location-sensitive medical image few-shot", "proxy implemented", "channel/plot/patch-specific local features and artifact-family scoring", "vlm_raw_plus_local", "Supports our local-position feature design; direct attention network deferred.")

    add("Green learning", "PixelHop 2019", "successive subspace / feedforward local features", "proxy implemented", "green_pca_prototype / green_pca_logistic", "green_pca_prototype", "Green proxy is useful but not strongest; supports non-gradient small-data feature design.")
    add("Green learning", "PixelHop++ 2020", "small successive-subspace model", "proxy implemented", "Green/PCA feature selection + prototype/logistic", "green_pca_logistic", "Small feedforward feature extraction is included as proxy.")
    add("Green learning", "Interpretable CNN via Feedforward Design 2018", "PCA/feedforward interpretable CNN", "proxy implemented", "Green/PCA local descriptors", "green_pca_prototype", "Supports interpretable feature extraction rather than end-to-end deep training.")

    add("Data augmentation", "Iwana & Uchida 2021", "time-series augmentation survey", "implemented family tests", "raw augmentation, DTW/T-SMOTE, fusion augmentation", "dtw_boundary_fusion", "Augmentation helps some branches but does not exceed proposed uncertain-band fusion.")
    add("Data augmentation", "T-SMOTE 2022", "temporal-oriented minority oversampling", "proxy implemented", "DTW/T-SMOTE ExtraTrees", "dtw_tsmote_extratrees", "DTW/T-SMOTE improves raw branch AUPRC vs some baselines but remains below 0.8.")
    add("Data augmentation", "SMOTE 2002", "synthetic minority oversampling", "feature-space proxy implemented", "smote_patch_logistic", "smote_patch_logistic", "Plain feature-space SMOTE is weaker; time-series-aware augmentation is preferred.")
    add("Data augmentation", "DTWSSE 2021", "DTW + Siamese encoder augmentation", "partial proxy implemented", "DTW boundary/T-SMOTE + relation proxy; no Siamese encoder yet", "dtw_tsmote_extratrees", "Keep as future stronger augmentation; current partial proxy below proposed fusion.")
    add("Validation/metric", "Saito & Rehmsmeier 2015", "PRC/AUPRC for imbalanced binary classification", "metric adopted", "AUPRC primary; AUROC secondary", "proposed_uncertain_band", "Justifies reporting PRC alongside ROC; not a model method.")
    add("Validation/metric", "Varoquaux 2018", "small-sample cross-validation instability", "validation interpretation adopted", "patient-aware 10-fold + random holdout proof-of-signal separation", "random_holdout_best", "Explains strict 10-fold variability; random holdout must not be claimed as final generalization.")
    add("EEG trial classifier", "EEG feature/statistical classifier literature family", "trial-level EEG classifier + patient score aggregation", "implemented on 114 cases", "training-fold trial classifier, patient-level max/top aggregation, mi_max score fusion", "eeg_trial_fusion_best_balanced", "Best balanced EEG trial fusion improves over mi_max but still falls just below 0.8/0.8; fixed-score fusion is marked exploratory/posthoc.")
    add("EEG aggregate classifier", "EEG spectral/statistical feature literature family", "patient-level EEG statistical/spectral features + LR fusion", "implemented on 114 cases", "non-leaky all-trial EEG aggregate LR + mi_max score fusion", "eeg_aggregate_fusion", "Aggregate EEG fusion is stricter/non-leaky and improves AUPRC over baseline, but remains below 0.8/0.8.")

    return rows


def implemented_result_rows(metrics: dict) -> List[dict]:
    order = [
        ("proposed_uncertain_band", "Proposed mi_max uncertain-band fusion"),
        ("raw_2d_fusion", "Raw + 2D ET8 grid fusion"),
        ("raw_initial", "Raw initial-segment branch"),
        ("dtw_tsmote_extratrees", "Raw DTW/T-SMOTE ExtraTrees"),
        ("green_pca_prototype", "Green/PCA prototype"),
        ("smote_patch_logistic", "SMOTE patch logistic"),
        ("patch_prototype", "Patch prototype"),
        ("relation_logistic", "Relation-learning proxy"),
        ("2d_cnn", "2D CNN PP/AR/RP/GAF"),
        ("1d_2d_cnn_ensemble", "1D+2D CNN ensemble"),
        ("1d_cnn", "1D CNN raw sequence"),
        ("artifact_family_suppressed", "Artifact-family suppressed"),
        ("eeg_aggregate_fusion", "EEG aggregate LR fusion"),
        ("eeg_aggregate_count_fusion", "EEG aggregate LR fusion + availability counts"),
        ("eeg_trial_fusion_best_balanced", "EEG trial-level classifier fusion, balanced best"),
        ("eeg_trial_fusion_best_auroc", "EEG trial-level classifier fusion, AUROC best"),
        ("eeg_trial_fusion_best_auprc", "EEG trial-level classifier fusion, AUPRC best"),
        ("random_holdout_best", "Random holdout best, proof-of-signal"),
    ]
    rows = []
    for key, label in order:
        row = metrics.get(key, {})
        if not row:
            continue
        rows.append(
            {
                "key": key,
                "method_label": label,
                "source": row.get("source", ""),
                "AUROC": row.get("AUROC", ""),
                "AUPRC": row.get("AUPRC", ""),
                "CM": row.get("CM", ""),
            }
        )
    return sorted(
        rows,
        key=lambda r: (
            1 if "random holdout" in r["source"].lower() else 0,
            -float(r["AUPRC"]) if r.get("AUPRC") else 0.0,
            -float(r["AUROC"]) if r.get("AUROC") else 0.0,
        ),
    )


def make_figures(rows: List[dict]) -> dict:
    fig_dir = OUT_DIR / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    strict = [r for r in rows if "random holdout" not in r["source"].lower()]
    strict = sorted(strict, key=lambda r: float(r["AUPRC"]), reverse=True)[:12]
    labels = [r["method_label"] for r in strict]
    y = list(range(len(strict)))
    plt.figure(figsize=(10.5, max(5, len(strict) * 0.42 + 1.2)))
    plt.barh([i + 0.18 for i in y], [float(r["AUROC"]) for r in strict], height=0.34, label="AUROC")
    plt.barh([i - 0.18 for i in y], [float(r["AUPRC"]) for r in strict], height=0.34, label="AUPRC")
    plt.axvline(0.8, color="gray", linestyle="--")
    plt.yticks(y, labels, fontsize=8)
    plt.xlabel("Metric")
    plt.title("Strict patient-aware 10-fold methods mapped to literature")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    plt.gca().invert_yaxis()
    path = fig_dir / "strict_literature_method_bar.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return {"strict_bar": str(path)}


def best_rows(result_rows: List[dict]) -> dict:
    strict = [r for r in result_rows if "random holdout" not in r["source"].lower()]
    strict_numeric = [
        r for r in strict
        if r.get("AUROC", "") not in ["", None] and r.get("AUPRC", "") not in ["", None]
    ]
    return {
        "balanced": max(
            strict_numeric,
            key=lambda r: (min(float(r["AUROC"]), float(r["AUPRC"])), float(r["AUPRC"]), float(r["AUROC"])),
        ) if strict_numeric else {},
        "auroc": max(strict_numeric, key=lambda r: float(r["AUROC"])) if strict_numeric else {},
        "auprc": max(strict_numeric, key=lambda r: float(r["AUPRC"])) if strict_numeric else {},
    }


def make_report(lit_rows: List[dict], result_rows: List[dict], figures: dict) -> None:
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("NSC 所有參考文獻方法比較報告", level=0)
    doc.add_paragraph(f"產出日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、總結", level=1)
    doc.add_paragraph(
        "本版將已下載 PDF 與前一版引用清單中的方法整理成 114 位資料上的 ROC/PRC 比較。"
        "可直接落地的文獻模型家族已用 patient-aware 10-fold 實跑；"
        "若原文模型需要大樣本深度訓練或資料型態不同，則以本案可重跑的 proxy 版本比較，"
        "並在表中標註為 proxy 或 future replication。"
    )
    best = best_rows(result_rows)
    balanced = best.get("balanced", {})
    auroc_best = best.get("auroc", {})
    auprc_best = best.get("auprc", {})
    doc.add_paragraph(
        f"目前 patient-aware 10-fold 的平衡最佳結果為 {balanced.get('method_label', '')}："
        f"AUROC {f3(balanced.get('AUROC', ''))}、AUPRC {f3(balanced.get('AUPRC', ''))}。"
        f"AUROC 最高結果為 {auroc_best.get('method_label', '')}：AUROC {f3(auroc_best.get('AUROC', ''))}、"
        f"AUPRC {f3(auroc_best.get('AUPRC', ''))}；"
        f"AUPRC 最高結果為 {auprc_best.get('method_label', '')}：AUROC {f3(auprc_best.get('AUROC', ''))}、"
        f"AUPRC {f3(auprc_best.get('AUPRC', ''))}。"
        "截至本版，嚴格/非洩漏 patient-aware 10-fold 尚無模型同時穩定達到 AUROC >= 0.8 與 AUPRC >= 0.8。"
        "random holdout proof-of-signal 仍只作代表性訊號證據，不作最終泛化宣稱。"
    )

    doc.add_heading("二、本案已實測/Proxy 方法結果", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["方法", "證據等級", "AUROC", "AUPRC", "CM", "對應文獻角色"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in result_rows:
        role = {
            "proposed_uncertain_band": "本案主方法；2D image-instance + raw uncertain-band fusion",
            "raw_2d_fusion": "1D/2D fusion 對照",
            "raw_initial": "1D raw baseline",
            "dtw_tsmote_extratrees": "T-SMOTE/DTW augmentation proxy",
            "green_pca_prototype": "PixelHop/Green/prototype proxy",
            "smote_patch_logistic": "SMOTE proxy",
            "patch_prototype": "Prototypical Network proxy",
            "relation_logistic": "Relation/Siamese proxy",
            "2d_cnn": "GAF/RP/2D CNN literature baseline",
            "1d_2d_cnn_ensemble": "1D+2D CNN ensemble baseline",
            "1d_cnn": "1D CNN baseline",
            "artifact_family_suppressed": "clinical image interpretability/artifact proxy",
            "eeg_aggregate_fusion": "EEG spectral/statistical aggregate + mi_max fusion",
            "eeg_aggregate_count_fusion": "EEG aggregate + availability-count sensitivity check",
            "eeg_trial_fusion_best_balanced": "EEG trial classifier + patient score aggregation; posthoc fixed fusion",
            "eeg_trial_fusion_best_auroc": "EEG trial fusion, AUROC-focused sensitivity result",
            "eeg_trial_fusion_best_auprc": "EEG trial fusion, AUPRC-focused sensitivity result",
            "random_holdout_best": "representative proof-of-signal only",
        }.get(r["key"], "")
        vals = [r["method_label"], r["source"], f3(r["AUROC"]), f3(r["AUPRC"]), r["CM"], role]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    if figures.get("strict_bar"):
        doc.add_paragraph("圖 1. 文獻方法對應之 strict patient-aware 10-fold 實測結果")
        doc.add_picture(figures["strict_bar"], width=Inches(6.4))

    doc.add_heading("三、所有參考文獻方法逐項比較", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["類別", "文獻", "文獻方法", "下載狀態", "本案狀態", "本案對應方法", "AUROC", "AUPRC", "結論"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in lit_rows:
        vals = [
            r["category"],
            r["reference"],
            r["literature_method"],
            r["downloaded_or_linked"],
            r["local_status"],
            r["local_proxy_or_exact_method"],
            r["AUROC"],
            r["AUPRC"],
            r["paper_conclusion"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、建議寫入 Paper 的比較設計", level=1)
    for item in [
        "主表：strict patient-aware 10-fold，列 proposed uncertain-band、raw-only、2D-only、raw+2D fusion、Green/PCA、prototype、relation proxy、SMOTE/T-SMOTE、1D CNN、2D CNN、1D+2D CNN。",
        "新增表：EEG CSV 參考模型家族，列 fold-local aggregate EEG fusion 與 trial-level EEG classifier fusion；trial-level fixed-score fusion需標註為 exploratory/posthoc。",
        "補充表：representative random holdout proof-of-signal，明確註明不是 final generalization claim。",
        "文獻方法表：每篇參考文獻都對應到 exact implementation、proxy implementation 或 future replication。",
        "討論：深度 CNN/metric/meta-learning 在本案小樣本下不作主模型；目前最接近 0.8/0.8 的方向是 mi_max 與 EEG trial/aggregate score-level fusion。",
        "限制：114 位個案不足以穩定支撐所有深度文獻方法的端到端復刻，新增個案後需重跑 frozen protocol。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("五、已下載 PDF 檢查", level=1)
    for r in lit_rows:
        if r["local_file"]:
            doc.add_paragraph(r["local_file"], style="List Bullet")

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = metric_lookup()
    lit = literature_rows(metrics)
    results = implemented_result_rows(metrics)
    figures = make_figures(results)
    write_csv(
        OUT_DIR / "all_literature_method_comparison.csv",
        lit,
        [
            "category",
            "reference",
            "literature_method",
            "downloaded_or_linked",
            "local_file",
            "local_status",
            "local_proxy_or_exact_method",
            "validation_level",
            "AUROC",
            "AUPRC",
            "CM",
            "paper_conclusion",
        ],
    )
    write_csv(
        OUT_DIR / "implemented_method_results.csv",
        results,
        ["key", "method_label", "source", "AUROC", "AUPRC", "CM"],
    )
    with (OUT_DIR / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(
            {
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "report": str(REPORT),
                "analysis_dir": str(OUT_DIR),
                "literature_rows": len(lit),
                "implemented_rows": len(results),
                "figures": figures,
            },
            f,
            ensure_ascii=False,
            indent=2,
        )
    make_report(lit, results, figures)
    print(json.dumps({"report": str(REPORT), "analysis_dir": str(OUT_DIR), "figures": figures}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
