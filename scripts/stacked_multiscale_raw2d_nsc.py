#!/usr/bin/env python3
"""Patient-aware group 10-fold raw multi-window + 2D multi-scale stacking.

This is a stricter score-level stacking experiment:

- raw branch: multi-window time-series statistics with train-fold-only
  DTW/T-SMOTE augmentation and ExtraTrees.
- 2D branch: multi-scale 4x4/8x8/16x16 patch features with fold-local
  imputation, variance filtering, feature selection, and ExtraTrees.
- stacking: outer-train-only inner out-of-fold branch scores train a logistic
  calibration/stacking model. The outer test fold is not used for branch
  training, augmentation, feature filtering, or stacker fitting.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from PIL import Image
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, VarianceThreshold, f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, confusion_matrix, precision_recall_curve, roc_auc_score, roc_curve
from sklearn.model_selection import StratifiedGroupKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from heuristic_rule_learning_nsc_dataset_images import (  # noqa: E402
    PLOT_NAMES,
    case_id_from_session,
    channel_group,
    collect_case_images,
    natural_key,
)
from raw_timeseries_dtw_tsmote_nsc_dataset_images import (  # noqa: E402
    dtw_tsmote_augment,
    extract_features,
    load_case_sequences,
)
from score_level_fusion_2d_raw_nsc import make_figures, metric_dict, write_csv  # noqa: E402


def safe_roc_auc(y_true: np.ndarray, score: np.ndarray) -> float:
    return float(roc_auc_score(y_true, score)) if len(np.unique(y_true)) == 2 else math.nan


def safe_auprc(y_true: np.ndarray, score: np.ndarray) -> float:
    return float(average_precision_score(y_true, score)) if len(np.unique(y_true)) == 2 else math.nan


def image_patch_features(path: Path, grid: int) -> np.ndarray:
    img = Image.open(path).convert("L").resize((128, 128))
    arr = np.asarray(img, dtype=np.float32) / 255.0
    ink = 1.0 - arr
    h, w = ink.shape
    patch_h = h // grid
    patch_w = w // grid
    feats = np.zeros((grid, grid), dtype=np.float32)
    for r in range(grid):
        for c in range(grid):
            patch = ink[r * patch_h : (r + 1) * patch_h, c * patch_w : (c + 1) * patch_w]
            feats[r, c] = float(patch.mean())
    return feats


def build_2d_multiscale_matrix(image_root: Path, subjects: List[str], grids: Tuple[int, ...] = (4, 8, 16)) -> Tuple[np.ndarray, List[str], dict]:
    cases = collect_case_images(image_root)
    subject_buckets: Dict[str, Dict[str, List[float]]] = defaultdict(lambda: defaultdict(list))
    image_count_by_subject: Dict[str, int] = defaultdict(int)
    for meta in cases.values():
        case_id = meta["case_id"]
        for path in meta["image_paths"]:
            if path.stem not in PLOT_NAMES:
                continue
            channel = path.parent.name
            group = channel_group(channel)
            plot = path.stem
            image_count_by_subject[case_id] += 1
            for grid in grids:
                patches = image_patch_features(path, grid)
                for r in range(grid):
                    for c in range(grid):
                        value = float(patches[r, c])
                        subject_buckets[case_id][f"g{grid}/ALL_CHANNELS/{plot}/r{r}c{c}"].append(value)
                        subject_buckets[case_id][f"g{grid}/{group}/{plot}/r{r}c{c}"].append(value)
                        subject_buckets[case_id][f"g{grid}/CHANNEL_{channel}/{plot}/r{r}c{c}"].append(value)

    subject_features: Dict[str, Dict[str, float]] = {}
    for subject, buckets in subject_buckets.items():
        subject_features[subject] = {k: float(np.mean(v)) for k, v in buckets.items() if v}

    missing = [s for s in subjects if s not in subject_features]
    if missing:
        raise ValueError(f"Missing 2D images for subjects: {missing[:10]}")
    feature_names = sorted({f for s in subjects for f in subject_features[s]})
    idx = {f: i for i, f in enumerate(feature_names)}
    X = np.full((len(subjects), len(feature_names)), np.nan, dtype=np.float32)
    for i, subject in enumerate(subjects):
        for f, value in subject_features[subject].items():
            X[i, idx[f]] = value
    meta = {
        "grids": list(grids),
        "feature_count": len(feature_names),
        "image_counts": {s: int(image_count_by_subject[s]) for s in subjects},
    }
    return X, feature_names, meta


def extract_multiwindow_features(X_seq: np.ndarray) -> np.ndarray:
    windows = [
        ("full", 0.0, 1.0),
        ("early25", 0.0, 0.25),
        ("mid25a", 0.25, 0.50),
        ("mid25b", 0.50, 0.75),
        ("late25", 0.75, 1.0),
        ("early50", 0.0, 0.50),
        ("late50", 0.50, 1.0),
    ]
    parts = []
    length = X_seq.shape[2]
    for _name, start_frac, end_frac in windows:
        start = int(round(start_frac * length))
        end = int(round(end_frac * length))
        end = max(end, start + 8)
        end = min(end, length)
        parts.append(extract_features(X_seq[:, :, start:end]))
    return np.concatenate(parts, axis=1).astype(np.float32)


def transform_with_fold_local_filter(
    X_train: np.ndarray,
    y_train: np.ndarray,
    X_train_original: np.ndarray,
    X_test: np.ndarray,
    top_k: int,
) -> Tuple[np.ndarray, np.ndarray, np.ndarray, dict]:
    imputer = SimpleImputer(strategy="median")
    Xtr = imputer.fit_transform(X_train)
    Xtr_orig = imputer.transform(X_train_original)
    Xte = imputer.transform(X_test)

    var = VarianceThreshold(threshold=1e-8)
    Xtr_v = var.fit_transform(Xtr)
    Xtr_orig_v = var.transform(Xtr_orig)
    Xte_v = var.transform(Xte)
    if Xtr_v.shape[1] == 0:
        Xtr_v, Xtr_orig_v, Xte_v = Xtr, Xtr_orig, Xte

    k = min(top_k, Xtr_v.shape[1])
    selector = SelectKBest(f_classif, k=k)
    with np.errstate(invalid="ignore", divide="ignore"):
        Xtr_s = selector.fit_transform(Xtr_v, y_train)
    Xtr_orig_s = selector.transform(Xtr_orig_v)
    Xte_s = selector.transform(Xte_v)
    stats = {
        "input_features": int(X_train.shape[1]),
        "after_variance_features": int(Xtr_v.shape[1]),
        "selected_features": int(k),
    }
    return Xtr_s, Xtr_orig_s, Xte_s, stats


def fit_raw_multiwindow_et(Xseq: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, top_k: int) -> Tuple[np.ndarray, np.ndarray, int, dict]:
    X_aug, y_aug, n_syn = dtw_tsmote_augment(Xseq[train_idx], y[train_idx], target_label=1, random_state=seed)
    Xtr = extract_multiwindow_features(X_aug)
    Xtr_orig = extract_multiwindow_features(Xseq[train_idx])
    Xte = extract_multiwindow_features(Xseq[test_idx])
    Xtr_s, Xtr_orig_s, Xte_s, stats = transform_with_fold_local_filter(Xtr, y_aug, Xtr_orig, Xte, top_k=top_k)
    model = ExtraTreesClassifier(
        n_estimators=420,
        random_state=seed,
        class_weight="balanced",
        max_features="sqrt",
        min_samples_leaf=2,
        n_jobs=-1,
    )
    model.fit(Xtr_s, y_aug)
    return model.predict_proba(Xtr_orig_s)[:, 1], model.predict_proba(Xte_s)[:, 1], n_syn, stats


def fit_2d_multiscale_et(X2d: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, top_k: int) -> Tuple[np.ndarray, np.ndarray, dict]:
    Xtr = X2d[train_idx]
    Xte = X2d[test_idx]
    Xtr_s, Xtr_orig_s, Xte_s, stats = transform_with_fold_local_filter(Xtr, y[train_idx], Xtr, Xte, top_k=top_k)
    model = ExtraTreesClassifier(
        n_estimators=520,
        random_state=seed,
        class_weight="balanced",
        max_features="sqrt",
        min_samples_leaf=2,
        n_jobs=-1,
    )
    model.fit(Xtr_s, y[train_idx])
    return model.predict_proba(Xtr_orig_s)[:, 1], model.predict_proba(Xte_s)[:, 1], stats


def stack_features(raw_score: np.ndarray, two_score: np.ndarray, kind: str) -> np.ndarray:
    raw_score = np.asarray(raw_score, dtype=float)
    two_score = np.asarray(two_score, dtype=float)
    if kind == "scores":
        return np.column_stack([raw_score, two_score])
    if kind == "scores_interaction":
        return np.column_stack([raw_score, two_score, raw_score * two_score, np.abs(raw_score - two_score)])
    raise ValueError(kind)


def fit_meta_model(raw_oof: np.ndarray, two_oof: np.ndarray, y_train: np.ndarray) -> Tuple[object, dict, List[dict]]:
    candidates = []
    for kind in ["scores", "scores_interaction"]:
        X_meta = stack_features(raw_oof, two_oof, kind)
        for c_value in [0.1, 0.25, 0.5, 1.0]:
            for class_weight in [None, "balanced"]:
                model = make_pipeline(
                    StandardScaler(),
                    LogisticRegression(max_iter=4000, C=c_value, class_weight=class_weight, solver="liblinear"),
                )
                model.fit(X_meta, y_train)
                score = model.predict_proba(X_meta)[:, 1]
                met = metric_dict(y_train, score)
                row = {
                    "meta_kind": kind,
                    "meta_C": c_value,
                    "meta_class_weight": "none" if class_weight is None else str(class_weight),
                    **met,
                }
                candidates.append((model, row))
    best_model, best_row = max(candidates, key=lambda item: (item[1]["AUPRC"], item[1]["AUROC"]))
    return best_model, best_row, [row for _model, row in candidates]


def make_report(report_path: Path, manifest: dict, rows: List[dict], fold_rows: List[dict], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC Raw Multi-window + 2D Multi-scale Stacking 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版採 patient-aware group 10-fold。raw branch 使用 full/early/middle/late/half 等多個時間窗抽取統計與頻域特徵，"
        "並只在 training fold 內做 DTW/T-SMOTE 增廣。2D branch 使用 4x4、8x8、16x16 多尺度 patch，"
        "包含 ALL_CHANNELS、channel group 與 channel-specific 特徵。"
    )
    doc.add_paragraph(
        "兩個 branch 均在各 fold 內完成 median imputation、variance filtering、SelectKBest 與 ExtraTrees 訓練。"
        "stacking 使用 outer training fold 的 inner out-of-fold branch scores 訓練 logistic calibration model，"
        "outer test fold 不參與任何前處理、模型訓練、augmentation 或 stacking 校準。"
    )

    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in rows:
        cells = table.add_row().cells
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、各折 stacking 與特徵統計", level=1)
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = ["fold", "meta_kind", "C", "class_weight", "synthetic", "raw_selected", "2d_var", "2d_selected", "test", "overlap"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in fold_rows:
        vals = [
            r["fold"],
            r["meta_kind"],
            r["meta_C"],
            r["meta_class_weight"],
            r["synthetic_samples"],
            r["raw_selected_features"],
            r["two_d_after_variance_features"],
            r["two_d_selected_features"],
            r["test_cases"],
            r["group_overlap"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、Stacking 圖表", level=1)
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))

    doc.add_heading("五、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(
        f"最佳方法為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。"
        "若 stacking 未優於單一 branch，代表目前 outer-train OOF 校準器尚未穩定學到兩個分支互補；"
        "若 stacking 優於單一 branch，則支持 raw multi-window 與 2D multi-scale patch 存在互補訊號。"
    )
    doc.add_paragraph(
        "本報告採 patient-aware group split，case_id/subject_id 不跨 train/test；因此比 random holdout 更適合作為個案層級泛化證據。"
        "但若仍未達 0.8/0.8，後續應聚焦於 fold-local 2D patch 錯誤分析、raw window selection、以及更穩定的 branch-level calibration。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_multiscale_raw2d_stacking_group10fold_20260519")
    parser.add_argument("--report", default="reports/NSC_raw_multiwindow_2D_multiscale_stacking_patient_aware_group10fold_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--raw-top-k", type=int, default=96)
    parser.add_argument("--two-d-top-k", type=int, default=160)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    subjects, Xseq, y, _meta = load_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.max_rows_per_csv,
    )
    X2d, feature_names_2d, meta_2d = build_2d_multiscale_matrix(Path(args.image_root).resolve(), subjects)
    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)

    raw_all = np.zeros(len(y), dtype=float)
    two_all = np.zeros(len(y), dtype=float)
    stacked_all = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    meta_rows_all: List[dict] = []

    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 100)
        raw_oof = np.zeros(len(train_idx), dtype=float)
        two_oof = np.zeros(len(train_idx), dtype=float)
        for inner_fold, (tr_local, va_local) in enumerate(
            inner.split(np.zeros(len(train_idx)), y[train_idx], groups=groups[train_idx]),
            start=1,
        ):
            tr_idx = train_idx[tr_local]
            va_idx = train_idx[va_local]
            _, raw_val, _n_syn, _raw_stats = fit_raw_multiwindow_et(Xseq, y, tr_idx, va_idx, args.random_state + fold * 1000 + inner_fold, args.raw_top_k)
            _, two_val, _two_stats = fit_2d_multiscale_et(X2d, y, tr_idx, va_idx, args.random_state + fold * 1000 + inner_fold, args.two_d_top_k)
            raw_oof[va_local] = raw_val
            two_oof[va_local] = two_val

        meta_model, best_meta, meta_candidates = fit_meta_model(raw_oof, two_oof, y[train_idx])
        for row in meta_candidates:
            row["outer_fold"] = fold
            meta_rows_all.append(row)

        raw_train, raw_test, n_syn, raw_stats = fit_raw_multiwindow_et(Xseq, y, train_idx, test_idx, args.random_state + fold, args.raw_top_k)
        two_train, two_test, two_stats = fit_2d_multiscale_et(X2d, y, train_idx, test_idx, args.random_state + fold, args.two_d_top_k)
        X_test_meta = stack_features(raw_test, two_test, best_meta["meta_kind"])
        stacked_test = meta_model.predict_proba(X_test_meta)[:, 1]

        raw_all[test_idx] = raw_test
        two_all[test_idx] = two_test
        stacked_all[test_idx] = stacked_test

        train_groups = set(groups[train_idx].tolist())
        test_groups = set(groups[test_idx].tolist())
        fold_rows.append(
            {
                "fold": fold,
                "meta_kind": best_meta["meta_kind"],
                "meta_C": best_meta["meta_C"],
                "meta_class_weight": best_meta["meta_class_weight"],
                "synthetic_samples": int(n_syn),
                "raw_input_features": raw_stats["input_features"],
                "raw_after_variance_features": raw_stats["after_variance_features"],
                "raw_selected_features": raw_stats["selected_features"],
                "two_d_input_features": two_stats["input_features"],
                "two_d_after_variance_features": two_stats["after_variance_features"],
                "two_d_selected_features": two_stats["selected_features"],
                "train_cases": int(len(train_idx)),
                "test_cases": int(len(test_idx)),
                "group_overlap": int(len(train_groups & test_groups)),
            }
        )

    rows = [
        {"method": "raw_multiwindow_dtw_tsmote_extratrees", **metric_dict(y, raw_all)},
        {"method": "2d_multiscale_varfilter_extratrees", **metric_dict(y, two_all)},
        {"method": "calibrated_score_level_stacking_raw_plus_2d", **metric_dict(y, stacked_all)},
    ]
    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)

    pred_rows = [
        {
            "subject_id": subject,
            "true_label": int(y[i]),
            "raw_multiwindow_score": float(raw_all[i]),
            "two_d_multiscale_score": float(two_all[i]),
            "stacked_score": float(stacked_all[i]),
        }
        for i, subject in enumerate(subjects)
    ]

    row_fields = ["method", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "stacked_multiscale_summary.csv", rows, row_fields)
    write_csv(out_dir / "stacked_multiscale_predictions.csv", pred_rows, ["subject_id", "true_label", "raw_multiwindow_score", "two_d_multiscale_score", "stacked_score"])
    write_csv(
        out_dir / "fold_stacking_details.csv",
        fold_rows,
        [
            "fold",
            "meta_kind",
            "meta_C",
            "meta_class_weight",
            "synthetic_samples",
            "raw_input_features",
            "raw_after_variance_features",
            "raw_selected_features",
            "two_d_input_features",
            "two_d_after_variance_features",
            "two_d_selected_features",
            "train_cases",
            "test_cases",
            "group_overlap",
        ],
    )
    write_csv(
        out_dir / "inner_meta_model_search.csv",
        meta_rows_all,
        ["outer_fold", "meta_kind", "meta_C", "meta_class_weight", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"],
    )

    stack_met = next(r for r in rows if r["method"] == "calibrated_score_level_stacking_raw_plus_2d")
    figures = make_figures(out_dir, y, stacked_all, stack_met)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold OOF stacking",
        "outer_splits": args.n_splits,
        "inner_splits": args.inner_splits,
        "sequence_length": args.sequence_length,
        "max_rows_per_csv": args.max_rows_per_csv,
        "raw_top_k": args.raw_top_k,
        "two_d_top_k": args.two_d_top_k,
        "two_d_feature_count": len(feature_names_2d),
        "two_d_meta": meta_2d,
        "best": rows[0],
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "stacked_multiscale_summary.csv"),
            "predictions": str(out_dir / "stacked_multiscale_predictions.csv"),
            "fold_details": str(out_dir / "fold_stacking_details.csv"),
            "meta_search": str(out_dir / "inner_meta_model_search.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": {k: str(v) for k, v in figures.items()},
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, fold_rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
