#!/usr/bin/env python3
"""SHAP-style explanation report for the NSC best EEG fusion method.

The validated model is `fusion_mi_max_LR_k256`, which uses a fold-local LR EEG
branch and an inner-fold selected fusion weight with the existing `mi_max`
baseline score. For interpretability, this script produces:

1. fold-local linear contribution statistics for the EEG branch;
2. a post-hoc all-data linear-SHAP surrogate using `mi_max + top EEG features`;
3. a feature-importance table with neuropsychological interpretation.

The surrogate is for explanation only and is not used as validation evidence.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
import warnings
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import shap
from docx import Document
from docx.shared import Inches
from sklearn.feature_selection import f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import average_precision_score, roc_auc_score
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

ROOT = Path(__file__).resolve().parents[1]
SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from nsc_eeg_csv_fusion_ablation import (  # noqa: E402
    build_subject_eeg_features,
    fit_model,
    inner_select_weight,
    load_base_scores,
    load_labels,
    natural_key,
    robust_unit_fit,
    select_topk,
)

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def safe_name(name: str) -> str:
    return name.replace("eeg_all/", "").replace("/", " / ")


def neuropsych_interpretation(feature: str, direction: str) -> str:
    f = feature.lower()
    if "mi_max" in f:
        return "既有 2D/raw 多實例模型的最高病例分數；代表影像或原始訊號中已有局部判別證據，是本案融合模型的主要風險來源。"
    if "duration" in f or "sample_rate" in f or "finite_ratio" in f or "/n/" in f or f.endswith("/n"):
        return "資料品質/可用性特徵，可能反映量測完整度、訊號缺失或紀錄品質；若重要性高，應優先做資料品質與缺失機制檢查。"
    if "band_delta" in f or "band_theta" in f:
        return (
            "低頻 delta/theta 變化常被視為皮質喚醒度下降、注意力效率降低、疲勞或神經功能變慢的候選訊號；"
            "若 SHAP 為正，表示低頻成分提高時推向 class 1；若為負，表示其提高較偏向 class 0。"
        )
    if "band_alpha" in f:
        return (
            "alpha 節律與安靜覺醒、注意力抑制/釋放及腦狀態穩定度有關；"
            "本資料沒有通道拓樸定位，因此只能解釋為全域 trial-level alpha 變異，而非特定腦區結論。"
        )
    if "band_beta" in f or "band_gamma" in f or "band_high" in f:
        return (
            "beta/gamma/high 頻段可能反映警覺度、認知負荷或動作/肌電 artifact。"
            "在穿戴或生理 CSV 中，高頻重要性需同時視為神經心理訊號與雜訊/肌電敏感指標，不能單獨作病理推論。"
        )
    if "spectral_entropy" in f:
        return (
            "spectral entropy 代表頻譜分散度與訊號複雜度；較高複雜度可對應較不穩定的腦/生理狀態，"
            "較低則可能代表節律集中或訊號單調化。"
        )
    if "fft_peak_frac" in f:
        return "fft peak fraction 表示單一主頻是否過度集中；可反映節律化程度，也可能受到設備或動作 artifact 影響。"
    if "spectral_centroid" in f or "spectral_edge" in f:
        return "頻譜重心/edge 反映能量往高頻或低頻偏移，可作為喚醒度、訊號快速變化或 artifact 負荷的綜合指標。"
    if "zero_cross" in f or "diff_" in f:
        return "零交叉率與差分特徵反映時間序列快速震盪與不穩定性；神經心理上可視為狀態波動或雜訊敏感指標。"
    if "seg" in f or "last_minus_first" in f or "trend" in f:
        return "分段與趨勢特徵反映 trial 內非平穩性；可對應疲勞、適應、喚醒度變化或量測狀態漂移。"
    if "std" in f or "iqr" in f or "mad" in f or "range" in f:
        return "變異度特徵反映 trial 內或跨 trial 的生理/腦狀態不穩定；在神經心理上可連到注意力維持與狀態調節，但也需排除 artifact。"
    if "mean" in f or "median" in f or "p25" in f or "p50" in f or "p75" in f:
        return "位置量數特徵反映該 EEG/生理摘要指標的整體水準；需結合原始訊號單位與通道來源才可作更精細解釋。"
    return "此特徵為統計或頻域摘要指標；目前可視為判別性候選 marker，需回看原始 trial 與臨床/神經心理量表才能做專業定論。"


def parse_domain(feature: str) -> str:
    f = feature.lower()
    if "mi_max" in f:
        return "融合基礎分數"
    if "band_delta" in f or "band_theta" in f:
        return "低頻節律/皮質喚醒"
    if "band_alpha" in f:
        return "alpha 節律/注意力狀態"
    if "band_beta" in f or "band_gamma" in f or "band_high" in f:
        return "高頻節律/肌電或警覺度"
    if "spectral" in f or "fft" in f:
        return "頻譜結構"
    if "diff_" in f or "zero_cross" in f:
        return "快速波動/非平穩性"
    if "seg" in f or "trend" in f or "last_" in f:
        return "trial 內趨勢"
    if "std" in f or "iqr" in f or "mad" in f or "range" in f:
        return "變異度/狀態穩定性"
    if "duration" in f or "sample_rate" in f or "finite_ratio" in f or "/n/" in f or f.endswith("/n"):
        return "資料品質"
    return "統計摘要"


def build_surrogate_shap(
    X_eeg: np.ndarray,
    y: np.ndarray,
    base_score: np.ndarray,
    feature_names: List[str],
    top_k: int,
    out_dir: Path,
) -> tuple[List[dict], Dict[str, str], dict]:
    imputer = SimpleImputer(strategy="median")
    X_imp = imputer.fit_transform(X_eeg)
    scores, _ = f_classif(X_imp, y)
    scores = np.nan_to_num(scores, nan=0.0, posinf=0.0, neginf=0.0)
    selected = np.argsort(scores)[::-1][: min(top_k, X_imp.shape[1])]
    selected_names = [feature_names[i] for i in selected]
    X_sel = X_imp[:, selected]
    X_combo = np.column_stack([base_score, X_sel])
    combo_names = ["mi_max"] + selected_names

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X_combo)
    model = LogisticRegression(max_iter=5000, C=0.25, class_weight="balanced", solver="liblinear")
    model.fit(X_scaled, y)
    pred = model.predict_proba(X_scaled)[:, 1]
    coef = model.coef_[0]
    # Exact linear-model contribution in standardized feature space. This is
    # equivalent to linear SHAP up to baseline centering for this explanatory
    # surrogate.
    shap_values = X_scaled * coef

    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    shap_path = fig_dir / "eeg_fusion_surrogate_shap_summary.png"
    plt.figure()
    shap.summary_plot(
        shap_values,
        X_combo,
        feature_names=[safe_name(n) for n in combo_names],
        max_display=20,
        show=False,
        plot_size=(9.5, 6.8),
    )
    plt.tight_layout()
    plt.savefig(shap_path, dpi=180, bbox_inches="tight")
    plt.close()

    mean_abs = np.mean(np.abs(shap_values), axis=0)
    signed_mean = np.mean(shap_values, axis=0)
    order = np.argsort(mean_abs)[::-1]
    rows: List[dict] = []
    for rank, j in enumerate(order[:40], start=1):
        name = combo_names[j]
        direction = "positive" if coef[j] > 0 else "negative"
        vals = X_combo[:, j]
        rows.append(
            {
                "rank": rank,
                "feature": name,
                "display_feature": safe_name(name),
                "domain": parse_domain(name),
                "mean_abs_shap": float(mean_abs[j]),
                "mean_signed_shap": float(signed_mean[j]),
                "coefficient": float(coef[j]),
                "direction": direction,
                "class0_mean": float(np.nanmean(vals[y == 0])),
                "class1_mean": float(np.nanmean(vals[y == 1])),
                "interpretation": neuropsych_interpretation(name, direction),
            }
        )

    bar_path = fig_dir / "eeg_fusion_top_feature_importance.png"
    top = rows[:20]
    plt.figure(figsize=(9.8, 6.4))
    plt.barh([r["display_feature"] for r in reversed(top)], [r["mean_abs_shap"] for r in reversed(top)])
    plt.xlabel("Mean absolute linear-SHAP contribution")
    plt.title("Top EEG Fusion Surrogate Features")
    plt.grid(axis="x", alpha=0.25)
    plt.tight_layout()
    plt.savefig(bar_path, dpi=180)
    plt.close()

    meta = {
        "surrogate_AUROC_in_sample": float(roc_auc_score(y, pred)),
        "surrogate_AUPRC_in_sample": float(average_precision_score(y, pred)),
        "top_k": int(top_k),
        "selected_feature_count": int(len(selected)),
        "note": "Post-hoc all-data explanation surrogate; not validation evidence.",
    }
    return rows, {"shap": str(shap_path), "importance": str(bar_path)}, meta


def fold_local_contributions(
    X_eeg: np.ndarray,
    y: np.ndarray,
    base_score: np.ndarray,
    feature_names: List[str],
    top_k: int,
    seed: int,
    n_splits: int,
    inner_splits: int,
) -> tuple[List[dict], List[dict], dict]:
    splitter = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=seed)
    contrib_by_feature: Dict[str, List[float]] = defaultdict(list)
    signed_by_feature: Dict[str, List[float]] = defaultdict(list)
    fold_rows: List[dict] = []
    oof = np.zeros(len(y), dtype=float)
    eeg_oof = np.zeros(len(y), dtype=float)
    for fold, (train_idx, test_idx) in enumerate(splitter.split(X_eeg, y), start=1):
        w_eeg, _inner = inner_select_weight(
            X_eeg,
            y,
            base_score,
            train_idx,
            top_k,
            "LR",
            seed + fold * 100,
            inner_splits,
            "min_metric",
        )
        imputer, selected = select_topk(X_eeg[train_idx], y[train_idx], top_k)
        Xtr = imputer.transform(X_eeg[train_idx])[:, selected]
        Xte = imputer.transform(X_eeg[test_idx])[:, selected]
        sc_test, model = fit_model("LR", Xtr, y[train_idx], Xte, seed + fold)
        sc_train, _ = fit_model("LR", Xtr, y[train_idx], Xtr, seed + fold + 7000)
        base_train_unit, base_test_unit = robust_unit_fit(base_score[train_idx], base_score[test_idx])
        eeg_train_unit, eeg_test_unit = robust_unit_fit(sc_train, sc_test)
        fused = (1.0 - w_eeg) * base_test_unit + w_eeg * eeg_test_unit
        oof[test_idx] = fused
        eeg_oof[test_idx] = eeg_test_unit

        scaler = model.named_steps["standardscaler"]
        lr = model.named_steps["logisticregression"]
        Xte_scaled = scaler.transform(Xte)
        local = Xte_scaled * lr.coef_[0]
        for col_pos, feat_idx in enumerate(selected):
            name = feature_names[int(feat_idx)]
            vals = np.abs(local[:, col_pos]) * float(w_eeg)
            signed = local[:, col_pos] * float(w_eeg)
            contrib_by_feature[name].extend(vals.tolist())
            signed_by_feature[name].extend(signed.tolist())
        fold_rows.append(
            {
                "fold": fold,
                "w_base": float(1.0 - w_eeg),
                "w_eeg": float(w_eeg),
                "test_cases": int(len(test_idx)),
                "mean_base_test_unit": float(np.mean(base_test_unit)),
                "mean_eeg_test_unit": float(np.mean(eeg_test_unit)),
            }
        )

    rows = []
    for name, vals in contrib_by_feature.items():
        arr = np.asarray(vals, dtype=float)
        signed = np.asarray(signed_by_feature[name], dtype=float)
        rows.append(
            {
                "feature": name,
                "display_feature": safe_name(name),
                "domain": parse_domain(name),
                "mean_abs_fold_local_contribution": float(np.mean(arr)),
                "mean_signed_fold_local_contribution": float(np.mean(signed)),
                "selection_count_or_test_contribution_count": int(len(arr)),
                "interpretation": neuropsych_interpretation(name, "positive" if np.mean(signed) > 0 else "negative"),
            }
        )
    rows.sort(key=lambda r: r["mean_abs_fold_local_contribution"], reverse=True)
    meta = {
        "oof_AUROC": float(roc_auc_score(y, oof)),
        "oof_AUPRC": float(average_precision_score(y, oof)),
        "eeg_branch_oof_AUROC": float(roc_auc_score(y, eeg_oof)),
        "eeg_branch_oof_AUPRC": float(average_precision_score(y, eeg_oof)),
    }
    return rows, fold_rows, meta


def make_weight_plot(out_dir: Path, fold_rows: List[dict]) -> str:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    path = fig_dir / "fusion_weight_by_fold.png"
    folds = [r["fold"] for r in fold_rows]
    w_base = [r["w_base"] for r in fold_rows]
    w_eeg = [r["w_eeg"] for r in fold_rows]
    plt.figure(figsize=(8.5, 4.2))
    plt.bar(folds, w_base, label="mi_max base weight")
    plt.bar(folds, w_eeg, bottom=w_base, label="EEG branch weight")
    plt.xlabel("Outer fold")
    plt.ylabel("Fusion weight")
    plt.ylim(0, 1.05)
    plt.title("Best Fusion Method: Inner-Fold Selected Weights")
    plt.legend(loc="lower right")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return str(path)


def make_report(
    report: Path,
    manifest: dict,
    shap_rows: List[dict],
    fold_rows: List[dict],
    fold_importance: List[dict],
    figures: Dict[str, str],
) -> None:
    doc = Document()
    doc.add_heading("NSC 114 位資料：本案最佳方法 SHAP 與神經心理特徵詮釋", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法定位", level=1)
    doc.add_paragraph(
        "本案目前 strict best 為 fusion_mi_max_LR_k256。正式效能仍以 patient-aware 10-fold OOF 結果為準。"
        "本報告補充兩種解釋：第一，fold-local EEG branch 線性貢獻；第二，mi_max + EEG top features 的 all-data linear-SHAP surrogate。"
    )
    doc.add_paragraph(
        "Surrogate 只用於解釋特徵方向與相對重要性，不作為新的效能宣稱。神經心理詮釋以候選機制描述，不能取代臨床診斷或腦區定位。"
    )

    doc.add_heading("二、融合權重", level=1)
    doc.add_picture(figures["weights"], width=Inches(6.2))
    doc.add_paragraph(
        "Inner-fold 選權重顯示，多數 fold 主要依賴 mi_max base score；EEG branch 在部分 fold 提供輔助增益。"
        "因此專業解釋應分成：整體融合層級以 mi_max 為主，EEG 特徵提供可解釋的神經生理輔助訊號。"
    )

    doc.add_heading("三、SHAP Summary Plot", level=1)
    doc.add_picture(figures["shap"], width=Inches(6.4))
    doc.add_picture(figures["importance"], width=Inches(6.4))

    doc.add_heading("四、Top 特徵重要性與神經心理詮釋", level=1)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Rank", "特徵", "領域", "Mean |SHAP|", "方向", "Class0 mean", "Class1 mean", "神經心理詮釋"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for row in shap_rows[:20]:
        cells = table.add_row().cells
        vals = [
            row["rank"],
            row["display_feature"],
            row["domain"],
            f"{row['mean_abs_shap']:.4f}",
            row["direction"],
            f"{row['class0_mean']:.4g}",
            f"{row['class1_mean']:.4g}",
            row["interpretation"],
        ]
        for cell, val in zip(cells, vals):
            cell.text = str(val)

    doc.add_heading("五、Fold-local EEG branch 重要性", level=1)
    doc.add_paragraph(
        f"OOF fusion AUROC={manifest['fold_local_meta']['oof_AUROC']:.3f}，"
        f"AUPRC={manifest['fold_local_meta']['oof_AUPRC']:.3f}；"
        f"EEG branch 單獨 OOF AUROC={manifest['fold_local_meta']['eeg_branch_oof_AUROC']:.3f}，"
        f"AUPRC={manifest['fold_local_meta']['eeg_branch_oof_AUPRC']:.3f}。"
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Rank", "特徵", "領域", "Fold-local mean |contribution|", "詮釋"]):
        cell.text = text
    for rank, row in enumerate(fold_importance[:15], start=1):
        cells = table.add_row().cells
        vals = [
            rank,
            row["display_feature"],
            row["domain"],
            f"{row['mean_abs_fold_local_contribution']:.4f}",
            row["interpretation"],
        ]
        for cell, val in zip(cells, vals):
            cell.text = str(val)

    doc.add_heading("六、專業限制", level=1)
    for item in [
        "本資料為 aggregate EEG CSV 特徵，缺少標準 10-20 channel 拓樸與任務事件標記，因此不能宣稱特定腦區機制。",
        "高頻 beta/gamma/high 重要性可能混合神經活動、肌電、動作與儀器 artifact，需搭配 artifact review。",
        "資料品質與缺失特徵若排在前面，應先檢查紀錄完整度與設備差異，避免把量測流程差異誤解成神經心理差異。",
        "SHAP surrogate 是 post-hoc 解釋模型，不取代 patient-aware 10-fold 的正式效能評估。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("七、輸出檔案", level=1)
    for key, value in manifest["outputs"].items():
        doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    report.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--eeg-root", default="eeg-csv-data-by-class")
    parser.add_argument("--base-predictions", default="analysis/nsc_uncertain_band_patch_refinement_20260520/uncertain_band_predictions.csv")
    parser.add_argument("--out-dir", default="analysis/nsc_eeg_shap_neuropsych_20260521")
    parser.add_argument("--report", default="reports/NSC_EEG_SHAP_feature_importance_neuropsych_20260521.docx")
    parser.add_argument("--cache-dir", default="analysis/nsc_eeg_shap_neuropsych_20260521/eeg_feature_cache")
    parser.add_argument("--top-k", type=int, default=256)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--random-state", type=int, default=20260520)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)

    labels = load_labels(Path(args.manifest).resolve())
    subjects = list(labels)
    y = np.asarray([labels[s] for s in subjects], dtype=int)
    X_eeg, feature_names, eeg_audit = build_subject_eeg_features(
        subjects,
        Path(args.eeg_root).resolve(),
        Path(args.cache_dir).resolve(),
        include_counts=False,
        class_aware_features=False,
    )
    base_scores = load_base_scores(Path(args.base_predictions).resolve(), subjects)
    base_score = base_scores["mi_max"]

    shap_rows, shap_figs, shap_meta = build_surrogate_shap(X_eeg, y, base_score, feature_names, args.top_k, out_dir)
    fold_importance, fold_rows, fold_meta = fold_local_contributions(
        X_eeg,
        y,
        base_score,
        feature_names,
        args.top_k,
        args.random_state,
        args.n_splits,
        args.inner_splits,
    )
    weights_fig = make_weight_plot(out_dir, fold_rows)

    write_csv(
        out_dir / "eeg_fusion_surrogate_shap_top_features.csv",
        shap_rows,
        [
            "rank",
            "feature",
            "display_feature",
            "domain",
            "mean_abs_shap",
            "mean_signed_shap",
            "coefficient",
            "direction",
            "class0_mean",
            "class1_mean",
            "interpretation",
        ],
    )
    write_csv(
        out_dir / "eeg_branch_fold_local_feature_importance.csv",
        fold_importance,
        [
            "feature",
            "display_feature",
            "domain",
            "mean_abs_fold_local_contribution",
            "mean_signed_fold_local_contribution",
            "selection_count_or_test_contribution_count",
            "interpretation",
        ],
    )
    write_csv(out_dir / "fusion_weight_by_fold.csv", fold_rows, ["fold", "w_base", "w_eeg", "test_cases", "mean_base_test_unit", "mean_eeg_test_unit"])

    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in Counter(y.tolist()).items()},
        "best_method_explained": "fusion_mi_max_LR_k256",
        "validation_note": "OOF metrics remain from patient-aware 10-fold; SHAP surrogate is explanation only.",
        "eeg_feature_count": len(feature_names),
        "top_k": args.top_k,
        "eeg_audit": eeg_audit,
        "surrogate_meta": shap_meta,
        "fold_local_meta": fold_meta,
        "outputs": {
            "analysis_dir": str(out_dir),
            "report": str(Path(args.report).resolve()),
            "surrogate_shap_top_features": str(out_dir / "eeg_fusion_surrogate_shap_top_features.csv"),
            "fold_local_feature_importance": str(out_dir / "eeg_branch_fold_local_feature_importance.csv"),
            "fusion_weight_by_fold": str(out_dir / "fusion_weight_by_fold.csv"),
            "shap_summary_plot": shap_figs["shap"],
            "feature_importance_plot": shap_figs["importance"],
            "fusion_weight_plot": weights_fig,
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(
        Path(args.report).resolve(),
        manifest,
        shap_rows,
        fold_rows,
        fold_importance,
        {"weights": weights_fig, **shap_figs},
    )
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
