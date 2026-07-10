#!/usr/bin/env python3
"""Sequential heuristic-rule learning for nsc_dataset_images.

This is intentionally not a neural classifier. It starts from two labelled
seed cases, builds a small interpretable patch-rule library, tests later cases
before revising rules, and records how many cases are covered without revision.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import numpy as np
from PIL import Image


PATCH_GRID = 8
PLOT_NAMES = {"PP", "AR", "RP", "GAF"}


@dataclass
class Rule:
    rule_id: str
    feature: str
    channel: str
    plot: str
    patch_row: int
    patch_col: int
    threshold: float
    direction_label1_high: bool
    weight: float
    mean_label0: float
    mean_label1: float
    support_label0: int
    support_label1: int
    coverage_cases: int
    description: str


def natural_key(text: str) -> Tuple:
    parts = re.split(r"(\d+)", text)
    return tuple(int(p) if p.isdigit() else p for p in parts)


def case_id_from_session(session_name: str) -> str:
    return session_name.split("_Sess", 1)[0]


def feature_key(channel: str, plot: str, row: int, col: int) -> str:
    return f"{channel}/{plot}/patch_r{row}c{col}"


def channel_group(channel: str) -> str:
    upper = channel.upper()
    if "RESP" in upper:
        return "GROUP_RESP"
    if "EKG" in upper or "PRIMARY_RAW" in upper or "IBI" in upper:
        return "GROUP_EKG_OR_PRIMARY"
    if "SC_" in upper or "SKIN" in upper:
        return "GROUP_SKIN_CONDUCTANCE"
    if "AUX" in upper:
        return "GROUP_AUX"
    return "GROUP_OTHER"


def parse_feature_key(key: str) -> Tuple[str, str, int, int]:
    channel, plot, patch = key.split("/")
    m = re.match(r"patch_r(\d+)c(\d+)", patch)
    if not m:
        raise ValueError(f"Bad patch feature key: {key}")
    return channel, plot, int(m.group(1)), int(m.group(2))


def image_patch_features(path: Path, grid: int = PATCH_GRID) -> np.ndarray:
    img = Image.open(path).convert("L").resize((128, 128))
    arr = np.asarray(img, dtype=np.float32) / 255.0
    ink = 1.0 - arr
    h, w = ink.shape
    patch_h = h // grid
    patch_w = w // grid
    feats = np.zeros((grid, grid), dtype=np.float32)
    for r in range(grid):
        for c in range(grid):
            patch = ink[r * patch_h : (r + 1) * patch_h, c * patch_w : (c + 1) * patch_w]
            feats[r, c] = float(patch.mean())
    return feats


def collect_case_images(root: Path) -> Dict[str, dict]:
    cases: Dict[str, dict] = {}
    for path in root.rglob("*"):
        if path.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
            continue
        rel = path.relative_to(root).parts
        if len(rel) < 4 or rel[0] not in {"0", "1"}:
            continue
        label = int(rel[0])
        session_name = rel[1]
        channel = rel[2]
        plot = path.stem
        if plot not in PLOT_NAMES:
            continue
        case_id = case_id_from_session(session_name)
        key = f"{label}:{case_id}"
        if key not in cases:
            cases[key] = {
                "case_id": case_id,
                "label": label,
                "sessions": set(),
                "image_paths": [],
                "channels": Counter(),
                "plots": Counter(),
            }
        elif cases[key]["label"] != label:
            raise ValueError(f"Case {case_id} appears with conflicting labels.")
        cases[key]["sessions"].add(session_name)
        cases[key]["image_paths"].append(path)
        cases[key]["channels"][channel] += 1
        cases[key]["plots"][plot] += 1

    for meta in cases.values():
        meta["sessions"] = sorted(meta["sessions"], key=natural_key)
        meta["image_paths"] = sorted(meta["image_paths"], key=lambda p: natural_key(str(p)))
    return cases


def build_case_features(cases: Dict[str, dict]) -> Dict[str, dict]:
    features: Dict[str, dict] = {}
    for key, meta in sorted(cases.items(), key=lambda kv: (kv[1]["label"], natural_key(kv[1]["case_id"]))):
        buckets: Dict[str, List[float]] = defaultdict(list)
        for path in meta["image_paths"]:
            rel = path.parts
            channel = path.parent.name
            plot = path.stem
            patches = image_patch_features(path)
            for r in range(PATCH_GRID):
                for c in range(PATCH_GRID):
                    value = float(patches[r, c])
                    buckets[feature_key(channel, plot, r, c)].append(value)
                    buckets[feature_key(channel_group(channel), plot, r, c)].append(value)
                    buckets[feature_key("ALL_CHANNELS", plot, r, c)].append(value)
        feat = {k: float(np.mean(v)) for k, v in buckets.items() if v}
        features[key] = {
            "case_id": meta["case_id"],
            "label": int(meta["label"]),
            "features": feat,
            "image_count": len(meta["image_paths"]),
            "session_count": len(meta["sessions"]),
            "sessions": list(meta["sessions"]),
            "channels": dict(meta["channels"]),
            "plots": dict(meta["plots"]),
        }
    return features


def choose_seed_cases(case_features: Dict[str, dict], seed0: str | None, seed1: str | None) -> Tuple[str, str]:
    def choose(label: int, requested: str | None) -> str:
        candidates = [k for k, v in case_features.items() if v["label"] == label]
        if requested:
            requested_keys = [k for k in candidates if case_features[k]["case_id"] == requested or k == requested]
            if not requested_keys:
                raise ValueError(f"Requested seed label {label} case not found: {requested}")
            return requested_keys[0]
        return sorted(
            candidates,
            key=lambda k: (-case_features[k]["image_count"], -case_features[k]["session_count"], natural_key(case_features[k]["case_id"])),
        )[0]

    return choose(0, seed0), choose(1, seed1)


def build_rules(case_features: Dict[str, dict], memory_keys: List[str], max_rules: int = 48, min_abs_diff: float = 0.015) -> List[Rule]:
    values0: Dict[str, List[float]] = defaultdict(list)
    values1: Dict[str, List[float]] = defaultdict(list)
    for key in memory_keys:
        row = case_features[key]
        target = values1 if row["label"] == 1 else values0
        for feat_key, value in row["features"].items():
            target[feat_key].append(value)

    candidates: List[Rule] = []
    common_features = sorted(set(values0) & set(values1))
    for feat in common_features:
        v0 = np.asarray(values0[feat], dtype=np.float32)
        v1 = np.asarray(values1[feat], dtype=np.float32)
        if len(v0) == 0 or len(v1) == 0:
            continue
        mean0 = float(v0.mean())
        mean1 = float(v1.mean())
        diff = mean1 - mean0
        if abs(diff) < min_abs_diff:
            continue
        pooled_sd = float(np.sqrt((v0.var() + v1.var()) / 2.0))
        pooled_sd = max(pooled_sd, 0.03)
        effect = diff / pooled_sd
        channel, plot, r, c = parse_feature_key(feat)
        threshold = (mean0 + mean1) / 2.0
        direction_high = diff > 0
        label1_phrase = "較深/較密" if direction_high else "較淺/較疏"
        coverage_cases = sum(1 for case in case_features.values() if feat in case["features"])
        candidates.append(
            Rule(
                rule_id="",
                feature=feat,
                channel=channel,
                plot=plot,
                patch_row=r,
                patch_col=c,
                threshold=threshold,
                direction_label1_high=direction_high,
                weight=float(min(abs(effect), 6.0)),
                mean_label0=mean0,
                mean_label1=mean1,
                support_label0=len(v0),
                support_label1=len(v1),
                coverage_cases=coverage_cases,
                description=f"{channel}/{plot} 第{r + 1}列第{c + 1}欄：label 1 相對 label 0 {label1_phrase}",
            )
        )

    def generality_rank(rule: Rule) -> int:
        if rule.channel == "ALL_CHANNELS":
            return 2
        if rule.channel.startswith("GROUP_"):
            return 1
        return 0

    candidates.sort(
        key=lambda x: (
            -x.coverage_cases,
            -generality_rank(x),
            -x.weight,
            x.channel,
            x.plot,
            x.patch_row,
            x.patch_col,
        )
    )
    selected = candidates[:max_rules]
    for i, rule in enumerate(selected, start=1):
        rule.rule_id = f"R{i:03d}"
    return selected


def classify(features: Dict[str, float], rules: List[Rule]) -> dict:
    used = []
    score = 0.0
    total_weight = 0.0
    for rule in rules:
        if rule.feature not in features:
            continue
        value = features[rule.feature]
        if rule.direction_label1_high:
            vote = 1.0 if value >= rule.threshold else -1.0
        else:
            vote = 1.0 if value <= rule.threshold else -1.0
        weighted = vote * rule.weight
        score += weighted
        total_weight += rule.weight
        used.append(
            {
                "rule_id": rule.rule_id,
                "feature": rule.feature,
                "value": float(value),
                "threshold": float(rule.threshold),
                "vote": int(vote),
                "weight": float(rule.weight),
            }
        )
    if total_weight == 0:
        return {"pred_label": None, "score": 0.0, "confidence": 0.0, "used_rules": 0, "used_rule_details": used}
    norm_score = score / total_weight
    return {
        "pred_label": 1 if norm_score >= 0 else 0,
        "score": float(norm_score),
        "confidence": float(abs(norm_score)),
        "used_rules": len(used),
        "used_rule_details": used[:10],
    }


def write_csv(path: Path, rows: Iterable[dict], fieldnames: List[str]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def make_figures(out_dir: Path, trials: List[dict], final_rules: List[Rule]) -> Dict[str, Path]:
    import matplotlib.pyplot as plt

    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, Path] = {}

    steps = [t["step"] for t in trials]
    cumulative_success = np.cumsum([int(t["no_revision_success"]) for t in trials])
    cumulative_revisions = np.cumsum([1 if t["action"] == "revise_rules_with_case_feedback" else 0 for t in trials])
    plt.figure(figsize=(8, 4.5))
    plt.plot(steps, cumulative_success, marker="o", label="No-revision successes")
    plt.plot(steps, cumulative_revisions, marker="o", label="Rule revisions")
    plt.xlabel("Sequential test step")
    plt.ylabel("Cumulative count")
    plt.title("Heuristic learning trajectory")
    plt.grid(alpha=0.25)
    plt.legend()
    paths["trajectory"] = fig_dir / "heuristic_learning_trajectory.png"
    plt.tight_layout()
    plt.savefig(paths["trajectory"], dpi=180)
    plt.close()

    top = final_rules[:12]
    if top:
        labels = [f"{r.rule_id} {r.channel}/{r.plot} r{r.patch_row + 1}c{r.patch_col + 1}" for r in top]
        weights = [r.weight for r in top]
        plt.figure(figsize=(9, 5))
        y = np.arange(len(top))
        plt.barh(y, weights)
        plt.yticks(y, labels, fontsize=8)
        plt.gca().invert_yaxis()
        plt.xlabel("Rule weight")
        plt.title("Top final heuristic rules")
        paths["top_rules"] = fig_dir / "top_final_rules.png"
        plt.tight_layout()
        plt.savefig(paths["top_rules"], dpi=180)
        plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, trials: List[dict], final_rules: List[Rule], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC Dataset Images 啟發式學習規則庫逐案測試報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法定位", level=1)
    doc.add_paragraph(
        "本報告採用啟發式學習概念：在個案數仍不足以支撐穩定深度模型時，"
        "先以兩個已知標籤個案建立可讀規則庫，後續個案先用既有規則分類；"
        "若分類正確且信心足夠，記為不用修規則即成功，若錯誤或不確定，才將該個案回饋給規則庫修訂。"
    )
    doc.add_paragraph(
        "因此，本報告呈現的是規則庫逐案吸收回饋的探索性證據，不宣稱已完成新個案泛化模型。"
    )
    doc.add_paragraph(
        "特徵設計採三層規則：精確通道、通道群組與全通道摘要。規則選擇先看可覆蓋的個案數，"
        "再看類別差異權重，避免只挑到差異很大但新個案常常沒有同名通道的規則。"
    )

    doc.add_heading("二、資料概況", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "項目"
    hdr[1].text = "數值"
    for label, value in [
        ("資料來源", manifest["dataset_root"]),
        ("總圖片數", str(manifest["total_images"])),
        ("總個案數", str(manifest["total_cases"])),
        ("label 0 個案數", str(manifest["cases_by_label"].get("0", 0))),
        ("label 1 個案數", str(manifest["cases_by_label"].get("1", 0))),
        ("起始 label 0 個案", manifest["seed_cases"]["label0"]),
        ("起始 label 1 個案", manifest["seed_cases"]["label1"]),
        ("逐案測試個案數", str(manifest["tested_cases"])),
    ]:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    doc.add_heading("三、逐案測試統計", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "指標"
    table.rows[0].cells[1].text = "結果"
    for label, value in [
        ("修規則前分類正確個案數", manifest["correct_before_revision_count"]),
        ("不用修規則即成功個案數", manifest["no_revision_success_count"]),
        ("需要修規則個案數", manifest["revision_count"]),
        ("不確定個案數", manifest["uncertain_count"]),
        ("修規則前正確率", f"{manifest['correct_before_revision_rate']:.3f}"),
        ("不用修規則成功率", f"{manifest['no_revision_success_rate']:.3f}"),
    ]:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    if "trajectory" in figures:
        doc.add_paragraph("圖 1. 啟發式學習逐案累積軌跡")
        doc.add_picture(str(figures["trajectory"]), width=Inches(6.2))

    doc.add_heading("四、逐案結果摘要", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["步驟", "個案", "真實", "預測", "信心", "規則數", "正確", "不確定", "動作"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for t in trials[:40]:
        cells = table.add_row().cells
        vals = [
            t["step"],
            t["case_id"],
            t["true_label"],
            t["pred_label"],
            f"{t['confidence']:.3f}",
            t["used_rules"],
            "是" if t["correct"] else "否",
            "是" if t["uncertain"] else "否",
            "成功免修" if t["no_revision_success"] else "修規則",
        ]
        for cell, val in zip(cells, vals):
            cell.text = str(val)
    if len(trials) > 40:
        doc.add_paragraph(f"註：完整逐案紀錄共 {len(trials)} 筆，請見 CSV/JSONL 輸出；本表僅列前 40 筆。")

    doc.add_heading("五、最後規則庫摘要", level=1)
    if "top_rules" in figures:
        doc.add_paragraph("圖 2. 最後規則庫權重最高的規則")
        doc.add_picture(str(figures["top_rules"]), width=Inches(6.2))
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["規則", "通道/圖型", "位置", "label 0 平均", "label 1 平均", "描述"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for rule in final_rules[:20]:
        cells = table.add_row().cells
        vals = [
            rule.rule_id,
            f"{rule.channel}/{rule.plot}",
            f"第{rule.patch_row + 1}列第{rule.patch_col + 1}欄",
            f"{rule.mean_label0:.3f}",
            f"{rule.mean_label1:.3f}",
            rule.description,
        ]
        for cell, val in zip(cells, vals):
            cell.text = val

    doc.add_heading("六、結論與後續", level=1)
    doc.add_paragraph(
        f"本次以兩個起始個案建立規則庫，後續 {manifest['tested_cases']} 個個案中，"
        f"{manifest['no_revision_success_count']} 個不用修規則即成功，"
        f"{manifest['revision_count']} 個需要修規則。"
    )
    doc.add_paragraph(
        "後續若有新個案進來，應固定使用同一流程：先用既有規則庫測試、紀錄是否成功，"
        "再決定是否修規則。當不用修規則即成功的比例穩定上升，才表示規則庫開始覆蓋更多樣本型態。"
    )

    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--dataset-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_dataset_images_heuristic_learning_20260519")
    parser.add_argument("--report", default="reports/NSC_dataset_images_啟發式學習規則庫逐案測試報告_20260519.docx")
    parser.add_argument("--seed-label0", default=None)
    parser.add_argument("--seed-label1", default=None)
    parser.add_argument("--confidence-threshold", type=float, default=0.12)
    parser.add_argument("--max-rules", type=int, default=48)
    args = parser.parse_args()

    root = Path(args.dataset_root).resolve()
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    cases = collect_case_images(root)
    case_features = build_case_features(cases)
    seed0, seed1 = choose_seed_cases(case_features, args.seed_label0, args.seed_label1)
    memory_keys = [seed0, seed1]
    observed_keys = [seed0, seed1]
    rules = build_rules(case_features, memory_keys, max_rules=args.max_rules)

    test_keys = [
        k
        for k in sorted(case_features, key=lambda kk: natural_key(case_features[kk]["case_id"]))
        if k not in {seed0, seed1}
    ]
    trials: List[dict] = []
    for step, key in enumerate(test_keys, start=1):
        row = case_features[key]
        pred = classify(row["features"], rules)
        uncertain = pred["pred_label"] is None or pred["confidence"] < args.confidence_threshold
        correct = (pred["pred_label"] == row["label"]) if pred["pred_label"] is not None else False
        no_revision_success = bool(correct and not uncertain)
        action = "no_revision_success" if no_revision_success else "revise_rules_with_case_feedback"
        trial = {
            "step": step,
            "case_key": key,
            "case_id": row["case_id"],
            "true_label": int(row["label"]),
            "pred_label": "" if pred["pred_label"] is None else int(pred["pred_label"]),
            "score": float(pred["score"]),
            "confidence": float(pred["confidence"]),
            "used_rules": int(pred["used_rules"]),
            "correct": int(correct),
            "uncertain": int(uncertain),
            "no_revision_success": int(no_revision_success),
            "action": action,
            "memory_cases_before": len(memory_keys),
            "image_count": row["image_count"],
            "session_count": row["session_count"],
        }
        trials.append(trial)
        observed_keys.append(key)
        if not no_revision_success:
            memory_keys = list(dict.fromkeys(observed_keys))
            rules = build_rules(case_features, memory_keys, max_rules=args.max_rules)

    final_rules = build_rules(case_features, memory_keys, max_rules=args.max_rules)
    fields = [
        "step",
        "case_id",
        "true_label",
        "pred_label",
        "score",
        "confidence",
        "used_rules",
        "correct",
        "uncertain",
        "no_revision_success",
        "action",
        "memory_cases_before",
        "image_count",
        "session_count",
    ]
    write_csv(out_dir / "heuristic_trials_summary.csv", trials, fields)
    with (out_dir / "heuristic_trials.jsonl").open("w", encoding="utf-8") as f:
        for t in trials:
            f.write(json.dumps(t, ensure_ascii=False) + "\n")
    with (out_dir / "heuristic_rule_library_final.json").open("w", encoding="utf-8") as f:
        json.dump([asdict(r) for r in final_rules], f, ensure_ascii=False, indent=2)
    with (out_dir / "case_feature_manifest.json").open("w", encoding="utf-8") as f:
        compact = {
            k: {
                "case_id": v["case_id"],
                "label": v["label"],
                "image_count": v["image_count"],
                "session_count": v["session_count"],
                "sessions": v["sessions"],
                "channels": v["channels"],
                "plots": v["plots"],
                "feature_count": len(v["features"]),
            }
            for k, v in case_features.items()
        }
        json.dump(compact, f, ensure_ascii=False, indent=2)

    total_images = sum(v["image_count"] for v in case_features.values())
    cases_by_label = Counter(str(v["label"]) for v in case_features.values())
    correct = sum(t["correct"] for t in trials)
    no_revision_success = sum(t["no_revision_success"] for t in trials)
    revisions = sum(1 for t in trials if t["action"] == "revise_rules_with_case_feedback")
    uncertain = sum(t["uncertain"] for t in trials)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "method_reference": "https://github.com/Trinkle23897/learning-beyond-gradients",
        "dataset_root": str(root),
        "total_images": total_images,
        "total_cases": len(case_features),
        "cases_by_label": dict(cases_by_label),
        "seed_cases": {
            "label0": case_features[seed0]["case_id"],
            "label1": case_features[seed1]["case_id"],
        },
        "tested_cases": len(trials),
        "correct_before_revision_count": int(correct),
        "correct_before_revision_rate": float(correct / len(trials)) if trials else math.nan,
        "no_revision_success_count": int(no_revision_success),
        "no_revision_success_rate": float(no_revision_success / len(trials)) if trials else math.nan,
        "revision_count": int(revisions),
        "uncertain_count": int(uncertain),
        "confidence_threshold": args.confidence_threshold,
        "max_rules": args.max_rules,
        "feature_design": "exact channel + channel group + all-channel patch summaries",
        "rule_selection": "coverage-first, then generality and effect-size weight",
        "final_rule_count": len(final_rules),
        "outputs": {
            "analysis_dir": str(out_dir),
            "trial_summary": str(out_dir / "heuristic_trials_summary.csv"),
            "trial_jsonl": str(out_dir / "heuristic_trials.jsonl"),
            "rule_library": str(out_dir / "heuristic_rule_library_final.json"),
            "report": str(Path(args.report).resolve()),
        },
    }
    figures = make_figures(out_dir, trials, final_rules)
    with (out_dir / "heuristic_learning_manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, trials, final_rules, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
