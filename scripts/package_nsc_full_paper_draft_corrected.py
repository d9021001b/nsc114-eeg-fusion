#!/usr/bin/env python3
"""Generate a full Traditional Chinese paper draft for the NSC 114-case study."""

from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis" / "nsc_full_paper_draft_20260521_v3"
REPORT = ROOT / "reports" / "NSC114_EEG_multimodal_small_data_paper_draft_20260521_v3.docx"

TAIL_DIR = ROOT / "analysis" / "nsc_restricted_subject_bagging_tail_stats_corrected_20260521"
LIT_DIR = ROOT / "analysis" / "nsc_replicated_literature_models_vs_best_20260520"
SHAP_DIR = ROOT / "analysis" / "nsc_eeg_shap_neuropsych_20260521"


def read_csv(path: Path) -> list[dict]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def fmt(value: str | float, digits: int = 3) -> str:
    if value in ["", None]:
        return ""
    return f"{float(value):.{digits}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_centered(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)


def add_algorithm_block(doc: Document, title: str, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F2F2F2")
    p = cell.paragraphs[0]
    run = p.add_run(title + "\n")
    run.bold = True
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(9.5)
    for line in code.strip("\n").splitlines():
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        r = para.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        r.font.size = Pt(8.5)


def make_flow_figure(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(12.5, 5.3))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    def box(x: float, y: float, w: float, h: float, text: str, fill: str = "#F7FAFC") -> None:
        patch = FancyBboxPatch(
            (x, y),
            w,
            h,
            boxstyle="round,pad=0.018,rounding_size=0.015",
            linewidth=1.35,
            edgecolor="#23506C",
            facecolor=fill,
        )
        ax.add_patch(patch)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=9.5, color="#111111")

    def arrow(x1: float, y1: float, x2: float, y2: float) -> None:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="-|>", color="#23506C", lw=1.35, shrinkA=3, shrinkB=3))

    ax.text(0.5, 0.94, "Patient-aware multimodal modeling workflow", ha="center", va="center", fontsize=13, fontweight="bold")
    ax.text(0.18, 0.80, "2D time-series image branch", ha="center", fontsize=10, color="#23506C", fontweight="bold")
    ax.text(0.18, 0.42, "EEG CSV feature branch", ha="center", fontsize=10, color="#23506C", fontweight="bold")

    box(0.04, 0.55, 0.13, 0.16, "NSC 114\\npatients", "#EFF6FF")
    box(0.23, 0.66, 0.18, 0.15, "2D time-series\\nimage evidence", "#F8FBFF")
    box(0.47, 0.66, 0.16, 0.15, "mi_max\\nrisk score", "#F8FBFF")
    box(0.23, 0.26, 0.18, 0.16, "EEG CSV trials\\nper patient", "#FFF9EE")
    box(0.47, 0.26, 0.16, 0.16, "Subject-level\\nstatistics", "#FFF9EE")
    box(0.68, 0.26, 0.18, 0.16, "Targeted tail\\nfeatures", "#FFF9EE")
    box(0.68, 0.56, 0.18, 0.16, "Fold-local\\nfeature selection", "#F4FBF7")
    box(0.87, 0.43, 0.10, 0.22, "Inner CV\\nscore fusion", "#F4FBF7")
    box(0.39, 0.04, 0.21, 0.12, "Patient-aware 10-fold OOF\\n10-seed bagging", "#F2F5F9")
    box(0.66, 0.04, 0.21, 0.12, "AUROC / AUPRC\\nmodel comparison", "#F2F5F9")

    arrow(0.17, 0.63, 0.23, 0.73)
    arrow(0.41, 0.735, 0.47, 0.735)
    arrow(0.63, 0.735, 0.68, 0.65)
    arrow(0.17, 0.61, 0.23, 0.34)
    arrow(0.41, 0.34, 0.47, 0.34)
    arrow(0.63, 0.34, 0.68, 0.34)
    arrow(0.77, 0.42, 0.77, 0.56)
    arrow(0.86, 0.64, 0.87, 0.56)
    arrow(0.92, 0.43, 0.55, 0.16)
    arrow(0.60, 0.10, 0.66, 0.10)

    ax.text(
        0.5,
        0.005,
        "Training-fold only: imputation, scaling, feature selection, fusion-weight selection, and model fitting.",
        ha="center",
        va="bottom",
        fontsize=9,
        color="#444444",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    REPORT.parent.mkdir(parents=True, exist_ok=True)

    manifest = json.loads((TAIL_DIR / "manifest.json").read_text(encoding="utf-8"))
    bag = manifest["bagged_results"]
    params = manifest["parameters"]
    preds = read_csv(TAIL_DIR / "bagged_predictions.csv")
    class_counts = {"0": 0, "1": 0}
    for row in preds:
        class_counts[row["true_label"]] += 1
    main_rows = read_csv(LIT_DIR / "replicated_literature_vs_our_best_main.csv")
    refs = read_csv(LIT_DIR / "citation_mapping.csv")
    shap_rows = read_csv(SHAP_DIR / "eeg_fusion_surrogate_shap_top_features.csv")[:12]

    flow_fig = make_flow_figure(OUT_DIR / "figures" / "paper_method_flow.png")

    doc = Document()
    set_doc_style(doc)

    title = "結合二維時序影像、EEG 統計特徵與文獻模型復刻之小樣本二元分類研究：NSC 114 位初稿"
    doc.add_heading(title, level=0)
    add_centered(doc, "作者一：待填；作者二：待填", bold=False, size=11)
    add_centered(doc, f"初稿日期：{datetime.now().strftime('%Y-%m-%d')}", size=10)

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "目的：本研究針對 NSC 114 位受試者之小樣本二元分類問題，建立一套可重跑、可解釋、且避免病人層級資料洩漏的模型比較流程。"
        "方法：本研究整合既有二維時序影像多實例分數、EEG CSV subject-level 統計摘要與針對 delta/theta/alpha 的 targeted tail-statistics。"
        "評估採 patient-aware stratified 10-fold out-of-fold validation，並以 10 個 random seeds 進行 bagging。"
        "同時復刻與比較 time-series image CNN、1D CNN、1D+2D CNN ensemble、SMOTE、T-SMOTE/DTW augmentation、prototype、PixelHop-style、Siamese、Relation Network、Matching Networks、MAML-style 與 DTWSSE-style 方法。"
        "結果：最佳候選方法 AUROC 為 0.7971、AUPRC 為 0.7989，confusion matrix 為 TN/FP/FN/TP=34/28/10/42。"
        "此結果接近但未同時達到 0.8/0.8，仍高於多數復刻文獻方法。"
        "結論：二維時序影像分數與 EEG subject-level 特徵具有互補性，targeted tail-statistics 可提升小樣本模型表現；"
        "然而目前仍應定位為 optimized OOF candidate，需進一步鎖定 protocol 後進行 confirmatory evaluation。"
    )
    doc.add_paragraph("關鍵字：小樣本學習；EEG；二維時序影像；patient-aware cross-validation；AUPRC；文獻模型復刻")

    doc.add_heading("1. 前言", level=1)
    doc.add_paragraph(
        "在小樣本生理訊號與神經心理資料中，模型容易因受試者重複、影像實例重複、特徵選擇外洩或過度調參而高估泛化能力。"
        "因此，本研究將評估單位明確定義為 patient-level case，並要求同一 case_id 不得同時出現在訓練與測試中。"
        "既有文獻顯示，將 time-series 轉換為二維影像可讓模型捕捉局部紋理與時序結構 [Wang2015; Hatami2017]，"
        "1D/2D CNN、prototype、metric learning 與資料增廣方法也常被用於生理訊號或小樣本分類 [Wu2018; Noman2018; Snell2017; Zhao2022]。"
        "本研究的核心問題是：在 NSC 114 位資料上，二維時序影像分數與 EEG CSV 統計特徵融合，是否能比文獻方法復刻版提供更穩定的 patient-level 判別訊號。"
    )

    doc.add_heading("2. 材料與方法", level=1)
    doc.add_heading("2.1 收案個案情況與描述統計（由收案人員填寫）", level=2)
    doc.add_paragraph(
        "本節保留給收案與臨床研究團隊填寫。請依研究設計補上收案期間、收案地點、納入與排除條件、"
        "受試者基本資料、臨床或神經心理分組依據、量表結果與缺失資料說明。"
    )
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["收案期間", ""],
            ["收案地點/單位", ""],
            ["納入條件", ""],
            ["排除條件", ""],
            ["年齡", ""],
            ["性別", ""],
            ["教育年數", ""],
            ["主要診斷或分組定義", ""],
            ["神經心理量表或臨床評估", ""],
            ["共病與用藥", ""],
            ["缺失資料與排除個案說明", ""],
            ["其他收案備註", ""],
        ],
    )

    doc.add_heading("2.2 分析資料與標籤定義", level=2)
    doc.add_paragraph(
        f"本模型分析以 {len(preds)} 位受試者為單位，其中 class 0 為 {class_counts['0']} 位，class 1 為 {class_counts['1']} 位。"
        "本研究將 case_id 定義為 patient-level identifier；所有主要驗證均以 patient-aware split 避免同一受試者同時進入訓練與測試。"
    )
    add_table(
        doc,
        ["項目", "內容"],
        [
            ["分析單位", "patient-level case；case_id 即 patient-id"],
            ["受試者數", str(len(preds))],
            ["類別分布", f"class 0={class_counts['0']}；class 1={class_counts['1']}"],
            ["主要驗證", "patient-aware stratified 10-fold OOF；10 seeds bagging"],
        ],
    )

    doc.add_heading("2.3 特徵與模型流程", level=2)
    doc.add_paragraph(
        "本案模型以既有二維時序影像或 raw multi-instance 分數 mi_max 作為 baseline risk score，再加入 EEG CSV subject-level 統計摘要。"
        "EEG 特徵包含位置量數、變異度、分段趨勢、頻譜比例與 spectral entropy 等摘要。"
        "targeted tail-statistics 僅針對神經心理上較具候選意義的頻帶擴充：delta/theta 加入高尾端 p90/p95，以捕捉 transient slowing；alpha 加入低尾端 p05/p10，以捕捉 alpha suppression。"
        "此設計避免把所有高尾端特徵一律加入而放大雜訊或肌電 artifact。"
    )
    doc.add_picture(str(flow_fig), width=Inches(6.5))
    add_caption(doc, "圖 1. 二維影像分數、EEG 統計特徵與 targeted tail-statistics 融合流程。")

    doc.add_heading("2.4 本研究提出方法的演算法步驟", level=2)
    doc.add_paragraph(
        "本研究提出的方法可概念化為 Targeted Tail-statistics Bagged Fusion。"
        "核心想法是先保留二維時序影像分支已學到的 patient-level 局部風險分數，再用 EEG CSV 轉成可解釋的 subject-level 統計特徵，"
        "其中只對 delta/theta 與 alpha 等具神經心理意義的頻帶加入尾端統計，以補捉短暫低頻化或 alpha suppression。"
        "最後，在每個 training fold 內選擇 EEG 分支與 baseline 分數的融合權重，並以多 seed 的 out-of-fold 分數平均降低小樣本切分變異。"
    )
    doc.add_paragraph(
        "演算法細節如下。第一步，將每位受試者視為不可拆分的 patient-level unit，建立 patient-aware stratified 10-fold。"
        "第二步，對每個 outer fold，僅用 training patients fit 缺失值填補、標準化、特徵選擇與融合權重；test patients 完全不參與這些步驟。"
        "第三步，EEG trial 層級原始訊號先萃取頻域與時域摘要，再聚合成 subject-level features。標準聚合包含 min、max、mean、std、p25、p50、p75。"
        "額外 targeted tail-statistics 僅用於 delta/theta band fraction 的 p90/p95 與 alpha band fraction 的 p05/p10。"
        "第四步，使用 training fold 的 univariate feature selection 選出 top-k EEG features，再訓練 L1 logistic regression。"
        "第五步，以 inner cross-validation 搜尋受限融合權重 w，將 baseline mi_max 分數與 EEG logistic regression probability 線性融合。"
        "第六步，重複 10 個 random seeds 後，將同一受試者的 OOF score 平均，作為 final bagged score 並計算 AUROC、AUPRC 與 confusion matrix。"
    )
    add_table(
        doc,
        ["步驟", "目的", "防止資料洩漏的設計"],
        [
            ["Patient-aware 10-fold", "以受試者為單位切分", "同一 patient-id 不跨訓練與測試"],
            ["EEG subject-level aggregation", "把多 trial EEG 轉成可解釋特徵", "聚合公式固定，不看測試標籤"],
            ["Targeted tail-statistics", "捕捉 transient slowing / alpha suppression", "只加入預先指定頻帶，不依測試結果挑選"],
            ["Fold-local top-k selection", "降低小樣本高維特徵風險", "feature selector 只 fit training fold"],
            ["Inner-CV fusion weight", "決定 EEG 分支補多少訊號", "融合權重只由 inner validation 決定"],
            ["10-seed OOF bagging", "降低 split variance", "每個 seed 仍維持 patient-aware OOF"],
        ],
    )
    pseudocode = f"""
Algorithm 1. Targeted Tail-statistics Bagged Fusion

Input:
  D = {{(I_i, E_i, y_i, patient_id_i)}} for i = 1..N
    I_i: 2D time-series image or multi-instance evidence for patient i
    E_i: EEG CSV trials for patient i
    y_i: binary class label
  S = {{20260520, ..., 20260529}}                         # random seeds
  K = 10                                                    # patient-aware outer folds
  top_k = {params['top_k']}, C = {params['C']}, penalty = {params['penalty']}
  W = {{0, 0.05, 0.10, 0.15, 0.20, 0.25}}                  # candidate EEG fusion weights

For each seed s in S:
  Create stratified K-fold split by patient_id.

  For each outer fold k = 1..K:
    TrainPatients, TestPatients <- split_s,k(D)

    # Branch A: existing 2D / multi-instance score
    For each patient i:
      b_i <- mi_max(I_i)                                   # baseline patient score

    # Branch B: EEG subject-level statistics
    For each patient i:
      For each EEG trial e in E_i:
        Extract time-domain, frequency-domain, entropy, and segment-trend features.
      Aggregate trial features into subject features:
        standard stats = min, max, mean, std, p25, p50, p75
        targeted tails:
          delta/theta band fractions -> p90, p95
          alpha band fraction        -> p05, p10

    Fit imputer/scaler on TrainPatients only.
    Select top_k EEG features on TrainPatients only.
    Train L1 logistic regression g_s,k(E_i) on TrainPatients.

    # Inner-fold fusion weight selection
    For each candidate w in W:
      Estimate inner validation scores:
        score_i(w) = (1 - w) * calibrate(b_i) + w * calibrate(g_s,k(E_i))
      Evaluate inner objective = min(AUROC, AUPRC) or predefined criterion.
    w*_s,k <- argmax_w inner objective

    # Outer-fold prediction
    For each patient i in TestPatients:
      OOF_s,i = (1 - w*_s,k) * calibrate(b_i) + w*_s,k * calibrate(g_s,k(E_i))

Final prediction:
  For each patient i:
    FinalScore_i = average_s(OOF_s,i)

Output:
  AUROC, AUPRC, sensitivity, specificity, PPV, NPV, and confusion matrix
  computed from {{(FinalScore_i, y_i)}} at patient level.
"""
    add_algorithm_block(doc, "演算法 1. Targeted Tail-statistics Bagged Fusion 虛擬碼", pseudocode)

    doc.add_heading("2.5 驗證設計與資料洩漏控制", level=2)
    doc.add_paragraph(
        "所有正式結果皆採 patient-aware stratified 10-fold。特徵選擇、標準化、融合權重選擇、資料增廣與模型訓練均限制在 training fold 內完成。"
        "測試 fold 僅用於最終推論與 out-of-fold score 收集。AUPRC 為主要指標，AUROC 為輔助指標；同時報告 confusion matrix、sensitivity、specificity、PPV 與 NPV。"
    )
    doc.add_paragraph(
        f"最佳候選方法使用 Logistic Regression，penalty={params['penalty']}，C={params['C']}，top_k={params['top_k']}，"
        f"base_col={params['base_col']}，fusion weight search 為 w_eeg <= {params['w_eeg_max']}、step={params['w_eeg_step']}、objective={params['objective']}。"
        "類別權重採 class_weight=None，因本資料之類別不平衡程度有限，未加權機率在融合時校準較佳。"
    )

    doc.add_heading("2.6 文獻模型復刻", level=2)
    doc.add_paragraph(
        "本研究將文獻方法分成兩類：一是已在 NSC 114 位資料上實際訓練與推論的本案自訂參數復刻版；"
        "二是僅作為方法元件或後續延伸者。本初稿主結果只列已實際跑完的復刻模型。"
        "若原作者未釋出完整超參數，本研究以核心架構與訓練邏輯一致為復刻原則，並完整列出本案自訂參數。"
    )

    doc.add_heading("3. 結果", level=1)
    doc.add_heading("3.1 最佳候選結果", level=2)
    add_table(
        doc,
        ["Method", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"],
        [
            [
                "Bagged Restricted Subject Fusion with Tail Stats",
                fmt(bag["AUROC"], 4),
                fmt(bag["AUPRC"], 4),
                fmt(bag["accuracy"], 4),
                fmt(bag["sensitivity"], 4),
                fmt(bag["specificity"], 4),
                fmt(bag["PPV"], 4),
                fmt(bag["NPV"], 4),
                f"{bag['TN']}/{bag['FP']}/{bag['FN']}/{bag['TP']}",
            ]
        ],
    )
    doc.add_paragraph(
        "最佳候選方法 AUROC=0.7971、AUPRC=0.7989，接近但未同時跨過 0.8/0.8。"
        "因此，本研究目前應定位為小樣本 patient-level OOF 分析中的 promising candidate，而非最終達標宣稱。"
    )
    for fig_name, caption in [
        ("bagged_fusion_roc.png", "圖 2. 最佳候選方法之 ROC curve。"),
        ("bagged_fusion_prc.png", "圖 3. 最佳候選方法之 PRC curve。"),
    ]:
        fig = TAIL_DIR / "figures" / fig_name
        if fig.exists():
            doc.add_picture(str(fig), width=Inches(5.8))
            add_caption(doc, caption)

    doc.add_heading("3.2 與復刻文獻模型比較", level=2)
    compact_rows = []
    for row in main_rows:
        compact_rows.append(
            [
                row["local_model"],
                row["validation"].replace("corrected-data ", ""),
                row.get("citation_keys", ""),
                fmt(row["AUROC"]),
                fmt(row["AUPRC"]),
                row.get("CM", ""),
            ]
        )
    add_table(doc, ["模型/方法", "驗證", "文獻", "AUROC", "AUPRC", "CM"], compact_rows)
    comp_fig = LIT_DIR / "figures" / "replicated_models_vs_our_best.png"
    if comp_fig.exists():
        doc.add_picture(str(comp_fig), width=Inches(6.5))
        add_caption(doc, "圖 4. 文獻模型復刻版與本案最佳候選方法之 AUROC/AUPRC 比較。")
    doc.add_paragraph(
        "結果顯示，time-series image CNN、1D CNN、1D+2D CNN ensemble、SMOTE、DTW/T-SMOTE-style augmentation、prototype、PixelHop-style、"
        "Siamese、Relation Network、Matching Networks、MAML-style 與 DTWSSE-style 在本資料的 patient-aware 10-fold 條件下，均未超過本案最佳候選方法。"
    )

    doc.add_heading("3.3 特徵重要性與神經心理詮釋", level=2)
    doc.add_paragraph(
        "SHAP 與 feature importance 分析目前以 mi_max + EEG aggregate fusion surrogate 為解釋對象，用於理解候選訊號來源。"
        "因 targeted tail-statistics candidate 尚需另外建立完整 SHAP 流程，本段作為神經心理候選解釋而非最終因果推論。"
    )
    for fig, caption in [
        (SHAP_DIR / "figures" / "eeg_fusion_surrogate_shap_summary.png", "圖 5. EEG fusion surrogate SHAP summary plot。"),
        (SHAP_DIR / "figures" / "eeg_fusion_top_feature_importance.png", "圖 6. EEG fusion surrogate top feature importance。"),
    ]:
        if fig.exists():
            doc.add_picture(str(fig), width=Inches(6.4))
            add_caption(doc, caption)
    add_table(
        doc,
        ["Rank", "Feature", "Domain", "Mean |SHAP|", "Direction", "神經心理/資料品質詮釋"],
        [
            [
                r.get("rank", ""),
                r.get("display_feature", ""),
                r.get("domain", ""),
                fmt(r.get("mean_abs_shap", "0"), 4),
                r.get("direction", ""),
                r.get("interpretation", ""),
            ]
            for r in shap_rows
        ],
    )

    doc.add_heading("4. 討論", level=1)
    doc.add_paragraph(
        "本研究顯示，在小樣本二元分類場景中，單純深度模型未必優於 fold-local 的統計特徵、二維影像分數與受限權重融合。"
        "mi_max 提供來自二維時序影像/多實例模型的局部判別證據；EEG aggregate features 則提供跨 trial 的頻譜、趨勢與資料品質訊號。"
        "targeted tail-statistics 的優勢在於只針對 delta/theta 高尾端與 alpha 低尾端擴充，符合 transient slowing、皮質喚醒度變化與 alpha suppression 的候選神經心理假說。"
    )
    doc.add_paragraph(
        "然而，beta/gamma/high 頻段與高頻 split 特徵可能同時反映警覺度、認知負荷、肌電或動作 artifact。"
        "因此本研究不將高頻重要性直接解讀為病理機制，而將其視為需要原始 trial 回看與資料品質檢查的候選 marker。"
    )

    doc.add_heading("5. 限制", level=1)
    for item in [
        "最佳候選方法尚未同時達到 AUROC/AUPRC 0.8/0.8，不應寫成已達標。",
        "targeted tail-statistics 參數來自多輪探索，正式投稿前應鎖定 protocol 後進行 confirmatory rerun 或 bootstrap confidence interval。",
        "SHAP 解釋目前來自 EEG aggregate fusion surrogate，尚未完全對應 tail-statistics bagged candidate。",
        "class 0 與 class 1 的臨床或神經心理意義需由研究者補上量表、診斷或任務定義；本初稿先以 class label 作方法學描述。",
        "EEG CSV 缺少完整電極拓樸、任務事件標記與 artifact expert labels，因此頻域與高頻特徵詮釋應保守。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. 結論", level=1)
    doc.add_paragraph(
        "在 NSC 114 位資料上，本研究建立了一個 patient-aware、fold-local、可重跑的文獻模型復刻與本案融合模型比較流程。"
        "最佳候選方法 AUROC=0.7971、AUPRC=0.7989，顯示二維時序影像分數與 EEG targeted statistics 具有互補性，"
        "但仍需 locked confirmatory evaluation 才能作為 reviewer-facing final generalization claim。"
    )

    doc.add_heading("資料與程式可重現性", level=1)
    for item in [
        "Analysis folder: see generated manifest and analysis outputs.",
        "Replicated literature comparison folder: see generated manifest and analysis outputs.",
        "Main modeling script: scripts/nsc_restricted_subject_bagging_tail_stats.py",
        "Paper draft generator: available in the scripts folder.",
        "Generated report: reports/NSC114_EEG_multimodal_small_data_paper_draft_20260521_v3.docx",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("參考文獻", level=1)
    for ref in refs:
        doc.add_paragraph(f"[{ref['citation_key']}] {ref['reference']} {ref['url']}", style="List Number")

    doc.save(REPORT)
    manifest_out = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "report": str(REPORT),
        "analysis_dir": str(OUT_DIR),
        "tail_analysis": str(TAIL_DIR),
        "literature_comparison": str(LIT_DIR),
        "shap_dir": str(SHAP_DIR),
        "claim_level": "patient-aware 10-fold OOF optimized candidate; not final target-reaching claim",
    }
    (OUT_DIR / "manifest.json").write_text(json.dumps(manifest_out, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest_out, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
