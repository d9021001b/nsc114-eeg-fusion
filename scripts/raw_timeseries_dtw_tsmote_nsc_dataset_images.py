#!/usr/bin/env python3
"""Raw time-series DTW-neighbor T-SMOTE-style augmentation.

This script connects bio-addict/csv-data raw physiological CSV files to the
NSC patient-level modelling task. In each train fold only, it:

1. resamples raw time-series per case into fixed-length channel-group sequences;
2. extracts fold-local raw-series features;
3. optionally creates minority-class synthetic raw sequences using DTW-neighbor
   interpolation inside the training fold;
4. compares no-augmentation vs DTW/T-SMOTE-style augmentation under patient-level
   stratified 10-fold validation.

The implementation is deliberately conservative: no test-fold sequence is used
for augmentation, scaling, PCA, feature selection, or nearest-neighbor search.
"""

from __future__ import annotations

import argparse
import csv
import json
import math
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    average_precision_score,
    confusion_matrix,
    precision_recall_curve,
    roc_auc_score,
    roc_curve,
)
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler


GROUPS = ["EKG_PRIMARY", "RESP", "SC_AUX"]


def write_csv(path: Path, rows: Iterable[dict], fieldnames: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def natural_key(text: str) -> Tuple:
    parts = re.split(r"(\d+)", str(text))
    return tuple(int(p) if p.isdigit() else p for p in parts)


def signal_group(signal: str) -> str:
    upper = signal.upper()
    if "RESP" in upper:
        return "RESP"
    if "SC_" in upper or "AUX" in upper:
        return "SC_AUX"
    if "EKG" in upper or "PRIMARY_RAW" in upper or "IBI" in upper:
        return "EKG_PRIMARY"
    return "OTHER"


def value_column(df: pd.DataFrame) -> str:
    for col in ["value", "ibi_ms", "resp_period_ms"]:
        if col in df.columns:
            return col
    numeric = [c for c in df.columns if c != "time_sec" and pd.api.types.is_numeric_dtype(df[c])]
    if not numeric:
        raise ValueError("No value-like numeric column found")
    return numeric[-1]


def robust_resample(times: np.ndarray, values: np.ndarray, length: int) -> np.ndarray:
    mask = np.isfinite(times) & np.isfinite(values)
    times = times[mask]
    values = values[mask]
    if len(values) == 0:
        return np.zeros(length, dtype=np.float32)
    if len(values) == 1 or np.nanmax(times) <= np.nanmin(times):
        return np.full(length, float(values[0]), dtype=np.float32)
    order = np.argsort(times)
    times = times[order]
    values = values[order]
    _, unique_idx = np.unique(times, return_index=True)
    times = times[unique_idx]
    values = values[unique_idx]
    grid = np.linspace(times.min(), times.max(), length)
    seq = np.interp(grid, times, values).astype(np.float32)
    return seq


def normalize_sequence(seq: np.ndarray) -> np.ndarray:
    med = np.nanmedian(seq)
    iqr = np.nanpercentile(seq, 75) - np.nanpercentile(seq, 25)
    scale = iqr if iqr > 1e-6 else np.nanstd(seq)
    scale = scale if scale > 1e-6 else 1.0
    return ((seq - med) / scale).astype(np.float32)


def load_labels(manifest_path: Path) -> Dict[str, int]:
    labels: Dict[str, int] = {}
    with manifest_path.open("r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            labels[str(row["subject_id"])] = int(row["label"])
    return labels


def load_case_sequences(csv_dir: Path, manifest_path: Path, length: int, max_rows_per_csv: int) -> Tuple[List[str], np.ndarray, np.ndarray, Dict[str, dict]]:
    labels = load_labels(manifest_path)
    grouped: Dict[str, Dict[str, List[np.ndarray]]] = defaultdict(lambda: defaultdict(list))
    file_counts: Dict[str, Counter] = defaultdict(Counter)
    pattern = re.compile(r"^(\d+)_Sess\d+_(.+)\.csv$")
    for path in sorted(csv_dir.glob("*.csv"), key=lambda p: natural_key(p.name)):
        m = pattern.match(path.name)
        if not m:
            continue
        subject, signal = m.group(1), m.group(2)
        if subject not in labels:
            continue
        group = signal_group(signal)
        if group == "OTHER":
            continue
        try:
            header = pd.read_csv(path, nrows=0)
            cols = list(header.columns)
            read_cols = ["time_sec"]
            for candidate in ["value", "ibi_ms", "resp_period_ms"]:
                if candidate in cols:
                    read_cols.append(candidate)
                    break
            if len(read_cols) == 1:
                read_cols = cols[:3]
            df = pd.read_csv(path, usecols=read_cols, nrows=max_rows_per_csv)
            col = value_column(df)
            seq = robust_resample(df["time_sec"].to_numpy(dtype=float), df[col].to_numpy(dtype=float), length)
            grouped[subject][group].append(normalize_sequence(seq))
            file_counts[subject][group] += 1
        except Exception:
            continue

    subjects = sorted(grouped.keys(), key=natural_key)
    X_seq = np.zeros((len(subjects), len(GROUPS), length), dtype=np.float32)
    meta: Dict[str, dict] = {}
    for i, subject in enumerate(subjects):
        for g, group in enumerate(GROUPS):
            seqs = grouped[subject].get(group, [])
            if seqs:
                X_seq[i, g, :] = np.mean(np.stack(seqs), axis=0)
        meta[subject] = {
            "label": labels[subject],
            "file_counts": dict(file_counts[subject]),
            "total_files": int(sum(file_counts[subject].values())),
        }
    y = np.asarray([labels[s] for s in subjects], dtype=int)
    return subjects, X_seq, y, meta


def extract_features(X_seq: np.ndarray) -> np.ndarray:
    rows = []
    for case in X_seq:
        feats = []
        for seq in case:
            dif = np.diff(seq)
            x = np.arange(len(seq), dtype=float)
            slope = float(np.polyfit(x, seq, 1)[0]) if len(seq) > 1 else 0.0
            centered = seq - np.mean(seq)
            autocorr = float(np.dot(centered[:-1], centered[1:]) / (np.dot(centered, centered) + 1e-9)) if len(seq) > 2 else 0.0
            fft = np.abs(np.fft.rfft(seq))
            dom = float(np.argmax(fft[1:]) + 1) / max(len(fft), 1) if len(fft) > 1 else 0.0
            hist, _ = np.histogram(seq, bins=16, density=True)
            hist = hist / (hist.sum() + 1e-9)
            entropy = float(-(hist * np.log(hist + 1e-9)).sum())
            feats.extend(
                [
                    float(np.mean(seq)),
                    float(np.std(seq)),
                    float(np.min(seq)),
                    float(np.max(seq)),
                    float(np.percentile(seq, 10)),
                    float(np.percentile(seq, 25)),
                    float(np.percentile(seq, 50)),
                    float(np.percentile(seq, 75)),
                    float(np.percentile(seq, 90)),
                    float(np.mean(np.abs(dif))) if len(dif) else 0.0,
                    float(np.std(dif)) if len(dif) else 0.0,
                    float(np.sum(np.abs(dif))) if len(dif) else 0.0,
                    slope,
                    autocorr,
                    dom,
                    entropy,
                ]
            )
        rows.append(feats)
    return np.asarray(rows, dtype=np.float32)


def dtw_distance(a: np.ndarray, b: np.ndarray) -> float:
    # Sakoe-Chiba band keeps this cheap and temporal-aware.
    n = len(a)
    m = len(b)
    band = max(abs(n - m), int(0.15 * max(n, m)))
    inf = 1e18
    prev = np.full(m + 1, inf, dtype=np.float64)
    curr = np.full(m + 1, inf, dtype=np.float64)
    prev[0] = 0.0
    for i in range(1, n + 1):
        lo = max(1, i - band)
        hi = min(m, i + band)
        curr.fill(inf)
        for j in range(lo, hi + 1):
            cost = float((a[i - 1] - b[j - 1]) ** 2)
            curr[j] = cost + min(prev[j], curr[j - 1], prev[j - 1])
        prev, curr = curr, prev
    return float(np.sqrt(prev[m]))


def flattened_for_dtw(case_seq: np.ndarray, reduced_length: int = 32) -> np.ndarray:
    parts = []
    for seq in case_seq:
        x = np.linspace(0, len(seq) - 1, reduced_length)
        parts.append(np.interp(x, np.arange(len(seq)), seq))
    return np.concatenate(parts).astype(np.float32)


def dtw_tsmote_augment(
    X_train_seq: np.ndarray,
    y_train: np.ndarray,
    target_label: int = 1,
    random_state: int = 0,
) -> Tuple[np.ndarray, np.ndarray, int]:
    rng = np.random.default_rng(random_state)
    counts = Counter(y_train.tolist())
    majority = max(counts.values())
    target_count = counts[target_label]
    n_to_make = max(0, majority - target_count)
    if n_to_make == 0:
        return X_train_seq, y_train, 0
    target_idx = np.where(y_train == target_label)[0]
    if len(target_idx) < 2:
        return X_train_seq, y_train, 0
    reduced = [flattened_for_dtw(X_train_seq[i]) for i in target_idx]
    dist = np.full((len(target_idx), len(target_idx)), np.inf, dtype=float)
    for i in range(len(target_idx)):
        for j in range(i + 1, len(target_idx)):
            d = dtw_distance(reduced[i], reduced[j])
            dist[i, j] = dist[j, i] = d
    synthetic = []
    for _ in range(n_to_make):
        anchor_local = int(rng.integers(0, len(target_idx)))
        neighbor_order = np.argsort(dist[anchor_local])
        neighbor_local = int(neighbor_order[0] if neighbor_order[0] != anchor_local else neighbor_order[1])
        a = X_train_seq[target_idx[anchor_local]]
        b = X_train_seq[target_idx[neighbor_local]]
        lam = float(rng.uniform(0.25, 0.75))
        synth = a + lam * (b - a)
        # small temporal jitter, still derived only from train fold.
        shift = int(rng.integers(-3, 4))
        synth = np.roll(synth, shift=shift, axis=1)
        noise = rng.normal(0, 0.015, size=synth.shape).astype(np.float32)
        synthetic.append((synth + noise).astype(np.float32))
    X_aug = np.concatenate([X_train_seq, np.stack(synthetic)], axis=0)
    y_aug = np.concatenate([y_train, np.full(len(synthetic), target_label, dtype=int)])
    return X_aug, y_aug, len(synthetic)


def metrics(y_true: np.ndarray, score: np.ndarray, pred: np.ndarray) -> dict:
    tn, fp, fn, tp = confusion_matrix(y_true, pred, labels=[0, 1]).ravel()
    return {
        "AUROC": float(roc_auc_score(y_true, score)),
        "AUPRC": float(average_precision_score(y_true, score)),
        "accuracy": float((tp + tn) / max(tp + tn + fp + fn, 1)),
        "sensitivity": float(tp / max(tp + fn, 1)),
        "specificity": float(tn / max(tn + fp, 1)),
        "PPV": float(tp / max(tp + fp, 1)),
        "NPV": float(tn / max(tn + fn, 1)),
        "TN": int(tn),
        "FP": int(fp),
        "FN": int(fn),
        "TP": int(tp),
    }


def run_model(X_train_feat: np.ndarray, y_train: np.ndarray, X_test_feat: np.ndarray, model_name: str, seed: int) -> Tuple[np.ndarray, np.ndarray]:
    if model_name == "logistic":
        model = make_pipeline(
            SimpleImputer(strategy="median"),
            SelectKBest(f_classif, k=min(48, X_train_feat.shape[1])),
            StandardScaler(),
            LogisticRegression(max_iter=4000, C=0.25, class_weight="balanced", solver="liblinear"),
        )
    elif model_name == "extratrees":
        model = make_pipeline(
            SimpleImputer(strategy="median"),
            SelectKBest(f_classif, k=min(48, X_train_feat.shape[1])),
            ExtraTreesClassifier(
                n_estimators=240,
                random_state=seed,
                class_weight="balanced",
                max_features="sqrt",
                min_samples_leaf=2,
                n_jobs=-1,
            ),
        )
    else:
        raise ValueError(model_name)
    with np.errstate(invalid="ignore"):
        model.fit(X_train_feat, y_train)
    score = model.predict_proba(X_test_feat)[:, 1]
    pred = (score >= 0.5).astype(int)
    return score, pred


def make_figures(out_dir: Path, y_true: np.ndarray, score: np.ndarray, pred: np.ndarray, title: str, met: dict) -> Dict[str, Path]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths = {}
    fpr, tpr, _ = roc_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={met['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray", linewidth=1)
    plt.title(f"Raw time-series ROC: {title}")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = fig_dir / "best_raw_timeseries_roc.png"
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()

    precision, recall, _ = precision_recall_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={met['AUPRC']:.3f}")
    plt.axhline(float(y_true.mean()), linestyle="--", color="gray", linewidth=1, label=f"prevalence={y_true.mean():.3f}")
    plt.title(f"Raw time-series PRC: {title}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = fig_dir / "best_raw_timeseries_prc.png"
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()

    cm = confusion_matrix(y_true, pred, labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title(f"Raw time-series CM: {title}")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = fig_dir / "best_raw_timeseries_confusion_matrix.png"
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], best: dict, figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC 原始 Time-series DTW/T-SMOTE Augmentation 初版報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法定位", level=1)
    doc.add_paragraph(
        "本報告將 bio-addict/csv-data 原始生理訊號 CSV 接入 patient-level 10-fold 實驗。"
        "每一 fold 只使用 training fold 建立 resampling、特徵萃取與 DTW-neighbor synthetic augmentation；test fold 不參與 augmentation。"
    )
    doc.add_paragraph(
        "本版實作為 raw time-series DTW-neighbor T-SMOTE-style augmentation：先以 DTW 在 label 1 training sequences 中尋找近鄰，"
        "再以時間序列插值與小幅 temporal jitter 產生合成序列。"
    )

    doc.add_heading("二、資料概況", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "內容"
    for k, v in [
        ("CSV 資料夾", manifest["csv_dir"]),
        ("可對應個案數", manifest["total_cases"]),
        ("label 0 / label 1", f"{manifest['cases_by_label'].get('0', 0)} / {manifest['cases_by_label'].get('1', 0)}"),
        ("序列長度", manifest["sequence_length"]),
        ("每個 CSV 讀取列數上限", manifest["max_rows_per_csv"]),
        ("通道群組", ", ".join(GROUPS)),
        ("驗證方式", manifest["validation"]),
    ]:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)

    doc.add_heading("三、結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "模型", "合成樣本", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "CM"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for r in rows:
        cells = table.add_row().cells
        vals = [
            r["augmentation"],
            r["model"],
            r["synthetic_samples"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        for cell, val in zip(cells, vals):
            cell.text = str(val)

    doc.add_heading("四、最佳圖表", level=1)
    doc.add_paragraph(f"最佳方法：{best['augmentation']} + {best['model']}")
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))

    doc.add_heading("五、判讀", level=1)
    doc.add_paragraph(
        "本報告完成原始 CSV time-series 與 DTW-neighbor augmentation 的接軌。本版為初始片段讀取版，"
        "高頻大型 CSV 僅讀取前段固定列數以建立可重跑 baseline。若結果仍未達 0.8/0.8，"
        "代表僅靠目前 raw-series summary features 與簡化 DTW interpolation 尚不足；下一版可加入更完整的 T-SMOTE boundary sampling、"
        "DTW barycenter averaging、以及和 2D patch/green features 的 late fusion。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--out-dir", default="analysis/nsc_raw_timeseries_dtw_tsmote_20260519")
    parser.add_argument("--report", default="reports/NSC_raw_timeseries_DTW_TSMOTE_initial_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    args = parser.parse_args()

    csv_dir = Path(args.csv_dir).resolve()
    manifest_path = Path(args.manifest).resolve()
    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, X_seq, y, meta = load_case_sequences(csv_dir, manifest_path, args.sequence_length, args.max_rows_per_csv)
    methods = [
        ("none", "logistic"),
        ("none", "extratrees"),
        ("dtw_tsmote", "logistic"),
        ("dtw_tsmote", "extratrees"),
    ]
    splitter = StratifiedKFold(n_splits=10, shuffle=True, random_state=42)
    rows = []
    predictions = []
    best_payload = None
    for augmentation, model_name in methods:
        all_score = np.zeros(len(y), dtype=float)
        all_pred = np.zeros(len(y), dtype=int)
        synthetic_total = 0
        for fold, (train_idx, test_idx) in enumerate(splitter.split(X_seq, y), start=1):
            X_train_seq = X_seq[train_idx]
            y_train = y[train_idx]
            X_test_seq = X_seq[test_idx]
            if augmentation == "dtw_tsmote":
                X_train_aug, y_train_aug, n_syn = dtw_tsmote_augment(
                    X_train_seq,
                    y_train,
                    target_label=1,
                    random_state=args.random_state + fold,
                )
                synthetic_total += n_syn
            else:
                X_train_aug, y_train_aug = X_train_seq, y_train
            X_train_feat = extract_features(X_train_aug)
            X_test_feat = extract_features(X_test_seq)
            score, pred = run_model(X_train_feat, y_train_aug, X_test_feat, model_name, args.random_state + fold)
            all_score[test_idx] = score
            all_pred[test_idx] = pred
        met = metrics(y, all_score, all_pred)
        row = {
            "augmentation": augmentation,
            "model": model_name,
            "synthetic_samples": synthetic_total,
            **met,
        }
        rows.append(row)
        for i, subject in enumerate(subjects):
            predictions.append(
                {
                    "augmentation": augmentation,
                    "model": model_name,
                    "subject_id": subject,
                    "true_label": int(y[i]),
                    "score": float(all_score[i]),
                    "pred_label": int(all_pred[i]),
                }
            )
        if best_payload is None or (row["AUPRC"], row["AUROC"]) > (best_payload["row"]["AUPRC"], best_payload["row"]["AUROC"]):
            best_payload = {"row": row, "score": all_score.copy(), "pred": all_pred.copy()}

    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    row_fields = ["augmentation", "model", "synthetic_samples", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    pred_fields = ["augmentation", "model", "subject_id", "true_label", "score", "pred_label"]
    write_csv(out_dir / "raw_timeseries_dtw_tsmote_summary.csv", rows, row_fields)
    write_csv(out_dir / "raw_timeseries_predictions_long.csv", predictions, pred_fields)
    best = best_payload["row"]
    figures = make_figures(
        out_dir,
        y,
        best_payload["score"],
        best_payload["pred"],
        f"{best['augmentation']} + {best['model']}",
        best,
    )
    cases_by_label = {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))}
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "csv_dir": str(csv_dir),
        "manifest_path": str(manifest_path),
        "total_cases": len(subjects),
        "cases_by_label": cases_by_label,
        "sequence_length": args.sequence_length,
        "max_rows_per_csv": args.max_rows_per_csv,
        "validation": "patient-level stratified 10-fold",
        "augmentation": "raw time-series DTW-neighbor T-SMOTE-style, training-fold-only",
        "best": best,
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "raw_timeseries_dtw_tsmote_summary.csv"),
            "predictions": str(out_dir / "raw_timeseries_predictions_long.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": {k: str(v) for k, v in figures.items()},
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, best, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
