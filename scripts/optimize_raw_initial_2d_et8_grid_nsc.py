#!/usr/bin/env python3
"""Nested optimization for raw_initial + 2D ET8 + grid fusion."""

from __future__ import annotations

import argparse
import csv
import json
import sys
import warnings
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from sklearn.ensemble import ExtraTreesClassifier
from sklearn.feature_selection import SelectKBest, VarianceThreshold, f_classif
from sklearn.impute import SimpleImputer
from sklearn.metrics import confusion_matrix, precision_recall_curve, roc_curve
from sklearn.model_selection import StratifiedGroupKFold

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences  # noqa: E402
from score_level_fusion_2d_raw_nsc import build_2d_matrix, fit_raw_dtw_et, metric_dict, write_csv  # noqa: E402

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)


def config_grid() -> List[dict]:
    configs: List[dict] = []
    for top_k in [48, 64, 96, 128, 160, 192]:
        for min_leaf in [1, 2, 3]:
            configs.append(
                {
                    "config_id": f"k{top_k}_leaf{min_leaf}_sqrt_bal",
                    "top_k": top_k,
                    "min_samples_leaf": min_leaf,
                    "max_features": "sqrt",
                    "class_weight": "balanced",
                    "n_estimators": 360,
                }
            )
    configs.extend(
        [
            {
                "config_id": "k96_leaf2_log2_bal",
                "top_k": 96,
                "min_samples_leaf": 2,
                "max_features": "log2",
                "class_weight": "balanced",
                "n_estimators": 360,
            },
            {
                "config_id": "k128_leaf2_log2_bal",
                "top_k": 128,
                "min_samples_leaf": 2,
                "max_features": "log2",
                "class_weight": "balanced",
                "n_estimators": 360,
            },
            {
                "config_id": "k96_leaf2_sqrt_bal_sub",
                "top_k": 96,
                "min_samples_leaf": 2,
                "max_features": "sqrt",
                "class_weight": "balanced_subsample",
                "n_estimators": 360,
            },
            {
                "config_id": "k128_leaf2_sqrt_bal_sub",
                "top_k": 128,
                "min_samples_leaf": 2,
                "max_features": "sqrt",
                "class_weight": "balanced_subsample",
                "n_estimators": 360,
            },
        ]
    )
    return configs


def fit_2d_et8_config(X2d: np.ndarray, y: np.ndarray, train_idx: np.ndarray, test_idx: np.ndarray, seed: int, cfg: dict) -> Tuple[np.ndarray, np.ndarray, dict]:
    Xtr = X2d[train_idx]
    Xte = X2d[test_idx]
    imputer = SimpleImputer(strategy="median")
    Xtr_i = imputer.fit_transform(Xtr)
    Xte_i = imputer.transform(Xte)
    var = VarianceThreshold(threshold=1e-8)
    Xtr_v = var.fit_transform(Xtr_i)
    Xte_v = var.transform(Xte_i)
    if Xtr_v.shape[1] == 0:
        Xtr_v, Xte_v = Xtr_i, Xte_i
    k = min(int(cfg["top_k"]), Xtr_v.shape[1])
    selector = SelectKBest(f_classif, k=k)
    with np.errstate(invalid="ignore", divide="ignore"):
        Xtr_s = selector.fit_transform(Xtr_v, y[train_idx])
    Xte_s = selector.transform(Xte_v)
    model = ExtraTreesClassifier(
        n_estimators=int(cfg["n_estimators"]),
        random_state=seed,
        class_weight=cfg["class_weight"],
        max_features=cfg["max_features"],
        min_samples_leaf=int(cfg["min_samples_leaf"]),
        n_jobs=-1,
    )
    model.fit(Xtr_s, y[train_idx])
    stats = {
        "input_features": int(X2d.shape[1]),
        "after_variance_features": int(Xtr_v.shape[1]),
        "selected_features": int(k),
    }
    return model.predict_proba(Xtr_s)[:, 1], model.predict_proba(Xte_s)[:, 1], stats


def search_fusion(raw_score: np.ndarray, two_score: np.ndarray, y_true: np.ndarray, grid_step: float) -> Tuple[dict, List[dict]]:
    weights = np.round(np.arange(0.0, 1.0 + grid_step / 2.0, grid_step), 6)
    rows: List[dict] = []
    best = None
    for weight_2d in weights:
        fused = (1.0 - float(weight_2d)) * raw_score + float(weight_2d) * two_score
        met = metric_dict(y_true, fused)
        row = {"weight_raw": 1.0 - float(weight_2d), "weight_2d": float(weight_2d), **met}
        rows.append(row)
        if best is None or (met["AUPRC"], met["AUROC"]) > (best["AUPRC"], best["AUROC"]):
            best = row
    return best, rows


def make_figures(out_dir: Path, y_true: np.ndarray, score: np.ndarray, met: dict, fold_rows: List[dict], config_rows: List[dict]) -> Dict[str, Path]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, Path] = {}

    fpr, tpr, _ = roc_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={met['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Optimized raw+2D ET8 fusion ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    paths["roc"] = fig_dir / "optimized_fusion_roc.png"
    plt.tight_layout()
    plt.savefig(paths["roc"], dpi=180)
    plt.close()

    precision, recall, _ = precision_recall_curve(y_true, score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={met['AUPRC']:.3f}")
    plt.axhline(float(y_true.mean()), linestyle="--", color="gray", label=f"prevalence={y_true.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Optimized raw+2D ET8 fusion PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    paths["prc"] = fig_dir / "optimized_fusion_prc.png"
    plt.tight_layout()
    plt.savefig(paths["prc"], dpi=180)
    plt.close()

    cm = confusion_matrix(y_true, (score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Optimized raw+2D ET8 confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    paths["cm"] = fig_dir / "optimized_fusion_cm.png"
    plt.tight_layout()
    plt.savefig(paths["cm"], dpi=180)
    plt.close()

    weights = [r["weight_2d"] for r in fold_rows]
    labels = [str(r["fold"]) for r in fold_rows]
    plt.figure(figsize=(7.0, 4.2))
    plt.bar(labels, weights, color="#3b6ea8")
    plt.ylim(0, 1)
    plt.xlabel("Fold")
    plt.ylabel("Selected 2D weight")
    plt.title("Outer-fold selected fusion weights")
    plt.grid(axis="y", alpha=0.25)
    paths["weights"] = fig_dir / "selected_fusion_weights.png"
    plt.tight_layout()
    plt.savefig(paths["weights"], dpi=180)
    plt.close()

    top_configs = sorted(config_rows, key=lambda r: (r["selected_count"], r["mean_inner_AUPRC"]), reverse=True)[:10]
    plt.figure(figsize=(10, 4.8))
    plt.bar([r["config_id"] for r in top_configs], [r["mean_inner_AUPRC"] for r in top_configs], color="#7a8f3a")
    plt.xticks(rotation=35, ha="right", fontsize=8)
    plt.ylabel("Mean selected inner AUPRC")
    plt.title("Most frequently selected 2D configurations")
    plt.grid(axis="y", alpha=0.25)
    paths["configs"] = fig_dir / "selected_config_summary.png"
    plt.tight_layout()
    plt.savefig(paths["configs"], dpi=180)
    plt.close()
    return paths


def make_report(report_path: Path, manifest: dict, summary_rows: List[dict], fold_rows: List[dict], config_summary: List[dict], figures: Dict[str, Path]) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("NSC raw_initial + 2D ET8 Grid Fusion 優化報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")

    doc.add_heading("一、方法", level=1)
    doc.add_paragraph(
        "本版聚焦前一輪 ablation 最佳主線：raw_initial + 8x8 2D variance-filtered ExtraTrees + grid fusion。"
        "外層使用 patient-aware StratifiedGroupKFold 10-fold；每個 outer training fold 內再用 inner 5-fold 選擇 "
        "2D ExtraTrees 參數與 fusion weight，選定後才套用到 outer test。"
    )
    doc.add_paragraph(
        "調整範圍包含 2D SelectKBest top_k、ExtraTrees min_samples_leaf、max_features、class_weight，以及 0.05 間距的 fusion weight。"
        "raw_initial 分支維持前一版設定，避免同時更動太多因素。"
    )

    doc.add_heading("二、整體結果", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["方法", "AUROC", "AUPRC", "ACC", "Sensitivity", "Specificity", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in summary_rows:
        vals = [
            r["method"],
            f"{r['AUROC']:.3f}",
            f"{r['AUPRC']:.3f}",
            f"{r['accuracy']:.3f}",
            f"{r['sensitivity']:.3f}",
            f"{r['specificity']:.3f}",
            f"{r['PPV']:.3f}",
            f"{r['NPV']:.3f}",
            f"{r['TN']}/{r['FP']}/{r['FN']}/{r['TP']}",
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、各折選定參數", level=1)
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["fold", "config", "top_k", "leaf", "max_features", "class_weight", "w_2d", "inner_AUPRC", "overlap"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in fold_rows:
        vals = [
            r["fold"],
            r["config_id"],
            r["top_k"],
            r["min_samples_leaf"],
            r["max_features"],
            r["class_weight"],
            f"{r['weight_2d']:.2f}",
            f"{r['inner_AUPRC']:.3f}",
            r["group_overlap"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、參數選擇摘要", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for c, h in zip(table.rows[0].cells, ["config_id", "selected_count", "mean_inner_AUPRC", "mean_weight_2d"]):
        c.text = h
    for r in config_summary[:10]:
        vals = [r["config_id"], r["selected_count"], f"{r['mean_inner_AUPRC']:.3f}", f"{r['mean_weight_2d']:.2f}"]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("五、圖表", level=1)
    doc.add_picture(str(figures["roc"]), width=Inches(5.3))
    doc.add_picture(str(figures["prc"]), width=Inches(5.3))
    doc.add_picture(str(figures["cm"]), width=Inches(4.8))
    doc.add_picture(str(figures["weights"]), width=Inches(5.8))
    doc.add_picture(str(figures["configs"]), width=Inches(6.2))

    doc.add_heading("六、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(
        f"最佳輸出為 {best['method']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。"
        "此結果採 nested hyperparameter selection，可作為目前主線的 patient-aware 10-fold 證據。"
    )
    doc.add_paragraph(
        "若本版未達 0.8/0.8，下一步不建議再擴大高維 2D 特徵，而應針對錯誤案例做 patch-level error analysis，"
        "找出穩定出現在 false positive/false negative 的 channel/plot/patch，再建立更小、更有生理意義的 2D feature subset。"
    )
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_raw_initial_2d_et8_optimized_group10fold_20260519")
    parser.add_argument("--report", default="reports/NSC_raw_initial_2D_ET8_grid_fusion_optimized_group10fold_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--fusion-step", type=float, default=0.05)
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(
        Path(args.csv_dir).resolve(),
        Path(args.manifest).resolve(),
        args.sequence_length,
        args.max_rows_per_csv,
    )
    X2d, feature_names = build_2d_matrix(Path(args.image_root).resolve(), subjects)
    groups = np.asarray(subjects)
    configs = config_grid()

    raw_all = np.zeros(len(y), dtype=float)
    two_all = np.zeros(len(y), dtype=float)
    fusion_all = np.zeros(len(y), dtype=float)
    fold_rows: List[dict] = []
    inner_rows: List[dict] = []
    prediction_rows: List[dict] = []

    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)
    for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
        inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 100)
        raw_oof = np.zeros(len(train_idx), dtype=float)
        two_oof_by_config = {cfg["config_id"]: np.zeros(len(train_idx), dtype=float) for cfg in configs}
        for inner_fold, (tr_local, va_local) in enumerate(
            inner.split(np.zeros(len(train_idx)), y[train_idx], groups=groups[train_idx]),
            start=1,
        ):
            tr_idx = train_idx[tr_local]
            va_idx = train_idx[va_local]
            _raw_train, raw_val, _n_syn = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, args.random_state + fold * 100 + inner_fold)
            raw_oof[va_local] = raw_val
            for cfg in configs:
                _two_train, two_val, _stats = fit_2d_et8_config(
                    X2d,
                    y,
                    tr_idx,
                    va_idx,
                    args.random_state + fold * 1000 + inner_fold,
                    cfg,
                )
                two_oof_by_config[cfg["config_id"]][va_local] = two_val

        best = None
        for cfg in configs:
            best_fusion, candidates = search_fusion(raw_oof, two_oof_by_config[cfg["config_id"]], y[train_idx], args.fusion_step)
            for row in candidates:
                inner_rows.append({"outer_fold": fold, **cfg, **row})
            if best is None or (best_fusion["AUPRC"], best_fusion["AUROC"]) > (best["AUPRC"], best["AUROC"]):
                best = {**cfg, **best_fusion}

        selected_cfg = {k: best[k] for k in ["config_id", "top_k", "min_samples_leaf", "max_features", "class_weight", "n_estimators"]}
        _raw_train, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
        _two_train, two_test, stats = fit_2d_et8_config(X2d, y, train_idx, test_idx, args.random_state + fold * 10, selected_cfg)
        fused_test = best["weight_raw"] * raw_test + best["weight_2d"] * two_test
        raw_all[test_idx] = raw_test
        two_all[test_idx] = two_test
        fusion_all[test_idx] = fused_test
        train_groups = set(groups[train_idx].tolist())
        test_groups = set(groups[test_idx].tolist())
        fold_rows.append(
            {
                "fold": fold,
                **selected_cfg,
                "weight_raw": float(best["weight_raw"]),
                "weight_2d": float(best["weight_2d"]),
                "inner_AUROC": float(best["AUROC"]),
                "inner_AUPRC": float(best["AUPRC"]),
                "synthetic_samples": int(n_syn),
                **stats,
                "train_cases": int(len(train_idx)),
                "test_cases": int(len(test_idx)),
                "group_overlap": int(len(train_groups & test_groups)),
            }
        )

    for i, subject in enumerate(subjects):
        prediction_rows.append(
            {
                "subject_id": subject,
                "true_label": int(y[i]),
                "raw_score": float(raw_all[i]),
                "two_d_score": float(two_all[i]),
                "fusion_score": float(fusion_all[i]),
            }
        )

    summary_rows = [
        {"method": "optimized_nested_raw_initial_2d_et8_grid", **metric_dict(y, fusion_all)},
        {"method": "selected_2d_et8_branch_only", **metric_dict(y, two_all)},
        {"method": "raw_initial_branch_only", **metric_dict(y, raw_all)},
    ]
    summary_rows = sorted(summary_rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)

    config_summary: List[dict] = []
    for cfg in configs:
        selected = [r for r in fold_rows if r["config_id"] == cfg["config_id"]]
        if not selected:
            config_summary.append({"config_id": cfg["config_id"], "selected_count": 0, "mean_inner_AUPRC": 0.0, "mean_weight_2d": 0.0})
            continue
        config_summary.append(
            {
                "config_id": cfg["config_id"],
                "selected_count": len(selected),
                "mean_inner_AUPRC": float(np.mean([r["inner_AUPRC"] for r in selected])),
                "mean_weight_2d": float(np.mean([r["weight_2d"] for r in selected])),
            }
        )
    config_summary = sorted(config_summary, key=lambda r: (r["selected_count"], r["mean_inner_AUPRC"]), reverse=True)

    summary_fields = ["method", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"]
    write_csv(out_dir / "optimized_summary.csv", summary_rows, summary_fields)
    write_csv(out_dir / "optimized_predictions.csv", prediction_rows, ["subject_id", "true_label", "raw_score", "two_d_score", "fusion_score"])
    write_csv(
        out_dir / "outer_fold_selected_params.csv",
        fold_rows,
        [
            "fold",
            "config_id",
            "top_k",
            "min_samples_leaf",
            "max_features",
            "class_weight",
            "n_estimators",
            "weight_raw",
            "weight_2d",
            "inner_AUROC",
            "inner_AUPRC",
            "synthetic_samples",
            "input_features",
            "after_variance_features",
            "selected_features",
            "train_cases",
            "test_cases",
            "group_overlap",
        ],
    )
    write_csv(
        out_dir / "inner_config_fusion_search.csv",
        inner_rows,
        [
            "outer_fold",
            "config_id",
            "top_k",
            "min_samples_leaf",
            "max_features",
            "class_weight",
            "n_estimators",
            "weight_raw",
            "weight_2d",
            "AUROC",
            "AUPRC",
            "accuracy",
            "sensitivity",
            "specificity",
            "PPV",
            "NPV",
            "TN",
            "FP",
            "FN",
            "TP",
        ],
    )
    write_csv(out_dir / "selected_config_summary.csv", config_summary, ["config_id", "selected_count", "mean_inner_AUPRC", "mean_weight_2d"])

    fusion_met = next(r for r in summary_rows if r["method"] == "optimized_nested_raw_initial_2d_et8_grid")
    figures = make_figures(out_dir, y, fusion_all, fusion_met, fold_rows, config_summary)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_cases": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold config and fusion selection",
        "outer_splits": args.n_splits,
        "inner_splits": args.inner_splits,
        "config_count": len(configs),
        "fusion_step": args.fusion_step,
        "feature_count_2d": len(feature_names),
        "best": summary_rows[0],
        "optimized_fusion": fusion_met,
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "optimized_summary.csv"),
            "predictions": str(out_dir / "optimized_predictions.csv"),
            "outer_fold_selected_params": str(out_dir / "outer_fold_selected_params.csv"),
            "inner_search": str(out_dir / "inner_config_fusion_search.csv"),
            "selected_config_summary": str(out_dir / "selected_config_summary.csv"),
            "report": str(Path(args.report).resolve()),
            "figures": {k: str(v) for k, v in figures.items()},
        },
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, summary_rows, fold_rows, config_summary, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
