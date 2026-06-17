#!/usr/bin/env python3
"""Package literature deep-model comparison addendum for the NSC paper."""

from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "analysis" / "nsc_literature_deep_model_comparison_20260519"
REPORT = ROOT / "reports" / "NSC_paper_literature_deep_model_comparison_addendum_20260519.docx"


def read_csv_rows(path: Path) -> List[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: Iterable[dict], fields: List[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fields})


def current_method_rows() -> List[dict]:
    ablation = read_csv_rows(ROOT / "analysis" / "nsc_raw2d_ablation_group10fold_20260519" / "ablation_summary.csv")
    paper_methods = read_csv_rows(ROOT / "analysis" / "nsc_dataset_images_paper_method_comparison_20260519" / "method_comparison_summary.csv")
    deep_path = ROOT / "analysis" / "nsc_deep_baselines_group10fold_20260519" / "deep_baseline_summary.csv"
    picked = []
    keep_ablation = {
        "2d_et8_change_grid": "Current best hybrid raw+2D",
        "raw_initial": "1D raw statistical/tree branch",
        "raw_multiwindow": "1D raw multi-window branch",
        "2d_var_et8": "2D patch ExtraTrees branch",
        "raw_initial_plus_multiscale_grid": "2D multi-scale fusion proxy",
        "stacking_only_change": "Score-level calibrated stacking proxy",
    }
    for row in ablation:
        if row["method"] in keep_ablation:
            picked.append(
                {
                    "method": row["method"],
                    "role": keep_ablation[row["method"]],
                    "validation": "本案 patient-aware group 10-fold",
                    "AUROC": row["AUROC"],
                    "AUPRC": row["AUPRC"],
                    "status": "已實測",
                }
            )
    keep_paper = {
        "green_pca_prototype": "PixelHop/Green learning proxy",
        "relation_logistic": "Relation Network shallow proxy",
        "smote_patch_logistic": "Training-fold-only feature augmentation proxy",
        "patch_prototype": "Prototype classifier baseline",
    }
    for row in paper_methods:
        if row["method"] in keep_paper:
            picked.append(
                {
                    "method": row["method"],
                    "role": keep_paper[row["method"]],
                    "validation": "本案 patient-aware 10-fold proxy",
                    "AUROC": row["AUROC"],
                    "AUPRC": row["AUPRC"],
                    "status": "已實測",
                }
            )
    if deep_path.exists():
        keep_deep = {
            "2d_cnn_pp_ar_rp_gaf_maps": "Literature 2D CNN baseline",
            "1d_2d_cnn_score_ensemble": "Literature 1D+2D CNN score-fusion baseline",
            "1d_cnn_raw_sequence": "Literature 1D CNN raw-sequence baseline",
        }
        for row in read_csv_rows(deep_path):
            if row["method"] in keep_deep:
                picked.append(
                    {
                        "method": row["method"],
                        "role": keep_deep[row["method"]],
                        "validation": "本案 patient-aware group 10-fold deep baseline",
                        "AUROC": row["AUROC"],
                        "AUPRC": row["AUPRC"],
                        "status": "已實測",
                    }
                )
    return sorted(picked, key=lambda r: float(r["AUPRC"]), reverse=True)


def literature_rows() -> List[dict]:
    return [
        {
            "literature_model": "Time-series imaging + CNN / tiled CNN",
            "reference": "Wang and Oates, IJCAI 2015, Imaging Time-Series to Improve Classification and Imputation",
            "input_representation": "GAF/MTF/RP/time-series images",
            "paper_position": "核心文獻：支持將 1D physiological signals 轉為 2D images。",
            "our_mapping": "本案已用 PP/AR/RP/GAF 類 2D 圖與 patch-level features，並補跑 lightweight 2D CNN baseline。",
            "comparison_status": "列入文獻深度模型；本案以 2D ET8/Green proxy 與 2D CNN baseline 實測。",
            "paper_claim": "可作理論依據，不可直接拿文獻分數與本案 patient-aware 10-fold 比。",
        },
        {
            "literature_model": "RP + deep CNN",
            "reference": "Hatami, Gavet and Debayle, 2017, Classification of Time-Series Images Using Deep CNNs",
            "input_representation": "Recurrence plot image",
            "paper_position": "直接對應本案 2D RP/紋理圖。",
            "our_mapping": "本案使用 patch/Green/ET 對局部紋理做小樣本 proxy，並補跑 PP/AR/RP/GAF aggregated 2D CNN baseline。",
            "comparison_status": "已補跑 lightweight 2D CNN；尚非逐篇完全復刻 RP-CNN 架構。",
            "paper_claim": "可作為深度 2D image baseline 對照；本案小樣本下未優於主方法。",
        },
        {
            "literature_model": "GAF + Deep CNN",
            "reference": "Elmir et al., 2023, ECG classification using Deep CNN and Gramian Angular Field",
            "input_representation": "Gramian Angular Field image",
            "paper_position": "支持 GAF 類 2D physiological image representation。",
            "our_mapping": "本案有 GAF 圖，並納入 PP/AR/RP/GAF aggregated 2D CNN baseline；主方法仍為 2D patch ET/fusion。",
            "comparison_status": "已補跑 lightweight CNN baseline；可作後續 GAF-only CNN 復刻。",
            "paper_claim": "支持方法動機；本案 CNN baseline 不宣稱達到該文結果。",
        },
        {
            "literature_model": "1D CNN vs 2D CNN / transfer-initialized 2D CNN",
            "reference": "Wu et al., 2018, A Comparison of 1-D and 2-D Deep CNNs in ECG Classification",
            "input_representation": "raw ECG vs ECG image",
            "paper_position": "可用於論述 1D raw 與 2D image 應公平比較。",
            "our_mapping": "本案已比較 raw-only、2D-only、raw+2D fusion，並補跑 1D CNN、2D CNN 與 1D+2D CNN ensemble。",
            "comparison_status": "已列入 paper 方法比較架構與實測 baseline。",
            "paper_claim": "支持本案 report 2D-only/1D-only/fusion 三組，並呈現 CNN baseline 未優於主方法。",
        },
        {
            "literature_model": "Mix Time-Series Imaging + neural feature fusion",
            "reference": "Cai et al., 2022, Electrocardiogram Signal Classification Based on Mix Time-Series Imaging",
            "input_representation": "GAF + RP + tiling multi-channel image",
            "paper_position": "支持多種 2D representation 融合。",
            "our_mapping": "本案 multiscale/multi-plot 2D 嘗試未提升；需改成 error-guided feature subset。",
            "comparison_status": "列入文獻深度模型；本案 multi-scale proxy 已實測但未優於主線。",
            "paper_claim": "可說明為何嘗試 multi-scale/multi-plot，但本案小樣本下高維融合不穩。",
        },
        {
            "literature_model": "1D CNN and 2D CNN comparison",
            "reference": "Nakano and Sugiyama, 2022, Discriminating seismic events using 1D and 2D CNNs",
            "input_representation": "raw waveform and 2D representation",
            "paper_position": "跨領域支持 1D/2D CNN 互補比較。",
            "our_mapping": "本案 raw + 2D late fusion 是相同問題設定的 non-deep small-data version。",
            "comparison_status": "列入背景與討論。",
            "paper_claim": "支持融合設計，但不同領域結果不可直接比較。",
        },
        {
            "literature_model": "1D-CNN + 2D-CNN score-level ensemble",
            "reference": "Noman et al., 2018, Short-segment heart sound classification using an ensemble of deep CNNs",
            "input_representation": "raw heart sound + time-frequency feature map",
            "paper_position": "與本案 raw+2D score-level fusion 最接近。",
            "our_mapping": "本案已實測 score-level fusion，並補跑 1D+2D CNN score-level ensemble；目前最佳仍為 raw_initial + 2D ET8 grid fusion。",
            "comparison_status": "列為最重要深度對照方法，且已有 lightweight deep ensemble baseline。",
            "paper_claim": "可作為本案 fusion 架構的深度文獻對照；本案資料量下 non-deep fusion 較佳。",
        },
        {
            "literature_model": "Siamese / Relation Network few-shot classifier",
            "reference": "Few-shot metric/relation learning literature",
            "input_representation": "pairs/prototypes/relation features",
            "paper_position": "小樣本分類常見深度 baseline。",
            "our_mapping": "本案已做 relation_logistic proxy，AUPRC 0.605，未達主線。",
            "comparison_status": "列入補充比較；待新增病例後再跑 neural version。",
            "paper_claim": "目前不作主模型，避免小樣本過擬合。",
        },
    ]


def make_report(current_rows: List[dict], lit_rows: List[dict]) -> None:
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("NSC Paper 文獻深度模型比較補充報告", level=0)
    doc.add_paragraph(f"產出日期：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("一、結論", level=1)
    doc.add_paragraph(
        "Paper 需要納入文獻深度模型比較，但必須和本案 patient-aware group 10-fold 實測結果分開呈現。"
        "文獻 CNN/深度融合方法多在較大資料集或不同切分條件下報告，因此可作為方法對照與研究動機，"
        "不可直接把文獻分數與本案 AUROC/AUPRC 並列成同一實驗結果。"
    )
    doc.add_paragraph(
        "本案已在 torch conda environment 補跑 1D CNN、2D CNN 與 1D+2D CNN score-level ensemble；"
        "同時也完成 patient-aware 10-fold 下的 1D raw、2D patch/Green/ExtraTrees、Relation proxy、SMOTE proxy 與 score-level fusion。"
    )

    doc.add_heading("二、本案已實測方法", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["方法", "角色", "驗證", "AUROC", "AUPRC", "狀態"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in current_rows:
        vals = [r["method"], r["role"], r["validation"], f"{float(r['AUROC']):.3f}", f"{float(r['AUPRC']):.3f}", r["status"]]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("三、Paper 應列入之文獻深度模型", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["文獻深度模型", "參考文獻", "輸入形式", "Paper 定位", "本案對應", "比較狀態", "可寫宣稱"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for r in lit_rows:
        vals = [
            r["literature_model"],
            r["reference"],
            r["input_representation"],
            r["paper_position"],
            r["our_mapping"],
            r["comparison_status"],
            r["paper_claim"],
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)

    doc.add_heading("四、建議寫入 Paper 的比較架構", level=1)
    doc.add_paragraph("建議在 Methods/Experiments 加入以下比較層級：")
    for item in [
        "Classical small-data baselines：patch prototype、logistic、ExtraTrees。",
        "Green / PixelHop-style proxy：Green/PCA prototype 或 Green/PCA logistic。",
        "Metric-learning proxy：relation_logistic；端到端 Relation/Siamese Network 列為 deep baseline/future replication。",
        "Literature deep image models：GAF/RP/MTF + CNN、1D CNN vs 2D CNN、1D+2D CNN ensemble。",
        "Proposed small-data method：raw_initial + 2D ET8 + score-level grid fusion。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("五、建議正式論述", level=1)
    doc.add_paragraph(
        "本研究參考 time-series imaging 與深度 CNN 文獻，將生理時序資料轉換為 PP/AR/RP/GAF 等 2D 特徵圖，"
        "並比較 1D raw、2D patch 與 late-fusion 表徵。然而在本案 114 位個案的小樣本條件下，"
        "端到端深度模型需要更多個案與更穩定的 nested validation；因此本研究主方法採用可解釋、資料效率較高的 "
        "patch-level feature selection、tree ensemble 與 score-level fusion。文獻深度模型列為對照與後續復刻方向。"
    )

    doc.add_heading("六、已下載文獻 PDF", level=1)
    refs = sorted((ROOT / "references" / "nsc_2d_vs_1d_timeseries_20260519").glob("*.pdf"))
    for ref in refs:
        doc.add_paragraph(str(ref), style="List Bullet")

    doc.save(REPORT)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    current = current_method_rows()
    lit = literature_rows()
    write_csv(
        OUT_DIR / "current_patient_aware_results_for_paper.csv",
        current,
        ["method", "role", "validation", "AUROC", "AUPRC", "status"],
    )
    write_csv(
        OUT_DIR / "literature_deep_model_comparison_for_paper.csv",
        lit,
        ["literature_model", "reference", "input_representation", "paper_position", "our_mapping", "comparison_status", "paper_claim"],
    )
    make_report(current, lit)
    print(
        json_like(
            {
                "report": str(REPORT),
                "analysis_dir": str(OUT_DIR),
                "current_results": str(OUT_DIR / "current_patient_aware_results_for_paper.csv"),
                "literature_table": str(OUT_DIR / "literature_deep_model_comparison_for_paper.csv"),
            }
        )
    )


def json_like(obj: dict) -> str:
    import json

    return json.dumps(obj, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
