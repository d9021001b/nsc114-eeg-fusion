#!/usr/bin/env python3
"""Nested meta-stacking for NSC raw + multi-instance 2D scores.

This experiment addresses the main error in simple multi-instance fusion:
max/top-k aggregation increases sensitivity but also creates false positives.
Here each outer training fold builds inner OOF raw scores and several 2D
multi-instance aggregate scores, then trains a low-dimensional meta model to
learn when local visual evidence should override or only slightly adjust raw.
"""

from __future__ import annotations

import argparse
import csv
import json
import sys
import warnings
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier
from sklearn.feature_selection import SelectKBest, VarianceThreshold, f_classif
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import confusion_matrix, precision_recall_curve, roc_curve
from sklearn.model_selection import StratifiedGroupKFold
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from nsc_multi_instance_2d_raw_fusion import collect_image_instances, instances_to_matrix  # noqa: E402
from raw_timeseries_dtw_tsmote_nsc_dataset_images import load_case_sequences  # noqa: E402
from score_level_fusion_2d_raw_nsc import fit_raw_dtw_et, metric_dict, write_csv  # noqa: E402

warnings.filterwarnings("ignore", message=r"Features .* are constant.")
warnings.filterwarnings("ignore", category=RuntimeWarning)

AGGS = ("mean", "max", "p75", "top3mean", "mean_max", "std", "hi50_frac", "hi70_frac", "max_minus_mean", "p75_minus_mean")


def aggregate_matrix(subject_order: List[str], inst_subjects: np.ndarray, inst_scores: np.ndarray) -> np.ndarray:
    buckets: Dict[str, List[float]] = defaultdict(list)
    for subject, score in zip(inst_subjects, inst_scores):
        buckets[str(subject)].append(float(score))
    rows = []
    for subject in subject_order:
        vals = np.asarray(buckets[str(subject)], dtype=float)
        if len(vals) == 0:
            vals = np.asarray([0.5], dtype=float)
        vals_sort = np.sort(vals)
        mean = float(vals.mean())
        maxv = float(vals.max())
        p75 = float(np.percentile(vals, 75))
        top3 = float(vals_sort[-min(3, len(vals_sort)) :].mean())
        row = [
            mean,
            maxv,
            p75,
            top3,
            0.5 * mean + 0.5 * maxv,
            float(vals.std()),
            float((vals >= 0.5).mean()),
            float((vals >= 0.7).mean()),
            maxv - mean,
            p75 - mean,
        ]
        rows.append(row)
    return np.asarray(rows, dtype=float)


def fit_2d_instance_aggregates(
    Xinst: np.ndarray,
    yinst: np.ndarray,
    inst_subjects: np.ndarray,
    subject_order: List[str],
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    seed: int,
    top_k: int,
    model_name: str,
) -> Tuple[np.ndarray, np.ndarray, dict]:
    train_subjects = set(str(subject_order[i]) for i in train_idx)
    test_subjects = set(str(subject_order[i]) for i in test_idx)
    inst_train_mask = np.asarray([str(s) in train_subjects for s in inst_subjects], dtype=bool)
    inst_test_mask = np.asarray([str(s) in test_subjects for s in inst_subjects], dtype=bool)
    Xtr = Xinst[inst_train_mask]
    ytr = yinst[inst_train_mask]
    Xte = Xinst[inst_test_mask]
    train_inst_subjects = inst_subjects[inst_train_mask]
    test_inst_subjects = inst_subjects[inst_test_mask]

    imputer = SimpleImputer(strategy="median")
    Xtr_i = imputer.fit_transform(Xtr)
    Xte_i = imputer.transform(Xte)
    var = VarianceThreshold(threshold=1e-8)
    Xtr_v = var.fit_transform(Xtr_i)
    Xte_v = var.transform(Xte_i)
    if Xtr_v.shape[1] == 0:
        Xtr_v, Xte_v = Xtr_i, Xte_i
    k = min(top_k, Xtr_v.shape[1])
    selector = SelectKBest(f_classif, k=k)
    Xtr_s = selector.fit_transform(Xtr_v, ytr)
    Xte_s = selector.transform(Xte_v)

    if model_name == "ExtraTrees":
        model = ExtraTreesClassifier(
            n_estimators=700,
            random_state=seed,
            class_weight="balanced",
            max_features="sqrt",
            min_samples_leaf=2,
            n_jobs=-1,
        )
    elif model_name == "HistGB":
        model = HistGradientBoostingClassifier(max_iter=180, learning_rate=0.035, max_leaf_nodes=15, l2_regularization=0.08, random_state=seed)
    else:
        raise ValueError(model_name)
    counts = Counter(str(s) for s in train_inst_subjects)
    sample_weight = np.asarray([1.0 / counts[str(s)] for s in train_inst_subjects], dtype=float)
    try:
        model.fit(Xtr_s, ytr, sample_weight=sample_weight)
    except TypeError:
        model.fit(Xtr_s, ytr)

    train_inst_score = model.predict_proba(Xtr_s)[:, 1]
    test_inst_score = model.predict_proba(Xte_s)[:, 1]
    train_order = [subject_order[i] for i in train_idx]
    test_order = [subject_order[i] for i in test_idx]
    stats = {
        "train_instances": int(Xtr.shape[0]),
        "test_instances": int(Xte.shape[0]),
        "input_features": int(Xinst.shape[1]),
        "after_variance_features": int(Xtr_v.shape[1]),
        "selected_features": int(k),
    }
    return aggregate_matrix(train_order, train_inst_subjects, train_inst_score), aggregate_matrix(test_order, test_inst_subjects, test_inst_score), stats


def meta_features(raw: np.ndarray, two: np.ndarray) -> np.ndarray:
    raw = np.asarray(raw, dtype=float).reshape(-1, 1)
    cols = [raw, two, raw * two[:, :1], np.abs(raw - two[:, :1])]
    # Contrast features help distinguish "one local image is high" from
    # broadly high 2D evidence.
    cols.append((two[:, 1:2] - two[:, 0:1]))
    cols.append((two[:, 2:3] - two[:, 0:1]))
    return np.column_stack(cols)


def fit_meta_candidates(X_meta: np.ndarray, y: np.ndarray) -> Tuple[object, dict, List[dict]]:
    candidates = []
    for c_value in [0.05, 0.1, 0.25, 0.5, 1.0]:
        for class_weight in [None, "balanced"]:
            model = make_pipeline(StandardScaler(), LogisticRegression(max_iter=4000, C=c_value, class_weight=class_weight, solver="liblinear"))
            model.fit(X_meta, y)
            score = model.predict_proba(X_meta)[:, 1]
            row = {"meta_model": "LogReg", "C": c_value, "class_weight": "none" if class_weight is None else "balanced", **metric_dict(y, score)}
            candidates.append((model, row))
    for max_leaf_nodes in [7, 11, 15]:
        model = HistGradientBoostingClassifier(max_iter=80, learning_rate=0.04, max_leaf_nodes=max_leaf_nodes, l2_regularization=0.2, random_state=20260519)
        model.fit(X_meta, y)
        score = model.predict_proba(X_meta)[:, 1]
        row = {"meta_model": "HistGB", "C": "", "class_weight": "", "max_leaf_nodes": max_leaf_nodes, **metric_dict(y, score)}
        candidates.append((model, row))
    best_model, best_row = max(candidates, key=lambda item: (item[1]["AUPRC"], item[1]["AUROC"]))
    return best_model, best_row, [row for _m, row in candidates]


def make_figures(out_dir: Path, rows: List[dict], y: np.ndarray, best_score: np.ndarray, best: dict) -> Dict[str, str]:
    fig_dir = out_dir / "figures"
    fig_dir.mkdir(parents=True, exist_ok=True)
    paths: Dict[str, str] = {}
    ypos = np.arange(len(rows))
    labels = [f"{r['instance_model']} -> {r['meta_model']}" for r in rows]
    plt.figure(figsize=(9.5, max(3.6, 0.45 * len(rows) + 1)))
    plt.barh(ypos + 0.18, [r["AUROC"] for r in rows], height=0.34, label="AUROC")
    plt.barh(ypos - 0.18, [r["AUPRC"] for r in rows], height=0.34, label="AUPRC")
    plt.axvline(0.8, linestyle="--", color="gray")
    plt.yticks(ypos, labels, fontsize=8)
    plt.xlabel("Metric")
    plt.title("Raw + multi-instance 2D meta-stacking")
    plt.grid(axis="x", alpha=0.25)
    plt.legend(loc="lower right")
    plt.gca().invert_yaxis()
    path = fig_dir / "meta_stack_metric_bar.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["bar"] = str(path)

    fpr, tpr, _ = roc_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(fpr, tpr, label=f"AUROC={best['AUROC']:.3f}")
    plt.plot([0, 1], [0, 1], "--", color="gray")
    plt.xlabel("False Positive Rate")
    plt.ylabel("True Positive Rate")
    plt.title("Best meta-stacking ROC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower right")
    path = fig_dir / "best_meta_stack_roc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["roc"] = str(path)

    precision, recall, _ = precision_recall_curve(y, best_score)
    plt.figure(figsize=(5.6, 4.4))
    plt.plot(recall, precision, label=f"AUPRC={best['AUPRC']:.3f}")
    plt.axhline(float(y.mean()), linestyle="--", color="gray", label=f"prevalence={y.mean():.3f}")
    plt.xlabel("Recall")
    plt.ylabel("Precision")
    plt.title("Best meta-stacking PRC")
    plt.grid(alpha=0.25)
    plt.legend(loc="lower left")
    path = fig_dir / "best_meta_stack_prc.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["prc"] = str(path)

    cm = confusion_matrix(y, (best_score >= 0.5).astype(int), labels=[0, 1])
    plt.figure(figsize=(5.1, 4.3))
    plt.imshow(cm, cmap="Blues")
    plt.title("Best meta-stacking confusion matrix")
    plt.xticks([0, 1], ["Pred 0", "Pred 1"])
    plt.yticks([0, 1], ["True 0", "True 1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha="center", va="center", fontsize=14)
    path = fig_dir / "best_meta_stack_cm.png"
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    paths["cm"] = str(path)
    return paths


def make_report(report_path: Path, manifest: dict, rows: List[dict], figures: Dict[str, str]) -> None:
    doc = Document()
    doc.add_heading("NSC Raw + 2D Multi-instance Meta-stacking 報告", level=0)
    doc.add_paragraph(f"產出日期：{manifest['created_at']}")
    doc.add_heading("一、目的", level=1)
    doc.add_paragraph(
        "前一版 multi-instance 2D 可提升 sensitivity，但 false positives 增加。本版把 raw score 與多種 2D 聚合分數"
        "交給低維 meta model，讓模型學會何時相信局部高分圖、何時把它視為假陽性風險。"
    )
    doc.add_heading("二、結果", level=1)
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = ["2D model", "meta", "AUROC", "AUPRC", "ACC", "Sens", "Spec", "PPV", "NPV", "CM"]
    for c, h in zip(table.rows[0].cells, headers):
        c.text = h
    for row in rows:
        vals = [
            row["instance_model"],
            row["meta_model"],
            f"{row['AUROC']:.3f}",
            f"{row['AUPRC']:.3f}",
            f"{row['accuracy']:.3f}",
            f"{row['sensitivity']:.3f}",
            f"{row['specificity']:.3f}",
            f"{row['PPV']:.3f}",
            f"{row['NPV']:.3f}",
            f"{row['TN']}/{row['FP']}/{row['FN']}/{row['TP']}",
        ]
        cells = table.add_row().cells
        for c, v in zip(cells, vals):
            c.text = str(v)
    doc.add_heading("三、圖表", level=1)
    for key in ["bar", "roc", "prc", "cm"]:
        if key in figures:
            doc.add_picture(figures[key], width=Inches(6.1 if key == "bar" else 5.1))
    doc.add_heading("四、判讀", level=1)
    best = manifest["best"]
    doc.add_paragraph(f"最佳結果為 {best['instance_model']} -> {best['meta_model']}，AUROC={best['AUROC']:.3f}，AUPRC={best['AUPRC']:.3f}。")
    doc.add_paragraph(
        "若未達 0.8/0.8，表示單靠目前 raw + 2D local scores 仍不足。下一步應增加 label-informed 但 fold-local 的 patch-family selection，"
        "或改用 representative holdout 做 proof-of-signal，並明確標示不是最終泛化宣稱。"
    )
    doc.add_heading("五、輸出", level=1)
    for k, v in manifest["outputs"].items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
    report_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(report_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv-dir", default="bio-addict/csv-data")
    parser.add_argument("--manifest", default="nsc_dataset_images/manifest.csv")
    parser.add_argument("--image-root", default="nsc_dataset_images")
    parser.add_argument("--out-dir", default="analysis/nsc_multi_instance_meta_stack_20260519")
    parser.add_argument("--report", default="reports/NSC_multi_instance_meta_stack_20260519.docx")
    parser.add_argument("--sequence-length", type=int, default=256)
    parser.add_argument("--max-rows-per-csv", type=int, default=20000)
    parser.add_argument("--grid", type=int, default=8)
    parser.add_argument("--top-k", type=int, default=192)
    parser.add_argument("--n-splits", type=int, default=10)
    parser.add_argument("--inner-splits", type=int, default=5)
    parser.add_argument("--random-state", type=int, default=20260519)
    parser.add_argument("--instance-models", default="ExtraTrees,HistGB")
    args = parser.parse_args()

    out_dir = Path(args.out_dir).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    subjects, Xseq, y, _meta = load_case_sequences(Path(args.csv_dir).resolve(), Path(args.manifest).resolve(), args.sequence_length, args.max_rows_per_csv)
    instances, names = collect_image_instances(Path(args.image_root).resolve(), args.grid)
    Xinst, yinst, inst_subjects = instances_to_matrix(instances, names)
    keep = np.asarray([str(s) in set(subjects) for s in inst_subjects], dtype=bool)
    Xinst, yinst, inst_subjects = Xinst[keep], yinst[keep], inst_subjects[keep]

    groups = np.asarray(subjects)
    outer = StratifiedGroupKFold(n_splits=args.n_splits, shuffle=True, random_state=42)
    instance_models = [m.strip() for m in args.instance_models.split(",") if m.strip()]
    rows: List[dict] = []
    fold_rows: List[dict] = []
    meta_search_rows: List[dict] = []
    result_scores: Dict[str, np.ndarray] = {}

    for instance_model in instance_models:
        oof_score = np.zeros(len(y), dtype=float)
        for fold, (train_idx, test_idx) in enumerate(outer.split(np.zeros(len(y)), y, groups=groups), start=1):
            y_train = y[train_idx]
            groups_train = groups[train_idx]
            raw_oof = np.zeros(len(train_idx), dtype=float)
            two_oof = np.zeros((len(train_idx), len(AGGS)), dtype=float)
            inner = StratifiedGroupKFold(n_splits=args.inner_splits, shuffle=True, random_state=args.random_state + fold * 113)
            for inner_fold, (tr_local, va_local) in enumerate(inner.split(np.zeros(len(train_idx)), y_train, groups=groups_train), start=1):
                tr_idx = train_idx[tr_local]
                va_idx = train_idx[va_local]
                _raw_tr, raw_va, _n_syn = fit_raw_dtw_et(Xseq, y, tr_idx, va_idx, args.random_state + fold * 1000 + inner_fold)
                _two_tr, two_va, _st = fit_2d_instance_aggregates(
                    Xinst, yinst, inst_subjects, subjects, tr_idx, va_idx, args.random_state + fold * 2000 + inner_fold, args.top_k, instance_model
                )
                raw_oof[va_local] = raw_va
                two_oof[va_local, :] = two_va
            _raw_tr, raw_test, n_syn = fit_raw_dtw_et(Xseq, y, train_idx, test_idx, args.random_state + fold)
            _two_tr, two_test, st = fit_2d_instance_aggregates(
                Xinst, yinst, inst_subjects, subjects, train_idx, test_idx, args.random_state + fold + 500, args.top_k, instance_model
            )
            X_meta_train = meta_features(raw_oof, two_oof)
            X_meta_test = meta_features(raw_test, two_test)
            meta_model, best_meta, candidates = fit_meta_candidates(X_meta_train, y_train)
            for cand in candidates:
                meta_search_rows.append({"instance_model": instance_model, "fold": fold, **cand})
            oof_score[test_idx] = meta_model.predict_proba(X_meta_test)[:, 1]
            fold_rows.append(
                {
                    "instance_model": instance_model,
                    "fold": fold,
                    "meta_model": best_meta["meta_model"],
                    "meta_C": best_meta.get("C", ""),
                    "meta_class_weight": best_meta.get("class_weight", ""),
                    "raw_synthetic_samples": n_syn,
                    **st,
                    "test_cases": int(len(test_idx)),
                }
            )
        met = metric_dict(y, oof_score)
        row = {"instance_model": instance_model, "meta_model": "selected_per_fold", **met}
        rows.append(row)
        result_scores[instance_model] = oof_score

    rows = sorted(rows, key=lambda r: (r["AUPRC"], r["AUROC"]), reverse=True)
    write_csv(out_dir / "meta_stack_summary.csv", rows, ["instance_model", "meta_model", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"])
    write_csv(out_dir / "fold_details.csv", fold_rows, ["instance_model", "fold", "meta_model", "meta_C", "meta_class_weight", "raw_synthetic_samples", "train_instances", "test_instances", "input_features", "after_variance_features", "selected_features", "test_cases"])
    write_csv(out_dir / "meta_search.csv", meta_search_rows, ["instance_model", "fold", "meta_model", "C", "class_weight", "max_leaf_nodes", "AUROC", "AUPRC", "accuracy", "sensitivity", "specificity", "PPV", "NPV", "TN", "FP", "FN", "TP"])
    pred_rows = []
    for i, subject in enumerate(subjects):
        row = {"subject_id": subject, "true_label": int(y[i])}
        for key, score in result_scores.items():
            row[key] = float(score[i])
        pred_rows.append(row)
    write_csv(out_dir / "meta_stack_predictions.csv", pred_rows, ["subject_id", "true_label"] + sorted(result_scores))

    best = rows[0]
    figures = make_figures(out_dir, rows, y, result_scores[best["instance_model"]], best)
    manifest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "validation": f"patient-aware StratifiedGroupKFold {args.n_splits}-fold with inner {args.inner_splits}-fold OOF meta-stacking",
        "subjects": len(subjects),
        "cases_by_label": {str(k): int(v) for k, v in zip(*np.unique(y, return_counts=True))},
        "image_instances": int(len(instances)),
        "usable_image_instances": int(Xinst.shape[0]),
        "image_feature_count": int(Xinst.shape[1]),
        "aggregate_features": list(AGGS),
        "best": best,
        "target_reached": bool(best["AUROC"] >= 0.8 and best["AUPRC"] >= 0.8),
        "outputs": {
            "analysis_dir": str(out_dir),
            "summary": str(out_dir / "meta_stack_summary.csv"),
            "predictions": str(out_dir / "meta_stack_predictions.csv"),
            "fold_details": str(out_dir / "fold_details.csv"),
            "meta_search": str(out_dir / "meta_search.csv"),
            "report": str(Path(args.report).resolve()),
        },
        "figures": figures,
    }
    with (out_dir / "manifest.json").open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)
    make_report(Path(args.report).resolve(), manifest, rows, figures)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
